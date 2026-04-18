"""Generate Day 3 materials: Do Now + Student Packet + Teacher Packet.

Day 3 = Multiplicity. The payoff day. On Day 2 students discovered that
g(x) = x(x - 4)^2 TOUCHES the x-axis at x = 4 instead of crossing. Today we
name that phenomenon (multiplicity) and build a rule: odd multiplicity =
cross, even multiplicity = touch and turn.

Budget: 55 minutes usable (short F-period on Tuesday). 65 min on Monday A.
Built for the 55-min version; slack goes to Explore phase.

Phases (total 55 min):
  0-5     Do Now    : Solo paper - predict touch/cross for 3 factored forms [DOK 2]
  5-7     Do Now    : Blooket login                                         [-]
  7-14    Do Now    : Blooket game (new multiplicity CSV)                   [DOK 1]
  14-22   Launch    : Name MULTIPLICITY, Even vs Odd T-chart, Frayer        [DOK 2]
  22-37   Explore   : Desmos compare (x-1)^n for n=1,2,3,4 + Try It         [DOK 2-3]
  37-47   Explore   : Tonya error analysis (DOK 3 from revised packet)      [DOK 3]
  47-51   Share/Sum : Synthesis + self-reflection                           [DOK 2]
  51-55   Exit      : CER - label behavior at each zero                     [DOK 2]
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL

TABLE_STYLE = "Table Grid"

OBJECTIVES = [
    ("Math Objective",
     "Determine the multiplicity of each zero of a polynomial in factored "
     "form, and use multiplicity to predict whether the graph crosses, "
     "touches, or flattens through the x-axis at each zero."),
    ("Language Objective",
     "Students will use the frame \u201cThe multiplicity is ___, so the graph "
     "___ at x = ___\u201d to justify graph behavior."),
    ("Essential Understanding",
     "A repeated factor creates a zero of higher multiplicity. Even "
     "multiplicity causes the graph to touch and turn; odd multiplicity "
     "causes it to cross."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Students formalize yesterday\u2019s discovery. They read a factored "
     "polynomial, count the multiplicity of each zero, and predict graph "
     "behavior before verifying in Desmos."),
    ("Essential Question",
     "What does a repeated factor do to the graph of a polynomial, and how "
     "can you predict the behavior at every zero just by looking at the "
     "factored form?"),
    ("Materials",
     "Student laptops with Desmos; Day 2 exit ticket (returned); Blooket "
     "access; pencils. No chart paper."),
]


def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if i == 0 and bold_first:
            run.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_two_col(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = TABLE_STYLE
    t.columns[0].width = Inches(1.8)
    t.columns[1].width = Inches(4.7)
    for i, (label, content) in enumerate(rows):
        set_cell(t.rows[i].cells[0], label, bold_first=True)
        if isinstance(content, str):
            content = [content]
        set_cell(t.rows[i].cells[1], content)
    doc.add_paragraph()


def add_callout(doc, title, lines):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.add_run(title).bold = True
    for line in lines:
        cell.add_paragraph(line)
    doc.add_paragraph()


def header_line(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


# ==================================================================
# DO NOW SHEET
# ==================================================================

def build_do_now(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    header_line(doc, "Day 3  |  DO NOW   \u2014   Touch or Cross?", size=13)

    p = doc.add_paragraph()
    r = p.add_run("Name: _____________________________________        Date: _____________")
    r.font.size = Pt(11)
    doc.add_paragraph()

    add_callout(doc, "\u270d\ufe0f  INSTRUCTIONS", [
        "Work SILENTLY. Pencil only. Do NOT open Desmos.",
        "Remember g(x) = x(x \u2212 4)\u00b2 yesterday? The graph TOUCHED the "
        "x-axis at x = 4 instead of crossing. Today you\u2019ll predict which "
        "zeros cross and which touch \u2014 just from the factored form.",
        "Time: 5 minutes.",
    ])

    add_callout(doc, "\U0001F52E  For EACH zero, predict: CROSS or TOUCH?", [
        "(You don\u2019t need a rule yet \u2014 use your eyes and your gut. "
        "Look at how many times each factor appears.)",
    ])

    add_two_col(doc, [
        ("1.  f(x) = (x + 2)(x \u2212 3)", [
            "At x = \u22122 the graph will: _________________",
            "At x = 3 the graph will: _________________",
        ]),
        ("2.  h(x) = (x \u2212 1)\u00b2 (x + 5)", [
            "At x = 1 the graph will: _________________",
            "At x = \u22125 the graph will: _________________",
        ]),
        ("3.  k(x) = x\u00b3 (x \u2212 2)\u00b2", [
            "At x = 0 the graph will: _________________",
            "At x = 2 the graph will: _________________",
        ]),
        ("4.  Pattern guess:", [
            "When the factor appears an EVEN number of times, the graph",
            "__________________________________________ at that zero.",
            "When the factor appears an ODD number of times, the graph",
            "__________________________________________ at that zero.",
        ]),
    ])

    add_callout(doc, "\u2705  HOW YOU WILL KNOW YOU\u2019RE DONE", [
        "\u2022  All six predictions written in words.",
        "\u2022  You tried to write a pattern in #4 \u2014 even a rough guess.",
        "\u2022  You have NOT opened Desmos or the packet yet.",
    ])

    doc.save(path)


# ==================================================================
# STUDENT PACKET
# ==================================================================

LAUNCH_STUDENT = [
    "Yesterday you saw g(x) = x(x \u2212 4)\u00b2 touch the x-axis at x = 4.",
    "Today we name it. The word is MULTIPLICITY.",
    "",
    "Definition (copy in your own words):",
    "The MULTIPLICITY of a zero is the number of times its factor appears "
    "in the factored form of the polynomial.",
    "",
    "Example: f(x) = (x \u2212 4)\u00b2 (x + 1)",
    "\u2022  Zero at x = 4: factor (x \u2212 4) appears ____ times.  Multiplicity = ____",
    "\u2022  Zero at x = \u22121: factor (x + 1) appears ____ times. Multiplicity = ____",
    "",
    "FRAYER \u2014 MULTIPLICITY (fill in the four boxes):",
    "Definition: ________________________________________",
    "Characteristics: ____________________________________",
    "Example: ___________________________________________",
    "Non-example: _______________________________________",
]

EVEN_ODD_STUDENT = [
    "EVEN vs. ODD MULTIPLICITY \u2014 T-chart",
    "(Fill in after the Desmos investigation below.)",
    "",
    "ODD multiplicity (1, 3, 5, ...)  |  EVEN multiplicity (2, 4, 6, ...)",
    "-----------------------------------|-----------------------------------",
    "Behavior at the zero:              |  Behavior at the zero:",
    "______________________________     |  ______________________________",
    "______________________________     |  ______________________________",
    "Sketch a picture:                  |  Sketch a picture:",
    "                                   |",
    "Sentence frame:",
    "\u201cThe multiplicity is ____, so the graph ____ at x = ____.\u201d",
]

DESMOS_EXPLORE_STUDENT = [
    "Investigate in Desmos. Type each and observe x = 1.",
    "",
    "(a)  y = (x \u2212 1)\u00b9        At x = 1 the graph: ____________",
    "(b)  y = (x \u2212 1)\u00b2        At x = 1 the graph: ____________",
    "(c)  y = (x \u2212 1)\u00b3        At x = 1 the graph: ____________",
    "(d)  y = (x \u2212 1)\u2074        At x = 1 the graph: ____________",
    "",
    "What pattern do you see as the multiplicity increases?",
    "_______________________________________________",
    "_______________________________________________",
    "",
    "BONUS: multiplicity 3 vs. multiplicity 1 \u2014 both cross. What\u2019s "
    "DIFFERENT about how they cross?",
    "_______________________________________________",
    "(Hint: the graph gets ____________ right at the zero.)",
]

TRY_IT_STUDENT = [
    "For each polynomial, find the zeros AND the multiplicity of each, then",
    "write what the graph does (cross / touch / flatten-through).",
    "",
    "A.  f(x) = (x + 3)(x \u2212 2)\u00b2",
    "    Zeros & multiplicities: ___________________________",
    "    Behavior at each:       ___________________________",
    "",
    "B.  g(x) = x\u00b2 (x \u2212 5)\u00b3",
    "    Zeros & multiplicities: ___________________________",
    "    Behavior at each:       ___________________________",
    "",
    "C.  h(x) = (x + 1)\u2074 (x \u2212 4)",
    "    Zeros & multiplicities: ___________________________",
    "    Behavior at each:       ___________________________",
    "",
    "D.  CHALLENGE: Write a polynomial whose graph CROSSES at x = 0, "
    "TOUCHES at x = 2, and FLATTENS THROUGH at x = \u22121.",
    "    f(x) = _________________________________",
]

ERROR_ANALYSIS_STUDENT = [
    "Tonya sketches f(x) = (x + 1)\u00b2 (x \u2212 3).",
    "She draws the graph CROSSING the x-axis at x = \u22121 and CROSSING at x = 3.",
    "",
    "(a)  What did Tonya get RIGHT?",
    "_______________________________________________",
    "_______________________________________________",
    "",
    "(b)  What did Tonya get WRONG? Use the word \u201cmultiplicity.\u201d",
    "_______________________________________________",
    "_______________________________________________",
    "",
    "(c)  Use CER format. Defend the CORRECT behavior at x = \u22121.",
    "Claim:    ________________________________________",
    "Evidence: ________________________________________",
    "           ________________________________________",
    "Reasoning:________________________________________",
    "           ________________________________________",
]

SHARE_SUMMARY_STUDENT = [
    "Synthesis:",
    "A repeated factor creates a zero of higher ____________. The parity "
    "of the multiplicity (even or odd) tells you whether the graph "
    "________ or ________ at that zero.",
    "",
    "Quick self-check \u2014 circle one for each (Criteria for Success):",
    "1.  I can state the multiplicity of each zero from factored form.   \u2713    partly    not yet",
    "2.  I can predict cross vs. touch from the multiplicity.            \u2713    partly    not yet",
    "3.  I can explain WHY even multiplicity causes a touch.             \u2713    partly    not yet",
    "4.  I can spot and correct a multiplicity error in someone else\u2019s graph.  \u2713    partly    not yet",
    "",
    "Preview: tomorrow we handle polynomials whose roots aren\u2019t all real.",
]

EXIT_STUDENT = [
    "For p(x) = (x + 2)\u00b3 (x \u2212 1)\u00b2 (x \u2212 4),",
    "",
    "(a)  List each zero and its multiplicity.",
    "_______________________________________________",
    "",
    "(b)  For EACH zero, predict the behavior (cross / touch / flatten-through).",
    "_______________________________________________",
    "_______________________________________________",
    "",
    "(c)  In ONE sentence, explain how the multiplicity told you each behavior.",
    "Use the frame: \u201cThe multiplicity is ___, so the graph ___ at x = ___.\u201d",
    "_______________________________________________",
    "_______________________________________________",
]


def build_student(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    header_line(doc, "Day 3  |  The Magic of Multiplicity", size=13)

    add_two_col(doc, OBJECTIVES)
    add_two_col(doc, FRAMEWORK_HEADER)

    add_callout(doc, "\U0001F501  RECALL FROM DAY 2", [
        "Yesterday you factored g(x) = x\u00b3 \u2212 8x\u00b2 + 16x into "
        "x(x \u2212 4)\u00b2 and saw the graph TOUCH the x-axis at x = 4 "
        "instead of crossing. We said \u201cwe\u2019ll name it tomorrow.\u201d",
        "Today is tomorrow. The word is MULTIPLICITY.",
    ])

    add_callout(doc, "\U0001F4DD  CER REMINDER", [
        "CLAIM \u2014 one sentence answer.",
        "EVIDENCE \u2014 specific numbers, factors, or graph behavior.",
        "REASONING \u2014 the WHY, using \u201cbecause\u201d or \u201csince.\u201d",
    ])

    add_callout(doc, "\U0001F4C4  DO NOW \u2014 SEPARATE SHEET", [
        "You received the Touch-or-Cross prediction sheet on entry.",
        "Hold onto it \u2014 we\u2019ll compare predictions during the Launch "
        "synthesis after the Blooket.",
    ])

    add_callout(doc, "\U0001F3AE  Blooket Warm-Up   [Do Now B+C]  [DOK 1]   (2 + 7 min)", [
        "Teacher displays the code. Log in (2 min). Game runs 7 minutes.",
        "Topics: reading multiplicity from factored form, even vs. odd parity, "
        "matching graphs to equations, and two recall questions from Day 2.",
    ])

    add_callout(doc, "\U0001F4A1  Launch \u2014 Name the Thing   [Launch]  [DOK 2]   (8 min)",
                LAUNCH_STUDENT)

    add_two_col(doc, [(
        "\U0001F50D  Explore \u2014 Desmos: (x \u2212 1)\u207f   [Explore]  [DOK 2\u20133]   (15 min)",
        DESMOS_EXPLORE_STUDENT,
    )])

    add_two_col(doc, [(
        "\U0001F4CA  Even vs. Odd T-Chart",
        EVEN_ODD_STUDENT,
    )])

    add_two_col(doc, [(
        "\u270F\uFE0F  Try It \u2014 Zeros + Multiplicities + Behavior",
        TRY_IT_STUDENT,
    )])

    add_two_col(doc, [(
        "\U0001F50E  Error Analysis \u2014 Tonya\u2019s Graph   [Explore]  [DOK 3]   (10 min)",
        ERROR_ANALYSIS_STUDENT,
    )])

    add_callout(doc,
        "\U0001F501  SHARE / SUMMARY   [Share/Summary]  [DOK 2]   (4 min)",
        SHARE_SUMMARY_STUDENT)

    add_callout(doc,
        "\U0001F4E4  EXIT TICKET \u2014 CER   [Exit Ticket]  [DOK 2]   (4 min)",
        EXIT_STUDENT)

    doc.save(path)


# ==================================================================
# TEACHER PACKET
# ==================================================================

CRITERIA_FOR_SUCCESS = [
    "Student-facing, revisited during Share/Summary:",
    "1.  I can state the multiplicity of each zero from factored form.",
    "2.  I can predict cross vs. touch from the multiplicity.",
    "3.  I can explain WHY even multiplicity produces a touch.",
    "4.  I can spot and correct a multiplicity error in another student\u2019s graph.",
    "Formative assessment evidence:",
    "\u2022  Do Now sheets (scan for pattern guesses in #4).",
    "\u2022  Blooket per-question percentages (DOK 1).",
    "\u2022  Try It A-C answers circulated + photographed.",
    "\u2022  Tonya error analysis written CER (DOK 3, collected).",
    "\u2022  Share/Summary self-rating.",
    "\u2022  CER exit ticket (DOK 2).",
]

DO_NOW_KEY = [
    "Expected predictions (accept any clear \u201ctouch\u201d / \u201ccross\u201d answer):",
    "1.  f(x) = (x + 2)(x \u2212 3)",
    "    \u2022  x = \u22122: CROSS  (multiplicity 1)",
    "    \u2022  x = 3:  CROSS  (multiplicity 1)",
    "2.  h(x) = (x \u2212 1)\u00b2 (x + 5)",
    "    \u2022  x = 1:  TOUCH  (multiplicity 2)",
    "    \u2022  x = \u22125: CROSS  (multiplicity 1)",
    "3.  k(x) = x\u00b3 (x \u2212 2)\u00b2",
    "    \u2022  x = 0:  CROSS (flattened)  (multiplicity 3)",
    "    \u2022  x = 2:  TOUCH  (multiplicity 2)",
    "4.  Pattern guess target:",
    "    \u2022  EVEN \u2192 touches and turns",
    "    \u2022  ODD  \u2192 crosses",
    "DIAGNOSTIC SCAN as students hand sheets in:",
    "\u2022  Students who wrote \u201ccross\u201d on every zero in #2-#3 "
    "\u2192 cold-call during Launch to surface + correct publicly.",
    "\u2022  Students who stated the even/odd rule in #4 \u2192 have them share "
    "first in Launch. Peer-voiced rule beats teacher-voiced rule.",
]

LAUNCH_KEY = [
    "Name the word: MULTIPLICITY = number of times a factor appears.",
    "f(x) = (x \u2212 4)\u00b2 (x + 1):  x = 4 \u2192 mult 2;  x = \u22121 \u2192 mult 1.",
    "",
    "FRAYER answer-bank (accept any reasonable phrasing):",
    "\u2022  Definition: how many times a factor repeats in factored form.",
    "\u2022  Characteristics: counted per zero; always a positive integer; "
    "equals the exponent on the factor.",
    "\u2022  Example: (x \u2212 5)\u00b3 has a zero of multiplicity 3 at x = 5.",
    "\u2022  Non-example: a coefficient (like the 3 in 3x) is NOT multiplicity.",
    "",
    "TEACHER MOVE \u2014 bridge from yesterday:",
    "Write g(x) = x(x \u2212 4)\u00b2 on the board. Ask: \u201cWhich zero had "
    "the weird behavior yesterday?\u201d (x = 4). \u201cHow many times is "
    "(x \u2212 4) in the factored form?\u201d (twice). \u201cThat\u2019s what "
    "we mean by multiplicity 2. And the graph touched.\u201d Now the word is "
    "attached to a shared memory, not a definition.",
]

DESMOS_EXPLORE_KEY = [
    "Target observations at x = 1:",
    "(a)  (x \u2212 1)\u00b9  \u2192 clean straight-through cross (linear slope at zero)",
    "(b)  (x \u2212 1)\u00b2  \u2192 TOUCHES (parabolic kiss, turns back)",
    "(c)  (x \u2212 1)\u00b3  \u2192 CROSSES but FLATTENS at the zero (S-curve)",
    "(d)  (x \u2212 1)\u2074  \u2192 TOUCHES and flattens harder than (x-1)\u00b2",
    "",
    "Pattern: EVEN \u2192 touch/turn; ODD \u2192 cross; higher multiplicity "
    "\u2192 flatter at the zero.",
    "",
    "BONUS (why higher odd mult flattens):",
    "Near x = 1, (x-1) is small. Cubing a small number makes it tiny. So "
    "for the graph to reach the same y-value on each side, it has to stay "
    "near the x-axis longer. The flatten-through behavior is the curve "
    "\u201csticking\u201d at the axis.",
    "",
    "TEACHER MOVE: don\u2019t volunteer the explanation for the BONUS. Let "
    "partners wrestle with it. If no one gets there, preview: \u201cCalculus "
    "will explain the flattening. For now, the PATTERN is the rule.\u201d",
]

TRY_IT_KEY = [
    "A.  f(x) = (x + 3)(x \u2212 2)\u00b2",
    "    \u2705 x = \u22123, mult 1 \u2192 CROSS",
    "    \u2705 x = 2, mult 2  \u2192 TOUCH",
    "B.  g(x) = x\u00b2 (x \u2212 5)\u00b3",
    "    \u2705 x = 0, mult 2  \u2192 TOUCH",
    "    \u2705 x = 5, mult 3  \u2192 CROSS (flattened)",
    "C.  h(x) = (x + 1)\u2074 (x \u2212 4)",
    "    \u2705 x = \u22121, mult 4 \u2192 TOUCH (very flat)",
    "    \u2705 x = 4, mult 1  \u2192 CROSS",
    "D.  CHALLENGE (answers vary; sample target):",
    "    Cross at x=0 (odd mult), touch at x=2 (even mult), flatten-through at x=\u22121 (odd mult \u2265 3).",
    "    \u2705 f(x) = x(x \u2212 2)\u00b2 (x + 1)\u00b3   (any leading coefficient).",
    "    Common student answer missing the flatten: f(x) = x(x-2)\u00b2(x+1) \u2014 the x+1 needs exponent \u2265 3 for flatten-through.",
]

ERROR_ANALYSIS_KEY = [
    "Tonya sketched f(x) = (x + 1)\u00b2 (x \u2212 3) with BOTH zeros crossing.",
    "",
    "(a)  RIGHT: x = 3 does cross (multiplicity 1, odd).",
    "(b)  WRONG: at x = \u22121 the factor (x + 1) is squared \u2192 multiplicity 2 "
    "(even) \u2192 the graph should TOUCH and TURN, not cross.",
    "",
    "SAMPLE CER:",
    "CLAIM: The graph should touch the x-axis at x = \u22121, not cross.",
    "EVIDENCE: The factor (x + 1) appears with an exponent of 2, so the "
    "zero at x = \u22121 has multiplicity 2. Multiplicity 2 is even.",
    "REASONING: Even multiplicity means the graph touches and turns at that "
    "zero because the sign of the factor doesn\u2019t change across x = \u22121 "
    "(both (x+1)\u00b2 just before and just after are positive), so the "
    "overall sign of f(x) can\u2019t flip there.",
    "",
    "Why DOK 3: students critique a peer\u2019s reasoning and justify the "
    "correct alternative with domain-specific language (multiplicity, even, "
    "sign change).",
    "ELL look-for: \u201cbecause\u201d / \u201csince\u201d in Reasoning; "
    "\u201ctouch\u201d / \u201cturn\u201d instead of just \u201cbounce.\u201d",
    "Release valve: if time tight, verbalize CER with partner now; written CER \u2192 homework.",
]

SHARE_SUMMARY_KEY = [
    "Sentence completion: A repeated factor creates a zero of higher "
    "MULTIPLICITY. Parity tells you whether the graph TOUCHES (even) or "
    "CROSSES (odd).",
    "",
    "Self-rating scan targets:",
    "\u2022  #1 state multiplicity: \u2713 for most (direct from Launch).",
    "\u2022  #2 cross vs. touch: \u2713 or partly. Strong formative signal.",
    "\u2022  #3 WHY even = touch: expect \u201cpartly.\u201d Sign-chart reasoning is hard.",
    "\u2022  #4 spot/correct error: based on Tonya task success.",
    "",
    "Teacher script (4 min):",
    "1.  Essential Question callback (45s): \u201cWhat does a repeated factor do?\u201d",
    "2.  Language Objective check (45s): \u201cWho used \u2018the multiplicity is ___, so the graph ___\u2019?\u201d",
    "3.  Self-rating circle (60s).",
    "4.  Preview Day 4 (30s): \u201cWhat if a quadratic factor won\u2019t factor over the reals? Complex zeros tomorrow.\u201d",
    "5.  Transition (30s): exit ticket next.",
]

EXIT_KEY = [
    "p(x) = (x + 2)\u00b3 (x \u2212 1)\u00b2 (x \u2212 4)",
    "",
    "(a)  Zeros and multiplicities:",
    "\u2022  x = \u22122, mult 3",
    "\u2022  x = 1,  mult 2",
    "\u2022  x = 4,  mult 1",
    "",
    "(b)  Behaviors:",
    "\u2022  x = \u22122 \u2192 CROSS (flattened, odd \u2265 3)",
    "\u2022  x = 1  \u2192 TOUCH (even)",
    "\u2022  x = 4  \u2192 CROSS (clean, odd = 1)",
    "",
    "(c)  Sample sentences:",
    "\u201cThe multiplicity is 3, so the graph crosses (with a flatten) at x = \u22122.\u201d",
    "\u201cThe multiplicity is 2, so the graph touches at x = 1.\u201d",
    "\u201cThe multiplicity is 1, so the graph crosses at x = 4.\u201d",
    "",
    "Why DOK 2: applying the rule to a NEW polynomial; not constructing a "
    "novel argument (that was error analysis).",
    "Scoring: 3 zeros \u00d7 (mult + behavior) = 6 items + frame in (c). "
    "Full credit = all three behaviors correct AND frame used.",
]

IEP_MODS = [
    "\u2022  Frayer organizer reduces blank-page anxiety; four boxes scope the task.",
    "\u2022  Do Now has predictions in CHECKBOX style \u2014 low language load.",
    "\u2022  Desmos exploration: partner with a strong-grapher student; IEP "
    "student owns recording behavior at x = 1 in writing.",
    "\u2022  Tonya error analysis: verbal CER with partner satisfies task for "
    "processing-speed accommodations; written CER can move to homework.",
    "\u2022  Exit ticket is structured per-zero \u2014 students work down a "
    "list rather than writing continuous prose.",
]

ELL_SUPPORTS = [
    "\u2022  Sentence frame: \u201cThe multiplicity is ___, so the graph ___ at x = ___.\u201d",
    "\u2022  \u201cTouches / bounces / kisses / turns\u201d all accepted for even-mult behavior.",
    "\u2022  \u201cCrosses / goes through / cuts through\u201d all accepted for odd-mult behavior.",
    "\u2022  Frayer model scaffolds the new vocabulary word.",
    "\u2022  Predict-before-verify: students commit in writing before speaking.",
    "\u2022  T-chart is visual \u2014 EVEN and ODD columns stay side-by-side as a reference.",
    "\u2022  Cognitive load concentrated on ONE new word today (multiplicity). No other new vocab.",
]


def build_teacher(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    header_line(doc, "TEACHER EDITION", size=12)
    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    header_line(doc, "Day 3  |  The Magic of Multiplicity", size=13)

    add_two_col(doc, OBJECTIVES + [("CCSS", "F-IF.C.7c, A-APR.B.3, MP.3, MP.7")])
    add_two_col(doc, FRAMEWORK_HEADER)

    add_callout(doc, "\U0001F4CB  DOK GUIDE FOR EVALUATORS", [
        "Framework: Do Now \u2192 Launch \u2192 Explore \u2192 Share/Summary \u2192 Exit Ticket. Every phase carries [Framework] + [DOK] tags.",
        "[DOK 1] Blooket warm-up [Do Now C].",
        "[DOK 2] Touch/Cross predictions [Do Now A]; Launch naming; Desmos investigation; Try It; Share/Summary; CER exit ticket.",
        "[DOK 3] Tonya error analysis [Explore]: students critique and defend with multiplicity language.",
        "DOK 3 placed with partner + teacher support. Exit Ticket intentionally DOK 2.",
    ])

    add_callout(doc, "\u2705  CRITERIA FOR SUCCESS  /  FORMATIVE ASSESSMENT", CRITERIA_FOR_SUCCESS)

    add_callout(doc, "\u26a0\ufe0f  PACING NOTE \u2014 FITS 55-MIN PERIOD", [
        "Tuesday F-period = 55 min. Monday A / Tuesday A = 65 min. Built for "
        "the 55-min version. On 65-min days, extra 10 min goes to:",
        "\u2022  Desmos exploration (deeper discussion of the flatten)",
        "\u2022  Try It D challenge (more student attempts, partner critique)",
        "\u2022  Written CER on error analysis",
        "RELEASE VALVE: if short, Tonya written CER \u2192 homework; preserve verbal partner justification.",
        "NO extra Blooket or fluff work \u2014 use slack for depth, not new tasks.",
    ])

    # Do Now A
    add_callout(doc, "[Do Now A]  [DOK 2] Touch or Cross? (Solo Paper)", [
        "Students receive the Do Now sheet on entry. Silent, 5 min, pencil only.",
        "This is DOK 2, not DOK 1 \u2014 students apply observation from "
        "yesterday to NEW factored forms and attempt a rule in #4. "
        "Cognitive demand is right for post-break Tuesday.",
    ])
    add_callout(doc, "\u2705  DO NOW / ANSWER KEY + DIAGNOSTIC", DO_NOW_KEY)

    # Blooket
    add_callout(doc, "[Do Now B+C]  [\u2014 / DOK 1] Blooket \u2014 Targeted Multiplicity", [
        "NEW CSV: Blooket_Day3_Multiplicity.csv. Import before class.",
        "B: 2-min login window.",
        "C: 7-min game. 20 questions on reading multiplicity from factored "
        "form, cross/touch identification, matching graphs, error spots, "
        "plus 2 recall questions from Day 2.",
        "Diagnostic: note lowest-percent questions. If students miss "
        "\u201cmultiplicity of (x-3)\u2074,\u201d address directly during Launch.",
    ])

    # Launch
    add_callout(doc, "[Launch]  [DOK 2] Name the Thing", [
        "Framework: Launch, 5-15 min. Purpose: introduce the word, tie to "
        "yesterday\u2019s discovery, build Frayer.",
        "Script: reference g(x) = x(x-4)\u00b2 on board \u2192 \u201c(x-4) "
        "appears twice \u2192 multiplicity 2 \u2192 graph touched.\u201d "
        "Now the word is attached to shared memory.",
    ])
    add_callout(doc, "\u2705  LAUNCH / ANSWER KEY", LAUNCH_KEY)

    # Desmos explore
    add_callout(doc, "[Explore]  [DOK 2-3] Desmos Investigation (x-1)\u207f", [
        "Framework: Explore. Students own the discovery. Teacher circulates only.",
        "Partners graph (x-1)\u00b9 through (x-1)\u2074, record behavior at x=1, "
        "then fill the Even-vs-Odd T-chart. Bonus pushes toward \u201cwhy does "
        "higher odd mult flatten?\u201d \u2014 productive struggle, no rescue.",
    ])
    add_callout(doc, "\U0001F50D  DESMOS / ANSWER KEY", DESMOS_EXPLORE_KEY)

    # Try It
    add_callout(doc, "\u270F\uFE0F  TRY IT / ANSWER KEY", TRY_IT_KEY)

    # Error analysis
    add_callout(doc, "[Explore]  [DOK 3] Tonya Error Analysis", [
        "Framework: Explore (continued) OR synthesis push before Share. "
        "10 min. This is the revised-packet DOK 3 task.",
        "Students must name Tonya\u2019s error using academic vocabulary "
        "(multiplicity, even, touch) and defend with CER.",
    ])
    add_callout(doc, "\U0001F50E  ERROR ANALYSIS / ANSWER KEY + CER sample", ERROR_ANALYSIS_KEY)

    # Share/Summary
    add_callout(doc, "[Share/Summary]  [DOK 2] Synthesis + Reflection", [
        "Framework: 5-15 min. Synthesis, return to objectives, self-rating.",
        "4-min version fits the 55-min period. Do not cut \u2014 use release valve elsewhere.",
    ])
    add_callout(doc, "\u2705  SHARE / SUMMARY / ANSWER KEY + DIAGNOSTIC", SHARE_SUMMARY_KEY)

    # Exit Ticket
    add_callout(doc, "[Exit Ticket]  [DOK 2] Full Multi-Zero Prediction", [
        "Framework: Exit Ticket shows mastery relative to objective.",
        "p(x) = (x+2)\u00b3 (x-1)\u00b2 (x-4). Three zeros, three multiplicities, "
        "three behaviors. 4 min tight; pre-printed structure reduces writing load.",
    ])
    add_callout(doc, "\u2705  EXIT TICKET / ANSWER KEY", EXIT_KEY)

    # Closing
    add_callout(doc, "\U0001F4F7  WHAT TO COLLECT", [
        "1.  Do Now sheets (scan #4 pattern guesses).",
        "2.  Exit tickets (DOK 2, every student).",
        "3.  Tonya error analysis page (DOK 3 evidence) \u2014 photo or collect.",
        "4.  Share/Summary self-ratings \u2014 quick scan.",
    ])

    add_callout(doc, "\U0001F440  LOOK-FORS DURING WALKTHROUGH", [
        "[Do Now A] Students writing predictions silently; teacher scanning on entry.",
        "[Do Now C] Blooket dashboard open; teacher reading diagnostics.",
        "[Launch] Frayer built collaboratively; bridge to yesterday made explicit.",
        "[Explore Desmos] Partners recording in writing; teacher NOT lecturing.",
        "[Explore Tonya] CER structure visible; \u201cbecause\u201d / \u201csince\u201d in reasoning.",
        "[Share/Summary] Self-ratings honest; Essential Question called back.",
        "[Exit] Sentence frame used in (c); multiplicity vocabulary present.",
    ])

    add_callout(doc, "\U0001F9E9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS", IEP_MODS)
    add_callout(doc, "\U0001F30D  SUPPORTS FOR ELL STUDENTS", ELL_SUPPORTS)

    doc.save(path)


if __name__ == "__main__":
    build_do_now("Day_3_Do_Now.docx")
    build_student("Day_3_Student_Packet.docx")
    build_teacher("Day_3_Teacher_Packet.docx")
    print("Built Day_3_Do_Now.docx, Day_3_Student_Packet.docx, Day_3_Teacher_Packet.docx")
