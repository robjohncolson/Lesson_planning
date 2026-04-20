"""Day 1 DOK-3 supplement handout for Period A Monday.

Context: Day 1 was taught without the Gallery Walk / "The Leap" closers.
Students built sign charts for a cubic and a quartic on chart paper. This
one-page handout is an optional end-of-period DOK-3 transfer task \u2014
breathable, self-contained, all info on the page.

The student should be able to pick up the sheet and work without the
teacher explaining anything new.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL

TABLE_STYLE = "Table Grid"


def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if i == 0 and bold_first:
            run.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_callout(doc, title, lines):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.add_run(title).bold = True
    for line in lines:
        cell.add_paragraph(line)
    doc.add_paragraph()


def header_line(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def blank_lines(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.add_run("______________________________________________________________")


def build(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    header_line(doc, "Day 1  |  The Leap  \u2014  Predict Without Graphing   [DOK 3]", size=13)

    p = doc.add_paragraph()
    r = p.add_run("Name: _____________________________________     Date: _____________")
    r.font.size = Pt(11)
    doc.add_paragraph()

    add_callout(doc, "\U0001F680  THE SET-UP (all info is right here)", [
        "You\u2019ve been building sign charts from FACTORED functions.  Here\u2019s a new one.",
        "You will NOT compute anything.  You will predict, using only what you already know.",
        "",
        "    f(x) = x (x \u2212 2) (x + 3) (x\u00b2 + 1)",
        "",
        "What you already know from today:",
        "  \u2022  Zero Product Property:  each factor = 0 gives one zero of f(x).",
        "  \u2022  x\u00b2 + (any positive number) is ALWAYS POSITIVE \u2014 it never equals zero "
        "and never flips the sign of f(x).",
        "  \u2022  Highest power \u2192 end behavior.  Even \u2192 same direction on both sides; "
        "odd \u2192 opposite directions.",
        "",
        "Hint:  x \u00b7 x \u00b7 x \u00b7 x\u00b2  multiplies to  x\u2075.  That tells you the end behavior "
        "without graphing.",
    ])

    header_line(doc, "1.  Real zeros.  Which values of x make f(x) = 0?", size=12)
    p = doc.add_paragraph()
    p.add_run("List them.  Also: does x\u00b2 + 1 contribute any real zero?  Why / why not?").italic = True
    blank_lines(doc, 3)
    doc.add_paragraph()

    header_line(doc, "2.  How many intervals will the sign chart have?", size=12)
    p = doc.add_paragraph()
    p.add_run("Remember: N real zeros \u2192 N + 1 intervals on the number line.").italic = True
    blank_lines(doc, 1)
    doc.add_paragraph()

    header_line(doc, "3.  End behavior without graphing.", size=12)
    p = doc.add_paragraph()
    p.add_run("The highest power of f(x) is ______ (count the x\u2019s across all factors, "
              "counting x\u00b2 as two).  So f(x) behaves like ______.").italic = True
    p = doc.add_paragraph()
    p.add_run("Far left:  f(x) goes ______.         Far right:  f(x) goes ______.")
    doc.add_paragraph()

    header_line(doc, "4.  Predict the sign chart.", size=12)
    p = doc.add_paragraph()
    p.add_run("Put the real zeros on this number line in order, then write +  or \u2212  "
              "in each interval.  You are predicting, not computing \u2014 use the pattern "
              "from today\u2019s work.").italic = True
    # A blank number line
    p = doc.add_paragraph()
    p.add_run("\u2190\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500"
              "\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500"
              "\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500"
              "\u2500\u2500\u2500\u2500\u2500\u2500\u2192").font.size = Pt(14)
    p = doc.add_paragraph()
    p.add_run("Zeros (in order): _______     _______     _______").font.size = Pt(11)
    p = doc.add_paragraph()
    p.add_run("Sign in each interval:  ______   ______   ______   ______").font.size = Pt(11)
    doc.add_paragraph()

    header_line(doc, "5.  Justify \u2014 one sentence.", size=12)
    p = doc.add_paragraph()
    p.add_run("Pick ONE thing about your sign chart that was \u201chard to figure out\u201d "
              "and explain WHY your answer is right, using evidence from f(x).  Use "
              "\u201cbecause\u201d or \u201csince\u201d.").italic = True
    blank_lines(doc, 4)
    doc.add_paragraph()

    add_callout(doc, "\U0001F50E  ANSWER KEY  (TEACHER USE \u2014 remove before copying for students)", [
        "1.  Real zeros: x = 0, x = 2, x = \u22123.  x\u00b2 + 1 contributes no real zero "
        "because x\u00b2 = \u22121 has no real solution; it\u2019s always \u2265 1 > 0.",
        "2.  4 intervals (3 real zeros \u2192 4 intervals).",
        "3.  Highest power is 5 (x \u00b7 x \u00b7 x \u00b7 x\u00b2 = x\u2075). Positive leading coefficient.",
        "    Far left: DOWN.   Far right: UP.  (odd degree, positive leading coeff)",
        "4.  Order of zeros on the number line: \u22123, 0, 2.",
        "    Signs, left to right:  \u2212 , + , \u2212 , + .",
        "    (At x = \u22124 in far left: (\u2212)(\u2212)(\u2212)(+) = \u2212, matches DOWN.",
        "     x\u00b2 + 1 is always + so it never flips a sign.)",
        "5.  Strong responses cite:",
        "    \u2022  \u201cx\u00b2 + 1 doesn\u2019t change the signs because it\u2019s always positive.\u201d",
        "    \u2022  \u201cThe highest power is 5 (odd), so the ends go opposite directions.\u201d",
        "    \u2022  \u201cThe four sign-chart signs alternate because each real zero is a simple crossing.\u201d",
        "WHY THIS IS DOK 3:  students transfer the sign-chart procedure to a NEW function "
        "with a non-factorable factor (x\u00b2 + 1) and predict without computing \u2014 they "
        "must plan, justify, and generalize from the day\u2019s pattern.",
    ])

    doc.save(path)


if __name__ == "__main__":
    build("Day_1_Period_A_Supplement_The_Leap.docx")
    print("Built Day_1_Period_A_Supplement_The_Leap.docx")
