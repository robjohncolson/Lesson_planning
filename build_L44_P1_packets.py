"""Generate Lesson 4-4 Period 1 materials — single-period quick treatment.

Single-period lesson — condensed Klimsara. If 45-min Wed F variant, cut Explore
by 8 min (drop q16 + q26).

NOTE: This is a COMPRESSED lesson — two Examples are modeled in Launch (Ex 3 + Ex 4),
which deviates from the standard Klimsara one-example-only rule. Adding AND
subtracting with unlike denominators are two distinct skills on the assessment
and both warrant a modeled example.

Design: Adding & Subtracting Rational Expressions + DOK-3 resistance capstone
(Topic 4 LEHS Q#15 prep). This single-period treatment covers LCM-based
combination of rational expressions and ends with Practice #32 (electrical
resistance parallel-circuit reciprocal sum), which rehearses the structural
closure-and-LCM reasoning tested on Topic 4 LEHS Q#15.

Phases (55 min base; Period F 65-min):
  0–  5   Do Now    : Model & Discuss — 1/2 + 1/3 conflict       [DOK 1]
  5– 18   Launch    : teacher models Example 3 (add unlike) +
                      Example 4 (subtract unlike)                [DOK 1-2]  [COMPRESSED]
 18– 50   Explore   : Try-It 3/4 + Practice #12/#16/#26 + ⭐ DOK-3 q32 [DOK 2-3]
 50– 55   Share/Sum : DOK-3 pair presents resistance argument +
                      bridge to Topic 4 LEHS Q#15                [DOK 2]
 55– 58   Exit      : summary recap fill-in                      [DOK 1-2]
 (back)   Optional: Practice #34 (compound-fraction multi-select) — fast-finisher
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_ID   = "4-4-savvas-model-discuss-lesson-4-4-launch"
LAUNCH_IDS  = [
    "4-4-savvas-example-3-lesson-4-4",   # Ex 3: add unlike denominators
    "4-4-savvas-example-4-lesson-4-4",   # Ex 4: subtract unlike denominators
]
EXPLORE_IDS = [
    "4-4-savvas-try-it-3-lesson-4-4",
    "4-4-savvas-try-it-4-lesson-4-4",
    "4-4-savvas-q12",
    "4-4-savvas-q16",
    "4-4-savvas-q26",
    "4-4-savvas-q32",    # ⭐ DOK-3 spine — electrical resistance reciprocal addition
]
DOK3_ID       = "4-4-savvas-q32"
REINFORCE_IDS = ["4-4-savvas-q34"]   # compound-fraction equivalence multi-select

_ALL_IDS = [DO_NOW_ID] + LAUNCH_IDS + EXPLORE_IDS + REINFORCE_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 4-4 · Quick"
LESSON_TITLE = "Adding & Subtracting Rational Expressions + Resistance DOK-3 (Q#15 Prep)"

OBJECTIVES = [
    ("Math Objective",
     "Add and subtract rational expressions with like and unlike denominators; "
     "identify the LCM of polynomial denominators; simplify results."),
    ("Language Objective",
     "Explain why adding rational expressions requires common denominators, "
     "using the word 'LCM' and referring to the analogy with fractions."),
    ("Essential Understanding",
     "Rational expressions behave like rational numbers — closed under addition "
     "and subtraction, and requiring a common denominator before you can combine them."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Fluency adding/subtracting rational expressions; understanding closure "
     "under these operations (Topic 4 assessment Q#15 prep)."),
    ("Essential Question",
     "When you add two rational expressions with different denominators, why "
     "must you find the LCM first?"),
    ("Materials",
     "Student packet, pencil, calculator. DOK-3 Practice #32 (electrical "
     "resistance) is the capstone and rehearses the structural reasoning tested "
     "on Topic 4 LEHS Q#15."),
]


# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.0):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    ps.render_prompt(prompt_p, q["prompt"], font_size=11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            run = line.add_run(f"  ({chr(65+i)})  " + ps.latex_to_unicode(a))
            run.font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, for_teacher=False):
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, LESSON_TITLE,
                  subtitle=LESSON_LABEL + (" · Teacher Edition" if for_teacher else ""))
    t = doc.add_table(rows=len(OBJECTIVES), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(OBJECTIVES):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    t = doc.add_table(rows=len(FRAMEWORK_HEADER), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(FRAMEWORK_HEADER):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()


# ── Do Now ───────────────────────────────────────────────────────────────

def build_do_now(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, "Do Now — " + LESSON_TITLE, subtitle=LESSON_LABEL)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from 4-3 to adding rational expressions",
        dok="1",
        minutes=5,
        teacher_does=[
            "Hand out at door.",
            "Circulate 3 min.",
            "Cold-call 1 student: \"Yesterday we multiplied rational expressions. "
            "Today we ADD. What new move will adding require that multiplying didn't?\"",
        ],
        students_do=[
            "Read the Model & Discuss prompt (two students add 1/2 + 1/3 in "
            "conflicting ways).",
            "Decide which is right and why.",
        ],
        questions_to_ask=[
            "What's 1/2 + 1/3? What did you do to get there?",
            "Why can't you just add the numerators and the denominators?",
        ],
        adult_role="Solo teacher: circulate, time 5 min, pull sheets.",
    )

    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label="Do Now — Model & Discuss: 1/2 + 1/3 conflict",
                    space_after_inches=2.0)

    add_callout(doc, "\U0001f3af  SENTENCE FRAME",
                ["“To add 1/2 + 1/3, I first find a common denominator, which is ___. "
                 "Then 1/2 = ___/6 and 1/3 = ___/6, so the sum is ___. "
                 "The student who got ___ was wrong because ___.”"],
                color_hex="FFF2CC")

    doc.save(path)


# ── Student packet ───────────────────────────────────────────────────────

def build_student(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=False)

    ps.teal_section(doc, "LAUNCH — Adding & Subtracting Rational Expressions "
                         "(teacher models Ex 3 + Ex 4)", [])
    add_callout(doc, "\U0001f4d8  RULES — ADD/SUBTRACT RATIONAL EXPRESSIONS "
                     "(keep printed on your page)",
                [
                    "1. LIKE denominators: combine numerators, keep the denominator, simplify.",
                    "2. UNLIKE denominators:",
                    "   (a) factor each denominator,",
                    "   (b) find the LCM,",
                    "   (c) rewrite each fraction with the LCM as denominator,",
                    "   (d) combine numerators,",
                    "   (e) simplify.",
                    "3. The LCM of two polynomial denominators includes every DISTINCT "
                    "factor, taken to its highest power across both.",
                    "4. Always check whether the final expression can be simplified "
                    "by cancelling a common factor.",
                ],
                color_hex="D9EAD3")

    add_callout(doc, "⚠️  COMPRESSED LESSON — TWO EXAMPLES IN LAUNCH",
                [
                    "This is a single-period quick treatment. Teacher models BOTH "
                    "Example 3 (add unlike denominators) AND Example 4 (subtract "
                    "unlike denominators) during Launch.",
                    "Watch carefully on Example 4 — when you subtract, you must "
                    "distribute the minus sign across every term of the second "
                    "numerator. This is the #1 error source.",
                ],
                color_hex="F4CCCC")
    doc.add_paragraph()

    ps.teal_section(doc, "EXPLORE — Think-Pair-Share (32 min)",
                    ["Work with a partner. Move in order. Practice #32 "
                     "(Electrical Resistance) is the big one — plan ~10 min for it.",
                     "45-min Wed F variant: skip #16 and #26."])

    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 3 — Add rational expressions (unlike denominators)",
        "Try It 4 — Subtract rational expressions (unlike denominators)",
        "Practice #12 — Like denominators",
        "Practice #16 — Unlike denominators, factor first",
        "Practice #26 — Subtract with sign-distribution trap",
        "Practice #32  ⭐ DOK-3 ELECTRICAL RESISTANCE — Topic 4 LEHS Q#15 prep",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=label, space_after_inches=1.5)

    ps.teal_section(doc, "SHARE / SUMMARY (5 min)", [])
    ps.sentence_frame_box(doc, [
        "Pair share (Resistance): “For the parallel circuit, 1/R_total = 1/R₁ + 1/R₂. "
        "I found a common denominator of ___, added the reciprocals to get "
        "1/R_total = ___, then flipped to get R_total = ___. "
        "The ratio R_A/R_B simplified to ___ because ___.”",
        "Whole class: one pair reads their resistance solution. Class signals "
        "thumbs up/sideways.",
    ])

    add_callout(doc, "\U0001f52d  BRIDGE TO TOPIC 4 LEHS Q#15",
                ["The resistance move — find LCM, combine reciprocals into a single "
                 "rational expression — is the structural reasoning tested on "
                 "Topic 4 LEHS Q#15 (operations/closure). You have now rehearsed it."],
                color_hex="FCE5CD")

    ps.summary_exit_box(doc, "EXIT TICKET — Summary Recap",
                        [
                            "To add rational expressions with unlike denominators, I must "
                            "first find the _____ of the denominators.",
                            "Practice #32 (resistance) showed me that reciprocal-sum "
                            "problems (like 1/R₁ + 1/R₂) use the same algebraic steps "
                            "as adding rational expressions because _____.",
                            "One biggest thing I learned today: ________________________",
                        ])

    doc.add_page_break()
    ps.teal_section(doc, "OPTIONAL REINFORCEMENT (not collected)",
                    ["If you finish Explore early. #34 is a compound-fraction "
                     "equivalence multi-select that rehearses closure under "
                     "addition and subtraction."])
    reinforce = qb.get_for_packet(REINFORCE_IDS)
    for i, q in enumerate(reinforce, 1):
        bank_item_block(doc, q,
                        label=f"Optional #{i}  (Savvas Practice #34 — compound-fraction multi-select)",
                        space_after_inches=1.8)

    doc.save(path)


# ── Teacher packet ───────────────────────────────────────────────────────

def build_teacher(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=True)

    add_callout(doc, "\U0001f4cc  PRE-FLIGHT — SINGLE-PERIOD COMPRESSED LESSON + DOK-3 CAPSTONE",
                [
                    "This is a single-period quick treatment of Lesson 4-4. TWO examples "
                    "are modeled in Launch (deviation from standard Klimsara one-example "
                    "rule — adding and subtracting are distinct assessment skills). "
                    "Release promptly at minute 18.",
                    "Today's capstone is Practice #32 (electrical resistance parallel "
                    "circuit). Students must (a) recognize that 1/R_total = 1/R_A + 1/R_B "
                    "is a sum of rational expressions, (b) build a common denominator, "
                    "(c) form the ratio of two such resistances as a single simplified "
                    "rational expression. Same structural reasoning as Topic 4 LEHS Q#15.",
                    "Pace: Try-It 3 (4 min) → Try-It 4 (4 min) → #12 (3 min) → "
                    "#16 (4 min) → #26 (5 min) → #32 (10 min). ~30 min total Explore "
                    "(2 min buffer for transitions).",
                    "45-min Wed F variant: cut #16 + #26. Never cut #32.",
                ],
                color_hex="FFD9E0")

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from 4-3 to adding rational expressions",
        dok="1",
        minutes=5,
        teacher_does=[
            "Hand out at door.",
            "Circulate 3 min.",
            "Cold-call 1 student on the 1/2 + 1/3 conflict resolution.",
        ],
        students_do=[
            "Read conflict and decide which student's method is correct.",
            "Justify with the word 'common denominator'.",
        ],
        questions_to_ask=[
            "What's the LCM of 2 and 3?",
            "Why can't you add 1/2 + 1/3 by adding numerators and denominators?",
        ],
        adult_role="Solo teacher: circulate silently. No hints.",
    )
    add_callout(doc, "\U0001f3a4  BRIDGE SCRIPT (read aloud after Do Now)",
                ["\"Yesterday (Lesson 4-3) you practiced MULTIPLYING and DIVIDING "
                 "rational expressions. Today you'll ADD and SUBTRACT them — which "
                 "requires one new move that multiplication didn't: finding a common "
                 "denominator first. The Model & Discuss prompt shows two students "
                 "adding 1/2 + 1/3 in conflicting ways. Which is right?\""],
                color_hex="CFE2F3")
    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label=f"Do Now ({q['id']})", space_after_inches=0.3)
    add_callout(doc, "✅  ANSWER KEY",
                [q.get("notes") or "(verify)"],
                color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Launch",
        framework_tag="teacher models Example 3 + Example 4 [COMPRESSED — 2 examples]",
        dok="1-2",
        minutes=13,
        teacher_does=[
            "Project Example 3 (add unlike denominators). Think aloud: "
            "factor each denominator, identify LCM, rewrite each fraction.",
            "\"The LCM takes every distinct factor at its highest power.\"",
            "Transition to Example 4 without stopping — announce: "
            "\"Subtracting looks identical UNTIL the distribute step. Watch the minus sign.\"",
            "Project Example 4. Think aloud: after combining numerators, "
            "distribute the negative across EVERY term.",
            "Flag the Rules card — students copy into margin.",
            "30-sec pause: \"Turn to partner: where does the minus sign "
            "trap catch students? What do you do to avoid it?\"",
            "Do NOT solve Try-Its or Practice items — release at minute 18.",
        ],
        students_do=[
            "Watch Example 3 silently. Note the LCM-building procedure.",
            "Watch Example 4. Mark the sign-distribution step.",
            "During pause: explain sign trap to partner using Rule 2(d).",
            "Copy Rules card into margin.",
        ],
        questions_to_ask=[
            "What's the LCM of (x-2) and (x+3)?",
            "How do we rewrite 5/(x-2) so it has denominator (x-2)(x+3)?",
            "When we subtract, where does the minus sign apply — to the first term only, "
            "or every term of the second numerator?",
            "After combining, what's the LAST thing you check before you finish?",
        ],
        adult_role="Project both examples back-to-back. Release promptly at minute 18.",
    )
    example_items = qb.get_for_packet(LAUNCH_IDS)
    add_callout(doc, "\U0001f3a4  TEACHER SCRIPT",
                [
                    "1. \"Watch me add two rational expressions with different denominators.\"",
                    "2. \"Step 1: Factor each denominator. Step 2: Find the LCM.\"",
                    "3. \"Step 3: Rewrite each fraction with the LCM as denominator — "
                    "multiply top and bottom by the missing factors.\"",
                    "4. \"Step 4: Combine numerators. Step 5: Simplify.\"",
                    "5. \"Now Example 4 — SUBTRACTION. Same setup, but watch the "
                    "minus sign distribute across EVERY term of the second numerator.\"",
                    "6. \"Today's DOK-3 task uses these moves on a parallel-circuit "
                    "resistance formula. Same algebraic move as Topic 4 LEHS Q#15.\"",
                    "7. Release to Explore.",
                ],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Explore",
        framework_tag="TPS + DOK-3 driver",
        dok="2-3",
        minutes=32,
        teacher_does=[
            "5 laps across 32 min.",
            "Lap 1 (4 min): Try-It 3 — add unlike. Press pairs on LCM construction.",
            "Lap 2 (4 min): Try-It 4 — subtract unlike. Watch for sign-distribution errors.",
            "Lap 3 (3 min): #12 — like denominators (quick confidence build).",
            "Lap 4 (4+5 min): #16 + #26. #26 is the sign-trap item — press \"where "
            "does the minus go?\" on every pair.",
            "Lap 5 (10 min): ⭐ #32 Electrical Resistance. Students build common "
            "denominators for reciprocal sums, then form the ratio.",
            "Do NOT give #32 CER answers. Pairs present argument at Share.",
        ],
        students_do=[
            "6 items in order: Try-It 3 → Try-It 4 → #12 → #16 → #26 → #32.",
            "For #32: BEFORE computing, write out 1/R_total = 1/R_A + 1/R_B "
            "for each circuit and identify the LCM.",
            "Justify the ratio form with sentence frame.",
        ],
        questions_to_ask=[
            "Try-It 3: what's the LCM of the two denominators?",
            "Try-It 4: where does the minus sign distribute?",
            "#12: like denominators — what's the one step you can skip compared to unlike?",
            "#16: factor both denominators first. Now what's the LCM?",
            "#26: watch the subtraction. Where do students typically lose the sign?",
            "#32: what's the LCM for 1/R₁ + 1/R₂? How do you flip to get R_total?",
            "#32: how do you form the ratio R_A/R_B as a SINGLE simplified rational expression?",
        ],
        adult_role="Solo teacher: 5 laps. On #32, press the LCM-of-reciprocals move "
                   "and the final single-expression form — those are the "
                   "assessment-critical moves.",
    )
    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 3", "Try It 4", "Practice #12", "Practice #16",
        "Practice #26",
        "Practice #32 ⭐ DOK-3 ELECTRICAL RESISTANCE",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=f"{label} ({q['id']})", space_after_inches=0.3)
        if "DOK-3" in label or "#32" in label:
            add_callout(doc, "⭐ DOK-3 ELECTRICAL RESISTANCE — Practice #32 "
                             "(Topic 4 LEHS Q#15 prep)",
                        [
                            "q32 is the electrical resistance problem. \"Two circuits, "
                            "A and B. You're asked for the ratio of their resistances, "
                            "where each resistance is expressed as a sum of reciprocals "
                            "(parallel resistors). The mathematical move is: build a "
                            "common denominator, add the reciprocals, then form the ratio.\"",
                            "CER expectation: (C) ratio as a single simplified rational "
                            "expression; (E) the algebra steps — common denominator, add "
                            "fractions, simplify; (R) the resistance-sum formula "
                            "1/R_total = 1/R₁ + 1/R₂ for parallel circuits, with the "
                            "observation that the ratio's form depends only on the "
                            "constituent resistances.",
                            "If students get stuck: remind them parallel resistance "
                            "combines via reciprocal sums (same structural move as adding "
                            "rational expressions). Ask: what's the LCM of the two "
                            "denominators? Once they have that, the rest is algebra.",
                            "Sentence frame required for ELL.",
                        ],
                        color_hex="FFD9E0")
        else:
            add_callout(doc, "✅  ANSWER / ANTICIPATED TALK",
                        [q.get("notes") or "(no answer key — verify before class)"],
                        color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Share / Summary",
        framework_tag="academic conversation + Topic 4 LEHS Q#15 bridge",
        dok="2",
        minutes=5,
        teacher_does=[
            "Cold-call 1 pair to present their resistance solution using sentence frame.",
            "Class signals thumbs. Write the LCM step on board.",
            "Topic 4 bridge: \"This is the structural reasoning on Topic 4 LEHS Q#15. "
            "You have now rehearsed it.\"",
        ],
        students_do=[
            "Presenting pair reads their solution and cites the LCM step.",
            "Rest of class signals and writes takeaway in margin.",
        ],
        questions_to_ask=[
            "Summarize: how did you get from 1/R₁ + 1/R₂ to a single rational expression?",
            "Why does the ratio simplify the way it does?",
        ],
        adult_role="Cold-call a pair that struggled on #32 — hold them accountable "
                   "to the LCM step.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit Ticket",
        framework_tag="summary recap (not graded)",
        dok="1-2",
        minutes=3,
        teacher_does=[
            "Collect at door.",
            "Scan stem 2 (reciprocal-sum connection). Plan re-teach if many "
            "students can't explain the parallel to adding rational expressions.",
        ],
        students_do=["3 sentence stems, honest."],
        questions_to_ask=[
            "Did students connect LCM to the addition procedure?",
            "Did students recognize the structural parallel between parallel "
            "circuits and rational-expression addition?",
        ],
        adult_role="Collect. No grade.",
    )
    add_callout(doc, "\U0001f4cb  EXIT TICKET — Student Form",
                [
                    "To add rational expressions with unlike denominators, I must "
                    "first find the _____ of the denominators.",
                    "Practice #32 (resistance) showed me that reciprocal-sum "
                    "problems (like 1/R₁ + 1/R₂) use the same algebraic steps "
                    "as adding rational expressions because _____.",
                    "One biggest thing I learned today: ________________________",
                ],
                color_hex="FFF2CC")

    doc.add_page_break()
    add_callout(doc, "\U0001f550  PERIOD A / PERIOD F — 65-MIN CLASS NOTES",
                [
                    "Both sections should have 65 min for this lesson. Use the "
                    "extra 10 min on Explore #32 (Resistance) — the richer the "
                    "LCM argument, the better the Topic 4 LEHS Q#15 payoff.",
                    "45-min Wed F variant: cut #16 + #26 from Explore. Never cut #32.",
                    "If finished early: hand out Practice #34 "
                    "(compound-fraction multi-select) as fast-finisher.",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "\U0001f9e9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS",
                [
                    "• Provide the Add/Subtract Rational Expressions Rules card as a "
                    "sticker/cut-out.",
                    "• For #32, provide the setup: \"1/R_total = 1/R₁ + 1/R₂ for each "
                    "circuit. Find a common denominator for each side.\"",
                    "• Accept verbal explanation of the LCM step in lieu of written steps.",
                ],
                color_hex="E2F0D9")
    add_callout(doc, "\U0001f30d  SUPPORTS FOR ELL STUDENTS",
                [
                    "• Sentence frames on student page (don’t skip).",
                    "• Word cards: \"LCM\" = least common multiple — the smallest "
                    "expression both denominators divide into; \"rational expression\" "
                    "= a fraction of polynomials; \"closure\" = if you add two of them, "
                    "you get another one of the same kind.",
                    "• \"Parallel circuit\" = a circuit where the current splits "
                    "across multiple resistors.",
                    "• Pair ELL students with English-dominant partners for TPS.",
                ],
                color_hex="DEEBF7")

    doc.save(path)


# ── Main ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_do_now("L44_P1_Do_Now.docx")
    build_student("L44_P1_Student_Packet.docx")
    build_teacher("L44_P1_Teacher_Packet.docx")
    print("Built L44_P1_Do_Now.docx, L44_P1_Student_Packet.docx, L44_P1_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L44_P1_Visuals_Checklist.md",
                              title="L44 P1 — Visuals Checklist")
