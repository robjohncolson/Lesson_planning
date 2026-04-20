"""Generate Lesson 4-1 Period 3 materials — Klimsara-adapted, no Blooket.

Design: Reciprocal function + translations = the **assessment-critical** period.
Topic 4 8-Q assessment items #3 (asymptotes) and #5 (translation description)
both come from Ex 4-5 content. P3 is pure DOK-2 mastery (no DOK-3 driver — P2
owned the spine). Students leave P3 able to identify asymptotes and describe
translations of the reciprocal function.

Phases (55 min base; Period F 65-min):
  0–  5   Do Now    : Practice #13 — generalize k = xy (P2 bridge)            [DOK 1]
  5– 17   Launch    : teacher models Example 4 + Example 5 back-to-back       [DOK 2]
 17– 52   Explore   : Try-It 4 + #11 + #18 + Try-It 5 + ⭐ #19 + #21          [DOK 2]
                     (#19 is assessment-critical: intercepts + asymptotes)
 52– 57   Share/Sum : Concept Summary return + LEHS 8-Q assessment preview   [DOK 2]
 57– 58   Exit      : pre-assessment confidence check                        [DOK 1-2]
 (back)   Optional : Practice #25 (SAT/ACT MC) — extra stretch
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_ID   = "4-1-savvas-q13"
EXPLORE_IDS = [
    "4-1-savvas-try-it-4-lesson-4-1-matches-examp",
    "4-1-savvas-q11",        # Error Analysis (y=5/x)
    "4-1-savvas-q18",        # y = -2/x
    "4-1-savvas-try-it-5-lesson-4-1-matches-examp",
    "4-1-savvas-q19",        # ⭐ ASSESSMENT-CRITICAL: g(x) = 1/(x-2)+6 + intercepts
    "4-1-savvas-q21",        # video games modeling
]
REINFORCE_IDS = ["4-1-savvas-q25"]

_ALL_IDS = [DO_NOW_ID] + EXPLORE_IDS + REINFORCE_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 4-1 · Period 3"
LESSON_TITLE = "Reciprocal Function + Translations — Assessment-Critical"

OBJECTIVES = [
    ("Math Objective",
     "Graph the reciprocal function f(x) = 1/x and its translations g(x) = "
     "1/(x−h) + k. Identify vertical and horizontal asymptotes, state domain "
     "and range, and find intercepts where they exist."),
    ("Language Objective",
     "Students will use the frame \u201cThis graph is the parent f(x) = 1/x "
     "translated ___ units ___. The asymptotes are at x = ___ and y = ___.\u201d "
     "to describe any translated reciprocal function."),
    ("Essential Understanding",
     "The reciprocal function has a vertical asymptote at x = 0 and horizontal "
     "asymptote at y = 0. Translations shift the asymptotes: g(x) = 1/(x−h) + k "
     "has asymptotes at x = h and y = k. The point (h, k) is the new \u201ccenter\u201d."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Students will identify asymptotes and describe translations for 6 "
     "reciprocal-function items. This is Topic 4 assessment-prep material."),
    ("Essential Question",
     "How do the values of h and k in g(x) = 1/(x − h) + k determine the graph's "
     "asymptotes, domain, range, and intercepts?"),
    ("Materials",
     "Pencils; student packet; graphing calculator (optional, for verification)."),
]


# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.0):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    prompt_p.add_run(q["prompt"]).font.size = Pt(11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            line.add_run(f"  ({chr(65+i)})  {a}").font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, for_teacher=False):
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 3, LESSON_TITLE,
                  subtitle=LESSON_LABEL + (" · Teacher Edition" if for_teacher else ""))
    t = doc.add_table(rows=len(OBJECTIVES), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(OBJECTIVES):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    t = doc.add_table(rows=len(FRAMEWORK_HEADER), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(FRAMEWORK_HEADER):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()


# ── Do Now ───────────────────────────────────────────────────────────────

def build_do_now(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 3, "Do Now — " + LESSON_TITLE, subtitle=LESSON_LABEL)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from P2",
        dok="1",
        minutes=5,
        teacher_does=["Hand out at door.", "Circulate 3 min.",
                      "Cold-call: \"What are the two forms of inverse variation?\""],
        students_do=["Write k in terms of x and y.",
                     "List both forms: y = k/x AND xy = k."],
        questions_to_ask=[
            "What's the product form of inverse variation?",
            "Today we graph this — predict: what does y = 1/x LOOK like?",
        ],
        adult_role="Solo teacher: circulate silently.",
    )

    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label="Do Now — Generalize: k in terms of x and y",
                    space_after_inches=2.0)

    add_callout(doc, "🔭  PREDICTION (spend 30 sec)",
                ["Before we graph y = 1/x together, sketch what you THINK it looks like. (No wrong answers — just predict.) We'll compare to the real thing in Launch."],
                color_hex="FCE5CD")

    doc.save(path)


# ── Student packet ───────────────────────────────────────────────────────

def build_student(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=False)

    ps.teal_section(doc, "LAUNCH — Reciprocal Function + Translations (teacher models Ex 4 + Ex 5)", [])
    add_callout(doc, "📘  RECIPROCAL FUNCTION — PARENT f(x) = 1/x",
                [
                    "• Vertical asymptote: x = 0",
                    "• Horizontal asymptote: y = 0",
                    "• Domain: {x | x ≠ 0}     • Range: {y | y ≠ 0}",
                    "• Graph: hyperbola in Quadrants I and III (two branches, no intersection)",
                ],
                color_hex="D9EAD3")
    add_callout(doc, "📘  TRANSLATED RECIPROCAL — g(x) = 1/(x − h) + k",
                [
                    "• Vertical asymptote shifts to:  x = h",
                    "• Horizontal asymptote shifts to: y = k",
                    "• Domain: {x | x ≠ h}     • Range: {y | y ≠ k}",
                    "• The point (h, k) is the \u201cnew center\u201d — where the asymptotes cross.",
                    "• Example: g(x) = 1/(x − 3) + 2 → asymptotes x = 3 and y = 2.",
                    "",
                    "To find INTERCEPTS:",
                    "  • y-intercept: set x = 0. y = 1/(0 − h) + k = −1/h + k.",
                    "  • x-intercept: set y = 0. Solve 0 = 1/(x − h) + k → x = h − 1/k.",
                ],
                color_hex="FFF2CC")
    doc.add_paragraph()

    ps.teal_section(doc, "EXPLORE — Think-Pair-Share (35 min)",
                    ["Work with a partner. Practice #19 is assessment-critical (asymptotes + intercepts) — plan ~8 min for it."])

    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 4 — Graph g(x) = 10/x (scaled parent)",
        "Practice #11 — Error Analysis: y = 5/x",
        "Practice #18 — Graph y = −2/x (reflected parent)",
        "Try It 5 — Graph g(x) = 1/(x+2) − 4 (translated)",
        "Practice #19  ⭐ ASSESSMENT CRITICAL — g(x) = 1/(x−2) + 6 + intercepts",
        "Practice #21 — Video game storage (modeled)",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=label, space_after_inches=1.5)

    ps.teal_section(doc, "SHARE / SUMMARY (5 min)", [])
    ps.sentence_frame_box(doc, [
        "Pair share (#19): \u201cThis graph is the parent f(x) = 1/x translated ___ units ___ and ___ units ___. The asymptotes are x = ___ and y = ___. The intercepts are ___.\u201d",
        "Whole class: one pair reads #19. Class signals thumbs up/sideways.",
    ])

    add_callout(doc, "📊  CONCEPT SUMMARY (your reference for Topic 4 assessment)",
                [
                    "INVERSE VARIATION (y = k/x):",
                    "  — WORDS: As one variable increases, the other decreases; their product k is constant.",
                    "  — ALGEBRA: y = k/x where k ≠ 0",
                    "  — GRAPH: parent hyperbola, asymptotes x = 0 and y = 0",
                    "",
                    "RECIPROCAL FUNCTION TRANSLATIONS (y = a/(x − h) + k):",
                    "  — WORDS: The reciprocal function can shift left/right (h) and up/down (k).",
                    "  — ALGEBRA: y = a/(x − h) + k",
                    "  — GRAPH: asymptotes shift to x = h and y = k; new \u201ccenter\u201d at (h, k).",
                ],
                color_hex="DEEBF7")

    ps.summary_exit_box(doc, "EXIT TICKET — Pre-Assessment Confidence Check",
                        [
                            "I can identify asymptotes of y = 1/(x−h) + k.     Confidence: 1 / 2 / 3 / 4 / 5",
                            "I can describe a translation from f(x) = 1/x to g(x).     Confidence: 1 / 2 / 3 / 4 / 5",
                            "One thing I want to review before Topic 4 assessment: _____________________________________________________.",
                        ])

    doc.add_page_break()
    ps.teal_section(doc, "OPTIONAL REINFORCEMENT (not collected)",
                    ["Extra stretch for fast finishers. Try without a calculator first."])
    reinforce = qb.get_for_packet(REINFORCE_IDS)
    for i, q in enumerate(reinforce, 1):
        bank_item_block(doc, q, label=f"Optional #{i}  (Savvas SAT/ACT style)",
                        space_after_inches=1.8)

    doc.save(path)


# ── Teacher packet ───────────────────────────────────────────────────────

def build_teacher(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=True)

    add_callout(doc, "📌  PRE-FLIGHT — P3 IS ASSESSMENT-CRITICAL",
                [
                    "Topic 4 8-Q assessment Q#3 (asymptotes) and Q#5 (translation description) both come from Ex 4 + Ex 5 content. P3 is where students earn those 2 of 8 points.",
                    "Practice #19 is THE rehearsal item: it has asymptotes + intercepts + implicit translation description.",
                    "No DOK-3 driver today (P2 owned the spine with #26). P3 is pure DOK-2 mastery.",
                    "Pace: Try-It 4 (5) → #11 (5) → #18 (6) → Try-It 5 (6) → #19 (8) → #21 (5). Total ~35 min.",
                    "If pace slows: cut #21 (modeling) first. Never cut #19.",
                ],
                color_hex="FFD9E0")

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from P2",
        dok="1",
        minutes=5,
        teacher_does=["Hand out at door.", "Circulate 3 min.",
                      "Cold-call for two forms (y = k/x and xy = k)."],
        students_do=["Write k in both forms.", "Sketch prediction of y = 1/x."],
        questions_to_ask=[
            "What ARE the two forms of inverse variation?",
            "Based on yesterday: predict what y = 1/x looks like.",
        ],
        adult_role="Solo teacher: circulate silently.",
    )
    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label=f"Do Now ({q['id']})", space_after_inches=0.3)
    add_callout(doc, "✅  ANSWER KEY",
                [q.get("notes") or "(verify)"],
                color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Launch",
        framework_tag="teacher models Ex 4 + Ex 5 back-to-back",
        dok="2",
        minutes=12,
        teacher_does=[
            "FIRST 6 min: Example 4 — graph f(x) = 1/x. Emphasize asymptote behavior (approaching but never reaching). Show table of x → 0 and x → ±∞.",
            "NEXT 5 min: Example 5 — graph g(x) = 1/(x−3) + 2. Identify (h, k) = (3, 2). Asymptotes shift accordingly.",
            "30 sec: point at packet rule cards. \"Mark them. These are your assessment reference.\"",
            "Do NOT solve Try-Its. Release at minute 17.",
        ],
        students_do=[
            "Watch Ex 4. Compare your Do Now prediction to the real graph.",
            "Watch Ex 5. Note: h shifts horizontal, k shifts vertical. (h, k) is the new center.",
            "Copy BOTH rule cards into margins.",
        ],
        questions_to_ask=[
            "Why does y = 1/x never equal 0? (Because 1/x = 0 has no solution.)",
            "Why does x ≠ 0 in the domain? (Division by zero.)",
            "For g(x) = 1/(x−3) + 2: what's h? What's k? Where do asymptotes cross?",
            "What happens as x → 3 from the right? From the left?",
        ],
        adult_role="Project both examples. Use a visible graphing tool (Desmos). Think aloud.",
    )
    add_callout(doc, "🎤  TEACHER SCRIPT",
                [
                    "1. \u201cYesterday we saw y = k/x. Today we graph it and name its features.\u201d",
                    "2. Example 4: \u201cAs x gets huge (1000, 10000), y gets tiny (0.001, 0.0001) — approaching ZERO but never reaching. That's a HORIZONTAL ASYMPTOTE at y = 0.\u201d",
                    "3. \u201cAs x gets close to 0, y explodes — POSITIVE infinity from the right, NEGATIVE infinity from the left. That's a VERTICAL ASYMPTOTE at x = 0.\u201d",
                    "4. Example 5: \u201cNow watch TRANSLATION. g(x) = 1/(x−3) + 2. Subtract 3 inside → shift RIGHT 3. Add 2 outside → shift UP 2. New asymptotes: x = 3, y = 2.\u201d",
                    "5. \u201cThe POINT (h, k) is the center of the translated graph. For g(x), center is (3, 2).\u201d",
                    "6. Release.",
                ],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Explore",
        framework_tag="TPS · pure DOK-2 mastery · assessment rehearsal",
        dok="2",
        minutes=35,
        teacher_does=[
            "6 laps across 35 min.",
            "Laps 1-3: Try-It 4 + #11 + #18. Parent-form scaled/reflected. (5-6 min each.)",
            "Laps 4-5: Try-It 5 + #19 — translated forms. #19 is the assessment rehearsal — press on BOTH asymptotes AND intercepts.",
            "Lap 6: #21 — modeling. If time.",
            "Do not give answers. Use sentence frame to press justification.",
        ],
        students_do=[
            "6 items in order. Parent forms first (Try-It 4, #11, #18), then translations (Try-It 5, #19, #21).",
            "For #19: BEFORE graphing, state (h, k), asymptotes, and domain/range.",
            "Justify with sentence frame.",
        ],
        questions_to_ask=[
            "What's the center of the translation? (h, k)",
            "Which asymptote is which? (x = h vertical; y = k horizontal.)",
            "For #11: what y-value does y = 5/x actually give at x = 2? (2.5 — the labeled 4 is wrong.)",
            "For #18: how does the negative flip the graph? (Quadrants II and IV instead of I and III.)",
            "For #19: what are the asymptotes? What are the intercepts?",
        ],
        adult_role="Solo teacher: 6 laps. Press BOTH asymptotes AND intercepts on #19. This IS the assessment rehearsal.",
    )
    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 4", "Practice #11", "Practice #18", "Try It 5",
        "Practice #19 ⭐ ASSESSMENT-CRITICAL", "Practice #21",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=f"{label} ({q['id']})", space_after_inches=0.3)
        add_callout(doc, "✅  ANSWER / ANTICIPATED TALK",
                    [q.get("notes") or "(no answer key — verify before class)"],
                    color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Share / Summary",
        framework_tag="Concept Summary + LEHS 8-Q preview",
        dok="2",
        minutes=5,
        teacher_does=[
            "Cold-call 1 pair to present #19 (asymptotes + intercepts).",
            "Project Concept Summary (from student packet).",
            "\"Topic 4 8-Q assessment: expect 2 of 8 questions on THIS content — asymptotes and translation descriptions. You just rehearsed both.\"",
        ],
        students_do=[
            "Presenting pair uses sentence frame for #19.",
            "Rest of class signals and annotates their Concept Summary.",
        ],
        questions_to_ask=[
            "What are the TWO big ideas from Lesson 4-1?",
            "If you see g(x) = a/(x−h) + k on the assessment, what's the FIRST thing you identify?",
        ],
        adult_role="Frame today as assessment rehearsal. Students leave confident.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit Ticket",
        framework_tag="pre-assessment confidence check",
        dok="1-2",
        minutes=3,
        teacher_does=[
            "Collect at door.",
            "Tonight: scan confidence scores. If averages <3 on asymptotes or translations, insert a 5-min Do Now recap when the 8-Q lands.",
        ],
        students_do=[
            "Rate own confidence 1-5 on two skills.",
            "Name one thing to review.",
        ],
        questions_to_ask=[
            "Did any student score <3 on either skill?",
            "What specific content did students name wanting to review?",
        ],
        adult_role="Collect. No grade. Use as data for future assessment prep.",
    )
    add_callout(doc, "📋  EXIT TICKET — Student Form",
                [
                    "I can identify asymptotes of y = 1/(x−h) + k.     Confidence: 1 / 2 / 3 / 4 / 5",
                    "I can describe a translation from f(x) = 1/x to g(x).     Confidence: 1 / 2 / 3 / 4 / 5",
                    "One thing I want to review before Topic 4 assessment: ___.",
                ],
                color_hex="FFF2CC")

    doc.add_page_break()
    add_callout(doc, "🕐  PERIOD A / PERIOD F — 65-MIN CLASS NOTES",
                [
                    "Both sections should have 65 min. Use extra 10 min on Explore #19 (assessment-critical).",
                    "If finished early: Practice #25 (SAT/ACT MC) as fast-finisher stretch.",
                    "Alternative: project Practice #11 as a second Error Analysis prompt — more asymptote practice.",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "🧩  MODIFICATIONS / SUPPORT FOR IEP STUDENTS",
                [
                    "• Provide both Rule cards as half-sheet cut-outs on their desks.",
                    "• For #19, provide a pre-axed grid (labeled x and y from -4 to 8) so they don't spend time drawing axes.",
                    "• Accept verbal asymptote identification in lieu of written.",
                ],
                color_hex="E2F0D9")
    add_callout(doc, "🌍  SUPPORTS FOR ELL STUDENTS",
                [
                    "• Sentence frames on student page (don't skip).",
                    "• Word cards: \u201casymptote\u201d = line the graph approaches but never touches.",
                    "• \u201creciprocal\u201d = the fraction flipped; 1/x is the reciprocal of x.",
                    "• \u201ctranslation\u201d = shift of the whole graph (no rotation or resize).",
                    "• Pair ELL students with English-dominant partners for TPS.",
                ],
                color_hex="DEEBF7")

    doc.save(path)


# ── Main ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_do_now("L41_P3_Do_Now.docx")
    build_student("L41_P3_Student_Packet.docx")
    build_teacher("L41_P3_Teacher_Packet.docx")
    print("Built L41_P3_Do_Now.docx, L41_P3_Student_Packet.docx, L41_P3_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L41_P3_Visuals_Checklist.md",
                              title="L41 P3 — Visuals Checklist")
