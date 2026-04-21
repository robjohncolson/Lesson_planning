"""Generate Lesson 5-1 Period 1 materials — single-period quick treatment.

Single-period lesson — condensed Klimsara. If 45-min Wed F variant, cut Explore
by 8 min (drop q37 + q48).

NOTE: This is a COMPRESSED lesson — two Examples are modeled in Launch (Ex 1 + Ex 3),
which deviates from the standard Klimsara one-example-only rule. The content density
of nth roots AND rational exponents together warrants the extra Launch model.

Design: nth Roots, Rational Exponents + DOK-3 cylinder capstone (Topic 5 Item 1A prep).
This single-period treatment covers the lesson essentials and ends with Practice #50
(cylindrical milk container), which rehearses the exact algebraic move tested on the
Topic 5 assessment Item 1A: derive radius from volume using rational exponents, then
rationalize the denominator.

Phases (55 min base; Period F 65-min):
  0–  5   Do Now    : Explore & Reason — y=x² graph, all (x,y) pairs [DOK 1]
  5– 18   Launch    : teacher models Example 1 (real nth roots) + Example 3
                      (evaluate rational exponents) [DOK 1-2]  [COMPRESSED: 2 examples]
 18– 50   Explore   : Try-It 1/3 + Practice #28/#30/#37/#48 + ⭐ DOK-3 q50 (Cylinder) [DOK 2-3]
 50– 55   Share/Sum : DOK-3 pair presents cylinder argument + bridge to Topic 5 [DOK 2]
 55– 58   Exit      : summary recap fill-in                              [DOK 1-2]
 (back)   Optional: Practice #49 (SAT/ACT cube-root item) — fast-finisher stretch
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_ID   = "5-1-savvas-model-discuss-lesson-5-1-launch"
LAUNCH_IDS  = [
    "5-1-savvas-example-1-lesson-5-1",   # Ex 1: real nth roots
    "5-1-savvas-example-3-lesson-5-1",   # Ex 3: evaluate rational exponents
]
EXPLORE_IDS = [
    "5-1-savvas-try-it-1-lesson-5-1",
    "5-1-savvas-try-it-3-lesson-5-1",
    "5-1-savvas-q28",
    "5-1-savvas-q30",
    "5-1-savvas-q37",
    "5-1-savvas-q48",
    "5-1-savvas-q50",        # ⭐ DOK-3 spine — Cylinder Modeling
]
DOK3_ID       = "5-1-savvas-q50"
REINFORCE_IDS = ["5-1-savvas-q49"]   # SAT/ACT ∛4096x^18y^30 stretch

_ALL_IDS = [DO_NOW_ID] + LAUNCH_IDS + EXPLORE_IDS + REINFORCE_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 5-1 · Quick"
LESSON_TITLE = "nth Roots, Rational Exponents + Cylinder DOK-3 (Item 1A Prep)"

OBJECTIVES = [
    ("Math Objective",
     "Relate nth roots to rational exponents, simplify radicals with variables, "
     "evaluate rational-exponent expressions, and use nth roots to solve "
     "geometric-modeling problems."),
    ("Language Objective",
     "Explain why (x^(1/n))^n = x, using radical and rational-exponent notation "
     "interchangeably."),
    ("Essential Understanding",
     "Rational exponents ARE radicals — x^(m/n) = ⁿ√(x^m) = (ⁿ√x)^m. "
     "Choose whichever form makes the computation easier."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Fluency with nth roots and rational exponents; apply both to volume/radius "
     "modeling problems (Topic 5 assessment Item 1A prep)."),
    ("Essential Question",
     "When is rational-exponent form easier than radical form, and how do we "
     "rationalize denominators in cube-root expressions?"),
    ("Materials",
     "Student packet, pencil, calculator. DOK-3 Practice #50 (cylinder) is the "
     "lesson's capstone and rehearses the Topic 5 Assessment Item 1A move (derive "
     "variable from volume using rational exponents, then rationalize)."),
]


# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.0):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    prompt_p.add_run(q["prompt"]).font.size = Pt(11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            line.add_run(f"  ({chr(65+i)})  {a}").font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, for_teacher=False):
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, LESSON_TITLE,
                  subtitle=LESSON_LABEL + (" · Teacher Edition" if for_teacher else ""))
    t = doc.add_table(rows=len(OBJECTIVES), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(OBJECTIVES):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    t = doc.add_table(rows=len(FRAMEWORK_HEADER), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(FRAMEWORK_HEADER):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()


# ── Do Now ───────────────────────────────────────────────────────────────

def build_do_now(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, "Do Now — " + LESSON_TITLE, subtitle=LESSON_LABEL)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge to nth roots launch",
        dok="1",
        minutes=5,
        teacher_does=["Hand out at door.", "Circulate 3 min.",
                      "Cold-call 1 student: \"For y = x², which (x,y) pairs make y = 9? "
                      "What does that tell you about square roots?\""],
        students_do=["Explore the y = x² graph — identify all (x, y) pairs for "
                     "given y values.",
                     "Connect: if y = x², what is x in terms of y?"],
        questions_to_ask=[
            "If y = 16, how many x values satisfy y = x²?",
            "What happens when y is negative — can you find a real x?",
        ],
        adult_role="Solo teacher: circulate, time 5 min, pull sheets.",
    )

    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label="Do Now — Explore & Reason: y = x² graph",
                    space_after_inches=2.0)

    add_callout(doc, "\U0001f3af  SENTENCE FRAME",
                ["“For y = x², the input x = ___ gives output y = ___. "
                 "This means the square root of ___ is ___. "
                 "There are ___ real values of x because ___.”"],
                color_hex="FFF2CC")

    doc.save(path)


# ── Student packet ───────────────────────────────────────────────────────

def build_student(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=False)

    ps.teal_section(doc, "LAUNCH — nth Roots + Rational Exponents (teacher models Ex 1 + Ex 3)", [])
    add_callout(doc, "\U0001f4d8  nth ROOTS ↔ RATIONAL EXPONENTS (keep printed on your page)",
                [
                    "nth ROOTS ↔ RATIONAL EXPONENTS",
                    "1. x^(1/n) = ⁿ√x  (principal nth root).",
                    "2. x^(m/n) = (ⁿ√x)^m = ⁿ√(x^m).",
                    "3. To rationalize an nth-root denominator: multiply num/denom by ⁿ√(x^(n-1)).",
                    "4. Negative nth roots: only odd n gives a real negative root; "
                    "even n is undefined for negative radicands in Reals.",
                ],
                color_hex="D9EAD3")

    add_callout(doc, "⚠️  COMPRESSED LESSON — TWO EXAMPLES IN LAUNCH",
                [
                    "This is a single-period quick treatment. Teacher models BOTH Example 1 "
                    "(real nth roots) AND Example 3 (evaluate rational exponents) during Launch.",
                    "Pay attention to the switch between radical form and rational-exponent form "
                    "in Example 3 — that conversion is the core skill for the DOK-3 task.",
                ],
                color_hex="F4CCCC")
    doc.add_paragraph()

    ps.teal_section(doc, "EXPLORE — Think-Pair-Share (32 min)",
                    ["Work with a partner. Move in order. Practice #50 (Cylinder Modeling) "
                     "is the big one — plan ~10 min for it."])

    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 1 — Real nth roots",
        "Try It 3 — Evaluate rational exponents",
        "Practice #28 — nth roots with variables",
        "Practice #30 — Rational exponent form",
        "Practice #37 — Simplify: mixed radical/rational",
        "Practice #48 — Multi-step rational exponent",
        "Practice #50  ⭐ DOK-3 CYLINDER MODELING — Topic 5 Assessment Item 1A prep",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=label, space_after_inches=1.5)

    ps.teal_section(doc, "SHARE / SUMMARY (5 min)", [])
    ps.sentence_frame_box(doc, [
        "Pair share (Cylinder): “I isolated r from V = πr²h by substituting "
        "h = 2r, giving V = ___. Solving for r, I got r = ___. "
        "In rational-exponent form, r = (V/(2π))^(1/3). "
        "To rationalize I multiplied by ___ to get ___.”",
        "Whole class: one pair reads their cylinder solution. Class signals thumbs up/sideways.",
    ])

    add_callout(doc, "\U0001f52d  BRIDGE TO TOPIC 5 ASSESSMENT",
                ["The cylinder move — isolate a variable, convert to rational-exponent form, "
                 "rationalize — is exactly what the Topic 5 Assessment Item 1A asks. "
                 "You have now rehearsed it."],
                color_hex="FCE5CD")

    ps.summary_exit_box(doc, "EXIT TICKET — Summary Recap",
                        [
                            "x^(m/n) is the same as ___ in radical form.",
                            "To rationalize a cube-root denominator I ___.",
                            "One biggest thing I learned today: ________________________",
                        ])

    doc.add_page_break()
    ps.teal_section(doc, "OPTIONAL REINFORCEMENT (not collected)",
                    ["If you finish Explore early. #49 is an SAT/ACT cube-root item "
                     "that extends rational-exponent fluency with large radicands."])
    reinforce = qb.get_for_packet(REINFORCE_IDS)
    for i, q in enumerate(reinforce, 1):
        bank_item_block(doc, q, label=f"Optional #{i}  (Savvas Practice #49 — SAT/ACT ∛4,096x¹⁸y³⁰)",
                        space_after_inches=1.8)

    doc.save(path)


# ── Teacher packet ───────────────────────────────────────────────────────

def build_teacher(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=True)

    add_callout(doc, "\U0001f4cc  PRE-FLIGHT — SINGLE-PERIOD COMPRESSED LESSON + DOK-3 CAPSTONE",
                [
                    "This is a single-period quick treatment of Lesson 5-1. TWO examples are "
                    "modeled in Launch (deviation from standard Klimsara one-example rule — "
                    "content density requires it). Release promptly at minute 18.",
                    "Today's capstone is Practice #50 (cylindrical milk container). Students must "
                    "(a) isolate r from V = πr²h with h = 2r, giving r = (V/(2π))^(1/3), "
                    "(b) substitute into lateral surface formula S = 2πrh = 4πr², "
                    "(c) revisit r to count stackable containers against the 20 ft cargo height. "
                    "Same algebraic move as pyramid assessment Item 1A.",
                    "Pace: Try-It 1 (4 min) → Try-It 3 (4 min) → #28 (3 min) → "
                    "#30 (3 min) → #37 (4 min) → #48 (4 min) → #50 (10 min). "
                    "~32 min total Explore.",
                    "45-min Wed F variant: cut q37 + q48. Never cut #50.",
                ],
                color_hex="FFD9E0")

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge to nth roots launch",
        dok="1",
        minutes=5,
        teacher_does=["Hand out at door.", "Circulate 3 min.",
                      "Cold-call 1 student for y = x² pair identification and root connection."],
        students_do=["Identify (x, y) pairs on y = x².",
                     "Connect output y to the concept of nth root."],
        questions_to_ask=[
            "If y = 16, what are all real x values?",
            "What happens for negative y on y = x²?",
        ],
        adult_role="Solo teacher: circulate silently. No hints.",
    )
    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label=f"Do Now ({q['id']})", space_after_inches=0.3)
    add_callout(doc, "✅  ANSWER KEY",
                [q.get("notes") or "(verify)"],
                color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Launch",
        framework_tag="teacher models Example 1 + Example 3 [COMPRESSED — 2 examples]",
        dok="1-2",
        minutes=13,
        teacher_does=[
            "Project Example 1. Think aloud: index n of root, principal root, "
            "even-index domain restriction.",
            "\"For even n, the radicand must be non-negative. For odd n, any real works.\"",
            "Transition to Example 3 without stopping — announce: "
            "\"Now watch the same concept in rational-exponent form.\"",
            "Project Example 3. Think aloud: x^(m/n) ↔ (ⁿ√x)^m.",
            "Flag the Rules card — students copy into margin.",
            "30-sec pause: \"Turn to partner: rewrite your answer in the other notation.\"",
            "Do NOT solve Try-Its or Practice items — release at minute 18.",
        ],
        students_do=[
            "Watch Example 1 silently. Note the even-index rule.",
            "Watch Example 3. Copy the x^(m/n) ↔ radical conversion.",
            "During pause: convert partner’s answer to the other notation.",
            "Copy Rules card into margin.",
        ],
        questions_to_ask=[
            "What is the index n in ⁴√81? What principal root does it give?",
            "Why can’t we take the square root of a negative number in Reals?",
            "Rewrite 8^(2/3) in radical form. Which form is easier to evaluate?",
            "If x^(1/n) = ⁿ√x, what is (x^(1/n))^n equal to?",
        ],
        adult_role="Project both examples back-to-back. Release promptly at minute 18.",
    )
    example_items = qb.get_for_packet(LAUNCH_IDS)
    add_callout(doc, "\U0001f3a4  TEACHER SCRIPT",
                [
                    "1. “Watch me find a real nth root. The index tells me which root.”",
                    "2. “Step 1: Identify the index n. Even index → radicand ≥ 0 in Reals.”",
                    "3. “Step 2: Find the principal (non-negative) root.”",
                    "4. “Now Example 3 — same idea, rational-exponent form.”",
                    "5. “x^(m/n) = (ⁿ√x)^m. I can go radical → exponent or exponent → radical — "
                    "pick the easier path.”",
                    "6. “Today’s DOK-3 task uses this to derive a radius from a volume formula, "
                    "then rationalize. That’s the exact move on your Topic 5 Assessment.”",
                    "7. Release to Explore.",
                ],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Explore",
        framework_tag="TPS + DOK-3 driver",
        dok="2-3",
        minutes=32,
        teacher_does=[
            "5 laps across 32 min.",
            "Lap 1 (4 min): Try-It 1 — nth roots. Press pairs on index and principal root.",
            "Lap 2 (4 min): Try-It 3 — rational exponents. Watch for notation confusion.",
            "Lap 3 (3+3 min): #28 + #30. Press conversion between forms.",
            "Lap 4 (4+4 min): #37 + #48. Multi-step — confirm form choice.",
            "Lap 5 (10 min): ⭐ #50 Cylinder Modeling. Students set up V = πr²h with h = 2r, "
            "solve for r in rational-exponent form, then rationalize.",
            "Do NOT give #50 answers. Pairs present argument at Share.",
        ],
        students_do=[
            "7 items in order: Try-It 1 → Try-It 3 → #28 → #30 → #37 → #48 → #50.",
            "For #50: BEFORE computing, write out V = πr²(2r) = 2πr³ and plan the "
            "rational-exponent isolation step.",
            "Justify rationalization step with sentence frame.",
        ],
        questions_to_ask=[
            "Try-It 1: what is the index? Is the radicand non-negative?",
            "#28: how do you simplify a variable inside an nth root with exponent?",
            "#30: convert to radical form first — which form is easier to evaluate?",
            "#37: do you simplify inside or outside the radical first?",
            "#48: what is (x^(1/3))^3? Use that to check your answer.",
            "#50: after substituting h = 2r, what formula do you get? "
            "Now isolate r using rational exponents.",
            "#50: to rationalize ∛(V/(2π)), what do you multiply numerator and denominator by?",
        ],
        adult_role="Solo teacher: 5 laps. On #50, press the substitution h = 2r and "
                   "the rationalization step — those are the assessment-critical moves.",
    )
    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 1", "Try It 3", "Practice #28", "Practice #30",
        "Practice #37", "Practice #48",
        "Practice #50 ⭐ DOK-3 CYLINDER MODELING",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=f"{label} ({q['id']})", space_after_inches=0.3)
        if "DOK-3" in label or "#50" in label:
            add_callout(doc, "⭐ DOK-3 CYLINDER MODELING — Practice #50 "
                            "(Topic 5 Assessment Item 1A prep)",
                        [
                            "Students must (a) isolate r from V = πr²h with h = 2r, "
                            "giving r = (V/(2π))^(1/3), (b) substitute into lateral surface "
                            "formula S = 2πrh = 4πr², (c) revisit r to count stackable "
                            "containers against the 20 ft cargo height. Same algebraic move as "
                            "pyramid assessment Item 1A.",
                            "Key criteria: (1) correct substitution h = 2r, "
                            "(2) rational-exponent form for r, "
                            "(3) rationalization shown with multiplication by ∛((2π)²).",
                        ],
                        color_hex="FFD9E0")
        else:
            add_callout(doc, "✅  ANSWER / ANTICIPATED TALK",
                        [q.get("notes") or "(no answer key — verify before class)"],
                        color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Share / Summary",
        framework_tag="academic conversation + Topic 5 bridge",
        dok="2",
        minutes=5,
        teacher_does=[
            "Cold-call 1 pair to present their cylinder solution using sentence frame.",
            "Class signals thumbs. Write r = (V/(2π))^(1/3) on board.",
            "Topic 5 bridge: \"This is the exact move on your Topic 5 Assessment Item 1A. "
            "You have now rehearsed it.\"",
        ],
        students_do=[
            "Presenting pair reads their solution and cites the rationalization step.",
            "Rest of class signals and writes takeaway in margin.",
        ],
        questions_to_ask=[
            "Summarize: how did you convert V = 2πr³ to r = (V/(2π))^(1/3)?",
            "Why did you need to rationalize? What did you multiply by?",
        ],
        adult_role="Cold-call a pair that struggled on #50 — hold them accountable "
                   "to the rationalization step.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit Ticket",
        framework_tag="summary recap (not graded)",
        dok="1-2",
        minutes=3,
        teacher_does=[
            "Collect at door.",
            "Scan stem 2 (rationalize). Plan re-teach if many students blank on "
            "the ⁿ√(x^(n-1)) move.",
        ],
        students_do=["3 sentence stems, honest."],
        questions_to_ask=[
            "Did students connect x^(m/n) to radical form correctly?",
            "Did students explain the rationalization multiplier?",
        ],
        adult_role="Collect. No grade.",
    )
    add_callout(doc, "\U0001f4cb  EXIT TICKET — Student Form",
                [
                    "x^(m/n) is the same as ___ in radical form.",
                    "To rationalize a cube-root denominator I ___.",
                    "One biggest thing I learned today: ________________________",
                ],
                color_hex="FFF2CC")

    doc.add_page_break()
    add_callout(doc, "\U0001f550  PERIOD A / PERIOD F — 65-MIN CLASS NOTES",
                [
                    "Both sections should have 65 min for this lesson. Use the extra 10 min "
                    "on Explore #50 (Cylinder Modeling) — the richer the rationalization "
                    "argument, the better the Topic 5 Assessment prep payoff.",
                    "45-min Wed F variant: cut q37 + q48 from Explore. Never cut #50.",
                    "If finished early: hand out Practice #49 (SAT/ACT cube-root stretch).",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "\U0001f9e9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS",
                [
                    "• Provide the nth Roots ↔ Rational Exponents Rules card as a sticker/cut-out.",
                    "• For #50, provide the setup: V = πr²h with h = 2r → V = 2πr³. "
                    "Students solve from that point.",
                    "• Accept verbal rationalization explanation in lieu of written steps.",
                ],
                color_hex="E2F0D9")
    add_callout(doc, "\U0001f30d  SUPPORTS FOR ELL STUDENTS",
                [
                    "• Sentence frames on student page (don’t skip).",
                    "• Word cards: “index” = the small number outside the radical; "
                    "“rational exponent” = fraction exponent; "
                    "“rationalalize” = remove the radical from the denominator.",
                    "• “Principal root” = the non-negative root (even index only).",
                    "• Pair ELL students with English-dominant partners for TPS.",
                ],
                color_hex="DEEBF7")

    doc.save(path)


# ── Main ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_do_now("L51_P1_Do_Now.docx")
    build_student("L51_P1_Student_Packet.docx")
    build_teacher("L51_P1_Teacher_Packet.docx")
    print("Built L51_P1_Do_Now.docx, L51_P1_Student_Packet.docx, L51_P1_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L51_P1_Visuals_Checklist.md",
                              title="L51 P1 — Visuals Checklist")
