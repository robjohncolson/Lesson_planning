"""Generate Lesson 5-5 Period 1 materials — Klimsara-adapted, single-period quick-hit.

Design: Function Operations (add/subtract/multiply/divide/compose) — compressed
single-period treatment. DOK-3 spine is Practice #32 (music store double discount)
which is a STANDALONE reasoning item and does NOT overlap with the Topic 5 assessment
(function composition is not assessed). Accept #32 on its own merits.

Single-period lesson — compressed. 45-min Wed F variant: cut q34 + q28 from Explore.

Phases (55 min base; Period F 65-min):
  0–  5   Do Now    : Launch model-discuss — profit = revenue − cost [DOK 1]
  5– 17   Launch    : teacher models Example 1 (add/subtract) + Example 4 (compose) [DOK 2]
 17– 50   Explore   : Try-Its + Practice items + ⭐ DOK-3 q32 (music store) [DOK 2-3]
 50– 55   Share/Sum : pair presents #32 composition order argument + lesson recap [DOK 2]
 55– 58   Exit      : summary recap fill-in                                    [DOK 1-2]
 (back)   Optional: Practice #35 (SAT/ACT f(g(5))) — fast-finisher stretch
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_ID   = "5-5-savvas-model-discuss-lesson-5-5-launch"
LAUNCH_IDS  = [
    "5-5-savvas-example-1-lesson-5-5",   # add/subtract
    "5-5-savvas-example-4-lesson-5-5",   # compose
]
EXPLORE_IDS = [
    "5-5-savvas-try-it-1-lesson-5-5",
    "5-5-savvas-try-it-4-lesson-5-5",
    "5-5-savvas-q20",
    "5-5-savvas-q21",
    "5-5-savvas-q24",
    "5-5-savvas-q28",
    "5-5-savvas-q34",
    "5-5-savvas-q32",    # ⭐ DOK-3 spine — ORDER MATTERS music store
]
DOK3_ID       = "5-5-savvas-q32"
REINFORCE_IDS = ["5-5-savvas-q35"]    # SAT/ACT f(g(5)) fast-finisher

_ALL_IDS = [DO_NOW_ID] + LAUNCH_IDS + EXPLORE_IDS + REINFORCE_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 5-5 · Quick"
LESSON_TITLE = "Function Operations + Composition Order DOK-3 (Standalone)"

OBJECTIVES = [
    ("Math Objective",
     "Combine functions via +, −, ×, ÷, and composition; track how domain "
     "restrictions change; recognize that composition is non-commutative and explain "
     "when the order of operations changes the outcome."),
    ("Language Objective",
     "Justify in writing which order of a two-step discount gives a better price, "
     "using function composition."),
    ("Essential Understanding",
     "f∘g ≠ g∘f in general — the ORDER of composition matters and often changes "
     "the answer, which matters for real-world decisions."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Add, subtract, multiply, divide, and compose functions; recognize "
     "non-commutativity of composition."),
    ("Essential Question",
     "Why does the ORDER of a two-step discount or two-function composition change "
     "the result, and when does it not?"),
    ("Materials",
     "Student packet, pencil. DOK-3 Practice #32 (music store) is the lesson’s "
     "standalone reasoning capstone — no Topic 5 assessment overlap, treated on "
     "its own merits."),
]


# ── Helpers ─────────────────────────────────────────────────────────────────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.0):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    ps.render_prompt(prompt_p, q["prompt"], font_size=11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            run = line.add_run(f"  ({chr(65+i)})  " + ps.latex_to_unicode(a))
            run.font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, for_teacher=False):
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, LESSON_TITLE,
                  subtitle=LESSON_LABEL + (" · Teacher Edition" if for_teacher else ""))
    t = doc.add_table(rows=len(OBJECTIVES), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(OBJECTIVES):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    t = doc.add_table(rows=len(FRAMEWORK_HEADER), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(FRAMEWORK_HEADER):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()


# ── Do Now ─────────────────────────────────────────────────────────────────────────

def build_do_now(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, "Do Now — " + LESSON_TITLE, subtitle=LESSON_LABEL)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="lesson hook",
        dok="1",
        minutes=5,
        teacher_does=["Hand out at door.", "Circulate 3 min.",
                      "Cold-call 1 student: \"What does profit = revenue − cost look like as a function rule?\""],
        students_do=["Read the model-discuss prompt.",
                     "Write the profit function using function subtraction notation."],
        questions_to_ask=[
            "If R(x) is revenue and C(x) is cost, how do you write profit?",
            "Is this a new function? What would we call it?",
        ],
        adult_role="Solo teacher: circulate, time 5 min, pull sheets.",
    )

    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label="Do Now — Model & Discuss: Profit as a Function",
                    space_after_inches=2.0)

    add_callout(doc, "\U0001f3af  SENTENCE FRAME",
                ["“If R(x) = ___ and C(x) = ___, then profit P(x) = R(x) − C(x) = ___. "
                 "This combined function gives us profit for any value x = ___.”"],
                color_hex="FFF2CC")

    doc.save(path)


# ── Student packet ─────────────────────────────────────────────────────────────────────────

def build_student(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=False)

    ps.teal_section(doc, "LAUNCH — Function Operations + Composition (teacher models Ex 1 & Ex 4)", [])
    add_callout(doc, "\U0001f4d8  FUNCTION OPERATIONS RULES (keep printed on your page)",
                [
                    "FUNCTION OPERATIONS",
                    "1. (f±g)(x) = f(x)±g(x); (fg)(x) = f(x)·g(x); (f/g)(x) = f(x)/g(x), g(x)≠0.",
                    "2. Domain of combined function = intersection of domains, minus any new excluded values (e.g., g(x)=0 for quotient).",
                    "3. Composition: (f∘g)(x) = f(g(x)) — apply g FIRST, then f.",
                    "4. f∘g ≠ g∘f in general. ORDER MATTERS. Always check with a test value if unsure.",
                ],
                color_hex="D9EAD3")

    add_callout(doc, "⚠️  COMPOSITION IS NOT COMMUTATIVE",
                [
                    "(f∘g)(x) = f(g(x))  ≠  (g∘f)(x) = g(f(x))  in general.",
                    "Step 1: identify which function is applied FIRST (the inner function).",
                    "Step 2: substitute the entire inner function as the input to the outer function.",
                    "Step 3: simplify. Then check: would reversing the order change the answer?",
                ],
                color_hex="F4CCCC")
    doc.add_paragraph()

    ps.teal_section(doc, "EXPLORE — Think-Pair-Share (33 min)",
                    ["Work with a partner. Move in order. Practice #32 (Music Store) is the big one — plan ~12 min for it.",
                     "45-min Wed F variant: skip q34 and q28."])

    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 1 — Add/subtract functions",
        "Try It 4 — Compose functions",
        "Practice #20 — Function operations",
        "Practice #21 — Function operations",
        "Practice #24 — Domain of combined function",
        "Practice #28 — Composition evaluation",
        "Practice #34 — Mixed operations",
        "Practice #32  ⭐ DOK-3 ORDER-MATTERS — Music Store Double Discount",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=label, space_after_inches=1.5)

    ps.teal_section(doc, "SHARE / SUMMARY (5 min)", [])
    ps.sentence_frame_box(doc, [
        "Pair share (#32 Music Store): “Applying ___ first gives ___, then ___ gives ___. "
        "If I reverse the order I get ___, so (f∘g)(x) = ___ ≠ g∘f = ___.”",
        "Whole class: one pair presents their order-matters argument. Class signals thumbs up/sideways.",
    ])

    add_callout(doc, "\U0001f52d  LESSON WRAP",
                ["This is a standalone lesson — composition order is the key idea to carry forward. "
                 "It appears in real decisions: discounts, conversions, multi-step processes."],
                color_hex="FCE5CD")

    ps.summary_exit_box(doc, "EXIT TICKET — Summary Recap",
                        [
                            "f∘g means ___ is applied ___.",
                            "An example where f∘g ≠ g∘f is ___.",
                            "For a double discount, the better order for the consumer is ___ because ___.",
                            "One biggest thing I learned today: ________________________",
                        ])

    doc.add_page_break()
    ps.teal_section(doc, "OPTIONAL REINFORCEMENT (not collected)",
                    ["If you finish Explore early. #35 is a SAT/ACT-style item requiring you to evaluate f(g(5))."])
    reinforce = qb.get_for_packet(REINFORCE_IDS)
    for i, q in enumerate(reinforce, 1):
        bank_item_block(doc, q, label=f"Optional #{i}  (Savvas Practice, SAT/ACT f(g(5)))",
                        space_after_inches=1.8)

    doc.save(path)


# ── Teacher packet ─────────────────────────────────────────────────────────────────────────

def build_teacher(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=True)

    add_callout(doc, "\U0001f4cc  PRE-FLIGHT — STANDALONE DOK-3, NOT ASSESSED",
                [
                    "Today’s conceptual core is Practice #32 (Music Store Double Discount). Students must set up TWO different compositions: ($5 off)∘(15% off) and (15% off)∘($5 off), apply to $90, compute both, AND justify in writing which order is better for the consumer.",
                    "The justification in Part C is the DOK-3 move. Accepted answers: (a) $72.25, (b) $71.50, (c) 15% first is better (the $5 saved on a smaller base is worth less than 15% off the larger base).",
                    "NOTE: Function composition is NOT on the Topic 5 assessment. Accept #32 on its own merits as a rich reasoning item.",
                    "Pace: Try-It 1 (4 min) → Try-It 4 (4 min) → #20/#21 (4 min) → #24 (3 min) → #28/#34 (4 min) → #32 (12 min). ~35 min total Explore.",
                    "45-min Wed F variant: cut #34 and #28. Never cut #32.",
                ],
                color_hex="FFD9E0")

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="lesson hook",
        dok="1",
        minutes=5,
        teacher_does=["Hand out at door.", "Circulate 3 min.",
                      "Cold-call 1 student for profit function notation."],
        students_do=["Write P(x) = R(x) − C(x) from the model-discuss prompt.",
                     "Connect function subtraction to the algebraic definition."],
        questions_to_ask=[
            "What operation on functions gives us profit?",
            "Could we combine functions with other operations? (Yes — preview of today.)",
        ],
        adult_role="Solo teacher: circulate silently. No hints.",
    )
    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label=f"Do Now ({q['id']})", space_after_inches=0.3)
    add_callout(doc, "✅  ANSWER KEY",
                [q.get("notes") or "(verify)"],
                color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Launch",
        framework_tag="teacher models Example 1 (add/subtract) + Example 4 (compose)",
        dok="2",
        minutes=12,
        teacher_does=[
            "Project Example 1. Think aloud: (f+g)(x) = f(x) + g(x). Domain = intersection.",
            "Project Example 4. Think aloud: (f∘g)(x) = f(g(x)) — apply g FIRST.",
            "Flag the Function Operations Rules card — print it in student minds.",
            "30-sec pause: “What’s the difference between f(g(x)) and g(f(x))? Check with x=2.” (partner check)",
            "Do NOT solve Try-Its or Practice items — release at minute 17.",
        ],
        students_do=[
            "Watch Example 1 silently; copy combined-function notation.",
            "Watch Example 4 silently; note the inner/outer function structure.",
            "During pause: test f(g(2)) vs g(f(2)) with partner.",
            "Copy Rules card into margin.",
        ],
        questions_to_ask=[
            "In (f∘g)(x), which function do I apply first?",
            "Does f(g(x)) always equal g(f(x))? When might it?",
            "If g(x) = 0 for some x, what happens to (f/g)(x) at that point?",
        ],
        adult_role="Project. Think aloud. Release promptly at minute 17.",
    )
    example_items = qb.get_for_packet(LAUNCH_IDS)
    add_callout(doc, "\U0001f3a4  TEACHER SCRIPT",
                [
                    "1. “Watch me add two functions. The domain of the combined function is the intersection of the two original domains.”",
                    "2. “Now composition: (f∘g)(x) means I substitute g(x) ENTIRELY into f. g goes in first.”",
                    "3. “I’ll test x=2: f(g(2)) vs g(f(2)). Watch what happens — the results are different.”",
                    "4. “This is the big idea today: ORDER MATTERS in composition. It’s like putting on your shoes before your socks vs. socks before shoes.”",
                    "5. “Today’s DOK-3 (#32) asks you to decide which order of a double discount is better for the consumer. You’ll set up both compositions and compare.”",
                    "6. Release to Explore.",
                ],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Explore",
        framework_tag="TPS + DOK-3 driver",
        dok="2-3",
        minutes=33,
        teacher_does=[
            "5 laps across 33 min.",
            "Lap 1 (4 min): Try-It 1 — add/subtract. Press pairs to use function notation.",
            "Lap 2 (4 min): Try-It 4 — compose. Watch for inner/outer confusion.",
            "Lap 3 (4 min): #20/#21 — operations. Confirm domain intersections.",
            "Lap 4 (4 min): #24/#28/#34 — domain + evaluation. Watch for #24 domain restrictions.",
            "Lap 5 (12 min): ⭐ #32 Music Store. Students set up BOTH orders, compute both, and justify Part C in writing.",
            "Do NOT give #32 Part C answers. Pairs present argument to each other; one pair presents at Share.",
        ],
        students_do=[
            "8 items in order: Try-It 1 → Try-It 4 → #20 → #21 → #24 → #28 → #34 → #32.",
            "For #32: BEFORE computing, identify which function represents which discount and set up both compositions.",
            "Justify Part C with sentence frame.",
        ],
        questions_to_ask=[
            "Try-It 1: is (f+g)(x) the same as f(x+g)? (No — watch the notation.)",
            "Try-It 4: which value goes in FIRST — the inner function output or x?",
            "#24: what domain values are excluded, and why?",
            "#32: what does the $5-off function look like? The 15%-off function?",
            "#32: apply each to $90. Which gives a lower price?",
            "#32 Part C: WHY is one order better? (The 15% acts on a larger base amount.)",
        ],
        adult_role="Solo teacher: 5 laps. Press COMPOSITION ORDER justification on #32 Part C.",
    )
    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 1", "Try It 4", "Practice #20", "Practice #21",
        "Practice #24", "Practice #28", "Practice #34",
        "Practice #32 ⭐ DOK-3 ORDER-MATTERS — Music Store Double Discount",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=f"{label} ({q['id']})", space_after_inches=0.3)
        if "DOK-3" in label or "#32" in label:
            add_callout(doc, "⭐ DOK-3 ORDER-MATTERS — Practice #32 (standalone reasoning — NOT on Topic 5 assessment)",
                        [
                            "Students must set up TWO different compositions: ($5 off)∘(15% off) and (15% off)∘($5 off), apply to $90, compute both, AND justify in writing which order is better for the consumer.",
                            "The justification in Part C is the DOK-3 move.",
                            "Accepted answers: (a) $72.25, (b) $71.50, (c) 15% first is better (the $5 saved on a smaller base is worth less than 15% off the larger base).",
                            "Anchor to Ex 4/Ex 6 (composition modeling). Sentence frame required for ELL.",
                        ],
                        color_hex="FFD9E0")
        else:
            add_callout(doc, "✅  ANSWER / ANTICIPATED TALK",
                        [q.get("notes") or "(no answer key — verify before class)"],
                        color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Share / Summary",
        framework_tag="academic conversation + lesson wrap",
        dok="2",
        minutes=5,
        teacher_does=[
            "Cold-call 1 pair to present their #32 Part C justification using sentence frame.",
            "Class signals thumbs. Write the composition-order definition on board: f∘g ≠ g∘f.",
            "Lesson wrap: “Composition order matters — real decisions (discounts, unit conversions, multi-step processes) depend on it.”",
        ],
        students_do=[
            "Presenting pair reads their Part C and shows both computations.",
            "Rest of class signals and writes takeaway in margin.",
        ],
        questions_to_ask=[
            "Why does 15% off first give a better price than $5 off first?",
            "Can you think of a case where the order of two functions WOULDN’T matter?",
        ],
        adult_role="Cold-call a pair that struggled on #32 Part C — hold them to the justification structure.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit Ticket",
        framework_tag="summary recap (not graded)",
        dok="1-2",
        minutes=3,
        teacher_does=[
            "Collect at door.",
            "Scan stem 3 (discount order). Plan review if many students missed the direction of the inequality.",
        ],
        students_do=["4 sentence stems, honest."],
        questions_to_ask=[
            "Did students correctly identify which order is better and give a mathematical reason?",
        ],
        adult_role="Collect. No grade.",
    )
    add_callout(doc, "\U0001f4cb  EXIT TICKET — Student Form",
                [
                    "f∘g means ___ is applied ___.",
                    "An example where f∘g ≠ g∘f is ___.",
                    "For a double discount, the better order for the consumer is ___ because ___.",
                    "One biggest thing I learned today: ________________________",
                ],
                color_hex="FFF2CC")

    doc.add_page_break()
    add_callout(doc, "\U0001f550  PERIOD A / PERIOD F — 65-MIN CLASS NOTES",
                [
                    "Both sections should have 65 min. Use the extra 10 min on Explore #32 (Music Store) — the richer the written justification, the better the DOK-3 payoff.",
                    "45-min Wed F variant: cut #34 and #28 from Explore. Never cut #32.",
                    "If finished early: hand out Practice #35 (SAT/ACT f(g(5))) as fast-finisher.",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "\U0001f9e9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS",
                [
                    "• Provide the Function Operations Rules card as a sticker/cut-out.",
                    "• For #32, provide a starter: “Let D(x) = x − 5 (the $5-off function) and P(x) = 0.85x (the 15%-off function). Now find P(D(90)) and D(P(90)).”",
                    "• Accept verbal Part C justification in lieu of written paragraph.",
                ],
                color_hex="E2F0D9")
    add_callout(doc, "\U0001f30d  SUPPORTS FOR ELL STUDENTS",
                [
                    "• Sentence frames on student page (don’t skip).",
                    "• Word cards: “composition” = apply one function’s OUTPUT as the next function’s INPUT; “inner function” = applied first; “outer function” = applied second.",
                    "• “Order matters” = changing the sequence changes the answer.",
                    "• Pair ELL students with English-dominant partners for TPS.",
                ],
                color_hex="DEEBF7")

    doc.save(path)


# ── Main ─────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_do_now("L55_P1_Do_Now.docx")
    build_student("L55_P1_Student_Packet.docx")
    build_teacher("L55_P1_Teacher_Packet.docx")
    print("Built L55_P1_Do_Now.docx, L55_P1_Student_Packet.docx, L55_P1_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L55_P1_Visuals_Checklist.md",
                              title="L55 P1 — Visuals Checklist")
