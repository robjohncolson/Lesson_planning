"""Generate Lesson 4-1 Period 1 materials — Klimsara-adapted, no Blooket.

Design: Introduce inverse variation. Period 1 lays the conceptual foundation
(xy = k, y = k/x). Assessment items for 4-1 (#3 asymptotes, #5 translations)
come from Examples 4-5, which land in P2-P3. P1 is foundation, not assessment-
critical.

Phases (55 min base; Period F 65-min on Fri; Period A 65-min on Mon):
  0–  5   Do Now    : 4-1 Model & Discuss Part A — notice-and-wonder on two  [DOK 2]
                     144-area rectangles (compressed, 5 min max)
  5– 17   Launch    : teacher models Example 1 (identify IV) + Example 2     [DOK 1-2]
                     (write IV equation)
 17– 50   Explore   : Try-It 1 (both tables) + Try-It 2 (both parts) +       [DOK 1-2]
                     Model & Discuss Parts B + C as TPS extension
 50– 55   Share/Sum : sentence-frame consolidation on xy = k                 [DOK 2]
 55– 58   Exit      : summary recap fill-in                                  [DOK 1-2]
 (back)   Optional Reinforcement: (once #9, #14 ingested — currently empty)

TODO BEFORE FRIDAY: ingest Savvas Practice #9 and #14 so Explore has 1-2 more
TPS items. Current bank: 6 4-1 content items (Launch + 2 Examples + 2 Try-Its
+ Concept). Teacher will need to stretch Try-Its and Model & Discuss to fill
Explore's 33 min.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_ID      = "4-1-savvas-model-discuss-lesson-4-1-launch"
# Explore sequence: Try-It 1 -> #9 (Ex 1) -> #14 (Ex 1) -> Try-It 2 -> #13 (Ex 2) -> #16 (Ex 2) -> #10 (Construct Args).
EXPLORE_IDS = [
    "4-1-savvas-try-it-1-lesson-4-1-matches-examp",
    "4-1-savvas-q9",
    "4-1-savvas-q14",
    "4-1-savvas-try-it-2-lesson-4-1-matches-examp",
    "4-1-savvas-q13",
    "4-1-savvas-q16",
    "4-1-savvas-q10",
]
REINFORCE_IDS  = ["4-1-savvas-q15"]

_ALL_IDS = [DO_NOW_ID] + EXPLORE_IDS + REINFORCE_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 4-1 · Period 1"
LESSON_TITLE = "Inverse Variation — Introduction"

OBJECTIVES = [
    ("Math Objective",
     "Identify when two variables vary inversely by testing whether their "
     "product xy is constant, then write the inverse-variation equation "
     "y = k/x and use it to predict unknown values."),
    ("Language Objective",
     "Students will use the frame \u201cAs ___ increases, ___ decreases; their "
     "product is ___, so the equation is y = ___/x.\u201d to explain an "
     "inverse-variation relationship."),
    ("Essential Understanding",
     "Two variables vary inversely when their product is a constant k. The "
     "equation y = k/x is the algebraic form; xy = k is the structural form. "
     "Graphically, the relationship is a hyperbola (the reciprocal function)."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Students will distinguish inverse variation from direct variation using "
     "product vs ratio tests, and write the equation y = k/x from two given "
     "values. All student-facing items are Savvas-sourced (bank IDs in teacher "
     "edition)."),
    ("Essential Question",
     "How are inverse variations related to the reciprocal function?"),
    ("Materials",
     "Pencils; student packet; calculator (optional). No Chromebooks required "
     "for Period 1."),
]


# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.0):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    prompt_p.add_run(q["prompt"]).font.size = Pt(11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            line.add_run(f"  ({chr(65+i)})  {a}").font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, for_teacher=False):
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, LESSON_TITLE,
                  subtitle=LESSON_LABEL + (" · Teacher Edition" if for_teacher else ""))
    t = doc.add_table(rows=len(OBJECTIVES), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(OBJECTIVES):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    t = doc.add_table(rows=len(FRAMEWORK_HEADER), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(FRAMEWORK_HEADER):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()


# ── Do Now ───────────────────────────────────────────────────────────────

def build_do_now(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, "Do Now — " + LESSON_TITLE, subtitle=LESSON_LABEL)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="notice & wonder",
        dok="2",
        minutes=5,
        teacher_does=["Project the two rectangles on screen or sketch on board.",
                      "Don't answer questions — this is exploration.",
                      "Collect sheets at 5 min sharp."],
        students_do=["Look at the two rectangles. Both have area = 144 sq units.",
                     "Write 1 thing you NOTICE and 1 thing you WONDER.",
                     "Bonus: sketch a third rectangle with area = 144."],
        questions_to_ask=[
            "What is the same about both rectangles?",
            "What is different?",
            "What would length × width equal for ANY rectangle with area 144?",
        ],
        adult_role="Solo teacher: circulate silently. No hints. Students must struggle productively for a moment.",
    )

    add_callout(doc, "🔲  TWO RECTANGLES — Both have area = 144 square units",
                [
                    "Rectangle 1: length = 72, width = 2",
                    "Rectangle 2: length = 24, width = 6",
                    "",
                    "A) Notice: Write something you notice about these two rectangles.",
                    "",
                    "B) Wonder: Write something you wonder about rectangles with area 144.",
                    "",
                    "C) Sketch: Draw a third rectangle with area 144. Label its length and width.",
                ],
                color_hex="FFF2CC")

    doc.save(path)


# ── Student packet ───────────────────────────────────────────────────────

def build_student(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=False)

    ps.teal_section(doc, "LAUNCH — Inverse Variation (teacher models on screen)", [])
    add_callout(doc, "📘  INVERSE VARIATION RULES (keep printed on your page)",
                [
                    "• Two variables vary INVERSELY when their product is a constant k:",
                    "     xy = k      or equivalently      y = k/x",
                    "• The constant k is the CONSTANT OF VARIATION.",
                    "• As one variable INCREASES, the other DECREASES proportionally.",
                    "• Example: if x = 10 and y = 3 vary inversely, then k = xy = 30.",
                    "   Equation: y = 30/x. When x = −6: y = 30/(−6) = −5.",
                ],
                color_hex="D9EAD3")

    add_callout(doc, "⚠️  COMMON ERROR",
                [
                    "Do NOT write y/x = k — that is DIRECT variation.",
                    "Inverse variation multiplies, direct variation divides.",
                    "Test: compute xy for several points. If constant → inverse.",
                ],
                color_hex="F4CCCC")
    doc.add_paragraph()

    ps.teal_section(doc, "EXPLORE — Think-Pair-Share (33 min)",
                    ["Work with a partner. Use the Inverse Variation Rules above for every problem. Move through items in order — each builds on the last."])

    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 1 — Identify inverse variation",
        "Practice #9 — Generalize",
        "Practice #14 — Another table check",
        "Try It 2 — Write the equation",
        "Practice #13 — Generalize: write k in terms of x and y",
        "Practice #16 — Apply the equation",
        "Practice #10 — Construct Arguments (extension)",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=label, space_after_inches=1.4)

    # Share / Summary
    ps.teal_section(doc, "SHARE / SUMMARY (5 min)", [])
    ps.sentence_frame_box(doc, [
        "Pair share: \u201cAs ___ increases, ___ decreases. Their product is ___, so the equation is y = ___/x.\u201d",
        "Whole class: one pair reads their sentence. Class signals thumbs up/sideways.",
    ])

    ps.summary_exit_box(doc, "EXIT TICKET — Summary Recap",
                        [
                            "Today I learned that _____________________________________________________.",
                            "The difference between inverse and direct variation is _____________________________________________________.",
                            "One thing I'm still unsure about is _____________________________________________________.",
                        ])

    doc.add_page_break()
    ps.teal_section(doc, "OPTIONAL REINFORCEMENT (not collected)",
                    ["If you finish Explore early or want extra practice before Period 2."])

    # Model & Discuss extension (demoted from Explore to optional)
    add_callout(doc, "🔭  OPTIONAL · Model & Discuss Parts B + C",
                [
                    "Return to the two rectangles from the Do Now (both have area 144). With your partner:",
                    "",
                    "B) Use Structure. Considering rectangles with area 144, what happens to the WIDTH as the LENGTH increases?",
                    "",
                    "C) Examine at least five more pairs of rectangles with area 144. How would you describe the relationship between their lengths and widths — in your own words?",
                ],
                color_hex="CFE2F3")

    reinforce = qb.get_for_packet(REINFORCE_IDS)
    for i, q in enumerate(reinforce, 1):
        bank_item_block(doc, q, label=f"Optional #{i}  (Savvas Practice)",
                        space_after_inches=1.8)

    doc.save(path)


# ── Teacher packet ───────────────────────────────────────────────────────

def build_teacher(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=True)

    add_callout(doc, "📌  PRE-FLIGHT NOTE (read before class)",
                [
                    "Explore is 7 bank items: Try-It 1 → #9 → #14 → Try-It 2 → #13 → #16 → #10 (extension). Paced for ~4-5 min per item.",
                    "If pairs finish early, push them to the Optional Reinforcement: Model & Discuss Parts B + C (verbal/TPS), then Practice #15.",
                    "If pace slows, cut #10 (Construct Arguments) — it's the DOK-2 stretch; the core DOK-1 arc (Try-It 1 → #9/#14 → Try-It 2 → #13/#16) is the minimum viable set.",
                ],
                color_hex="FFD9E0")

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="notice & wonder",
        dok="2",
        minutes=5,
        teacher_does=["Project the two rectangles (or use printed Do Now sheet).",
                      "Circulate silently. Do NOT teach during this phase.",
                      "Collect at 5 min sharp."],
        students_do=["Write 1 notice + 1 wonder.",
                     "Sketch a third rectangle with area 144."],
        questions_to_ask=["(None during the phase — teacher harvests answers in Launch.)"],
        adult_role="Solo teacher: circulate silently. Resist the urge to give hints.",
    )
    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label=f"Do Now ({q['id']})", space_after_inches=0.3)
    add_callout(doc, "✅  ANTICIPATED STUDENT RESPONSES",
                [
                    "NOTICE (typical): \"Both rectangles have the same area.\" \"One is long and thin, the other is shorter and wider.\" \"When the length is bigger, the width is smaller.\"",
                    "WONDER (typical): \"Can any rectangle have area 144?\" \"What's the rule?\" \"Is there a formula?\"",
                    "SKETCHES: students might draw 1 × 144, 4 × 36, 8 × 18, 12 × 12, 3 × 48, etc. All valid.",
                    "KEY MOVE in Launch: surface the insight that length × width = 144 for every rectangle — this is xy = k.",
                ],
                color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Launch",
        framework_tag="teacher models Example 1 + Example 2",
        dok="1-2",
        minutes=12,
        teacher_does=[
            "Bridge from Do Now: \"You all noticed length × width = 144. That's a RULE, and it has a name: inverse variation.\"",
            "Project Example 1 (Savvas Topic 4-1). Think aloud: for each table, compute xy for each column. Is it constant? → YES = inverse; NO = not.",
            "Project Example 2. Think aloud: given x = 10, y = 3, find k = xy = 30. Write y = 30/x. Substitute x = −6. Answer: y = −5.",
            "Pause 1× for 30-sec partner-check: \"What's k for your Do Now rectangles?\" (Answer: k = 144.)",
            "Do NOT solve Try-Its. Release promptly at minute 17.",
        ],
        students_do=[
            "Watch silently through Example 1.",
            "During partner-check: \"What's k for area = 144?\" (144)",
            "Copy the Inverse Variation Rules card into their page margin.",
            "Note the COMMON ERROR warning: y/x = k is DIRECT variation.",
        ],
        questions_to_ask=[
            "\"What's xy for column 1? Column 2? Column 3? Is it constant?\"",
            "\"If two variables have xy = 30 always, what is k?\"",
            "\"How do I write y in terms of x if xy = 30?\"",
        ],
        adult_role="Project, think aloud. Do NOT give students Try-Its to work until minute 17.",
    )
    add_callout(doc, "🎤  TEACHER SCRIPT",
                [
                    "1. \u201cFrom the Do Now, you noticed length × width = 144 for every rectangle. That's inverse variation — a pattern.\u201d",
                    "2. Example 1 walk-through: \u201cWatch me. For each column, I compute xy. If it's the same number for every column, the two variables vary inversely.\u201d",
                    "3. Example 2 walk-through: \u201cNow I'll write the EQUATION. k = xy, so if x = 10 and y = 3, then k = 30. The equation is y = k/x, so y = 30/x. When x = −6, y = 30/(−6) = −5.\u201d",
                    "4. Common-error flag: \u201cDon't mix up y/x = k (direct) with xy = k (inverse). Inverse multiplies.\u201d",
                    "5. Release: \u201cNow you try. Try-It 1 asks if two tables are inverse variation. Try-It 2 asks you to write the equation.\u201d",
                ],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Explore",
        framework_tag="Think-Pair-Share + Practice",
        dok="1-2",
        minutes=33,
        teacher_does=[
            "5 laps circulation across 33 min (~6-7 min each).",
            "Lap 1: Try-It 1 — press pairs to compute xy for every column.",
            "Lap 2: #9 + #14 — check that pairs notice NOT every table is inverse.",
            "Lap 3: Try-It 2 — press pairs to state k, then write the equation.",
            "Lap 4: #13 + #16 — quick check on generalization and application.",
            "Lap 5: #10 (Construct Arguments) — stretch. Cut if pace slows.",
            "Do not give answers. Use sentence frame to press justification.",
        ],
        students_do=[
            "Pairs move in order through 7 items: 2 Try-Its + 5 Practice.",
            "For each table (#14, #15 later): compute xy for EVERY column before answering.",
            "For Practice #10 (why can x=0 not be in the domain?): justify with the equation y = k/x.",
            "Justify with sentence frame before writing.",
        ],
        questions_to_ask=[
            "What's xy for each column? Is it constant?",
            "If xy is NOT constant, what kind of variation is it? (Check: is y/x constant = direct variation.)",
            "What's the value of k? How do you know?",
            "Can you write the equation y = k/x?",
            "For #10: if y = k/x and x = 0, what happens? (Division by zero = undefined.)",
        ],
        adult_role="Solo teacher: 5 laps. Press pairs to justify using the rule card. No answers.",
    )
    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 1",
        "Practice #9 — Generalize",
        "Practice #14",
        "Try It 2",
        "Practice #13 — Generalize",
        "Practice #16",
        "Practice #10 — Construct Arguments",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=f"{label} ({q['id']})", space_after_inches=0.3)
        add_callout(doc, "✅  ANSWER / ANTICIPATED TALK",
                    [q.get("notes") or "(no answer key — verify before class)"],
                    color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Share / Summary",
        framework_tag="academic conversation",
        dok="2",
        minutes=5,
        teacher_does=[
            "Cold-call 2 pairs to read their sentence-frame sentence.",
            "Write one student-generated sentence on board.",
            "Preview P2: \"Next class we'll solve real-world inverse-variation problems — like a bouzouki string.\"",
        ],
        students_do=[
            "Presenting pair reads their sentence.",
            "Class signals thumbs up/sideways.",
            "Note: next class applies this to a real situation.",
        ],
        questions_to_ask=[
            "What's a sentence that captures today's rule?",
            "What's the difference between inverse and direct variation?",
        ],
        adult_role="Cold-call 2 pairs. Keep it brief — 5 min.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit Ticket",
        framework_tag="summary recap (not graded)",
        dok="1-2",
        minutes=2,
        teacher_does=[
            "Collect at door.",
            "Scan #2 (inverse vs direct) — misconceptions tell you what to re-hit in P2 Do Now.",
        ],
        students_do=[
            "Complete 3 sentence stems.",
            "Be honest about confusion.",
        ],
        questions_to_ask=[
            "Did students correctly distinguish inverse (xy = k) from direct (y/x = k)?",
        ],
        adult_role="Collect. No grade.",
    )
    add_callout(doc, "📋  EXIT TICKET — Student Form",
                [
                    "Today I learned that ___.",
                    "The difference between inverse and direct variation is ___.",
                    "One thing I'm still unsure about is ___.",
                ],
                color_hex="FFF2CC")

    doc.add_page_break()
    add_callout(doc, "🕐  PERIOD A & F — 65-MIN CLASS NOTES",
                [
                    "Both sections get 65 min for P1 (Fri F / Mon A). Use the extra 10 min on Explore, NOT Launch.",
                    "Suggested add: project Savvas Practice #9 and #14 on screen (not in packet) for 2 more TPS rounds.",
                    "Or: extend Model & Discuss Parts B + C with \"find the equation in the form y = k/x for your three rectangles.\" (k = 144, so y = 144/x.)",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "🧩  MODIFICATIONS / SUPPORT FOR IEP STUDENTS",
                [
                    "• Provide the Inverse Variation Rules card as a sticker/cut-out on their desk.",
                    "• For Try-It 1, allow students to compute xy with a calculator for every column.",
                    "• Accept verbal sentence-frame completions in lieu of written.",
                ],
                color_hex="E2F0D9")
    add_callout(doc, "🌍  SUPPORTS FOR ELL STUDENTS",
                [
                    "• Sentence frames printed on student packet (don't skip).",
                    "• Word card: \u201cvaries inversely\u201d = as one goes up, the other goes down proportionally.",
                    "• \u201cconstant of variation\u201d = k = the fixed product xy.",
                    "• \u201cproduct\u201d = the result of multiplying two numbers.",
                    "• Pair ELL students with English-dominant partners for TPS.",
                ],
                color_hex="DEEBF7")

    doc.save(path)


# ── Main ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_do_now("L41_P1_Do_Now.docx")
    build_student("L41_P1_Student_Packet.docx")
    build_teacher("L41_P1_Teacher_Packet.docx")
    print("Built L41_P1_Do_Now.docx, L41_P1_Student_Packet.docx, L41_P1_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L41_P1_Visuals_Checklist.md",
                              title="L41 P1 — Visuals Checklist")
