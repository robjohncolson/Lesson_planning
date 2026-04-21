"""Generate Lesson 6-5 Period 1 materials — single-period quick treatment.

Single-period lesson — condensed Klimsara. If 45-min Wed F variant, drop q3
from Explore (cut one DOK-2 error-analysis item).

NOTE: This is a COMPRESSED lesson — one Example is modeled in Launch (Ex 3),
activating Product, Quotient, AND Power properties simultaneously via condensing.

Design: Properties of Logarithms + DOK-3 paired spine (Power Property proof
+ Richter generalization). Rehearses Form B Q13, Q14, Q16, and the DOK-3
generalization reasoning item.

Phases (50 min base; 45-min Wed-F variant drops q3):
  0–  7   Do Now    : 3 items — q12 (error analysis), q4 (expand), q5 (condense)
  7– 12   Launch    : teacher models Example 3 (condensing — activates all 3 properties)
 12– 40   Explore   : q19, q21, q34 (DOK-2 practice) + q3 (error analysis, cut on Wed-F)
                      + ⭐ DOK-3 PAIRED driver: q13 (Power Property proof) + q40 Part C
                      (Richter generalization)
 40– 45   Share/Sum : DOK-3 pair presents proof + Richter result + bridge to Topic 6 assess
 45– 50   Exit      : two-part custom — condense P=3 log t + log 4, then solve for t when P=2
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_IDS = [
    "6-5-savvas-q12",   # change-of-base numerator/denominator error analysis
    "6-5-savvas-q4",    # expand log_6(49/5) with Quotient Property
    "6-5-savvas-q5",    # condense 5 ln s + 6 ln t into single log
]
LAUNCH_IDS = [
    "6-5-savvas-example-3-lesson-6-5",   # Ex 3: write expressions as single logarithm
]
EXPLORE_IDS = [
    "6-5-savvas-q19",   # condense log_5 6 + 1/2 log_5 y
    "6-5-savvas-q21",   # condense 1/3 ln 27 - 3 ln(2y)
    "6-5-savvas-q34",   # solve 7^x = 100 via change of base
    "6-5-savvas-q3",    # Amanda's expand error on log_4(c^2 d^5) — cut on Wed-F
    "6-5-savvas-q13",   # ⭐ DOK-3: prove Power Property via exponent laws
    "6-5-savvas-q40",   # ⭐ DOK-3 paired: Richter generalization (Part C: 150x intensity)
]
DOK3_ID      = "6-5-savvas-q13"
DOK3_PAIR_ID = "6-5-savvas-q40"
REINFORCE_IDS = ["6-5-savvas-q40"]   # Part A — direct numerical Richter; fast-finisher

_ALL_IDS = DO_NOW_IDS + LAUNCH_IDS + EXPLORE_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 6-5 · Quick"
LESSON_TITLE = "Properties of Logarithms + DOK-3 Power Property Proof & Richter Generalization"

OBJECTIVES = [
    ("Math Objective",
     "Use the Product, Quotient, and Power Properties of Logarithms, plus the "
     "Change of Base Formula, to expand, condense, and evaluate logarithmic "
     "expressions and solve exponential equations."),
    ("Language Objective",
     "State each log property in 'words' form ('the log of a product equals "
     "the sum of the logs'), and justify a condensing step by naming which "
     "property applies."),
    ("Essential Understanding",
     "The log properties come directly from the exponent laws — every property "
     "of logarithms IS a property of exponents written in inverse form."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Fluency with all four log properties + change of base; apply to condense, "
     "expand, and solve; prove Power Property from first principles (DOK-3)."),
    ("Essential Question",
     "How do the properties of exponents become the properties of logarithms, "
     "and how do we use them to simplify or solve equations?"),
    ("Materials",
     "Student packet, pencil, calculator. DOK-3 paired driver (#13 proof + #40 "
     "Richter) is the lesson capstone and rehearses Form B DOK-3 generalization "
     "reasoning. Note: Form B Q8 and Q15 (geometric series) are NOT on the "
     "Topic 6 assessment and are NOT covered here."),
]

RULES_TEXT = [
    "PROPERTIES OF LOGARITHMS (all for b > 0, b ≠ 1; M, N > 0):",
    "  PRODUCT:   log_b (MN) = log_b M + log_b N",
    "  QUOTIENT:  log_b (M/N) = log_b M − log_b N",
    "  POWER:     log_b (M^n) = n · log_b M",
    "  IDENTITY:  log_b b = 1   log_b 1 = 0   b^(log_b x) = x",
    "CHANGE OF BASE:",
    "  log_b x = (log x) / (log b) = (ln x) / (ln b)",
    "SOLVING EXPONENTIAL EQUATIONS:",
    "  b^x = k  ⇒  x = log_b k = (log k) / (log b)",
]


# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.0):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    ps.render_prompt(prompt_p, q["prompt"], font_size=11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            run = line.add_run(f"  ({chr(65+i)})  " + ps.latex_to_unicode(a))
            run.font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, for_teacher=False):
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, LESSON_TITLE,
                  subtitle=LESSON_LABEL + (" · Teacher Edition" if for_teacher else ""))
    t = doc.add_table(rows=len(OBJECTIVES), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(OBJECTIVES):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    t = doc.add_table(rows=len(FRAMEWORK_HEADER), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(FRAMEWORK_HEADER):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()


# ── Do Now ───────────────────────────────────────────────────────────────

def build_do_now(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, "Do Now — " + LESSON_TITLE, subtitle=LESSON_LABEL)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge to log properties launch",
        dok="1-2",
        minutes=7,
        teacher_does=["Hand out at door.", "Circulate 4 min.",
                      "Cold-call 1 student on q12: 'Where is the error? Which rule did "
                      "the student violate?'"],
        students_do=["Analyze a change-of-base error (q12), expand using Quotient Property "
                     "(q4), and condense using Power + Product Properties (q5).",
                     "No calculator needed for q4/q5 — property fluency only."],
        questions_to_ask=[
            "q12: what does change-of-base say the denominator must be?",
            "q4: what is log_6(49/5) in expanded form? Which property applies?",
            "q5: what property lets you pull the coefficient 5 inside as an exponent?",
        ],
        adult_role="Solo teacher: circulate, time 7 min, pull sheets.",
    )

    items = qb.get_for_packet(DO_NOW_IDS)
    labels = [
        "Do Now #1 — Error Analysis: Change of Base",
        "Do Now #2 — Expand: log₆(49/5) with Quotient Property",
        "Do Now #3 — Condense: 5 ln s + 6 ln t",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=label, space_after_inches=1.6)

    add_callout(doc, "\U0001f3af  SENTENCE FRAME",
                ["“For Do Now #1, the error is ___ because change of base says "
                 "log_b x = (log ___) / (log ___). "
                 "For #2, log_6(49/5) = log_6 ___ − log_6 ___ (Quotient Property). "
                 "For #3, I used the ___ Property to write 5 ln s = ln(s^___), "
                 "then the ___ Property to combine.”"],
                color_hex="FFF2CC")

    doc.save(path)


# ── Student packet ───────────────────────────────────────────────────────

def build_student(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=False)

    ps.teal_section(doc, "LAUNCH — Example 3: Write as a Single Logarithm (teacher models)", [])
    add_callout(doc, "\U0001f4d8  PROPERTIES OF LOGARITHMS (keep printed on your page)",
                RULES_TEXT,
                color_hex="D9EAD3")

    add_callout(doc, "⚠️  COMPRESSED LESSON — ONE EXAMPLE IN LAUNCH, ALL THREE PROPERTIES ACTIVE",
                [
                    "Teacher models Example 3 (condensing into a single logarithm). "
                    "Condensing forces you to recognize Product, Quotient, AND Power Properties "
                    "at once — watch which property applies at each step.",
                    "The DOK-3 task (#13) asks you to prove the Power Property from scratch "
                    "using exponent laws. That proof IS on Form B. Plan 12 min for #13 + #40.",
                ],
                color_hex="F4CCCC")
    doc.add_paragraph()

    ps.teal_section(doc, "EXPLORE — Think-Pair-Share (28 min)",
                    ["Work with a partner. Move in order. #13 + #40 are the big ones — "
                     "plan ~12 min for the paired DOK-3 driver. On Wed-F 45-min, skip #3."])

    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Practice #19 — Condense: log₅ 6 + ½ log₅ y",
        "Practice #21 — Condense: ⅓ ln 27 − 3 ln(2y)",
        "Practice #34 — Solve: 7^x = 100 (change of base)",
        "Practice #3  — Amanda's expand error on log₄(c²d⁵)   [SKIP on Wed-F 45-min]",
        "Practice #13  ⭐ DOK-3 PROOF — Prove the Power Property: log_b(M^n) = n·log_b M",
        "Practice #40 Part C  ⭐ DOK-3 RICHTER — Generalize: one earthquake 150× greater; how much larger is its Richter magnitude?",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=label, space_after_inches=1.5)

    ps.teal_section(doc, "SHARE / SUMMARY (5 min)", [])
    ps.sentence_frame_box(doc, [
        "Pair share (Power Property proof): “Let x = log_b M. Then b^x = M. "
        "Raising both sides to the nth power: (b^x)^n = M^n, so b^(nx) = M^n. "
        "By the definition of log, log_b(M^n) = nx = n·log_b M. ✓”",
        "Richter: “R(150A₀) − R(A₀) = log(150A₀/S) − log(A₀/S) "
        "= log 150 ≈ 2.18. So the larger earthquake is about 2.18 Richter units greater.”",
        "Whole class: one pair reads the proof. Class signals thumbs up/sideways.",
    ])

    add_callout(doc, "\U0001f52d  BRIDGE TO TOPIC 6 ASSESSMENT",
                ["The Power Property proof (q13) and the Richter generalization (q40 Part C) "
                 "are the exact DOK-3 reasoning pattern on Form B. You have now rehearsed both. "
                 "Note: Form B Q8 and Q15 (geometric series) are NOT on the Topic 6 assessment."],
                color_hex="FCE5CD")

    ps.summary_exit_box(doc, "EXIT TICKET — Two-Part Summary",
                        [
                            "A biologist models a population as P = 3 log t + log 4.",
                            "(a) Write P as a single logarithm: P = log(___)",
                            "(b) Solve for t when P = 2: t = ___",
                            "Expected: (a) P = log(4t³)    (b) t = (25)^(1/3) ≈ 2.924",
                            "One biggest thing I learned today: ________________________",
                        ])

    doc.add_page_break()
    ps.teal_section(doc, "OPTIONAL REINFORCEMENT (not collected)",
                    ["If you finish #13 + #40 early. Part A of #40 is a direct numerical "
                     "Richter computation — builds toward the Part C generalization."])
    reinforce = qb.get_for_packet(REINFORCE_IDS)
    for i, q in enumerate(reinforce, 1):
        bank_item_block(doc, q, label=f"Optional  (Savvas Practice #40 Part A — direct Richter computation)",
                        space_after_inches=1.8)

    doc.save(path)


# ── Teacher packet ───────────────────────────────────────────────────────

def build_teacher(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=True)

    add_callout(doc, "\U0001f4cc  PRE-FLIGHT — SINGLE-PERIOD COMPRESSED LESSON + DOK-3 PAIRED CAPSTONE",
                [
                    "This is a single-period quick treatment of Lesson 6-5. ONE example is "
                    "modeled in Launch (Example 3: condensing). Release promptly at minute 12.",
                    "Today's DOK-3 driver is PAIRED: #13 (prove Power Property) + #40 Part C "
                    "(Richter generalization). Students must connect the proof to the application. "
                    "Both are Form B DOK-3 rehearsal items.",
                    "Pace: q19 (4 min) → q21 (4 min) → q34 (4 min) → q3 (4 min, skip Wed-F) "
                    "→ #13 + #40 paired (12 min). ~28 min total Explore.",
                    "45-min Wed-F variant: cut q3. Never cut #13 or #40.",
                    "SCOPE NOTE: Form B Q8 and Q15 (geometric series) are NOT covered in this "
                    "lesson and have been dropped from the Topic 6 assessment. Do NOT prep "
                    "students for those items.",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "\U0001f9e9  ASSESSMENT ALIGNMENT — Form B Rehearsal Mapping",
                [
                    "q19, q21 → Form B Q13 (condense logs) + Q16 (step-by-step condense token drag)",
                    "q34 → Form B Q14 (solve exponential equation with logs)",
                    "q3 → Form B Q3 (substitution/simplification in rational/log expressions)",
                    "q13 + q40 → DOK-3 rehearsal for Form B generalization reasoning",
                    "⚠️ Form B Q8 and Q15 (geometric series) are NOT on the Topic 6 "
                    "assessment and are NOT covered in this lesson.",
                ],
                color_hex="E2F0D9")

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge to log properties launch",
        dok="1-2",
        minutes=7,
        teacher_does=["Hand out at door.", "Circulate 4 min.",
                      "Cold-call 1 student on q12: change-of-base error."],
        students_do=["Error analysis (q12), expand with Quotient Property (q4), "
                     "condense with Power + Product Properties (q5)."],
        questions_to_ask=[
            "q12: what must the denominator be in change-of-base? Same base or same argument?",
            "q4: name the property. What does log_6(49/5) expand to?",
            "q5: Power Property first (coefficient → exponent), then Product Property to combine.",
        ],
        adult_role="Solo teacher: circulate silently. No hints on q12 — students must identify the error type.",
    )
    items_dn = qb.get_for_packet(DO_NOW_IDS)
    for q, label in zip(items_dn, [
        "Do Now #1 — Error Analysis: Change of Base",
        "Do Now #2 — Expand: log₆(49/5)",
        "Do Now #3 — Condense: 5 ln s + 6 ln t",
    ]):
        bank_item_block(doc, q, label=f"{label} ({q['id']})", space_after_inches=0.3)
        add_callout(doc, "✅  ANSWER KEY",
                    [q.get("notes") or "(verify before class)"],
                    color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Launch",
        framework_tag="teacher models Example 3 [condensing — all 3 properties active]",
        dok="1-2",
        minutes=5,
        teacher_does=[
            "Project Example 3. Think aloud: identify which property applies at each step.",
            "“Condensing forces you to recognize all three properties at once — "
            "Power first (coefficients become exponents), then Product/Quotient to merge.”",
            "Flag the Rules card — students copy POWER PROPERTY into margin.",
            "30-sec pause: ‘Turn to partner: what would expanding look like? Same properties, "
            "opposite direction.‘",
            "Do NOT solve any Explore items. Release at minute 12.",
        ],
        students_do=[
            "Watch Example 3 silently. Note the order: Power → Product/Quotient.",
            "Copy the Power Property rule into margin.",
            "During pause: describe expanding vs. condensing to partner.",
        ],
        questions_to_ask=[
            "Which property lets us move the coefficient n in front of log_b M?",
            "After applying Power Property, which property combines two logs into one?",
            "If I have log_b M + log_b N, what single log equals that? What if it’s a subtraction?",
            "How does Example 3 connect to exponent laws? (That’s the DOK-3 proof you’ll do.)",
        ],
        adult_role="Project Example 3 only. Release promptly at minute 12.",
    )
    example_items = qb.get_for_packet(LAUNCH_IDS)
    add_callout(doc, "\U0001f3a4  TEACHER SCRIPT",
                [
                    "1. “Watch me condense this expression into a single logarithm.”",
                    "2. “Step 1: Any coefficient in front of a log? → Power Property: move it up as exponent.”",
                    "3. “Step 2: Now I have log terms being added or subtracted.”",
                    "4. “Added logs → Product Property: log_b(MN). Subtracted logs → Quotient Property: log_b(M/N).”",
                    "5. “Result: one single logarithm. That’s condensing.”",
                    "6. “Today’s DOK-3 task asks you to PROVE why the Power Property works — "
                    "using nothing but the definition of logarithm and exponent laws.”",
                    "7. Release to Explore.",
                ],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Explore",
        framework_tag="TPS + DOK-3 paired driver",
        dok="2-3",
        minutes=28,
        teacher_does=[
            "4 laps across 28 min.",
            "Lap 1 (4+4 min): q19 + q21. Press: which property, in what order?",
            "Lap 2 (4+4 min): q34 + q3. q34: change of base for solving. q3: name the error type.",
            "Lap 3 (12 min): ⭐ #13 + #40 paired. DO NOT give the let-x substitution for #13 "
            "before 6 min in. For #40 Part C: press ‘write R(150A₀) and R(A₀) separately first.‘",
            "On Wed-F 45-min: skip q3. Move directly from q34 to #13 + #40.",
        ],
        students_do=[
            "6 items in order (5 on Wed-F): q19 → q21 → q34 → q3 (skip Wed-F) → #13 → #40 Part C.",
            "For #13: let x = log_b M. Rewrite b^x = M. Raise both sides to n. Apply log definition.",
            "For #40 Part C: compute R(150A₀) − R(A₀) using the log properties. Generalize.",
        ],
        questions_to_ask=[
            "q19: which property handles the ½ coefficient? Apply Power Property first.",
            "q21: what is ⅓ ln 27? Simplify before subtracting.",
            "q34: change of base — which base do you use on the calculator?",
            "q3: Amanda wrote log₄(c²d⁵) = 2·log₄ c · 5·log₄ d. "
            "Which property did she confuse? (Power ≠ Product of coefficients.)",
            "#13: if b^x = M and you raise both sides to n, what exponent rule gives b^(nx) = M^n?",
            "#40 Part C: write R(150A₀) − R(A₀) = log(150A₀/S) − log(A₀/S). "
            "Now apply the Quotient Property. What remains?",
        ],
        adult_role="Solo teacher: 4 laps. On #13, press the ‘let x = log_b M‘ substitution "
                   "and the exponent-law bridge. On #40 Part C, press the Quotient Property "
                   "simplification to isolate log 150.",
    )
    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Practice #19", "Practice #21", "Practice #34", "Practice #3",
        "Practice #13 ⭐ DOK-3 PROOF",
        "Practice #40 Part C ⭐ DOK-3 RICHTER",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=f"{label} ({q['id']})", space_after_inches=0.3)
        if "DOK-3" in label and "PROOF" in label:
            add_callout(doc, "⭐ DOK-3 PROOF — Practice #13",
                        [
                            "Students prove: log_b(M^n) = n·log_b M.",
                            "Expected argument: Let x = log_b M ⇒ b^x = M. "
                            "Raise both sides to n: (b^x)^n = M^n ⇒ b^(nx) = M^n. "
                            "By log definition: log_b(M^n) = nx = n·log_b M. ✓",
                            "Key criteria: (1) correct let-x substitution, "
                            "(2) exponent-power rule applied, "
                            "(3) final step cites log definition.",
                        ],
                        color_hex="FFD9E0")
        elif "DOK-3" in label and "RICHTER" in label:
            add_callout(doc, "⭐ DOK-3 RICHTER — Practice #40 Part C",
                        [
                            "Students must generalize: if one earthquake has intensity 150A₀, "
                            "how much larger is its Richter magnitude?",
                            "Expected: R(150A₀) − R(A₀) = log(150A₀/S) − log(A₀/S) "
                            "= log(150) ≈ 2.18.",
                            "Key criteria: (1) write both R values using the formula, "
                            "(2) apply Quotient Property to get log(150), "
                            "(3) evaluate log 150 ≈ 2.18 and interpret in context.",
                        ],
                        color_hex="FFD9E0")
        else:
            add_callout(doc, "✅  ANSWER / ANTICIPATED TALK",
                        [q.get("notes") or "(no answer key — verify before class)"],
                        color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Share / Summary",
        framework_tag="academic conversation + Topic 6 assess bridge",
        dok="2",
        minutes=5,
        teacher_does=[
            "Cold-call 1 pair to narrate the Power Property proof (#13) using sentence frame.",
            "Cold-call 1 pair to present the Richter generalization (#40 Part C).",
            "Class signals thumbs. Write log_b(M^n) = n·log_b M on board.",
            "Topic 6 bridge: “The Power Property proof and the Richter generalization "
            "are the exact DOK-3 pattern on Form B. You have now rehearsed both.”",
            "SCOPE NOTE: Form B Q8 and Q15 (geometric series) are NOT on the assessment.",
        ],
        students_do=[
            "Presenting pair reads proof and names each exponent-law step.",
            "Second pair reads Richter result: log 150 ≈ 2.18 with context sentence.",
            "Rest of class signals and writes takeaway in margin.",
        ],
        questions_to_ask=[
            "Summarize: which exponent law did you use to go from b^x to b^(nx)?",
            "Why does R(150A₀) − R(A₀) equal log 150 exactly? Which property did that?",
        ],
        adult_role="Cold-call a pair that struggled on #13 — hold them accountable "
                   "to the let-x substitution step.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit Ticket",
        framework_tag="two-part summary (not graded)",
        dok="1-2",
        minutes=5,
        teacher_does=[
            "Collect at door.",
            "Scan part (b): if students wrote t = 25 (not cube root), they dropped the exponent. "
            "Plan a 2-min re-teach before next Topic 6 item.",
        ],
        students_do=["Two-part exit: condense, then solve."],
        questions_to_ask=[
            "Did students apply Power Property correctly to get log(4t³)?",
            "Did students solve 10^2 = 4t³ → t³ = 25 → t = 25^(1/3)?",
        ],
        adult_role="Collect. No grade.",
    )
    add_callout(doc, "\U0001f4cb  EXIT TICKET — Student Form",
                [
                    "A biologist models a population as P = 3 log t + log 4.",
                    "(a) Write P as a single logarithm: P = log(___)",
                    "(b) Solve for t when P = 2: t = ___",
                    "Expected: (a) P = log(4t³)    (b) t = (25)^(1/3) ≈ 2.924",
                    "One biggest thing I learned today: ________________________",
                ],
                color_hex="FFF2CC")

    doc.add_page_break()
    add_callout(doc, "\U0001f550  PERIOD A / PERIOD F — 50-MIN CLASS NOTES",
                [
                    "Both sections have 50 min for this lesson. This is a compressed single-period "
                    "treatment with a paired DOK-3 driver (#13 proof + #40 Richter).",
                    "45-min Wed-F variant: cut q3 from Explore. Never cut #13 or #40.",
                    "If a pair finishes #13 + #40 early: hand out Practice #40 Part A "
                    "(direct numerical Richter computation).",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "\U0001f9e9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS",
                [
                    "• Provide the Properties of Logarithms Rules card as a sticker/cut-out.",
                    "• For #13, provide the first step: 'Let x = log_b M, so b^x = M.' "
                    "Students complete from there.",
                    "• For #40 Part C, provide: 'Write R(150A₀) = log(150A₀/S) and "
                    "R(A₀) = log(A₀/S). Now subtract.' Students apply Quotient Property.",
                    "• Accept verbal proof explanation in lieu of written algebraic steps.",
                ],
                color_hex="E2F0D9")
    add_callout(doc, "\U0001f30d  SUPPORTS FOR ELL STUDENTS",
                [
                    "• Sentence frames on student page (don’t skip).",
                    "• Word cards: ‘condense’ = write as one log; "
                    "‘expand’ = write as separate logs; "
                    "‘property’ = a rule that is always true.",
                    "• ‘Proof’ = show WHY the rule works, step by step, with algebra.",
                    "• Pair ELL students with English-dominant partners for TPS on #13.",
                ],
                color_hex="DEEBF7")

    doc.save(path)


# ── Main ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_do_now("L65_P1_Do_Now.docx")
    build_student("L65_P1_Student_Packet.docx")
    build_teacher("L65_P1_Teacher_Packet.docx")
    print("Built L65_P1_Do_Now.docx, L65_P1_Student_Packet.docx, L65_P1_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L65_P1_Visuals_Checklist.md",
                              title="L65 P1 — Visuals Checklist")
