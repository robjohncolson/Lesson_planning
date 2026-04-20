"""Shared docx formatting helpers for the daily lesson packet builders.

Adopts the visual language of `lesson35_8day.tex` (teal section labels,
dark-blue day banners, red-bordered exit box, light-gray sentence-frame
callout, running name/block header) into the python-docx toolchain so
the teacher prints DOCX \u2192 PDF and sees the same look.

Colors match the .tex definitions:
    tealheader  #0097A7
    lightgray   #F2F2F2
    warmred     #C0392B
    dayblue     #2B5C8A
"""
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


COLORS = {
    "teal":      "0097A7",
    "lightgray": "F2F2F2",
    "warmred":   "C0392B",
    "warmredlt": "FCEAE8",  # very pale red (title backdrop)
    "dayblue":   "2B5C8A",
    "white":     "FFFFFF",
    "tealpale":  "E6F5F7",  # 10% tint for generalize-box background
}

TABLE_STYLE = "Table Grid"


# ----------------------------------------------------------------------
# Low-level cell shading / border helpers (python-docx OxmlElement)
# ----------------------------------------------------------------------

def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, *, top=None, left=None, bottom=None, right=None,
                     default=None) -> None:
    """Set per-side borders. Each side is a dict like
    {"val": "single", "sz": 4, "color": "C0392B"} or None.
    If `default` is given, any side not specified gets that spec."""
    tc_pr = cell._tc.get_or_add_tcPr()
    # remove any existing tcBorders so we don't double up
    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    sides = {"top": top, "left": left, "bottom": bottom, "right": right}
    for side, spec in sides.items():
        if spec is None and default is not None:
            spec = default
        if spec is None:
            continue
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   spec.get("val", "single"))
        b.set(qn("w:sz"),    str(spec.get("sz", 4)))
        b.set(qn("w:space"), str(spec.get("space", 0)))
        b.set(qn("w:color"), spec.get("color", "auto"))
        borders.append(b)
    tc_pr.append(borders)


def _run(paragraph, text: str, *, bold=False, color=None, size=None,
         italic=False):
    r = paragraph.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


# ----------------------------------------------------------------------
# Day banner — dark-blue solid block with "Day N | Title" + gray subtitle
# ----------------------------------------------------------------------

def day_banner(doc, day_num: int, title: str, subtitle: str = "") -> None:
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    cell = t.rows[0].cells[0]
    cell.text = ""
    shade_cell(cell, COLORS["dayblue"])
    # thin border so the banner reads as a block
    set_cell_borders(cell, default={"val": "single", "sz": 4,
                                    "color": COLORS["dayblue"]})
    p = cell.paragraphs[0]
    _run(p, f"Day {day_num}  \u2502  {title}",
         bold=True, color=COLORS["white"], size=16)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if subtitle:
        p = doc.add_paragraph()
        _run(p, subtitle, italic=True, size=10, color="555555")
    doc.add_paragraph()


# ----------------------------------------------------------------------
# Teal section — 2-col table with teal-shaded left label, white bold,
# body lines on the right. Replaces the unstyled add_two_col pattern.
# ----------------------------------------------------------------------

def teal_section(doc, label: str, body_lines) -> None:
    if isinstance(body_lines, str):
        body_lines = [body_lines]
    t = doc.add_table(rows=1, cols=2)
    t.style = TABLE_STYLE
    t.columns[0].width = Inches(1.3)
    t.columns[1].width = Inches(5.2)
    left, right = t.rows[0].cells

    shade_cell(left, COLORS["teal"])
    left.text = ""
    p = left.paragraphs[0]
    _run(p, label, bold=True, color=COLORS["white"], size=11)
    left.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    right.text = ""
    for i, line in enumerate(body_lines):
        p = right.paragraphs[0] if i == 0 else right.add_paragraph()
        p.add_run(line)
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    doc.add_paragraph()


# ----------------------------------------------------------------------
# Summary-Exit box — red-bordered block with a titled top-left tab.
# Mirrors the .tex cerbox. Works for CER or for summary-style exits.
# ----------------------------------------------------------------------

def summary_exit_box(doc, title: str, lines) -> None:
    if isinstance(lines, str):
        lines = [lines]
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_borders(cell, default={"val": "single", "sz": 8,
                                    "color": COLORS["warmred"]})
    shade_cell(cell, COLORS["warmredlt"])
    cell.text = ""

    p = cell.paragraphs[0]
    _run(p, title, bold=True, color=COLORS["warmred"], size=11)
    for line in lines:
        cp = cell.add_paragraph()
        cp.add_run(line)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    doc.add_paragraph()


# ----------------------------------------------------------------------
# Sentence-frame callout — light-gray block with a 3-pt teal left accent.
# Mirrors the .tex generalizebox.
# ----------------------------------------------------------------------

def sentence_frame_box(doc, lines) -> None:
    if isinstance(lines, str):
        lines = [lines]
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade_cell(cell, COLORS["lightgray"])
    set_cell_borders(
        cell,
        left={"val": "single", "sz": 24, "color": COLORS["teal"]},
        top={"val": "nil"}, right={"val": "nil"}, bottom={"val": "nil"},
    )
    cell.text = ""
    p = cell.paragraphs[0]
    _run(p, "Sentence frame", bold=True, color=COLORS["teal"], size=10)
    for line in lines:
        cp = cell.add_paragraph()
        cp.add_run(line).italic = True
    doc.add_paragraph()


# ----------------------------------------------------------------------
# Running header — name / block fields on every page.
# ----------------------------------------------------------------------

def running_header_name_block(doc, label_right: str = "BLOCK") -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    # Tab stop to push right label to the right margin.
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5))
    _run(p, "FULL NAME: ________________________", size=9)
    p.add_run("\t")
    _run(p, f"{label_right}: __________", size=9)


# ----------------------------------------------------------------------
# Framework phase header with explicit "Teacher / Students / Questions
# to ask / Adult role" columns per DOKframework.txt.
# Used in teacher packets so evaluators scanning with the framework
# printout find the exact labels.
# ----------------------------------------------------------------------

def framework_phase_header(doc, *, phase: str, framework_tag: str,
                           dok: str, minutes,
                           teacher_does: list, students_do: list,
                           questions_to_ask: list, adult_role: str) -> None:
    title = f"[{framework_tag}]  [DOK {dok}]  {phase}"
    if minutes is not None:
        title += f"  ({minutes} min)"
    teal_section(doc, "PHASE", [title])

    rows = [
        ("Teachers will\u2026",    teacher_does),
        ("Students will\u2026",    students_do),
        ("Questions to ask",       questions_to_ask),
        ("Adult role",             [adult_role]),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = TABLE_STYLE
    t.columns[0].width = Inches(1.3)
    t.columns[1].width = Inches(5.2)
    for i, (label, lines) in enumerate(rows):
        left, right = t.rows[i].cells
        shade_cell(left, COLORS["tealpale"])
        left.text = ""
        _run(left.paragraphs[0], label, bold=True,
             color=COLORS["teal"], size=10)
        right.text = ""
        for j, line in enumerate(lines):
            p = right.paragraphs[0] if j == 0 else right.add_paragraph()
            p.add_run("\u2022 " + line if len(lines) > 1 else line)
    doc.add_paragraph()


# ---------------------------------------------------------------------
# Visuals checklist hook for packet builders.
#
# Call once per builder (after all docx files are saved) to emit a
# sidecar `<Day>_Visuals_Checklist.md` listing every embedded visual
# the packet relies on. Teacher reviews this before print to confirm
# photos/graphs/tables/maps are embedded cleanly (no answer-key leak,
# no multi-item crops).
#
# Usage (e.g. build_day6_packets.py __main__):
#     import qb
#     from packet_styles import emit_visuals_checklist
#     emit_visuals_checklist(_ALL_IDS, "Day_6_Visuals_Checklist.md",
#                            title="Day 6 - Visuals Checklist")
# ---------------------------------------------------------------------

def emit_visuals_checklist(packet_ids, out_path, *, title=None):
    """Write a visuals checklist for the given packet IDs.

    Thin wrapper around qb.write_visuals_checklist() so each builder
    can emit its own sidecar with one import. Prints a confirmation
    line if any visual items were found, else notes that none were.
    """
    import qb
    from pathlib import Path
    qb.write_visuals_checklist(
        packet_ids,
        out_path,
        title=title or f"{Path(out_path).stem} - Visuals Checklist",
    )
    rows = qb.visuals_for(packet_ids)
    n_flag = sum(1 for r in rows if r["needs_cleanup"])
    if rows:
        flag = f" ({n_flag} need cleanup)" if n_flag else ""
        print(f"  + {out_path}: {len(rows)} visual(s){flag}")
    else:
        print(f"  + {out_path}: (no visuals in this packet)")
