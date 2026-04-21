"""Generate Lesson 4-5 Period 1 materials — single-period quick treatment.

Single-period lesson — condensed Klimsara. If 45-min Wed F variant, cut Explore
by 8 min (drop q14 + q25).

NOTE: This is a COMPRESSED lesson — two Examples are modeled in Launch (Ex 1 + Ex 3),
which deviates from the standard Klimsara one-example-only rule. Solving a rational
equation (Ex 1) and identifying extraneous solutions (Ex 3) are two conceptually
distinct moves; compressed 2-example Launch justified.

Design: Solving Rational Equations + DOK-3 chemistry capstone (Topic 4 LEHS
Q#2 + Q#6 + Q#7 prep). This single-period treatment covers clearing
denominators, extraneous-solution checking, and work-rate/mixture applications,
and ends with Practice #33 (chemistry alcohol mixture performance task), which
rehearses Topic 4 LEHS Q#6 (solve rational equation) and Q#7 (work-rate
reasoning) simultaneously.

Phases (55 min base; Period F 65-min):
  0–  5   Do Now    : Model & Discuss — conflicting candidate solutions  [DOK 1]
  5– 18   Launch    : teacher models Example 1 (solve rational eq) +
                      Example 3 (extraneous identification)              [DOK 1-2]  [COMPRESSED]
 18– 50   Explore   : Try-It 1/3 + Practice #10/#14/#25 + ⭐ DOK-3 q33   [DOK 2-3]
 50– 55   Share/Sum : DOK-3 pair presents chemistry argument +
                      bridge to Topic 4 LEHS Q#2/Q#6/Q#7                 [DOK 2]
 55– 58   Exit      : summary recap fill-in                              [DOK 1-2]
 (back)   Optional: Practice #8 (work-rate validity reasoning) — fast-finisher
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_ID   = "4-5-savvas-model-discuss-lesson-4-5-launch"
LAUNCH_IDS  = [
    "4-5-savvas-example-1-lesson-4-5",   # Ex 1: solve rational equation
    "4-5-savvas-example-3-lesson-4-5",   # Ex 3: extraneous solution identification
]
EXPLORE_IDS = [
    "4-5-savvas-try-it-1-lesson-4-5",
    "4-5-savvas-try-it-3-lesson-4-5",
    "4-5-savvas-q10",
    "4-5-savvas-q14",
    "4-5-savvas-q25",
    "4-5-savvas-q33",    # ⭐ DOK-3 spine — chemistry alcohol mixture performance task
]
DOK3_ID       = "4-5-savvas-q33"
REINFORCE_IDS = ["4-5-savvas-q8"]    # work-rate validity reasoning (rehearses Q#7)

_ALL_IDS = [DO_NOW_ID] + LAUNCH_IDS + EXPLORE_IDS + REINFORCE_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 4-5 · Quick"
LESSON_TITLE = "Solving Rational Equations + Chemistry DOK-3 (Q#2/Q#6/Q#7 Prep)"

OBJECTIVES = [
    ("Math Objective",
     "Solve rational equations in one variable by clearing denominators, check "
     "for extraneous solutions, and apply rational equations to real-world "
     "rate and mixture problems."),
    ("Language Objective",
     "Justify why a candidate solution is extraneous, using the domain "
     "restrictions of the original rational equation."),
    ("Essential Understanding",
     "Multiplying both sides of a rational equation by an expression containing "
     "variables can introduce extraneous solutions — values that satisfy the "
     "cleared equation but violate the original domain."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Fluency solving rational equations; identifying extraneous solutions; "
     "applying to work-rate and mixture problems (Topic 4 LEHS Q#2, Q#6, Q#7 prep)."),
    ("Essential Question",
     "When you solve a rational equation by clearing denominators, why must you "
     "check every candidate solution against the original equation?"),
    ("Materials",
     "Student packet, pencil, calculator. DOK-3 Practice #33 (chemistry mixture "
     "performance task) rehearses Topic 4 LEHS Q#6 (solve rational equation) "
     "and Q#7 (work-rate reasoning)."),
]


# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.0):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    ps.render_prompt(prompt_p, q["prompt"], font_size=11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            run = line.add_run(f"  ({chr(65+i)})  " + ps.latex_to_unicode(a))
            run.font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, for_teacher=False):
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, LESSON_TITLE,
                  subtitle=LESSON_LABEL + (" · Teacher Edition" if for_teacher else ""))
    t = doc.add_table(rows=len(OBJECTIVES), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(OBJECTIVES):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    t = doc.add_table(rows=len(FRAMEWORK_HEADER), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(FRAMEWORK_HEADER):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()


# ── Do Now ───────────────────────────────────────────────────────────────

def build_do_now(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, "Do Now — " + LESSON_TITLE, subtitle=LESSON_LABEL)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from 4-4 to solving rational equations",
        dok="1",
        minutes=5,
        teacher_does=[
            "Hand out at door.",
            "Circulate 3 min.",
            "Cold-call 1 student: \"Two students solved a rational equation and "
            "got different candidate solutions. How can we decide which is valid?\"",
        ],
        students_do=[
            "Read the Model & Discuss prompt (two students solving a rational "
            "equation, conflicting candidates).",
            "Decide whether either candidate works and why.",
        ],
        questions_to_ask=[
            "If a candidate solution makes the ORIGINAL equation's denominator "
            "equal to zero, is it a real solution?",
            "Why might clearing denominators produce a value that doesn't work?",
        ],
        adult_role="Solo teacher: circulate, time 5 min, pull sheets.",
    )

    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label="Do Now — Model & Discuss: conflicting candidate solutions",
                    space_after_inches=2.0)

    add_callout(doc, "\U0001f3af  SENTENCE FRAME",
                ["“Student A got candidate x = ___ and Student B got x = ___. "
                 "When I check in the ORIGINAL equation, candidate ___ makes a "
                 "denominator zero, so it is extraneous. The valid solution is ___.”"],
                color_hex="FFF2CC")

    doc.save(path)


# ── Student packet ───────────────────────────────────────────────────────

def build_student(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=False)

    ps.teal_section(doc, "LAUNCH — Solving Rational Equations "
                         "(teacher models Ex 1 + Ex 3)", [])
    add_callout(doc, "\U0001f4d8  RULES — SOLVING RATIONAL EQUATIONS "
                     "(keep printed on your page)",
                [
                    "1. To solve a rational equation:",
                    "   (a) find the LCD of all fractions,",
                    "   (b) multiply BOTH SIDES by the LCD to clear fractions,",
                    "   (c) solve the resulting polynomial equation,",
                    "   (d) CHECK each candidate in the ORIGINAL equation — any "
                    "candidate that makes a denominator zero is extraneous.",
                    "2. Domain restrictions are set BEFORE clearing denominators. "
                    "A candidate that violates a restriction is extraneous even if "
                    "it algebraically satisfies the cleared equation.",
                    "3. Work-rate setup: if person A does a job in a hours and "
                    "person B in b hours, working together they complete 1/a + 1/b "
                    "of the job per hour, so combined time t satisfies "
                    "1/a + 1/b = 1/t.",
                    "4. Mixture setup: concentration × volume = amount of substance. "
                    "Build one fraction per mixture stream, then set up a rational "
                    "equation for the target concentration.",
                ],
                color_hex="D9EAD3")

    add_callout(doc, "⚠️  COMPRESSED LESSON — TWO EXAMPLES IN LAUNCH",
                [
                    "This is a single-period quick treatment. Teacher models BOTH "
                    "Example 1 (solve a rational equation) AND Example 3 "
                    "(identify the extraneous solution) during Launch.",
                    "Pay attention to Example 3 — the extraneous solution trap "
                    "is the #1 error on this lesson's assessment items. Always "
                    "check every candidate in the ORIGINAL equation.",
                ],
                color_hex="F4CCCC")
    doc.add_paragraph()

    ps.teal_section(doc, "EXPLORE — Think-Pair-Share (32 min)",
                    ["Work with a partner. Move in order. Practice #33 "
                     "(Chemistry Mixture) is the big one — plan ~10 min for it.",
                     "45-min Wed F variant: skip #14 and #25."])

    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 1 — Solve a rational equation",
        "Try It 3 — Identify extraneous solution",
        "Practice #10 — Solve + check",
        "Practice #14 — Extraneous trap",
        "Practice #25 — Work-rate application",
        "Practice #33  ⭐ DOK-3 CHEMISTRY MIXTURE — Topic 4 LEHS Q#6 + Q#7 prep",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=label, space_after_inches=1.5)

    ps.teal_section(doc, "SHARE / SUMMARY (5 min)", [])
    ps.sentence_frame_box(doc, [
        "Pair share (Chemistry): “I set up f(x) = (50·0.02 + 0.06x)/(50+x) = 0.05. "
        "After clearing the denominator I got ___. Solving for x gave x = ___ gallons. "
        "I verified with a calculator TABLE by entering y₁ = ___ and y₂ = ___.”",
        "Whole class: one pair reads their chemistry solution. Class signals "
        "thumbs up/sideways.",
    ])

    add_callout(doc, "\U0001f52d  BRIDGE TO TOPIC 4 LEHS Q#2, Q#6, Q#7",
                ["The chemistry move — set up rational equation, clear denominators, "
                 "solve, check — is exactly what Topic 4 LEHS Q#6 asks. The "
                 "domain-restriction reasoning is Q#2. The rate reasoning is Q#7. "
                 "You have now rehearsed all three."],
                color_hex="FCE5CD")

    ps.summary_exit_box(doc, "EXIT TICKET — Summary Recap",
                        [
                            "After clearing denominators in a rational equation, I must "
                            "check each candidate solution in the _____ equation to catch "
                            "_____ solutions.",
                            "In Practice #33 (chemistry), the equation came from setting "
                            "_____ equal to the target concentration 0.05.",
                            "One biggest thing I learned today: ________________________",
                        ])

    doc.add_page_break()
    ps.teal_section(doc, "OPTIONAL REINFORCEMENT (not collected)",
                    ["If you finish Explore early. #8 is a work-rate validity "
                     "reasoning item that rehearses Topic 4 LEHS Q#7 specifically."])
    reinforce = qb.get_for_packet(REINFORCE_IDS)
    for i, q in enumerate(reinforce, 1):
        bank_item_block(doc, q,
                        label=f"Optional #{i}  (Savvas Practice #8 — work-rate validity reasoning)",
                        space_after_inches=1.8)

    doc.save(path)


# ── Teacher packet ───────────────────────────────────────────────────────

def build_teacher(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=True)

    add_callout(doc, "\U0001f4cc  PRE-FLIGHT — SINGLE-PERIOD COMPRESSED LESSON + DOK-3 CAPSTONE",
                [
                    "This is a single-period quick treatment of Lesson 4-5. TWO "
                    "examples are modeled in Launch (deviation from standard Klimsara "
                    "one-example rule — solving AND extraneous checking are both "
                    "assessment-critical). Release promptly at minute 18.",
                    "Today's capstone is Practice #33 (chemistry alcohol mixture). "
                    "Students must (a) set up f(x) = (50·0.02 + 0.06x)/(50+x) = 0.05, "
                    "(b) clear the denominator and solve — 50·0.02 + 0.06x = 0.05(50+x) "
                    "→ 1 + 0.06x = 2.5 + 0.05x → 0.01x = 1.5 → x = 150 gallons, "
                    "(c) explain calculator TABLE verification. Same algebraic move "
                    "as Topic 4 LEHS Q#6; same reasoning as Q#7.",
                    "Pace: Try-It 1 (4 min) → Try-It 3 (4 min) → #10 (4 min) → "
                    "#14 (4 min) → #25 (6 min) → #33 (10 min). ~32 min total Explore.",
                    "45-min Wed F variant: cut #14 + #25. Never cut #33.",
                ],
                color_hex="FFD9E0")

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from 4-4 to solving rational equations",
        dok="1",
        minutes=5,
        teacher_does=[
            "Hand out at door.",
            "Circulate 3 min.",
            "Cold-call 1 student on the conflicting candidate solutions.",
        ],
        students_do=[
            "Read and decide whether either candidate is valid.",
            "Identify which (if any) is extraneous.",
        ],
        questions_to_ask=[
            "What does it mean for a candidate solution to be 'extraneous'?",
            "Why do we check candidates in the ORIGINAL equation, not the cleared one?",
        ],
        adult_role="Solo teacher: circulate silently. No hints.",
    )
    add_callout(doc, "\U0001f3a4  BRIDGE SCRIPT (read aloud after Do Now)",
                ["\"Topic 4 so far: inverse variation, reciprocal functions, "
                 "multiplying and adding rational expressions. Today: SOLVING "
                 "rational EQUATIONS. Same algebra, but with an equation to "
                 "solve and a new trap — extraneous solutions. The Model & "
                 "Discuss shows two students solving a rational equation and "
                 "getting different candidate solutions. Which (if any) is valid?\""],
                color_hex="CFE2F3")
    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label=f"Do Now ({q['id']})", space_after_inches=0.3)
    add_callout(doc, "✅  ANSWER KEY",
                [q.get("notes") or "(verify)"],
                color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Launch",
        framework_tag="teacher models Example 1 + Example 3 [COMPRESSED — 2 examples]",
        dok="1-2",
        minutes=13,
        teacher_does=[
            "Project Example 1 (solve rational equation). Think aloud: "
            "find LCD, multiply both sides, solve polynomial, check in original.",
            "\"Every candidate must be checked in the ORIGINAL equation — clearing "
            "denominators can introduce values that don't belong.\"",
            "Transition to Example 3 without stopping — announce: "
            "\"Now watch what happens when a candidate IS extraneous.\"",
            "Project Example 3. Think aloud: state domain restrictions FIRST, "
            "then solve; test each candidate against the restrictions.",
            "Flag the Rules card — students copy into margin.",
            "30-sec pause: \"Turn to partner: why does clearing denominators "
            "sometimes produce extraneous solutions? What's the mechanism?\"",
            "Do NOT solve Try-Its or Practice items — release at minute 18.",
        ],
        students_do=[
            "Watch Example 1 silently. Note the four-step procedure.",
            "Watch Example 3. Mark the domain-restriction identification step.",
            "During pause: explain mechanism to partner using Rule 1(d).",
            "Copy Rules card into margin.",
        ],
        questions_to_ask=[
            "What is the LCD in Example 1?",
            "After you clear denominators, what kind of equation do you get?",
            "In Example 3, what values are forbidden by the original denominators?",
            "Why does multiplying by an expression containing variables sometimes "
            "introduce extraneous roots?",
        ],
        adult_role="Project both examples back-to-back. Release promptly at minute 18.",
    )
    example_items = qb.get_for_packet(LAUNCH_IDS)
    add_callout(doc, "\U0001f3a4  TEACHER SCRIPT",
                [
                    "1. \"Watch me solve a rational equation. Step 1: find the LCD.\"",
                    "2. \"Step 2: Multiply BOTH SIDES by the LCD. This clears all fractions.\"",
                    "3. \"Step 3: Solve the resulting polynomial equation — you'll get "
                    "candidate solutions.\"",
                    "4. \"Step 4: CHECK each candidate in the ORIGINAL equation. Any "
                    "candidate that makes a denominator zero is extraneous — throw it out.\"",
                    "5. \"Example 3: watch what happens when a candidate IS extraneous. "
                    "I state the domain restrictions FIRST — before I even start "
                    "solving — so I know what's forbidden.\"",
                    "6. \"Today's DOK-3 task is a chemistry alcohol mixture. Same "
                    "procedure — set up rational equation, clear, solve, check. "
                    "Same algebraic move as Topic 4 LEHS Q#6.\"",
                    "7. Release to Explore.",
                ],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Explore",
        framework_tag="TPS + DOK-3 driver",
        dok="2-3",
        minutes=32,
        teacher_does=[
            "5 laps across 32 min.",
            "Lap 1 (4 min): Try-It 1 — solve. Press pairs on the LCD identification.",
            "Lap 2 (4 min): Try-It 3 — extraneous. Watch for students who skip the check.",
            "Lap 3 (4 min): #10 — solve + check. Confirm the check step.",
            "Lap 4 (4 min): #14 — extraneous trap. Press \"which candidate violates "
            "the original domain?\" on every pair.",
            "Lap 5 (6 min): #25 — work-rate. Use the 1/a + 1/b = 1/t setup.",
            "Lap 6 (10 min): ⭐ #33 Chemistry. Students set up the mixture equation, "
            "clear, solve for x, and explain TABLE verification.",
            "Do NOT give #33 Part A numeric answer. Pairs present argument at Share.",
        ],
        students_do=[
            "6 items in order: Try-It 1 → Try-It 3 → #10 → #14 → #25 → #33.",
            "For #33: BEFORE computing, write out f(x) = (50·0.02 + 0.06x)/(50+x) "
            "and set equal to 0.05. Ask: what would x mean in gallons?",
            "Justify TABLE verification with sentence frame.",
        ],
        questions_to_ask=[
            "Try-It 1: what's the LCD of the fractions?",
            "Try-It 3: which candidate makes the original denominator zero?",
            "#10: after you solve, WHY do you still need to check?",
            "#14: state the domain restrictions BEFORE you clear denominators.",
            "#25: how does 1/a + 1/b = 1/t capture the work-rate idea?",
            "#33 Part A: set up the mixture equation. What does x represent "
            "in gallons — does your answer make sense?",
            "#33 Part B: how do you use y₁ and y₂ in the TABLE to verify?",
        ],
        adult_role="Solo teacher: 6 laps. On #33, press the mixture-equation setup "
                   "AND the gallons reasonableness check — those are the "
                   "assessment-critical moves.",
    )
    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 1", "Try It 3", "Practice #10", "Practice #14",
        "Practice #25",
        "Practice #33 ⭐ DOK-3 CHEMISTRY MIXTURE",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=f"{label} ({q['id']})", space_after_inches=0.3)
        if "DOK-3" in label or "#33" in label:
            add_callout(doc, "⭐ DOK-3 CHEMISTRY MIXTURE — Practice #33 "
                             "(Topic 4 LEHS Q#6 + Q#7 prep)",
                        [
                            "q33 is the chemistry alcohol-mixture performance task: "
                            "50 gal of 2% alcohol solution, add x gal of 6% alcohol, "
                            "resulting concentration f(x) = (50·0.02 + 0.06x)/(50+x). "
                            "Part A: solve f(x) = 0.05 for x. Part B: explain how to "
                            "verify with a graphing calculator table.",
                            "\"The move Part A wants is: set up the rational equation, "
                            "clear the denominator, solve the resulting linear equation "
                            "for x. Students should recognize this as a rational equation "
                            "in disguise — same moves as Ex 1.\"",
                            "For Part B: enter y₁ = numerator, y₂ = 0.05(50+x), use "
                            "TABLE to find where they intersect.",
                            "If stuck: ask 'what would the answer x mean in gallons?' "
                            "before they compute, to anchor the reasonableness check.",
                            "Answer: 150 gallons.",
                            "Sentence frame required for ELL.",
                        ],
                        color_hex="FFD9E0")
        else:
            add_callout(doc, "✅  ANSWER / ANTICIPATED TALK",
                        [q.get("notes") or "(no answer key — verify before class)"],
                        color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Share / Summary",
        framework_tag="academic conversation + Topic 4 LEHS Q#2/Q#6/Q#7 bridge",
        dok="2",
        minutes=5,
        teacher_does=[
            "Cold-call 1 pair to present their chemistry solution using sentence frame.",
            "Class signals thumbs. Write x = 150 gallons on board.",
            "Topic 4 bridge: \"This is the exact move on Topic 4 LEHS Q#6, the "
            "domain reasoning is Q#2, and the rate reasoning is Q#7. You have "
            "now rehearsed all three.\"",
        ],
        students_do=[
            "Presenting pair reads their solution and cites the TABLE verification.",
            "Rest of class signals and writes takeaway in margin.",
        ],
        questions_to_ask=[
            "Summarize: how did you get from f(x) = 0.05 to x = 150?",
            "Why did 150 gallons make sense as an answer?",
        ],
        adult_role="Cold-call a pair that struggled on #33 — hold them accountable "
                   "to the setup step.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit Ticket",
        framework_tag="summary recap (not graded)",
        dok="1-2",
        minutes=3,
        teacher_does=[
            "Collect at door.",
            "Scan stem 1 (original / extraneous). Plan re-teach if many students "
            "can't articulate WHY the check is required.",
        ],
        students_do=["3 sentence stems, honest."],
        questions_to_ask=[
            "Did students connect the 'check in original' step to avoiding "
            "extraneous solutions?",
            "Did students correctly describe the mixture-equation setup?",
        ],
        adult_role="Collect. No grade.",
    )
    add_callout(doc, "\U0001f4cb  EXIT TICKET — Student Form",
                [
                    "After clearing denominators in a rational equation, I must "
                    "check each candidate solution in the _____ equation to catch "
                    "_____ solutions.",
                    "In Practice #33 (chemistry), the equation came from setting "
                    "_____ equal to the target concentration 0.05.",
                    "One biggest thing I learned today: ________________________",
                ],
                color_hex="FFF2CC")

    doc.add_page_break()
    add_callout(doc, "\U0001f550  PERIOD A / PERIOD F — 65-MIN CLASS NOTES",
                [
                    "Both sections should have 65 min for this lesson. Use the "
                    "extra 10 min on Explore #33 (Chemistry) — richer "
                    "reasonableness and TABLE-verification talk pays off on "
                    "Topic 4 LEHS Q#6/Q#7.",
                    "45-min Wed F variant: cut #14 + #25 from Explore. Never cut #33.",
                    "If finished early: hand out Practice #8 (work-rate validity "
                    "reasoning) as fast-finisher; rehearses Q#7 specifically.",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "\U0001f9e9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS",
                [
                    "• Provide the Solving Rational Equations Rules card as a "
                    "sticker/cut-out.",
                    "• For #33, provide the setup: \"f(x) = (50·0.02 + 0.06x)/(50+x) "
                    "= 0.05. Multiply both sides by (50+x).\" Students solve from there.",
                    "• Accept verbal TABLE-verification explanation in lieu of "
                    "written steps.",
                ],
                color_hex="E2F0D9")
    add_callout(doc, "\U0001f30d  SUPPORTS FOR ELL STUDENTS",
                [
                    "• Sentence frames on student page (don’t skip).",
                    "• Word cards: \"extraneous\" = an answer that looks right "
                    "algebraically but doesn't actually work in the original "
                    "equation; \"LCD\" = least common denominator; \"domain "
                    "restriction\" = a value that x is not allowed to be.",
                    "• \"Concentration\" = amount of substance per unit of mixture; "
                    "\"mixture\" = two solutions combined.",
                    "• Pair ELL students with English-dominant partners for TPS.",
                ],
                color_hex="DEEBF7")

    doc.save(path)


# ── Main ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_do_now("L45_P1_Do_Now.docx")
    build_student("L45_P1_Student_Packet.docx")
    build_teacher("L45_P1_Teacher_Packet.docx")
    print("Built L45_P1_Do_Now.docx, L45_P1_Student_Packet.docx, L45_P1_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L45_P1_Visuals_Checklist.md",
                              title="L45 P1 — Visuals Checklist")
