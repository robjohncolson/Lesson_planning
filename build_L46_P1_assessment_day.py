"""Generate Lesson 4-6 Period 1 (Topic 4 LEHS Assessment Day) wrapper materials.

The Topic 4 assessment itself is `a2topic4assess.docx` (not generated here).
LEHS uses 8 items from that document: Q#2, #3, #4, #5, #6, #7, #15, #19.

This script produces the surrounding class-period artifacts:

  - L46_P1_Do_Now.docx          — short 5-min review (4-5 Try-It 1, solve rational eq)
  - L46_P1_Student_Packet.docx  — cover sheet + policy + formula reference + tracker
  - L46_P1_Teacher_Packet.docx  — entry/timing/interventions/debrief
  - L46_P1_Visuals_Checklist.md — sidecar (likely empty; Do Now item has none)

Phases (55-min base period):
  0– 5  Entry / Do Now (4-5 Try-It 1, solve a rational equation)      [DOK 2]
  5–50  Assessment (8 items)                                          [varies]
 50–55  Early-finisher reflection prompt                              [DOK 2]
 55–58  Exit (turn in + 2-item reflection slip)                       [DOK 1]
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_IDS = ["4-5-savvas-try-it-1-lesson-4-5"]
_ALL_IDS = DO_NOW_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 4-6 · Period 1"
LESSON_TITLE = "Topic 4 Assessment Day (LEHS 8-Item)"

# 8-item LEHS subset → source-lesson map, with stems paraphrased from
# a2topic4assess.docx (for teacher reference only; NEVER surfaced to students).
ITEM_MAP = [
    ("Q#2",  "4-5", "Domain of a rational function (exclude zeros of the denominator).",
              "If stuck: factor the denominator; which x-values make it zero?"),
    ("Q#3",  "4-1", "Horizontal and vertical asymptotes of a rational/reciprocal graph.",
              "If stuck: H-asymptote from leading-coefficient ratio; V-asymptote from denominator = 0."),
    ("Q#4",  "4-3", "Simplify a rational expression AND state domain restrictions.",
              "If stuck: factor numerator and denominator; cancel; original-denominator zeros are still excluded."),
    ("Q#5",  "4-1", "Describe the translation from y = 1/x to a shifted reciprocal.",
              "If stuck: (x − h) shifts RIGHT h; + k shifts UP k."),
    ("Q#6",  "4-5", "Solve a rational equation.",
              "If stuck: clear denominators (multiply by LCD); check for extraneous solutions."),
    ("Q#7",  "4-5", "Work-rate word problem (Luke + Nora peeling carrots).",
              "If stuck: set up 1/a + 1/b = 1/t; convert rates to a common unit first."),
    ("Q#15", "4-4", "Evaluate/operate on rational expressions (closure-style).",
              "If stuck: common denominator FIRST, then combine numerators."),
    ("Q#19", "4-4", "Polynomial division — find length given area and width of a rectangle.",
              "If stuck: length = area ÷ width; use polynomial long division or synthetic division."),
]


def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.2):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    ps.render_prompt(prompt_p, q["prompt"], font_size=11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            run = line.add_run(f"  ({chr(65+i)})  " + ps.latex_to_unicode(a))
            run.font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


# ------------------------------------------------------------
# Do Now — 4-5 Try-It 1 (solve rational equation), 5 min
# ------------------------------------------------------------

def build_do_now(path):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.75); s.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, "Do Now — Quick Warm-Up (5 min)",
                  subtitle=LESSON_LABEL + " · Topic 4 Assessment")

    add_callout(doc, "🌤️  BEFORE THE ASSESSMENT",
                ["This is a 5-minute warm-up — NOT the assessment. "
                 "Solve cleanly. Check for extraneous solutions. "
                 "When the timer hits 0, clear your desk for the test."],
                color_hex="FFF2CC")

    for i, qid in enumerate(DO_NOW_IDS, 1):
        q = qb.get_for_packet([qid])[0]
        bank_item_block(doc, q, label=f"Warm-Up #{i}", space_after_inches=2.0)

    add_callout(doc, "📋  QUICK REMINDERS",
                ["• Clear denominators (multiply both sides by the LCD).",
                 "• Solve.",
                 "• Plug each solution back — discard any that make an original denominator 0."],
                color_hex="D9EAD3")

    doc.save(path)


# ------------------------------------------------------------
# Student Packet — cover sheet, policy, formula ref, checklist.
# No assessment items. Assessment lives in a2topic4assess.docx.
# ------------------------------------------------------------

def build_student_packet(path):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.75); s.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, "Topic 4 Assessment — Cover Sheet",
                  subtitle=LESSON_LABEL + " · Rational Functions & Inverse Variation")

    # Info block
    info = doc.add_table(rows=3, cols=2)
    info.style = TABLE_STYLE
    info.columns[0].width = Inches(1.6)
    info.columns[1].width = Inches(4.9)
    for i, (label, blank) in enumerate([
        ("NAME",   "_________________________________________________"),
        ("DATE",   "_________________________________________________"),
        ("PERIOD", "  [  A  ]     [  F  ]"),
    ]):
        left, right = info.rows[i].cells
        ps.shade_cell(left, "CFE2F3")
        set_cell(left,  label,  bold_first=True)
        set_cell(right, blank)
    doc.add_paragraph()

    add_callout(doc, "⏰  TIME POLICY (55-min period)",
                ["• 5 min  Do Now warm-up (on separate sheet)",
                 "• 45 min Assessment — 8 questions",
                 "• 5 min  Early-finisher reflection (if you finish before time)",
                 "• 3 min  Exit slip (turn in test + exit)",
                 "If you are stuck on an item > 3 min, circle it and move on. Come back if time."],
                color_hex="FFF2CC")

    add_callout(doc, "🧮  CALCULATOR POLICY",
                ["• Graphing calculator: ALLOWED",
                 "• Phone / smart-watch: NOT ALLOWED (place in caddy)",
                 "• Scratch paper: allowed (hand in with test)"],
                color_hex="D9EAD3")

    # Formula reference — ELL support
    add_callout(doc, "📘  FORMULA REFERENCE (you MAY use this during the test)",
                ["• Reciprocal / parent function:  f(x) = 1/x",
                 "• Translated reciprocal:         f(x) = 1/(x − h) + k   "
                     "(right h, up k)",
                 "• Inverse variation:             xy = k   ⇔   y = k/x",
                 "• Work-rate (two workers, shared job):   1/a + 1/b = 1/t",
                 "• Domain restriction rule:       Exclude any x that makes a "
                     "denominator = 0 (both before AND after simplifying).",
                 "• Asymptotes of f(x) = (ax+b)/(cx+d):  "
                     "V at cx+d = 0; H at y = a/c."],
                color_hex="CFE2F3")

    # Self-tracking checklist
    add_callout(doc, "☑️  ITEM TRACKER (check off as you finish each)",
                ["Q2 ☐    Q3 ☐    Q4 ☐    Q5 ☐    Q6 ☐    Q7 ☐    Q15 ☐    Q19 ☐",
                 "",
                 "Items you SKIPPED and want to come back to (write numbers):",
                 "______________________________________________________________"],
                color_hex="DEEBF7")

    add_callout(doc, "📣  ONE-SENTENCE PRE-TEST FOCUS",
                ["Which of the 4 lesson topics do you feel STRONGEST on right now? "
                 "(4-1 reciprocals · 4-3 simplifying · 4-4 operations/division · "
                 "4-5 solving / work-rate)",
                 "",
                 "I feel strongest on: ____________________________________________",
                 "",
                 "(This is just a focusing tool — not graded, not seen by anyone.)"],
                color_hex="FFF2CC")

    doc.save(path)


# ------------------------------------------------------------
# Teacher Packet — entry routine, minute-by-minute plan,
# per-item intervention cues, post-assessment debrief plan.
# ------------------------------------------------------------

def build_teacher_packet(path):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.5); s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)
    ps.day_banner(doc, 1, LESSON_TITLE,
                  subtitle=LESSON_LABEL + " · Teacher Packet")

    add_callout(doc, "🚪  ENTRY ROUTINE (pre-bell → 0:00)",
                ["• At the door: hand each student (1) Do Now sheet, (2) cover sheet, "
                     "(3) assessment packet (FACE DOWN — no peeking).",
                 "• Phones into caddy. Graphing calculators on desks.",
                 "• \"You have 5 minutes on the Do Now. The assessment starts when the timer hits 5.\"",
                 "• Start the visible class timer at 5:00 counting down."],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Entry / Do Now (4-5 Try-It 1)",
        framework_tag="bridge warm-up",
        dok="2",
        minutes=5,
        teacher_does=[
            "Circulate once. Answer no content questions.",
            "At 4:30: \"30 seconds. Finish your current step.\"",
            "At 5:00: \"Pencils down. Flip the assessment. Begin.\"",
        ],
        students_do=["Solve the Do Now rational equation.",
                     "Check for extraneous roots.",
                     "Stop when the timer hits 0."],
        questions_to_ask=["(None — priming, not teaching.)"],
        adult_role="Solo teacher: calm, low-talk. Do NOT re-teach.",
    )

    ps.framework_phase_header(
        doc,
        phase="Assessment — 8 items (Q#2, 3, 4, 5, 6, 7, 15, 19)",
        framework_tag="proctoring",
        dok="1–3 mix",
        minutes=45,
        teacher_does=[
            "Silent circulate every 10 min.",
            "At minute 20:  \"25 min remaining.\"",
            "At minute 35:  \"10 min remaining. Start moving to items you skipped.\"",
            "At minute 43:  \"2 minutes. Finalize.\"",
            "At minute 45:  \"Pencils down. Pass packets to the front — face down.\"",
            "Collect assessment packets. Check that each has a name.",
        ],
        students_do=[
            "Work independently. No partner talk.",
            "If stuck > 3 min on one item, circle and move on.",
            "Check your item tracker on the cover sheet as you go.",
        ],
        questions_to_ask=["(Silent proctoring — no content help of any kind.)"],
        adult_role="Solo teacher: proctor only. Redirect off-task quietly.",
    )

    # Item-by-item intervention cues — for teacher awareness ONLY.
    rows = [["Q#", "Lesson", "Stem (teacher reference)", "If-stuck cue (DO NOT give out)"]]
    for qlabel, lesson, stem, cue in ITEM_MAP:
        rows.append([qlabel, lesson, stem, cue])
    t = doc.add_table(rows=len(rows), cols=4)
    t.style = TABLE_STYLE
    t.columns[0].width = Inches(0.5)
    t.columns[1].width = Inches(0.7)
    t.columns[2].width = Inches(3.2)
    t.columns[3].width = Inches(2.5)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
                ps.shade_cell(cell, "CFE2F3")
    doc.add_paragraph()

    add_callout(doc, "🔇  INTERVENTION PROTOCOL (CRITICAL)",
                ["These cues are for YOUR awareness only. DO NOT say them during the test.",
                 "If a student freezes completely, the allowed response is: "
                     "\"Skip it, come back if time.\"",
                 "You may NOT hint, rephrase, confirm work, or discuss math. "
                     "This is a graded summative assessment."],
                color_hex="F7D7D7")

    ps.framework_phase_header(
        doc,
        phase="Early-Finisher Reflection",
        framework_tag="buffer / silent",
        dok="2",
        minutes=5,
        teacher_does=[
            "Hand early finishers a half-sheet with the prompt below.",
            "Finishers may start the prompt silently while others work.",
            "If class is still writing at minute 50, skip to Exit at 50.",
        ],
        students_do=[
            "Look back at the 8 items. Pick ONE you were unsure about.",
            "Write one sentence: what strategy did you try, and why?",
        ],
        questions_to_ask=["(Silent.)"],
        adult_role="Solo teacher: collect assessments from finishers as they go.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit",
        framework_tag="self-reflection",
        dok="1",
        minutes=3,
        teacher_does=[
            "Collect any remaining assessments. Confirm name on each.",
            "Hand out exit slip (or have students write on the back of cover sheet).",
            "Exit question (read aloud): "
                "\"(1) Which item felt hardest and why? "
                "(2) Which topic do you want to review tomorrow — "
                "4-1, 4-3, 4-4, or 4-5?\"",
            "Collect at the door.",
        ],
        students_do=["Answer the two exit questions honestly.",
                     "Turn in at the door."],
        questions_to_ask=["Which item felt hardest?",
                          "Which topic to review tomorrow?"],
        adult_role="Solo teacher: warm tone regardless of class vibe. "
                   "Do NOT give away answers.",
    )

    add_callout(doc, "📊  POST-ASSESSMENT DEBRIEF — RE-TEACH TRIGGERS",
                ["Tonight, tally misses per item. If > 30% of the class missed an item, "
                     "plan the next-day Do Now around that item's source lesson:",
                 "• Q#2 / Q#6 / Q#7 miss > 30% ⇒ re-teach 4-5 "
                     "(solving / work-rate). Biggest re-teach candidate.",
                 "• Q#3 / Q#5 miss > 30% ⇒ re-teach 4-1 "
                     "(reciprocal translations, H/V asymptotes).",
                 "• Q#4 miss > 30% ⇒ re-teach 4-3 "
                     "(simplifying + domain restrictions). Watch for the "
                     "\"cancel BUT original zero still excluded\" trap.",
                 "• Q#15 / Q#19 miss > 30% ⇒ re-teach 4-4 "
                     "(operations + polynomial division).",
                 "Cross-check with exit-slip \"topic to review\" tallies. "
                     "The two signals should agree — if they diverge, "
                     "student self-perception is off."],
                color_hex="DEEBF7")

    add_callout(doc, "🕐  PERIOD F (65-MIN) NOTE",
                ["Period F has 10 extra minutes. Use:",
                 "• 5 min extra on early-finisher reflection (do not extend assessment).",
                 "• 5 min soft preview: \"Next up: Topic 5 — trigonometric functions. "
                     "Think: what's a function that repeats forever?\"",
                 "Do NOT reteach assessment content. Period A doesn't get this extension."],
                color_hex="FFD9E0")

    doc.save(path)


if __name__ == "__main__":
    build_do_now("L46_P1_Do_Now.docx")
    build_student_packet("L46_P1_Student_Packet.docx")
    build_teacher_packet("L46_P1_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L46_P1_Visuals_Checklist.md",
                              title="L46 P1 (Topic 4 Assessment Day) — Visuals Checklist")
    print("Built L46_P1_Do_Now.docx, L46_P1_Student_Packet.docx, "
          "L46_P1_Teacher_Packet.docx, L46_P1_Visuals_Checklist.md")
