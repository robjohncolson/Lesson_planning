"""Generate Lesson 5-4 Period 3 materials — Klimsara-adapted, no Blooket.

Design: Radical Equations — Solve for Variable + Modeling = the
**assessment-critical** period. Topic 5 Performance Assessment Items 1A/1B
(pyramid volume → solve for x → evaluate) and Item 2 (table from rewritten
formula) are directly rehearsed by Practice #44 and Practice #10.
Ex 2 (Rewrite a Formula — Golden Gate cables) and q44 (table-completion)
are paired in Launch. P3 is pure DOK-2 mastery (no DOK-3 driver — P2
owned the spine).
Students leave P3 able to isolate any variable from a radical/
rational-exponent formula and evaluate the rewritten formula to complete
a table or answer a modeling question.

Phases (55 min base; Period F 65-min):
  0–  5   Do Now    : Practice #19 — Higher Order Thinking (extraneous sol'ns) [DOK 3]
  5– 17   Launch    : teacher models Example 2 (rewrite formula) + q44 (table)  [DOK 2]
 17– 52   Explore   : Try-It 2 + #10 + #25 + #26 + #27 + #39                  [DOK 2]
                     (#10 is Item 1A rehearsal; #44 paired in Launch is Item 2)
 52– 57   Share/Sum : Concept Summary return + Topic 5 Performance Task preview  [DOK 2]
 57– 58   Exit      : pre-assessment confidence check                           [DOK 1-2]
 (back)   Optional  : Practice #46 (Escape Velocity) — extra stretch
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_ID   = "5-4-savvas-q19"
LAUNCH_IDS  = [
    "5-4-savvas-example-2-lesson-5-4",   # Ex 2 — Rewrite a Formula (Golden Gate cables)
    "5-4-savvas-q44",                    # ⭐ ASSESSMENT REHEARSAL: table-completion (Assessment Item 2)
]
EXPLORE_IDS = [
    "5-4-savvas-try-it-2-lesson-5-4",
    "5-4-savvas-q10",                    # rewrite to isolate x — direct Item 1A rehearsal
    "5-4-savvas-q25",
    "5-4-savvas-q26",
    "5-4-savvas-q27",
    "5-4-savvas-q39",                    # BSA application
]
REINFORCE_IDS = ["5-4-savvas-q46"]

_ALL_IDS = [DO_NOW_ID] + LAUNCH_IDS + EXPLORE_IDS + REINFORCE_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 5-4 · Period 3"
LESSON_TITLE = "Radical Equations — Assessment-Critical: Solve for Variable + Modeling"

OBJECTIVES = [
    ("Math Objective",
     "Rewrite radical/rational-exponent equations to isolate a specified variable; "
     "evaluate such rewritten formulas at given values to produce tables and graphs."),
    ("Language Objective",
     "Explain in writing which variable is being isolated and why, using a "
     "formula’s context (volume, BSA, speed)."),
    ("Essential Understanding",
     "Solving a formula for a new variable is the SAME algebra as solving a radical "
     "equation — isolate the radical/rational-exponent term, raise to the reciprocal "
     "power, clean up. Works for both pure-algebra and real-world cases."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Rewrite formulas to isolate a specified variable (the Topic 5 assessment’s "
     "core skill); evaluate rewritten formulas at given V/x values."),
    ("Essential Question",
     "How do we extract ANY variable from a radical/rational-exponent formula, and "
     "why does the same procedure work for both algebra and real-world contexts?"),
    ("Materials",
     "Student packet, pencil, calculator for table evaluation. Practice #44 is the "
     "direct assessment rehearsal — treat it like the real test."),
]


# ── Helpers ────────────────────────────────────────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.0):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    ps.render_prompt(prompt_p, q["prompt"], font_size=11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            run = line.add_run(f"  ({chr(65+i)})  " + ps.latex_to_unicode(a))
            run.font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, for_teacher=False):
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 3, LESSON_TITLE,
                  subtitle=LESSON_LABEL + (" · Teacher Edition" if for_teacher else ""))
    t = doc.add_table(rows=len(OBJECTIVES), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(OBJECTIVES):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    t = doc.add_table(rows=len(FRAMEWORK_HEADER), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(FRAMEWORK_HEADER):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()


# ── Do Now ───────────────────────────────────────────────────

def build_do_now(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 3, "Do Now — " + LESSON_TITLE, subtitle=LESSON_LABEL)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from P2 — extraneous solutions (HOT)",
        dok="3",
        minutes=5,
        teacher_does=["Hand out at door.", "Circulate 3 min.",
                      "Cold-call: \"Why does squaring both sides sometimes introduce a solution that doesn't work?\""],
        students_do=["Identify which solution is extraneous and justify why.",
                     "Write a one-sentence explanation of the checking step."],
        questions_to_ask=[
            "How do you know which solution to discard?",
            "Today we rewrite formulas for a variable — predict: will extraneous solutions matter here too?",
        ],
        adult_role="Solo teacher: circulate silently.",
    )

    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label="Do Now — Higher Order Thinking: Extraneous Solutions",
                    space_after_inches=2.0)

    add_callout(doc, "\U0001f52d  PREDICTION (spend 30 sec)",
                ["Before we rewrite formulas together, predict: if you solved V = (1/3)πr²h for r, "
                 "would you expect an extraneous solution? Why or why not? (No wrong answers — "
                 "just reason it out.) We’ll revisit in Launch."],
                color_hex="FCE5CD")

    doc.save(path)


# ── Student packet ─────────────────────────────────────────────

def build_student(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=False)

    ps.teal_section(doc, "LAUNCH — Rewrite a Formula + Table Completion (teacher models Ex 2 + #44)", [])

    add_callout(doc, "\U0001f4d8  ISOLATE-A-VARIABLE RULES",
                [
                    "1. Identify which variable to isolate (look at the final answer form).",
                    "2. Strip away outer operations LAST-IN, FIRST-OUT (addition/subtraction last "
                    "→ multiplication/division → exponent/radical first to remove).",
                    "3. Raise both sides to the RECIPROCAL power to undo a rational exponent (m/n → n/m).",
                ],
                color_hex="D9EAD3")

    add_callout(doc, "\U0001f4d8  TABLE-EVALUATION RULES",
                [
                    "1. For each input value in the table, substitute into the rewritten formula.",
                    "2. Simplify rational exponents step-by-step (use calculator for irrational outputs).",
                    "3. Round to specified precision.",
                    "4. The table is a point-cloud of the formula’s graph — plot the points to see the shape.",
                ],
                color_hex="FFF2CC")

    doc.add_paragraph()

    ps.teal_section(doc, "EXPLORE — Think-Pair-Share (35 min)",
                    ["Work with a partner. Practice #10 rehearses Assessment Item 1A (isolate x). "
                     "Practice #44 (done in Launch) rehearses Assessment Item 2 (complete the table). "
                     "Plan ∼8 min for #39 (BSA application)."])

    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 2 — Rewrite a formula (partner practice)",
        "Practice #10 — Isolate x ⭐ Assessment Item 1A rehearsal",
        "Practice #25 — Solve for y, rationalize",
        "Practice #26 — Solve for y, rationalize",
        "Practice #27 — Solve for y, rationalize",
        "Practice #39 — BSA application (modeling)",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=label, space_after_inches=1.5)

    ps.teal_section(doc, "SHARE / SUMMARY (5 min)", [])
    ps.sentence_frame_box(doc, [
        "Pair share (#10 or #39): “To isolate ___, I first ___, then raise both sides to the power ___. "
        "The rewritten formula is ___. Plugging in ___ = ___ gives ___. In context this means ___.”",
        "Whole class: one pair reads their answer. Class signals thumbs up/sideways.",
    ])

    add_callout(doc, "\U0001f4ca  CONCEPT SUMMARY (your reference for Topic 5 assessment)",
                [
                    "ISOLATING A VARIABLE FROM A RADICAL/RATIONAL-EXPONENT FORMULA:",
                    "  — Identify the variable to isolate and the form of the final answer.",
                    "  — Undo outer operations first (add/sub), then multiply/divide, then remove the exponent/radical.",
                    "  — To undo x^(m/n): raise both sides to the power (n/m).",
                    "  — Simplify: distribute, combine like terms, solve for the target variable.",
                    "",
                    "COMPLETING A TABLE FROM A REWRITTEN FORMULA:",
                    "  — Use your rewritten formula (e.g., x = ∛(4V) / 2) as the rule for each row.",
                    "  — Substitute each V-value. Use a calculator for irrational outputs. Round as directed.",
                    "  — Check: does the table’s pattern match the shape you expect (growth/decay/root curve)?",
                    "  — Assessment Item 2 expects THIS process — from the rewritten formula to the table.",
                ],
                color_hex="DEEBF7")

    ps.summary_exit_box(doc, "EXIT TICKET — Pre-Assessment Confidence Check",
                        [
                            "To isolate a variable from a rational-exponent formula I ___.",
                            "When I evaluate the formula at a specific x, I ___.",
                            "One biggest thing I learned today: ________________________",
                        ])

    doc.add_page_break()
    ps.teal_section(doc, "OPTIONAL REINFORCEMENT (not collected)",
                    ["Extra stretch for fast finishers. Try without a calculator first."])
    reinforce = qb.get_for_packet(REINFORCE_IDS)
    for i, q in enumerate(reinforce, 1):
        bank_item_block(doc, q, label=f"Optional #{i}  (Escape Velocity — performance task)",
                        space_after_inches=1.8)

    doc.save(path)


# ── Teacher packet ──────────────────────────────────────────────

def build_teacher(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=True)

    add_callout(doc, "\U0001f4cc  PRE-FLIGHT — P3 IS ASSESSMENT-CRITICAL",
                [
                    "Topic 5 Performance Assessment Items 1A/1B (pyramid volume → solve for x → evaluate) "
                    "and Item 2 (complete the table) are directly rehearsed by Practice #10 and Practice #44.",
                    "Ex 2 (rewrite a formula — Golden Gate cables) and q44 (table-completion) are paired in "
                    "Launch — students see both ‘isolate a variable’ and ‘evaluate into a table’ back-to-back.",
                    "No DOK-3 driver today (P2 owned the spine). P3 is pure DOK-2 mastery.",
                    "Pace: Try-It 2 (5) → #10 (7) → #25/#26/#27 (5 each) → #39 (8). Total ∼35 min.",
                    "If pace slows: cut #27 first. Never cut #10 (Item 1A rehearsal).",
                ],
                color_hex="FFD9E0")

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from P2 — extraneous solutions (HOT)",
        dok="3",
        minutes=5,
        teacher_does=["Hand out at door.", "Circulate 3 min.",
                      "Cold-call for extraneous solution reasoning."],
        students_do=["Identify the extraneous solution and justify the check step.",
                     "One-sentence explanation of why squaring can introduce false solutions."],
        questions_to_ask=[
            "HOW do you know which solution to reject?",
            "Based on yesterday: predict — will rewriting a formula for a variable produce extraneous solutions?",
        ],
        adult_role="Solo teacher: circulate silently.",
    )
    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label=f"Do Now ({q['id']})", space_after_inches=0.3)
    add_callout(doc, "✅  ANSWER KEY",
                [q.get("notes") or "(verify)"],
                color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Launch",
        framework_tag="teacher models Ex 2 (rewrite formula) + #44 (table) back-to-back",
        dok="2",
        minutes=12,
        teacher_does=[
            "FIRST 6 min: Example 2 — Golden Gate suspension cable formula. Identify the target variable. "
            "Undo operations LAST-IN FIRST-OUT: subtract, then divide, then raise to reciprocal power. "
            "Emphasize: (m/n) → (n/m) flips the exponent.",
            "NEXT 5 min: Practice #44 — table-completion. Show the rewritten formula first, then "
            "substitute each V value row-by-row. Demonstrate calculator use for irrational outputs. "
            "Emphasize: this is EXACTLY Assessment Item 2.",
            "30 sec: point at packet rule cards. “Mark them. These are your assessment reference.”",
            "Do NOT solve Try-It 2 or Explore items. Release at minute 17.",
        ],
        students_do=[
            "Watch Ex 2. Copy each step. Circle the reciprocal-power move.",
            "Watch #44. Copy the table. Note which column is input (V) and which is output (x).",
            "Copy BOTH rule cards into margins.",
        ],
        questions_to_ask=[
            "Why do we raise to the RECIPROCAL power to undo x^(2/3)?",
            "In the table: as V increases, what happens to x? Does that make sense for volume?",
            "For Ex 2: what physical constraint rules out negative values of the variable?",
            "How is rewriting a formula different from solving an equation? (It’s the same algebra — "
            "just a letter, not a number, on the right side.)",
        ],
        adult_role="Project both examples. Think aloud on the reciprocal-power step.",
    )
    launch_items = qb.get_for_packet(LAUNCH_IDS)
    launch_labels = [
        "Example 2 — Rewrite a Formula (Golden Gate cables)",
        "Practice #44 — Table Completion ⭐ ASSESSMENT REHEARSAL (Item 2)",
    ]
    for label, lq in zip(launch_labels, launch_items):
        bank_item_block(doc, lq, label=f"{label} ({lq['id']})", space_after_inches=0.3)
        add_callout(doc, "✅  ANSWER KEY",
                    [lq.get("notes") or "(no answer key — verify before class)"],
                    color_hex="D9EAD3")

    add_callout(doc, "⭐  ASSESSMENT NOTE — Practice #44",
                [
                    "Practice #44 is the algebraic twin of Topic 5 Assessment Item 2 — complete the "
                    "table for x = ∛(4V)/2 at V values {0, 0.25, 2, 16, 32, 64, 128}.",
                    "Any student who can complete #44 can complete Assessment Item 2.",
                    "Teacher: walk through at least the first two rows of the table live so students "
                    "see the calculator workflow for irrational outputs.",
                ],
                color_hex="FFF2CC")

    add_callout(doc, "\U0001f3a4  TEACHER SCRIPT",
                [
                    "1. “Yesterday we solved radical equations — one variable, one answer. Today we solve for "
                    "A variable — we rewrite the whole formula.”",
                    "2. Ex 2: “Target is the length variable. What’s in the way? First the ‘+’, then the "
                    "coefficient, then the exponent. LAST-IN FIRST-OUT. Flip (2/3) → (3/2). Done.”",
                    "3. “#44: I already rewrote the formula. Now I just substitute. V=0 → x=0. "
                    "V=0.25 → x=∛(1)/2 = 0.5. Calculator for the rest.”",
                    "4. “Same procedure, new purpose. The math is identical whether it’s a number or "
                    "a table of numbers on the right.”",
                    "5. Release.",
                ],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Explore",
        framework_tag="TPS · pure DOK-2 mastery · assessment rehearsal",
        dok="2",
        minutes=35,
        teacher_does=[
            "6 laps across 35 min.",
            "Laps 1-2: Try-It 2 + #10. Isolate-a-variable practice. Press reciprocal-power step on every pair.",
            "Laps 3-5: #25 + #26 + #27 — solve for y with rationalization. Pattern recognition set.",
            "Lap 6: #39 — BSA application (modeling). Press interpretation: what does the variable mean?",
            "Do not give answers. Use sentence frame to press interpretation.",
        ],
        students_do=[
            "6 items in order. Isolate-a-variable first (Try-It 2, #10), then rationalization pattern "
            "(#25–#27), then application (#39).",
            "For #10: BEFORE solving, state which variable you’re isolating and what power will undo it.",
            "Justify with sentence frame.",
        ],
        questions_to_ask=[
            "What reciprocal power undoes x^(3/2)? What about x^(2/5)?",
            "For #25–#27: what’s the same in every problem? What changes?",
            "For #10: how does your answer connect to Assessment Item 1A?",
            "For #39: what do the units of BSA tell you about whether your answer is reasonable?",
            "After #10: could you now complete a table like #44 using your answer? Try one row.",
        ],
        adult_role="Solo teacher: 6 laps. Press the reciprocal-power step on every item. #10 IS the Item 1A rehearsal.",
    )
    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 2", "Practice #10 ⭐ ASSESSMENT ITEM 1A", "Practice #25",
        "Practice #26", "Practice #27", "Practice #39 — BSA Application",
    ]
    for label, qi in zip(labels, items):
        bank_item_block(doc, qi, label=f"{label} ({qi['id']})", space_after_inches=0.3)
        add_callout(doc, "✅  ANSWER / ANTICIPATED TALK",
                    [qi.get("notes") or "(no answer key — verify before class)"],
                    color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Share / Summary",
        framework_tag="Concept Summary + Topic 5 Performance Task preview",
        dok="2",
        minutes=5,
        teacher_does=[
            "Cold-call 1 pair to present #10 (isolate x) using the sentence frame.",
            "Project Concept Summary (from student packet).",
            "“Topic 5 Performance Assessment: Items 1A/1B ask you to rewrite a pyramid-volume formula "
            "for x, then evaluate. Item 2 asks you to complete a table. You just rehearsed BOTH.”",
        ],
        students_do=[
            "Presenting pair uses sentence frame for #10.",
            "Rest of class signals and annotates their Concept Summary.",
        ],
        questions_to_ask=[
            "What are the TWO big ideas from Lesson 5-4 P3?",
            "If you see a rational exponent on the assessment, what’s the FIRST thing you do?",
        ],
        adult_role="Frame today as assessment rehearsal. Students leave confident on isolating variables.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit Ticket",
        framework_tag="pre-assessment confidence check",
        dok="1-2",
        minutes=3,
        teacher_does=[
            "Collect at door.",
            "Tonight: scan for students who can’t complete the reciprocal-power step or the table row. "
            "Insert a 5-min Do Now recap when the Performance Assessment lands.",
        ],
        students_do=[
            "Complete the three fill-in stems.",
            "Write one biggest thing learned.",
        ],
        questions_to_ask=[
            "Did any student leave the reciprocal power blank?",
            "Did any student evaluate the formula correctly but skip the table interpretation?",
        ],
        adult_role="Collect. No grade. Use as data for assessment prep.",
    )
    add_callout(doc, "\U0001f4cb  EXIT TICKET — Student Form",
                [
                    "To isolate a variable from a rational-exponent formula I ___.",
                    "When I evaluate the formula at a specific x, I ___.",
                    "One biggest thing I learned today: ________________________",
                ],
                color_hex="FFF2CC")

    doc.add_page_break()
    add_callout(doc, "\U0001f550  PERIOD A / PERIOD F — 65-MIN CLASS NOTES",
                [
                    "Both sections should have 65 min. Use extra 10 min on #10 and #39 (assessment-critical).",
                    "If finished early: Practice #46 (Escape Velocity) as fast-finisher stretch.",
                    "Alternative: project #44 as a second table-completion prompt — more assessment rehearsal.",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "\U0001f9e9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS",
                [
                    "• Provide both Rule cards as half-sheet cut-outs on their desks.",
                    "• For #10, provide a scaffold showing the first two steps already completed so they "
                    "focus on the reciprocal-power move.",
                    "• Accept verbal formula-rewriting in lieu of written for #10.",
                ],
                color_hex="E2F0D9")
    add_callout(doc, "\U0001f30d  SUPPORTS FOR ELL STUDENTS",
                [
                    "• Sentence frames on student page (don’t skip).",
                    "• Word cards: “isolate” = get the variable alone; “reciprocal power” = flip the fraction in the exponent.",
                    "• “rewrite” = change what the formula solves for, without changing the math.",
                    "• “table” = list of (input, output) pairs — each row is one substitution.",
                    "• Pair ELL students with English-dominant partners for TPS.",
                ],
                color_hex="DEEBF7")

    doc.save(path)


# ── Main ─────────────────────────────────────────────────────

if __name__ == "__main__":
    build_do_now("L54_P3_Do_Now.docx")
    build_student("L54_P3_Student_Packet.docx")
    build_teacher("L54_P3_Teacher_Packet.docx")
    print("Built L54_P3_Do_Now.docx, L54_P3_Student_Packet.docx, L54_P3_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L54_P3_Visuals_Checklist.md",
                              title="L54 P3 — Visuals Checklist")
