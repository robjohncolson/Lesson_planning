"""Generate Day 2 materials: Do Now sheet + Student Packet + Teacher Packet.

Day 2 = Graphing from Factored Form (post-spring-break Monday).

DESIGN (revised after Day-1 evidence): single-DOK3 spine.
  Entrance ticket primes the DOK-3.  Blooket is pure rule-recall.  One
  self-contained DOK-3 driver (Reverse Engineering).  Summary exit ticket.
  The #13 \"touches at x=4\" discovery MOVES to Day 3 (its natural home
  with multiplicity).

Phases (57 min, fits 55-min Tue A with tight buffer; 65-min Tue F gets slack):
  0\u20135    Do Now A : Solo Sign Flip Prediction (2 prompts)     [DOK 2]
  5\u20137    Do Now B : Blooket login                             [\u2014]
  7\u201314   Do Now C : Blooket rule recall (habits of mind)      [DOK 1]
  14\u201322  Launch    : Sign Flip synthesis + Desmos verify       [DOK 2]
  22\u201347  Explore   : Reverse Engineering \u2014 the DOK-3 driver    [DOK 3]
  47\u201352  Share/Sum : Synthesis + self-reflection               [DOK 2]
  52\u201357  Exit      : Summary recap (NOT a CER)                 [DOK 1\u20132]
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL

TABLE_STYLE = "Table Grid"

OBJECTIVES = [
    ("Math Objective",
     "Identify the zeros of a function using factoring or synthetic division. "
     "Use zeros to graph a polynomial function."),
    ("Language Objective",
     "Students will use the frame \u201cThe zeros stay the same because ___. "
     "The graph changes because ___\u201d to justify their comparison."),
    ("Essential Understanding",
     "The zeros of a polynomial function can be determined using factoring or "
     "synthetic division. The zeros can be used to sketch its graph."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Given the zeros of a polynomial and one point on its graph, students "
     "build the complete equation and justify every step."),
    ("Essential Question",
     "How do the zeros of a polynomial \u2014 and a single point on the graph "
     "\u2014 determine its complete equation?"),
    ("Materials",
     "Student laptops with Desmos; blank paper; pencils; Blooket access "
     "(warm-up code on the board). No chart paper, no manipulatives."),
]


def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if i == 0 and bold_first:
            run.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_two_col(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = TABLE_STYLE
    t.columns[0].width = Inches(1.8)
    t.columns[1].width = Inches(4.7)
    for i, (label, content) in enumerate(rows):
        set_cell(t.rows[i].cells[0], label, bold_first=True)
        if isinstance(content, str):
            content = [content]
        set_cell(t.rows[i].cells[1], content)
    doc.add_paragraph()


def add_callout(doc, title, lines):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.add_run(title).bold = True
    for line in lines:
        cell.add_paragraph(line)
    doc.add_paragraph()


def header_line(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def blank_lines(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.add_run("_______________________________________________________________")


# ==================================================================
# DO NOW SHEET (handed out on entry) \u2014 SLIM
# Two prompts only. Breathable. Primes ZPP + end-behavior reasoning.
# ==================================================================

def build_do_now(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)

    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    header_line(doc, "Day 2  |  DO NOW  \u2014  Sign Flip", size=13)

    p = doc.add_paragraph()
    r = p.add_run("Name: _____________________________________     Date: _____________")
    r.font.size = Pt(11)
    doc.add_paragraph()

    add_callout(doc, "\u270d\ufe0f  5 minutes. Silent. Pencil only. No Desmos. No packet.", [
        "Compare these two functions:",
        "    f(x) = x(x \u2212 4)(x + 3)",
        "    g(x) = \u2212x(x \u2212 4)(x + 3)",
    ])

    # Prompt 1 \u2014 single open comparison with room to think
    header_line(doc, "1.  Before graphing, predict:", size=12)
    p = doc.add_paragraph()
    p.add_run("What STAYS THE SAME?  What CHANGES?  Write in whatever form helps you think \u2014 "
              "bullet points, sentences, a quick table. Be specific.").italic = True
    blank_lines(doc, 5)
    doc.add_paragraph()

    # Prompt 2 \u2014 sentence frame (the justification)
    header_line(doc, "2.  Finish both halves of this sentence:", size=12)
    p = doc.add_paragraph()
    p.add_run("\u201cThe zeros stay the same because ").bold = False
    doc.add_paragraph()
    blank_lines(doc, 2)
    p = doc.add_paragraph()
    p.add_run("and the graph changes because")
    blank_lines(doc, 2)
    p = doc.add_paragraph()
    p.add_run(".\u201d")

    doc.save(path)


# ==================================================================
# STUDENT PACKET (handed out after Blooket)
# One DOK-3 driver.  All info on the page.  Summary exit ticket.
# ==================================================================

# Self-contained Reverse Engineering page: every rule the student needs
# to pick up and run is printed right on the page.
REVERSE_ENG_STUDENT = [
    "THE PROBLEM",
    "Write the equation of a polynomial with zeros at  x = \u22122,  x = 1,  and "
    "x = 4  that passes through the point  (0, \u22128).",
    "",
    "THE TWO RULES YOU NEED (everything you need is right here)",
    "\u2022  RULE 1:  If  x = c  is a zero, then  (x \u2212 c)  is a factor.",
    "         (So a zero at x = \u22122 gives the factor (x + 2). Watch the sign.)",
    "\u2022  RULE 2:  The factors alone don\u2019t fix the equation \u2014 you still need "
    "a leading number  a.  Use the given point to find it.",
    "         General form:   f(x) = a \u00b7 (factor)(factor)(factor)",
    "",
    "STEP 1  \u2014  Turn the zeros into factors.",
    "    x = \u22122  \u2192  (________)",
    "    x =  1  \u2192  (________)",
    "    x =  4  \u2192  (________)",
    "So far:   f(x) = a (________)(________)(________)",
    "",
    "STEP 2  \u2014  Use the point (0, \u22128) to find  a.",
    "    Substitute  x = 0  and  f(x) = \u22128:",
    "    \u22128  =  a (________)(________)(________)",
    "    \u22128  =  a \u00b7 ________",
    "    a  =  ________",
    "",
    "STEP 3  \u2014  Write the final equation.",
    "    f(x) = ________________________________________",
    "",
    "STEP 4  \u2014  Verify in Desmos.",
    "    Graph your equation. Does it pass through (0, \u22128)?  "
    "Are the zeros at \u22122, 1, 4?",
    "",
    "STEP 5  \u2014  Talk to your partner.",
    "    Use the sentence frame from the Do Now:",
    "    \u201cThe zeros gave me the factors because ___.",
    "     The point gave me the leading number because ___.\u201d",
]

SHARE_SUMMARY_STUDENT = [
    "Self-check  \u2014  circle one for each:",
    "1.  I can turn a list of zeros into factors.               \u2713    partly    not yet",
    "2.  I can use a given point to find the leading number a.  \u2713    partly    not yet",
    "3.  I can predict what changes when f(x) becomes \u2212f(x).    \u2713    partly    not yet",
    "",
    "Preview: tomorrow we investigate what happens when a factor "
    "repeats (like (x \u2212 4) twice).",
]

EXIT_STUDENT = [
    "Today we discovered:",
    "",
    "1.  When f(x) becomes \u2212f(x), the zeros  ______________  "
    "and the graph  ______________ .",
    "",
    "2.  To find the leading number  a, I plug in a given  ______________  "
    "and solve for a.",
    "",
    "3.  One sentence \u2014 the most important thing I learned today:",
    "_____________________________________________________________________",
    "_____________________________________________________________________",
]


def build_student(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    header_line(doc, "Day 2  |  Graphing from Factored Form", size=13)

    add_two_col(doc, OBJECTIVES)
    add_two_col(doc, FRAMEWORK_HEADER)

    add_callout(doc, "\U0001F501  WELCOME BACK", [
        "Last time: zeros from factors, sign charts, sketches on chart paper.",
        "Today: work BACKWARD.  Given the zeros and one point, build the equation.",
        "Materials: laptop (Desmos), blank paper, pencil.",
    ])

    add_callout(doc, "\U0001F4C4  DO NOW  \u2014  done on the separate sheet", [
        "You got the Sign Flip sheet when you walked in.",
        "Keep it nearby \u2014 we\u2019ll return to your prediction during the Launch "
        "after the Blooket.",
    ])

    add_callout(doc, "\U0001F3AE  Blooket  \u2014  Rule Recall   [Do Now B+C]  [DOK 1]   (2 + 7 min)", [
        "Teacher puts the code on the screen. Log in (2 min), then 7-minute game.",
        "Focus: the RULES from last lesson \u2014 Zero Product Property, what "
        "x\u00b2 + 9 can and can\u2019t do, what the highest power tells you about end behavior.",
    ])

    add_callout(doc, "\U0001F504  Sign Flip Synthesis   [Launch]  [DOK 2]   (8 min)", [
        "Pull out your Do Now.",
        "Turn and Talk: what did you predict stayed the same?  What changed?",
        "Class synthesis + Desmos check.",
        "Sentence frame (use this aloud, then write it in STEP 5 below):",
        "    \u201cThe zeros stay the same because ___.  The graph changes because ___.\u201d",
    ])

    add_two_col(doc, [(
        "\U0001F527  Build the Equation (Reverse)   [Explore]  [DOK 3]   (25 min)",
        REVERSE_ENG_STUDENT,
    )])

    add_callout(doc,
        "\U0001F501  SHARE / SUMMARY   [Share/Summary]  [DOK 2]   (5 min)",
        SHARE_SUMMARY_STUDENT)

    add_callout(doc,
        "\U0001F4E4  EXIT TICKET  \u2014  SUMMARY   [Exit Ticket]  [DOK 1\u20132]   (5 min)",
        EXIT_STUDENT)

    doc.save(path)


# ==================================================================
# TEACHER PACKET
# ==================================================================

CRITERIA_FOR_SUCCESS = [
    "Student-facing, revisited during Share/Summary:",
    "1.  I can turn a list of zeros into factors (mind the signs).",
    "2.  I can use a given point on the graph to find the leading coefficient a.",
    "3.  I can predict what changes when f(x) becomes \u2212f(x).",
    "Formative evidence:",
    "\u2022  Do Now sheets \u2014 scan on entry for predictions.",
    "\u2022  Blooket dashboard (rule-recall items, DOK 1).",
    "\u2022  Partner justification during Reverse Engineering (circulate, listen).",
    "\u2022  Reverse Engineering page \u2014 collected or photographed (DOK 3 evidence).",
    "\u2022  Summary exit ticket.",
]

DO_NOW_KEY = [
    "Expected predictions (accept well-reasoned variants):",
    "STAYS THE SAME:",
    "\u2022  Zeros at x = \u22123, 0, 4 \u2014 ZPP still applies; the \u22121 out front never makes a factor zero.",
    "\u2022  The three x-intercepts.",
    "\u2022  Degree (still cubic).",
    "CHANGES:",
    "\u2022  The graph is a reflection of f(x) across the x-axis.",
    "\u2022  End behavior flips (was down-left/up-right; now up-left/down-right).",
    "\u2022  Sign chart: every sign flips.",
    "Strong sentence responses:",
    "\u201cThe zeros stay the same because the Zero Product Property doesn\u2019t care "
    "about a \u22121 in front, and the graph changes because multiplying by \u22121 "
    "reflects every y-value.\u201d",
    "DIAGNOSTIC: students who predict \u201czeros change\u201d are the cold-call "
    "during the Launch synthesis.",
]

BLOOKET_SCOPE = [
    "SCOPE (pure rule-recall \u2014 DOK 1 only).  No procedural factoring items.",
    "Habits of mind from Day 1 that today\u2019s work depends on:",
    "\u2022  Zero Product Property:  \u201cIf (x \u2212 c) is a factor, one zero is c.\u201d",
    "\u2022  Sign / factor direction:  \u201cIf one zero is x = \u22122, the factor is (x + 2), not (x \u2212 2).\u201d",
    "\u2022  x\u00b2 + c for c > 0 contributes NO real zeros (primes why some factors don\u2019t show up as x-intercepts).",
    "\u2022  Highest power \u2192 end behavior:  even \u2192 same direction both sides; odd \u2192 opposite.",
    "\u2022  Sign-flip primers:  zeros of \u2212f(x) are the same as zeros of f(x).",
    "DO NOT include Blooket items that require students to factor an expression \u2014 "
    "that\u2019s DOK 2 problem solving and doesn\u2019t belong on flashcards.",
    "Dashboard diagnostic: note the two lowest-percent items. If a rule is cold, "
    "60-second board reset before Launch.",
]

REVERSE_KEY = [
    "Target: zeros at x = \u22122, 1, 4 through (0, \u22128)",
    "\u2705 Step 1 \u2014 Factors: (x + 2), (x \u2212 1), (x \u2212 4)",
    "\u2705 Step 2 \u2014 Plug in (0, \u22128):",
    "   \u22128 = a(0 + 2)(0 \u2212 1)(0 \u2212 4) = a(2)(\u22121)(\u22124) = 8a",
    "   a = \u22121",
    "\u2705 Step 3 \u2014 Final:  f(x) = \u2212(x + 2)(x \u2212 1)(x \u2212 4)",
    "\u2705 Desmos check: passes (0, \u22128); zeros at \u22122, 1, 4; negative leading "
    "coefficient flips the end behavior compared to the Do Now\u2019s f(x).",
    "WHY THIS IS DOK 3:",
    "\u2022  No memorized procedure gets the student there in one step.",
    "\u2022  They must coordinate THREE moves (zeros\u2192factors, hidden  a , point "
    "constraint\u2192solve for  a ) in an order they choose.",
    "\u2022  They justify each move aloud with their partner.",
    "SAMPLE PARTNER JUSTIFICATION (what good talk sounds like):",
    "\u201cThe zeros told me the factors: \u22122 gives (x + 2), 1 gives (x \u2212 1), "
    "4 gives (x \u2212 4). The point told me  a  because plugging x = 0 gave \u22128 = 8a, "
    "so  a = \u22121.  It works because the Zero Product Property guarantees the "
    "factors produce zeros, and one point is enough to solve for the one unknown.\u201d",
]

SHARE_SUMMARY_KEY = [
    "Scan target for the three self-rating items:",
    "\u2022  #1 zeros \u2192 factors: \u2713 for most after the synthesis. \u201cNot yet\u201d = "
    "conference during next Do Now.",
    "\u2022  #2 point \u2192 a : \u201cpartly\u201d acceptable today (new skill).",
    "\u2022  #3 sign flip: \u2713 for most (explicit Do Now + Launch content).",
    "Teacher script (5 min flat):",
    "1.  Essential Question callback (60s).",
    "2.  Language Objective check (60s): \u201cWho used \u2018the zeros stay the "
    "same because...\u2019? Finish with what?\u201d",
    "3.  Students circle self-ratings (90s).  Teacher scans.",
    "4.  Preview Day 3 (30s): \u201cTomorrow: what happens when a factor appears twice?\u201d",
    "5.  Transition (30s): \u201cPencils up. Summary exit ticket.\u201d",
]

EXIT_KEY = [
    "Expected responses:",
    "1.  Zeros stay the SAME;  graph REFLECTS across the x-axis "
    "(accept: \u201cflips,\u201d \u201cupside down\u201d).",
    "2.  Plug in a given POINT (accept: \u201cy-value,\u201d \u201cy-intercept,\u201d "
    "\u201c(x, y)\u201d).",
    "3.  Open response. Look for:",
    "    \u2022  Naming the two-move structure (zeros\u2192factors, point\u2192a).",
    "    \u2022  Surfacing the Zero Product Property rule.",
    "    \u2022  Connecting to Day 1 (\u201cfactors \u2192 zeros\u201d going the other direction).",
    "INTENTIONALLY NOT A CER.  Walk-out under time pressure should recap, not generate "
    "new argument.  The DOK-3 generative work already happened in Reverse Engineering.",
    "If you want CER practice: assign Savvas Practice #6 as HW, or use it as Day 3 Do Now.",
]

IEP_MODS = [
    "\u2022  Chunked steps: each step numbered with its own whitespace block.",
    "\u2022  Desmos graphing reduces fine-motor sketch demand.",
    "\u2022  Partner roles during Reverse Engineering: \u201cfactor-finder\u201d (Steps 1, 3) "
    "and \u201cequation-checker\u201d (Steps 2, 4). Clear scope for processing-speed accommodations.",
    "\u2022  Release valve: if a pair is stuck after 15 min on Reverse Engineering, "
    "release a \u201chint card\u201d \u2014 show the a = \u22121 result and ask them to JUSTIFY why.",
    "\u2022  Blooket answers read aloud if an individual IEP requires audio support.",
    "\u2022  Preferential seating near board for Sign Flip synthesis.",
    "\u2022  Do Now sheet is deliberately breathable (2 prompts, large whitespace).",
]

ELL_SUPPORTS = [
    "\u2022  Sentence frame on Do Now sheet \u2014 also printed on the Reverse "
    "Engineering page for Step 5.",
    "\u2022  Predict-before-verify routine: students commit in writing before speaking.",
    "\u2022  No new academic vocabulary today (\u201cmultiplicity\u201d is deferred to Day 3).",
    "\u2022  Materials-light: no chart paper, no manipulatives.",
    "\u2022  Think-Pair-Share before any public share; writing before speaking.",
    "\u2022  Share/Summary self-rating uses \u2713 / partly / not-yet icons \u2014 "
    "keeps cognitive load on the math, not on producing English.",
    "\u2022  Summary exit ticket is two fill-in-the-blank + one sentence. Low writing load.",
]


def build_teacher(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    header_line(doc, "TEACHER EDITION", size=12)
    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    header_line(doc, "Day 2  |  Graphing from Factored Form", size=13)

    add_two_col(doc, OBJECTIVES + [("CCSS", "F-IF.C.7c, A-APR.B.3, MP.3")])
    add_two_col(doc, FRAMEWORK_HEADER)

    add_callout(doc, "\U0001F4CB  DESIGN \u2014  ONE DOK-3 SPINE", [
        "Do Now \u2192 Launch \u2192 ONE DOK-3 Explore \u2192 Share/Summary \u2192 Summary Exit.",
        "[DOK 1] Blooket rule recall.",
        "[DOK 2] Solo Sign Flip prediction; Launch synthesis; Share/Summary; summary exit.",
        "[DOK 3] Reverse Engineering \u2014 zeros + point \u2192 equation.  This is the sole "
        "DOK-3 task today and gets the full 25 minutes.",
        "Day 1 evidence: Day 1 took three class periods because it bundled two DOK-3 "
        "tasks (Gallery + Leap) after a DOK-2 group build.  Day 2 does NOT repeat that "
        "mistake.  The \u201c#13 touches at x = 4\u201d discovery moves to Day 3 (its natural "
        "home with Multiplicity).",
    ])

    add_callout(doc, "\u2705  CRITERIA FOR SUCCESS  /  FORMATIVE ASSESSMENT", CRITERIA_FOR_SUCCESS)

    add_callout(doc, "\u23f1\ufe0f  REALISTIC PACING", [
        "Total: 57 min.  Blocks: Tuesday A = 55 min (tight), Tuesday F = 65 min (8 min slack).",
        "  Do Now A (solo paper) ..... 5 min",
        "  Do Now B (Blooket login) .. 2 min",
        "  Do Now C (Blooket game) ... 7 min",
        "  Launch (Sign Flip) ........ 8 min",
        "  Explore (Reverse Eng) .... 25 min",
        "  Share/Summary ............. 5 min",
        "  Summary Exit .............. 5 min",
        "LIBRARY-CHARGER LOGISTICS: hand every student the DO NOW sheet as they walk in. "
        "Early arrivers work silently while the library run returns. When everyone\u2019s "
        "back, display the Blooket code.",
        "RELEASE VALVES (in order of preference):",
        "  1. Cut Launch synthesis from 8 \u2192 6 min (skip the extended Desmos verify).",
        "  2. Hint-card a stuck pair in Reverse Engineering (reveal a, ask for justification).",
        "  3. Share/Summary = 4 min (drop the preview line).",
        "  NEVER cut the Summary Exit. It\u2019s the only artifact that documents learning.",
    ])

    # ---- Do Now A ----
    add_callout(doc, "[Do Now A]  [DOK 2] Solo Paper \u2014 Sign Flip Prediction", [
        "Students receive the separate Do Now sheet on entry. Silent, 5 min, pencil only.",
        "TWO prompts only (breathable).  Primes Zero Product Property + end-behavior "
        "reasoning for the DOK-3 driver later.",
        "Purpose: productive use of the library-trip window + early written commitment "
        "before the synthesis.",
    ])
    add_callout(doc, "\u2705  DO NOW / ANSWER KEY", DO_NOW_KEY)

    # ---- Do Now B+C ----
    add_callout(doc, "[Do Now B+C]  [DOK 1] Blooket \u2014 Rule Recall ONLY", BLOOKET_SCOPE)

    # ---- Launch ----
    add_callout(doc, "[Launch]  [DOK 2] Sign Flip \u2014 Class Synthesis  (8 min)", [
        "Script:",
        "1.  \u201cPull out your Do Now.\u201d (15s)",
        "2.  Turn and Talk on predictions. (90s)",
        "3.  Cold-call a student who predicted \u201czeros change\u201d (from entry scan). (60s)",
        "4.  Desmos check \u2014 graph both functions. (90s)",
        "5.  Class synthesis: \u201cWhy are the zeros the same?\u201d Elicit ZPP. Record on board. (90s)",
        "6.  Transition to Reverse Engineering: \u201cYou just went from factors to graph. "
        "Next: you\u2019ll go from zeros to equation. Backward.\u201d (45s)",
        "Cold-call, don\u2019t take volunteers \u2014 post-break, low-hand students need re-engagement.",
    ])

    # ---- Explore ----
    add_callout(doc, "[Explore]  [DOK 3] Reverse Engineering  (25 min)", [
        "One driver.  All the information the student needs is printed on their packet page: "
        "both rules, the step-by-step scaffold, the verification prompt, the partner frame.",
        "Teacher role: circulate.  No board teaching.  Prompt-only \u2014 answer questions "
        "with questions (\u201cPlug in x = 0. Do you get \u22128?\u201d).",
        "Partner roles (IEP-friendly): \u201cfactor-finder\u201d (Steps 1, 3) + \u201cequation-checker\u201d (Steps 2, 4).",
    ])
    add_two_col(doc, [("\U0001F527  Reverse Engineering / ANSWER KEY", REVERSE_KEY)])
    add_callout(doc, "\U0001F3AF  TEACHER MOVES \u2014 Reverse Engineering", [
        "Common stall: Skip a. Write f(x) = (x+2)(x\u22121)(x\u22124). Prompt: \u201cPlug in x = 0. Do you get \u22128?\u201d",
        "Common error: Sign flip on a zero. x = \u22122 \u2192 (x \u2212 2). Prompt: \u201cSet your factor = 0. Does it give x = \u22122?\u201d",
        "Stuck-pair hint card: \u201ca = \u22121. Now justify WHY.\u201d (Preserves DOK 3 via justification.)",
        "Extension (fast finishers): \u201cWhat if the point were (0, 8) instead?\u201d  (a = 1.)",
        "Fast-finisher 2: \u201cWhat if the zeros were \u22122, 1, 4 and the point was (2, 0)? "
        "Can you find a?\u201d  (No \u2014 (2, 0) is on a zero, so a is not determined. Good conceptual catch.)",
    ])

    # ---- Share / Summary ----
    add_callout(doc, "[Share/Summary]  [DOK 2] Synthesis + Reflection  (5 min)", [
        "Framework: synthesis, return to objectives, self-rating.  NOT optional.",
    ])
    add_callout(doc,
        "\U0001F501  SHARE / SUMMARY (student-facing)",
        SHARE_SUMMARY_STUDENT)
    add_callout(doc, "\u2705  SHARE / SUMMARY / ANSWER KEY + DIAGNOSTIC", SHARE_SUMMARY_KEY)

    # ---- Exit Ticket ----
    add_callout(doc, "[Exit Ticket]  [DOK 1\u20132] Summary (NOT a CER)  (5 min)", [
        "Design intent: the exit ticket is a SUMMARY of what they learned today, not a "
        "new cognitive task.  Two fill-in-the-blank recap items + one open sentence.",
        "Why not CER here: full CER writing in 5 min is unrealistic and a DOK-3 task at "
        "the walk-out window is bad design.  The DOK-3 work is Reverse Engineering.",
    ])
    add_callout(doc,
        "\U0001F4E4  EXIT TICKET  \u2014  SUMMARY (student-facing)",
        EXIT_STUDENT)
    add_callout(doc, "\u2705  EXIT TICKET / ANSWER KEY", EXIT_KEY)

    # ---- Closers ----
    add_callout(doc, "\U0001F4F7  WHAT TO COLLECT", [
        "1.  Do Now sheets \u2014 collected on entry scan or during Launch synthesis.",
        "2.  Reverse Engineering page \u2014 photograph or collect (DOK-3 evidence).",
        "3.  Summary exit ticket (every student).",
        "4.  Share/Summary self-ratings \u2014 quick scan as packets come in.",
    ])
    add_callout(doc, "\U0001F440  LOOK-FORS DURING WALKTHROUGH", [
        "[Do Now A] Silent solo work? Predictions in writing?",
        "[Do Now C] Teacher monitoring Blooket dashboard diagnostically?",
        "[Launch] Student predictions referenced? Cold-call used to surface wrong ideas safely?",
        "[Explore] Partners justifying with rule language? Teacher circulating with prompts, not answers?",
        "[Share/Summary] Self-ratings honest? Essential Question + Language Objective called back?",
        "[Exit Ticket] All three items attempted? Sentence complete?",
    ])
    add_callout(doc, "\U0001F9E9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS", IEP_MODS)
    add_callout(doc, "\U0001F30D  SUPPORTS FOR ELL STUDENTS", ELL_SUPPORTS)

    doc.save(path)


if __name__ == "__main__":
    build_do_now("Day_2_Do_Now.docx")
    build_student("Day_2_Student_Packet.docx")
    build_teacher("Day_2_Teacher_Packet.docx")
    print("Built Day_2_Do_Now.docx, Day_2_Student_Packet.docx, Day_2_Teacher_Packet.docx")
