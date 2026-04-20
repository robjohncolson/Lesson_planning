"""Generate Combined Day 6 materials — Savvas-only polynomial inequalities day.

Design: polynomial inequalities (Lesson 3-5 Day 6). Every student-facing work
item traces to a Savvas bank ID. DOK 1-2 practice day — no DOK-3 driver
(Savvas's editorial stance puts DOK-3 on modeling days only).

Do Now bridge: 3-5-savvas-q21 connects Day 4-5 zeros → Day 6 inequalities
via the key insight "zeros split the number line."

Context: solo teacher, class of up to ten.

Phases (55 min):
  0– 5   Do Now A : 3-5-savvas-q21 — solve polynomial equation (bridge) [DOK 1]
  5– 7   Do Now B : Blooket login                                         [—]
  7–14   Do Now C : Blooket rule recall (6 min game)                      [DOK 1]
  14–15  Blooket Wrap : candy + lowest rule aloud                         [DOK 1]
  15–27  Launch    : sign-chart method — Savvas Example 6 / Try It 6a     [DOK 2]
  27–47  Practice  : Try It 6b + pick 2 of Practice #22/#23/#24           [DOK 2]
  47–51  Share/Sum : self-rating + one-sentence synthesis                  [DOK 2]
  51–56  Exit      : summary recap fill-in + one sentence                 [DOK 1-2]
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_ID    = "3-5-savvas-q21"
LAUNCH_ID    = "3-5-tryit-6a"
PRACTICE_IDS = ["3-5-tryit-6b",
                "3-5-savvas-q22", "3-5-savvas-q23", "3-5-savvas-q24"]
# EXIT_ID (3-5-lq-q3) is reserved for Day 8 quiz — not used on Day 6.
# Day 6 exit is a summary recap (fill-in + one sentence), not a problem.

# Fail loudly if any bank ID is missing — prevents silent placeholder renders.
_ALL_IDS = [DO_NOW_ID, LAUNCH_ID] + PRACTICE_IDS
qb.get_for_packet(_ALL_IDS)   # raises KeyError on first missing ID

OBJECTIVES = [
    ("Math Objective",
     "Solve polynomial inequalities by factoring the polynomial, identifying "
     "zeros, constructing a sign chart, and reading off the intervals where "
     "the inequality is satisfied."),
    ("Language Objective",
     "Students will use the frame \u201cThe zeros are ___, so the sign changes at "
     "___, and the solution is ___.\u201d to explain their sign-chart reasoning."),
    ("Essential Understanding",
     "The zeros of a continuous polynomial function divide the number line into "
     "intervals. Because the sign of the polynomial can only change at a zero of "
     "odd multiplicity, testing one point per interval determines the sign on "
     "the entire interval."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Students apply the sign-chart method to solve four types of cubic/quartic "
     "polynomial inequalities.  Every item is drawn from Savvas Lesson 3-5."),
    ("Essential Question",
     "How do the zeros of a polynomial divide the number line, and how does "
     "that structure let you solve any polynomial inequality?"),
    ("Materials",
     "Chromebooks (Blooket login); pencils; student packet. Desmos allowed on "
     "exit ticket to verify zeros only."),
]


# ── Low-level helpers (mirrored from build_day23_packets.py) ─────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if i == 0 and bold_first:
            run.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_two_col(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = TABLE_STYLE
    t.columns[0].width = Inches(1.8)
    t.columns[1].width = Inches(4.7)
    for i, (label, content) in enumerate(rows):
        set_cell(t.rows[i].cells[0], label, bold_first=True)
        if isinstance(content, str):
            content = [content]
        set_cell(t.rows[i].cells[1], content)
    doc.add_paragraph()


def add_callout(doc, title, lines):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.add_run(title).bold = True
    if isinstance(lines, str):
        lines = [lines]
    for line in lines:
        cell.add_paragraph(line)
    doc.add_paragraph()


def add_phase_callout(doc, phase_label, dok, minutes, title, body):
    """Callout leading with bold DOK badge + phase + time."""
    if isinstance(body, str):
        body = [body]
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.rows[0].cells[0]
    cell.text = ""
    # teal left-bar background
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), ps.COLORS["tealpale"])
    tc_pr.append(shd)
    ps.set_cell_borders(cell,
        left={"val": "single", "sz": 24, "color": ps.COLORS["teal"]},
        top={"val": "nil"}, right={"val": "nil"}, bottom={"val": "nil"})
    p = cell.paragraphs[0]
    r = p.add_run(f"[{dok}]  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(ps.COLORS["teal"])
    r = p.add_run(f"{phase_label}  ")
    r.bold = True
    r = p.add_run(f"({minutes} min)")
    r.font.color.rgb = RGBColor(0x55, 0x66, 0x77)
    p = cell.add_paragraph()
    p.add_run(title).bold = True
    for line in body:
        cell.add_paragraph(line)
    doc.add_paragraph()


def header_line(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def blank_lines(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.add_run("______________________________________________________________")


def _format_unicode(prompt):
    prompt = (prompt
              .replace("\u00e2\u2020\u2019", "\u2192")
              .replace("\u00e2\u20ac\u201d", "\u2014")
              .replace("\u00e2\u20ac\u201c", "\u2013"))
    SUPS = {"2": "\u00b2", "3": "\u00b3", "4": "\u2074",
            "5": "\u2075", "6": "\u2076"}
    out, i = [], 0
    while i < len(prompt):
        ch = prompt[i]
        if ch == "^" and i + 1 < len(prompt) and prompt[i+1] in SUPS:
            out.append(SUPS[prompt[i+1]])
            i += 2
            continue
        out.append("\u2212" if ch == "-" else ch)
        i += 1
    return "".join(out)


def _strip_stem(prompt):
    txt = _format_unicode(prompt)
    if ":" in txt:
        txt = txt.split(":", 1)[-1].strip()
    return txt.rstrip(".")


# ======================================================================
# DO NOW SHEET (Do Now A — bridge from Day 4-5)
# ======================================================================

def build_do_now(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin   = Inches(0.8)
        sec.right_margin  = Inches(0.8)

    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    header_line(doc, "Combined Day 6  |  DO NOW  \u2014  Solving a Polynomial Equation", size=13)

    p = doc.add_paragraph()
    r = p.add_run("Name: _____________________________________     Date: _____________")
    r.font.size = Pt(11)

    add_callout(doc, "[DOK 1]  \u2022  5 minutes  \u2022  Silent  \u2022  Pencil only  \u2022  No Desmos", [
        "Savvas Practice #21.  Find all solutions (real AND complex). Show your work.",
    ])

    q = qb.get_for_packet([DO_NOW_ID])[0]
    header_line(doc, "Savvas Practice #21:", size=12)
    p = doc.add_paragraph()
    p.add_run(_format_unicode(q["prompt"])).italic = True

    doc.add_paragraph()
    blank_lines(doc, 5)

    header_line(doc, "Your solutions:  x = _____________________________________", size=11)
    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run(
        "Bridge: these solutions are ZEROS of the polynomial.  Tomorrow we use "
        "zeros to slice the number line and solve INEQUALITIES.  Hold onto that "
        "idea."
    )
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x00, 0x97, 0xA7)

    doc.save(path)


# ======================================================================
# STUDENT PACKET
# ======================================================================

SIGN_CHART_RULES = [
    "RULE 1: Move everything to one side so the inequality reads  f(x) < 0  "
    "or  f(x) > 0.",
    "RULE 2: Factor f(x) completely.  Find all real zeros.",
    "RULE 3: Plot the zeros on a number line.  They divide it into intervals.",
    "RULE 4: Pick a TEST POINT in each interval.  Evaluate f at that point.",
    "RULE 5: Because f is continuous, sign only changes at zeros of ODD "
    "multiplicity.  Even-multiplicity zeros do NOT change sign.",
    "RULE 6: Shade intervals where the sign matches the inequality.  "
    "Zeros are included (\u2264, \u2265) or excluded (<, >).",
    "",
    "Sentence frame:  \u201cThe zeros are ___, so the sign changes at ___, "
    "and the solution is ___.\u201d",
]

LAUNCH_INVESTIGATION = [
    "SAVVAS EXAMPLE 6 / TRY IT 6a  \u2014  Walk Through the Sign-Chart Method",
    "",
    "Try It 6a:  Solve  2x\u00b3 + 12x\u00b2 + 12x < 0.",
    "",
    "Step 1 \u2014 Move to one side (already is). Factor out the GCF.",
    "    2x(x\u00b2 + 6x + 6) < 0",
    "",
    "Step 2 \u2014 Find zeros of  2x(x\u00b2 + 6x + 6).  Use Quadratic Formula on x\u00b2 + 6x + 6.",
    "    x = 0,   x = \u22123 \u2212 \u221a3 \u2248 \u22124.73,   x = \u22123 + \u221a3 \u2248 \u22121.27",
    "",
    "Step 3 \u2014 Plot zeros on the number line.  Label the four intervals.",
    "    ________|___________|__________|___________",
    "         \u22124.73       \u22121.27        0",
    "",
    "Step 4 \u2014 Test a point in each interval:",
    "    (pick a number that\u2019s easy to evaluate)",
    "    Interval 1 (x < \u22124.73):   test x = \u22125  \u2192  sign = ____",
    "    Interval 2 (\u22124.73 < x < \u22121.27):  test x = \u22123  \u2192  sign = ____",
    "    Interval 3 (\u22121.27 < x < 0):   test x = \u22121  \u2192  sign = ____",
    "    Interval 4 (x > 0):  test x = 1  \u2192  sign = ____",
    "",
    "Step 5 \u2014 Shade intervals where  2x(x\u00b2 + 6x + 6) < 0 (i.e., where sign = \u2212).",
    "    Solution:  ____________________________________________",
]

SHARE_SUMMARY_STUDENT = [
    "Self-check  \u2014  circle one for each:",
    "1.  I can factor a polynomial and identify all real zeros.            "
    "\u2713    partly    not yet",
    "2.  I can build a sign chart with the correct number of intervals.   "
    "\u2713    partly    not yet",
    "3.  I can read the solution set from the sign chart.                 "
    "\u2713    partly    not yet",
    "",
    "One-sentence synthesis:  \u201cThe key idea today was _______________________.\u201d",
    "___________________________________________________________________",
    "",
    "Preview Day 7:  synthesis + Savvas Performance Task #30 (Venetta, DOK 3).",
]


def _practice_lines():
    items = qb.get_for_packet(PRACTICE_IDS)
    lines = [
        "Use the sign-chart method: factor \u2192 zeros \u2192 number line \u2192 test points \u2192 shade.",
        "",
        "REQUIRED:  A.   PICK TWO of B, C, D.",
        "",
    ]
    labels = ["A.", "B.", "C.", "D."]
    for label, q in zip(labels, items):
        lines.append(f"{label}  {_format_unicode(q['prompt'])}")
        lines.append("    Zeros:  ____________________________________________________")
        lines.append("    Sign chart intervals:  ____________________________")
        lines.append("    Solution:  _________________________________________________")
        lines.append("")
    lines.append("(All items from Savvas Lesson 3-5.  "
                 "Bank IDs: " + ", ".join(PRACTICE_IDS) + ".)")
    return lines


def _exit_lines():
    # Summary recap exit — fill-in + one sentence.
    # Rule 4 (Summary exit, not CER): a 5-min walk-out task recaps today's learning.
    # 3-5-lq-q3 is reserved for Day 8 quiz.
    return [
        "Fill in each blank to recap what we learned today.",
        "",
        "1.  Zeros of a polynomial split the number line into ___________",
        "    where the sign _______________________________________.",
        "",
        "2.  On a sign chart, I test one point per ___________________ to determine",
        "    whether the inequality is satisfied.",
        "",
        "3.  The solution to a polynomial inequality is a(n) ___________________ of",
        "    intervals (those where the sign is correct).",
        "",
        "One sentence:",
        "  The biggest thing I learned today is ___________________________________",
        "  ______________________________________________________________________.",
    ]


def build_student(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin   = Inches(0.7)
        sec.right_margin  = Inches(0.7)

    ps.running_header_name_block(doc)
    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    ps.day_banner(doc, "6",
                  "Polynomial Inequalities (Savvas-anchored)",
                  "Sign-chart method: factor \u2192 zeros \u2192 number line \u2192 test points \u2192 solution.")

    add_two_col(doc, OBJECTIVES)
    add_two_col(doc, FRAMEWORK_HEADER)

    add_phase_callout(doc, "DO NOW", "DOK 1", "5",
        "Savvas Practice #21 \u2014 separate sheet",
        ["You got the equation sheet on entry. Hold onto it \u2014 "
         "we\u2019ll connect the zeros to today\u2019s number-line idea."])

    add_phase_callout(doc, "BLOOKET", "DOK 1", "2 + 6",
        "Rule Recall (on the board, not the leaderboard)",
        ["Rules: zero \u21d4 factor; sign chart; continuous polynomial; "
         "odd-multiplicity zero = sign change; even-multiplicity zero = "
         "same sign on both sides. No procedural factoring."])

    add_phase_callout(doc, "LAUNCH", "DOK 2", "12",
        "Sign-Chart Method \u2014 Savvas Example 6 / Try It 6a", [])
    add_callout(doc, "Sign-Chart Rules (keep these visible while you work):",
                SIGN_CHART_RULES)
    add_callout(doc, "", LAUNCH_INVESTIGATION)

    add_phase_callout(doc, "PRACTICE", "DOK 2", "20",
        "Savvas Try It 6b + Pick Two of Practice #22/#23/#24", [])
    add_callout(doc, "", _practice_lines())

    add_phase_callout(doc, "SHARE / SUMMARY", "DOK 2", "4",
        "Synthesis + self-rating",
        SHARE_SUMMARY_STUDENT)

    ps.summary_exit_box(doc,
        "[DOK 1\u20132]  EXIT TICKET  \u2014  5 min  \u2014  Summary Recap",
        _exit_lines())

    doc.save(path)


# ======================================================================
# TEACHER PACKET
# ======================================================================

CRITERIA_FOR_SUCCESS = [
    "Student-facing criteria (revisited during Share/Summary):",
    "1.  I can factor a polynomial and identify all real zeros.",
    "2.  I can build a sign chart with the correct number of intervals.",
    "3.  I can read the solution set from the sign chart.",
    "Formative evidence:",
    "\u2022  Do Now (Savvas Practice #21 equation \u2014 collect on entry or after Blooket).",
    "\u2022  Blooket per-student percentages on the dashboard.",
    "\u2022  Practice pages (Try It 6b + two of #22/#23/#24) \u2014 circulate + photograph.",
    "\u2022  Exit ticket (summary recap fill-in + one sentence, every student).",
]

BLOOKET_SCOPE = [
    "SCOPE (DOK 1 rule-recall only).  Rules today\u2019s work depends on:",
    "\u2022  Zero \u21d4 factor = 0 (Zero Product Property).",
    "\u2022  Sign of a polynomial can only change at a zero of odd multiplicity.",
    "\u2022  Even-multiplicity zero = same sign on both sides.",
    "\u2022  Continuous function \u2192 sign is constant on each interval between zeros.",
    "\u2022  Factor out GCF first; check Quadratic Formula if factor doesn\u2019t split.",
    "NO procedural inequality-solving items (those are DOK 2). "
    "Blooket CSV: regenerate via regen_blooket_csvs.py with tag blooket-pool day6.",
    "DISPLAY NOTE: project the Day 6 Slide 3 rule-recall list during play.",
    "Questions to ask (after game): \u201cWhich rule had the lowest percent? Say it now.\u201d",
]


def _do_now_key_lines():
    q = qb.get_for_packet([DO_NOW_ID])[0]
    correct = _format_unicode(str(q.get("correct", "")))
    return [
        f"[bank: {q['id']}, DOK {q.get('dok')}]",
        _format_unicode(q["prompt"]),
        "",
        f"\u2705  {correct}",
        "",
        "BRIDGE SCRIPT (30 sec at end of Do Now):",
        "  \u201cThese four solutions \u2014 x = 0, x = 3, x = 2i, x = \u22122i \u2014 are zeros of the",
        "  polynomial.  The real zeros (0 and 3) are the ones that sit on the number",
        "  line.  Today\u2019s question: what happens to the SIGN of the polynomial",
        "  between those zeros?  That\u2019s the key to solving inequalities.\u201d",
        "",
        "DIAGNOSTIC SCAN (on entry):",
        "\u2022  Students who wrote only x = 0 and x = 3 missed the complex roots.",
        "   Cold-call in Launch: \u201cAre complex zeros on the number line? Why not?\u201d",
        "\u2022  Students who wrote all four: acknowledge and move on.",
    ]


def _launch_key_lines():
    return [
        "SAVVAS EXAMPLE 6 / TRY IT 6a  \u2014  target walk-through",
        "",
        "  Solve  2x\u00b3 + 12x\u00b2 + 12x < 0.",
        "",
        "  Step 1: 2x(x\u00b2 + 6x + 6) < 0.  GCF = 2x.",
        "  Step 2: zeros at x = 0, x = \u22123 \u2212 \u221a3 (\u2248 \u22124.73), x = \u22123 + \u221a3 (\u2248 \u22121.27).",
        "          Quadratic Formula: x = (\u22126 \u00b1 \u221a(36\u221224)) / 2 = \u22123 \u00b1 \u221a3.",
        "  Step 3: Four intervals: x < \u22124.73 | \u22124.73 < x < \u22121.27 | \u22121.27 < x < 0 | x > 0.",
        "  Step 4: Test points \u2014 sign table:",
        "          x = \u22125 \u2192 2(\u22125)(pos) = NEG  \u2713  (< 0 \u2192 solution interval)",
        "          x = \u22123 \u2192 2(\u22123)(neg) = POS  \u2715",
        "          x = \u22121 \u2192 2(\u22121)(neg) = POS  \u2715",
        "          x =  1 \u2192 2(1)(pos)  = POS  \u2715",
        "  Step 5: Solution:  x < \u22123 \u2212 \u221a3  OR  \u22123 + \u221a3 < x < 0.",
        "          In interval notation: (\u2212\u221e, \u22123\u2212\u221a3) \u222a (\u22123+\u221a3, 0).",
        "",
        "GOTCHA: x = 0 has multiplicity 1 (odd) \u2014 SIGN CHANGES at x = 0.",
        "GOTCHA: the quadratic x\u00b2 + 6x + 6 does NOT factor over integers \u2014",
        "        students must use the Quadratic Formula. Warn them before they stall.",
        "",
        "LAUNCH STEP-BY-STEP (12 minutes, solo teacher):",
        "  0:00 \u2013 0:15  \u201cPull out your Do Now sheet.\u201d",
        "  0:15 \u2013 1:30  Turn-and-Talk: \u201cWhat real zeros did you find? Are they on the",
        "                number line? Why aren\u2019t the complex ones?\u201d",
        "  1:30 \u2013 2:30  Cold-call (low-stakes): \u201c[Name], which zeros go on the number",
        "                line?\u201d  \u2192 sincere \u201cthanks\u201d \u2192 \u201canyone want to build on that?\u201d",
        "  2:30 \u2013 5:00  Partners do Steps 1\u20132 on Try It 6a (factor + find zeros).",
        "                Teacher walks room. Check: is the GCF factored out?",
        "  5:00 \u2013 8:00  CLASS builds the number line together (Steps 3\u20134).",
        "                Teacher draws number line on the board; students fill in packet.",
        "  8:00 \u201310:00  Partners test one interval each (Step 5). Class shares.",
        " 10:00\u201312:00  NAME the method: \u201cThis is the SIGN CHART method.\u201d Write it on",
        "                the board. Students circle the rule box on the packet.",
        "",
        "Questions to ask (Savvas TE / Habits of Mind):",
        "  \u2022  \u201cDoes f(x) change sign between every pair of consecutive zeros?\u201d",
        "     Expected: No \u2014 even-multiplicity zeros keep the same sign on both sides.",
        "  \u2022  \u201cWhat information from Day 4-5 tells you the sign at very large x?\u201d",
        "     Expected: the leading coefficient + degree \u2192 end behavior.",
        "  \u2022  \u201cIf f(x) is positive for x > 3, what does that tell you about f(x) < 0?\u201d",
        "     Expected: f < 0 is false for x > 3; solution lives left of 3.",
    ]


def _practice_key_lines():
    items = qb.get_for_packet(PRACTICE_IDS)
    lines = []
    labels = ["A.", "B.", "C.", "D."]
    for label, q in zip(labels, items):
        lines.append(f"{label}  {_format_unicode(q['prompt'])}  [bank: {q['id']}, DOK {q.get('dok')}]")
        correct = _format_unicode(str(q.get("correct", "")))
        lines.append(f"    \u2705  {correct}")
        rationale = q.get("dok_rationale")
        if rationale:
            lines.append(f"    DOK rationale: {_format_unicode(rationale)}")
        lines.append("")
    lines += [
        "CIRCULATION PATTERN (solo teacher, ~10 students, ~5 pairs):",
        "  Lap 1 (min 0\u20135):   every pair has A started; zeros plotted on number line.",
        "  Lap 2 (min 5\u201310):  pairs choosing B vs C vs D; coach the sign-chart setup.",
        "  Lap 3 (min 10\u201315): pairs on their second chosen item; check sign-chart.",
        "  Lap 4 (min 15\u201320): fast-finishers extend; stuck pairs get one prompt, not answer.",
        "",
        "Questions to ask (NEVER deliver the answer):",
        "  \u2022  \u201cHow many intervals does your number line have?\u201d",
        "  \u2022  \u201cWhat\u2019s the sign of f at your test point?\u201d",
        "  \u2022  \u201cDoes this zero have odd or even multiplicity? Does the sign change?\u201d",
        "  \u2022  \u201cYour solution has a strict inequality \u2014 should zeros be included?\u201d",
    ]
    return lines


def _exit_summary_key_lines():
    # Summary-exit answer key.  No bank item — these are vocabulary recalls.
    return [
        "Fill-in answers:",
        "1.  \u201cintervals\u201d  /  \u201ccan only change at a zero of odd multiplicity\u201d",
        "    (accept: \u201csegments\u201d, \u201cregions\u201d for \u2018intervals\u2019)",
        "2.  \u201cinterval\u201d  (one representative test point per interval)",
        "3.  \u201cunion\u201d  (accept: \u201ccollection\u201d)",
        "",
        "One-sentence scan (monitor for misconceptions):",
        "\u2022  Look for vocabulary: zeros, sign chart, intervals, inequality.",
        "\u2022  Flag: students who write only \u2018we did sign charts\u2019 with no further content \u2192 "
        "target in Day 7 Launch.",
        "\u2022  Note: 3-5-lq-q3 is reserved for Day 8 quiz (not used here).",
    ]


IEP_MODS = [
    "\u2022  Launch is a step-by-step walk-through on the student packet (six numbered steps).",
    "\u2022  Sign-chart rules are printed on the packet \u2014 students reference at any time.",
    "\u2022  Practice is structured per item (zeros + chart + solution on three lines).",
    "\u2022  Choice in Practice: A required, pick any two of B/C/D. Reduces overwhelm.",
    "\u2022  Exit mirrors Launch structure \u2014 same verbs, same shape.",
    "\u2022  Desmos explicitly allowed on the exit to verify zeros.",
]

ELL_SUPPORTS = [
    "\u2022  Sentence frame printed on the Launch and rule box: \u201cThe zeros are ___,",
    "  so the sign changes at ___, and the solution is ___.\u201d",
    "\u2022  \u201cSign chart\u201d: bridge to \u201csign\u201d (positive/negative) vs \u201csign\u201d (drawing/symbol).",
    "\u2022  \u201cInterval\u201d: bridge to \u201cinterval in music\u201d = distance between two notes.",
    "\u2022  Visual number line on the packet keeps the spatial reasoning concrete.",
    "\u2022  Test-point strategy = \u201cpick one number and just check\u201d \u2014 name the strategy",
    "  explicitly; ELL students often know the arithmetic but not which move to make.",
]


def build_teacher(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin   = Inches(0.7)
        sec.right_margin  = Inches(0.7)

    header_line(doc, "TEACHER EDITION", size=12)
    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    ps.day_banner(doc, "6",
                  "Polynomial Inequalities (Savvas-anchored)",
                  "Solo teacher. Class of ~10. Sign-chart method. DOK framework "
                  "crosswalk + answer keys + Questions to ask per phase.")

    add_two_col(doc, OBJECTIVES + [("CCSS", "A-APR.B.3, F-IF.C.7c, MP.1, MP.6")])
    add_two_col(doc, FRAMEWORK_HEADER)

    add_callout(doc, "\U0001F4CB  DESIGN \u2014  SAVVAS-ONLY PRACTICE DAY (no DOK 3)", [
        "Do Now A (bridge) \u2192 Blooket (B+C) \u2192 Wrap \u2192 Launch \u2192 Practice \u2192 Share/Summary \u2192 Exit.",
        "[DOK 1] Blooket; Do Now #21 (equation).",
        "[DOK 2] Launch sign-chart walk-through; Practice Try It 6b + #22/#23/#24.",
        "[DOK 1\u20132] Exit = summary recap (fill-in + one sentence).  3-5-lq-q3 reserved for Day 8.",
        "[DOK 3] NOT on this day.  Savvas DOK-3 \u2192 Day 7 (Practice #30, Venetta deli).",
        "Every student-facing work item traces to a bank ID. No fabricated tasks.",
    ])

    add_callout(doc, "\u2705  CRITERIA FOR SUCCESS  /  FORMATIVE ASSESSMENT",
                CRITERIA_FOR_SUCCESS)

    add_callout(doc, "\u23f1\ufe0f  REALISTIC PACING (solo teacher, ~10 students)", [
        "Total: 55 min.  Fits 55-min Tue A; 65-min F has 10 min slack.",
        "  Do Now A (Practice #21) ........ 5 min",
        "  Do Now B (Blooket login) ....... 2 min",
        "  Do Now C (Blooket game) ........ 6 min",
        "  Blooket Wrap ................... 1 min",
        "  Launch (Example 6 / Try It 6a)..12 min",
        "  Practice (6b + 2 of #22/23/24)..20 min",
        "  Share/Summary .................. 4 min",
        "  Exit (summary recap) ........... 5 min",
        "RELEASE VALVES (ordered):",
        "  1. Cut Practice to A + ONE choice (save 7 min for stuck pairs).",
        "  2. Cut Launch from 12 \u2192 10 min (skip extended class share).",
        "  3. Cut Share/Summary to 3 min.",
        "  NEVER cut Exit.  Only artifact that documents learning.",
        "65-min F slack: all four Practice items; run ELL sentence-frame scaffold during Launch.",
    ])

    ps.framework_phase_header(doc,
        phase="Do Now A \u2014 Savvas Practice #21 (bridge equation)",
        framework_tag="Do Now A", dok="1", minutes=5,
        teacher_does=["Distribute sheets at the door. Walk one lap as students work.",
                      "Scan what they wrote: check if they included complex roots."],
        students_do=["Silent, pencil only. Solve the polynomial equation. Show work."],
        questions_to_ask=["(after) \u201cWhich of these solutions would you plot on a number line?\u201d",
                           "(bridge) \u201cThe real zeros divide the number line \u2014 what do you think is between them?\u201d"],
        adult_role="Silent circulation. Note the 2-3 students who missed complex roots (target for Launch cold-call).")

    add_callout(doc, "\u2705  DO NOW A / ANSWER KEY", _do_now_key_lines())

    ps.framework_phase_header(doc,
        phase="Do Now B+C \u2014 Blooket Login + Rule Recall",
        framework_tag="Do Now B/C", dok="1", minutes="2 + 6",
        teacher_does=["Display Blooket code. Load Day 6 rule-recall CSV."],
        students_do=["Open Chromebooks, enter code + nickname. Play 6-min Blooket game."],
        questions_to_ask=["(after wrap) \u201cWhich rule had the lowest percent? Everyone say it aloud.\u201d"],
        adult_role="Monitor dashboard. Note two lowest rules; use them in Blooket Wrap verbal repetition.")

    add_callout(doc, "[Do Now C]  BLOOKET SCOPE", BLOOKET_SCOPE)

    ps.framework_phase_header(doc,
        phase="Blooket Wrap \u2014 Candy + Lowest Rule",
        framework_tag="Blooket Wrap", dok="1", minutes=1,
        teacher_does=["Pass candy jar. Announce lowest rule. Class repeats aloud."],
        students_do=["Receive candy (if present). Repeat lowest-scored rule in unison."],
        questions_to_ask=["\u201cSay it with me: [lowest rule text].\u201d  One volunteer repeats it once more."],
        adult_role="Facilitate 60-second candy + verbal ritual. Keeps energy up before Launch.")

    ps.framework_phase_header(doc,
        phase="Launch \u2014 Sign-Chart Method (Example 6 / Try It 6a)",
        framework_tag="Launch", dok="2", minutes=12,
        teacher_does=["Walk through sign-chart steps on the board / packet.",
                      "Walk laps reading student sign charts."],
        students_do=["Graph observations in Desmos, then build sign chart for Try It 6a."],
        questions_to_ask=["Does f(x) change sign between every pair of zeros?",
                           "What does the leading coefficient tell you about the sign for large x?"],
        adult_role="Facilitator + circulation. Students build the chart; teacher prompts, not tells.")

    add_callout(doc, "\u2705  LAUNCH / ANSWER KEY + TEACHER SCRIPT", _launch_key_lines())

    ps.framework_phase_header(doc,
        phase="Practice \u2014 Try It 6b + Pick Two of #22/#23/#24",
        framework_tag="Practice", dok="2", minutes=20,
        teacher_does=["Circulate 4 laps. Prompt with questions, not answers.",
                      "Photograph practice pages as evidence."],
        students_do=["A (Try It 6b) required.  Choose two of B (#22), C (#23), D (#24). "
                     "Sign chart for each."],
        questions_to_ask=["How many zeros? How many intervals?",
                           "Is this zero odd or even multiplicity \u2014 does the sign change?",
                           "\u201cStrict inequality \u2014 are zeros included or excluded?\u201d"],
        adult_role="Solo circulation. 4 laps for ~5 pairs. Photo evidence replaces collecting papers.")

    add_callout(doc, "\u270F\uFE0F  PRACTICE / ANSWER KEY + CIRCULATION PATTERN",
                _practice_key_lines())

    ps.framework_phase_header(doc,
        phase="Share / Summary \u2014 Synthesis + self-rating",
        framework_tag="Share/Summary", dok="2", minutes=4,
        teacher_does=["Call on 2 students for sentence-frame usage.",
                      "Scan self-ratings as packets tilt up."],
        students_do=["Circle self-rating. Write one-sentence synthesis. Hear Day 7 preview."],
        questions_to_ask=["Which criterion is \u201cpartly\u201d? What\u2019s the specific blocker?",
                           "\u201cWho used the sentence frame? Read yours aloud.\u201d"],
        adult_role="Teacher script (4 min flat). Preview Day 7 DOK-3 driver (Venetta). Transition to exit.")

    add_callout(doc, "\U0001F501  SHARE / SUMMARY (student-facing)", SHARE_SUMMARY_STUDENT)

    ps.framework_phase_header(doc,
        phase="Exit Ticket \u2014 Summary Recap (fill-in + one sentence)",
        framework_tag="Exit Ticket", dok="1-2", minutes=5,
        teacher_does=["No questions asked during write time. Circulate silently.",
                      "Scan \u2018biggest thing\u2019 sentences at the door."],
        students_do=["Fill in the three blanks. Write one sentence: biggest thing learned today.",
                     "No notes, no Desmos."],
        questions_to_ask=["(post-collect) \u201cBiggest thing \u2014 one word. Go.\u201d"],
        adult_role="Silent collection at door. Skim criterion 2 self-ratings for Day 7 seating.")

    ps.summary_exit_box(doc, "\U0001F4E4  EXIT TICKET (student-facing)", _exit_lines())
    add_callout(doc, "\u2705  EXIT TICKET / ANSWER KEY", _exit_summary_key_lines())

    add_callout(doc, "\U0001F4F7  WHAT TO COLLECT", [
        "1.  Do Now sheets (Practice #21).",
        "2.  Practice pages (circulate + photograph during Practice phase).",
        "3.  Exit tickets (summary recap, every student).",
        "4.  Share/Summary self-ratings + one-sentence synthesis.",
    ])
    add_callout(doc, "\U0001F9E9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS", IEP_MODS)
    add_callout(doc, "\U0001F30D  SUPPORTS FOR ELL STUDENTS", ELL_SUPPORTS)

    doc.save(path)


if __name__ == "__main__":
    build_do_now("Day_6_Do_Now.docx")
    build_student("Day_6_Student_Packet.docx")
    build_teacher("Day_6_Teacher_Packet.docx")
    print("Built Day_6_Do_Now.docx, Day_6_Student_Packet.docx, Day_6_Teacher_Packet.docx")
    from packet_styles import emit_visuals_checklist
    emit_visuals_checklist(_ALL_IDS, "Day_6_Visuals_Checklist.md",
                           title="Day 6 - Visuals Checklist")
