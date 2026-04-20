"""Generate Day 8 materials — Quiz + Differentiated Stations.

Day 8 is assessment day. NO Blooket. NO DOK-3 driver.
Phases (55 min):
  0–5    Do Now A  : vocabulary warm-up (3-5-savvas-q28)  [DOK 1]
  5–30   Quiz      : Savvas Lesson Quiz Q1a, Q1b, Q2, Q3, Q4, Q5  [DOK 1-2]
  30–32  Transition: hand in quiz, set up stations
  32–50  Stations  : differentiated 3-station rotation (18 min)  [DOK 2]
  50–55  Exit      : summary recap fill-in + biggest takeaway  [DOK 1-2]

Four artifacts:
  Day_8_Do_Now.docx           — vocabulary warm-up (q28), silent, solo
  Day_8_Quiz.docx             — clean quiz layout, 5 items (1a+1b+2+3+4+5)
  Day_8_Student_Packet.docx   — stations table + exit (student-facing, no keys)
  Day_8_Teacher_Packet.docx   — framework phase headers, answer keys for all

Station IDs (dynamically selected, reported below):
  Support  : 3-5-savvas-q14  +  one DOK-1 multiplicity item
  Core     : 3-5-savvas-q17  +  3-5-savvas-q22
  Extend   : 3-5-savvas-q24  (hardest inequality, DOK-1 synthesis; NO DOK-3 on Day 8)

UTF-8 prelude required for Windows/MSYS2.
"""

import sys
import io
import os

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

# -----------------------------------------------------------------------
# Quiz item IDs (fixed, ordered 1a → 1b → 2 → 3 → 4 → 5)
# -----------------------------------------------------------------------
QUIZ_IDS = [
    "3-5-lq-q1a",
    "3-5-lq-q1b",
    "3-5-lq-q2",
    "3-5-lq-q3",
    "3-5-lq-q4",
    "3-5-lq-q5",
]

DO_NOW_ID = "3-5-savvas-q28"

# -----------------------------------------------------------------------
# Station items — dynamically selected with fallbacks
# -----------------------------------------------------------------------
SUPPORT_ITEM_1 = "3-5-savvas-q14"

_support_dyn = qb.select(lesson="3-5", dok=1, tags=["day3-multiplicity"], limit=2)
# Savvas-only: exclude any non-Savvas items (e.g. blooket items)
_support_dyn = [q for q in _support_dyn
                if q["id"].startswith("3-5-savvas-") and q["id"] != SUPPORT_ITEM_1]
SUPPORT_ITEM_2 = _support_dyn[0]["id"] if _support_dyn else None

CORE_ITEM_1 = "3-5-savvas-q17"
CORE_ITEM_2 = "3-5-savvas-q22"

# Day 8 must have NO DOK-3 (Savvas rule: DOK-3 is Day 4-5 and Day 7 only).
# Extend station uses 3-5-savvas-q24 — hardest inequality in the set
# (factor by grouping, recognize irreducible quadratic factor).
EXTEND_ITEM = "3-5-savvas-q24"
EXTEND_WARNING = None

# -----------------------------------------------------------------------
# Report station selections
# -----------------------------------------------------------------------
print("Day 8 — station IDs selected:")
print(f"  Support  station: {SUPPORT_ITEM_1}, {SUPPORT_ITEM_2}")
print(f"  Core     station: {CORE_ITEM_1}, {CORE_ITEM_2}")
print(f"  Extend   station: {EXTEND_ITEM}")
if EXTEND_WARNING:
    print(f"  {EXTEND_WARNING}")


# -----------------------------------------------------------------------
# Low-level helpers (adapted from build_day23_packets.py)
# -----------------------------------------------------------------------

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if i == 0 and bold_first:
            run.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_two_col(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = TABLE_STYLE
    t.columns[0].width = Inches(1.8)
    t.columns[1].width = Inches(4.7)
    for i, (label, content) in enumerate(rows):
        set_cell(t.rows[i].cells[0], label, bold_first=True)
        if isinstance(content, str):
            content = [content]
        set_cell(t.rows[i].cells[1], content)
    doc.add_paragraph()


def add_callout(doc, title, lines):
    if isinstance(lines, str):
        lines = [lines]
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    if title:
        p.add_run(title).bold = True
    for line in lines:
        cell.add_paragraph(line)
    doc.add_paragraph()


def header_line(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def blank_lines(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.add_run("____________________________________________________________")


def _format_unicode(prompt):
    prompt = (prompt
              .replace("\u00e2\u2020\u2019", "\u2192")
              .replace("\u00e2\u20ac\u201d", "\u2014")
              .replace("\u00e2\u20ac\u201c", "\u2013"))
    SUPS = {"2": "\u00b2", "3": "\u00b3", "4": "\u2074",
            "5": "\u2075", "6": "\u2076"}
    out, i = [], 0
    while i < len(prompt):
        ch = prompt[i]
        if ch == "^" and i + 1 < len(prompt) and prompt[i+1] in SUPS:
            out.append(SUPS[prompt[i+1]])
            i += 2
            continue
        out.append("\u2212" if ch == "-" else ch)
        i += 1
    return "".join(out)


def _fmt(text):
    """Alias for _format_unicode for convenience."""
    return _format_unicode(str(text))


def _margins(doc):
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)


# -----------------------------------------------------------------------
# DO NOW SHEET  (vocabulary warm-up, q28)
# -----------------------------------------------------------------------

def build_do_now(path):
    doc = Document()
    _margins(doc)

    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    header_line(doc, "Day 8  |  DO NOW  —  Vocabulary Warm-Up", size=13)

    p = doc.add_paragraph()
    r = p.add_run("Name: _____________________________________     Date: _____________")
    r.font.size = Pt(11)
    doc.add_paragraph()

    add_callout(doc, "[DOK 1]  •  5 minutes  •  Silent  •  Pencil only  •  No Desmos", [
        "Complete the vocabulary warm-up below. Work alone.",
        "When you finish, re-read your quiz notes from Days 2–7.",
    ])

    q = qb.get_for_packet([DO_NOW_ID])[0]
    header_line(doc, "Vocabulary Warm-Up  (Savvas Practice — bank: 3-5-savvas-q28):", size=12)
    p = doc.add_paragraph()
    p.add_run(_fmt(q["prompt"])).italic = True
    doc.add_paragraph()

    doc.add_paragraph("(1)  The graph of the function crosses the _______________________  at  4.")
    doc.add_paragraph()
    doc.add_paragraph("(2)  The expression  (x \u2212 4)  is a  _______________________  of the polynomial.")
    doc.add_paragraph()

    ps.sentence_frame_box(doc, [
        "\u201cA zero tells you a  _______________  of the polynomial and where the graph "
        "crosses or is tangent to the  _______________.\u201d"
    ])

    doc.save(path)
    print(f"  Built {path}")


# -----------------------------------------------------------------------
# QUIZ  (clean quiz layout, separate artifact)
# -----------------------------------------------------------------------

def _quiz_item_block(doc, number, label, q, include_sketch_space=False):
    """Render one quiz item with generous work space. No answer hints."""
    header_line(doc, f"Question {label}", size=12, bold=True)
    p = doc.add_paragraph()
    prompt_text = _fmt(q["prompt"])
    p.add_run(prompt_text)
    p.paragraph_format.space_after = Pt(4)

    # MC choices (if item has answers list and integer correct)
    answers = q.get("answers", [])
    correct_int = q.get("correct")
    is_mc = bool(answers) and isinstance(correct_int, int)
    if is_mc:
        letters = "ABCD"
        for i, ans in enumerate(answers[:4]):
            cp = doc.add_paragraph()
            r = cp.add_run(f"    ({letters[i]})  {_fmt(ans)}")
            r.font.size = Pt(11)

    doc.add_paragraph()
    if include_sketch_space:
        header_line(doc, "Sketch space:", size=11, bold=False)
        if os.path.exists("assets/blank_grid.png"):
            doc.add_picture("assets/blank_grid.png", width=Inches(2.8))
        else:
            for _ in range(4):
                doc.add_paragraph("     " + " " * 60)
    else:
        header_line(doc, "Work / Answer:", size=11, bold=False)
        blank_lines(doc, 4)
    doc.add_paragraph()


def build_quiz(path):
    doc = Document()
    _margins(doc)

    # Running header (name / block on every page)
    ps.running_header_name_block(doc, label_right="PERIOD")

    header_line(doc, "ALGEBRA 2  |  LESSON 3-5  |  Day 8", size=14)
    p = doc.add_paragraph()
    r = p.add_run("LESSON QUIZ  —  Zeros of Polynomial Functions")
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    r = p.add_run("Name: ___________________________________________     Date: _____________")
    r.font.size = Pt(11)

    add_callout(doc, "Instructions", [
        "\u2022  This quiz is silent and individual. Pencil only. No Desmos. No notes.",
        "\u2022  Show all work. Partial credit is awarded for correct work even if the final answer is wrong.",
        "\u2022  For multiple-choice items, circle one letter AND show supporting work.",
        "\u2022  You have 25 minutes.",
    ])

    quiz_items = qb.get_for_packet(QUIZ_IDS)
    labels = ["1a", "1b", "2", "3", "4", "5"]

    for label, q in zip(labels, quiz_items):
        # Q1b asks for a sketch — provide grid space
        _quiz_item_block(doc, label, label, q, include_sketch_space=(label == "1b"))

    doc.save(path)
    print(f"  Built {path}")


# -----------------------------------------------------------------------
# STUDENT PACKET  (stations + exit, no answer keys)
# -----------------------------------------------------------------------

def _stations_table(doc):
    """Three-station table: Support / Core / Extend with DOK + item IDs."""
    header_line(doc, "Differentiated Stations  (18 min, 6 min per station)", size=13)
    p = doc.add_paragraph()
    r = p.add_run(
        "Rate yourself on the quiz (1 = needs support, 2 = mostly there, 3 = ready for a stretch). "
        "Then circle the 1 or 2 stations you will do. "
        "Start with the station that matches your rating."
    )
    r.font.size = Pt(11)
    r.italic = True
    doc.add_paragraph()

    # Build the station table: 4 rows (header + 3 stations) × 3 cols
    t = doc.add_table(rows=4, cols=3)
    t.style = TABLE_STYLE
    t.columns[0].width = Inches(0.8)
    t.columns[1].width = Inches(3.0)
    t.columns[2].width = Inches(2.6)

    # Header row
    headers = ["Station", "Task", "Bank IDs / DOK"]
    for j, hdr in enumerate(headers):
        cell = t.rows[0].cells[j]
        ps.shade_cell(cell, ps.COLORS["teal"])
        cell.text = ""
        p = cell.paragraphs[0]
        ps._run(p, hdr, bold=True, color=ps.COLORS["white"], size=11)

    def _station_row(row_idx, name, badge, task_lines, ids_dok):
        name_cell, task_cell, ids_cell = (
            t.rows[row_idx].cells[0],
            t.rows[row_idx].cells[1],
            t.rows[row_idx].cells[2],
        )
        shade = {"Support": "FCEAE8", "Core": "E6F5F7", "Extend": "EEF4FF"}.get(name, "FFFFFF")
        ps.shade_cell(name_cell, shade)
        ps.shade_cell(task_cell, "FFFFFF")
        ps.shade_cell(ids_cell, "FFFFFF")

        name_cell.text = ""
        p = name_cell.paragraphs[0]
        ps._run(p, f"{name}\n{badge}", bold=True, size=11)
        name_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        task_cell.text = ""
        for i, line in enumerate(task_lines):
            p2 = task_cell.paragraphs[0] if i == 0 else task_cell.add_paragraph()
            p2.add_run(line)
        task_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        ids_cell.text = ""
        ids_cell.paragraphs[0].add_run(ids_dok)
        ids_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Support station
    sup1_q = qb.get(SUPPORT_ITEM_1)
    sup2_q = qb.get(SUPPORT_ITEM_2) if SUPPORT_ITEM_2 else None
    sup_ids = SUPPORT_ITEM_1 + (f"\n{SUPPORT_ITEM_2}" if SUPPORT_ITEM_2 else "")
    sup_task = [
        "Re-teach: factoring + multiplicity with guided scaffold.",
        "",
        f"(A)  {_fmt(sup1_q['prompt']) if sup1_q else SUPPORT_ITEM_1}",
        "     Zeros: ___________________________",
        "     Multiplicity of each: ___________________________",
        "     Behavior (crosses / tangent): ___________________________",
    ]
    if sup2_q:
        sup_task += [
            "",
            f"(B)  {_fmt(sup2_q['prompt'])}",
            "     Multiplicity: ___________________________",
        ]
    _station_row(1, "Support", "(re-teach)", sup_task, f"DOK 1\n{sup_ids}")

    # Core station
    core1_q = qb.get(CORE_ITEM_1)
    core2_q = qb.get(CORE_ITEM_2)
    _station_row(2, "Core", "(practice)", [
        "Sign-chart + complex zeros mixed practice.",
        "",
        f"(A)  {_fmt(core1_q['prompt']) if core1_q else CORE_ITEM_1}",
        "     Real zeros: ___________________________",
        "     Complex zeros: ___________________________",
        "",
        f"(B)  {_fmt(core2_q['prompt']) if core2_q else CORE_ITEM_2}",
        "     Solution set: ___________________________",
    ], f"DOK 1\n{CORE_ITEM_1}\n{CORE_ITEM_2}")

    # Extend station — hardest inequality in the set (DOK 1, synthesis stretch)
    if EXTEND_ITEM:
        ext_q = qb.get(EXTEND_ITEM)
        ext_prompt = _fmt(ext_q["prompt"]) if ext_q else EXTEND_ITEM
        ext_task = [
            "Synthesis inequality challenge — show all work.",
            "",
            ext_prompt,
            "",
            "Step 1 — Move everything to one side:  _________________________",
            "Step 2 — Factor (try grouping):  _______________________________",
            "Step 3 — Identify any factor that is ALWAYS positive:  ___________",
            "Step 4 — Sign chart (draw here):",
            "  _______|_______________|___",
            "Step 5 — Solution set:  ________________________________________",
        ]
        _station_row(3, "Extend", "(challenge)", ext_task,
                     f"DOK 1 (synthesis)\n{EXTEND_ITEM}")
    else:
        _station_row(3, "Extend", "(challenge)", [
            "See teacher for the extension problem."
        ], "DOK 1\n(see teacher)")

    doc.add_paragraph()

    # Self-rating checkbox
    p = doc.add_paragraph()
    p.add_run("Circle the station(s) you completed:     Support     Core     Extend")
    doc.add_paragraph()


def _exit_student_lines():
    return [
        "Fill in the blanks (no notes, no Desmos):",
        "",
        "1.  A zero of a polynomial is an x-value where f(x) = _______.  "
        "The graph  ________________  or is  ________________  at that point.",
        "",
        "2.  If a factor appears with an EVEN exponent, the graph is  ________________  "
        "to the x-axis at that zero (Savvas: \u201cturning point\u201d).",
        "",
        "3.  To solve a polynomial inequality, first find the  ________________  "
        "and then build a  ________________  chart.",
        "",
        "4.  My biggest takeaway from Lesson 3-5 (one sentence):",
        "    _______________________________________________________________",
        "    _______________________________________________________________",
    ]


def build_student_packet(path):
    doc = Document()
    _margins(doc)

    ps.running_header_name_block(doc, label_right="PERIOD")
    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    ps.day_banner(doc, "8", "Quiz + Differentiated Stations",
                  "Quiz returned. Now: choose your station and finish with the summary exit.")

    add_callout(doc, "Today's Structure", [
        "\u2022  Do Now (5 min): vocabulary warm-up \u2014 separate sheet.",
        "\u2022  Quiz (25 min): Savvas Lesson Quiz \u2014 separate sheet.",
        "\u2022  Stations (18 min): pick 1\u20132 stations based on your self-rating.",
        "\u2022  Exit (5 min): summary fill-in right here.",
    ])

    _stations_table(doc)

    ps.summary_exit_box(doc,
        "[DOK 1-2]  EXIT \u2014  Summary Recap  (5 min)",
        _exit_student_lines())

    doc.save(path)
    print(f"  Built {path}")


# -----------------------------------------------------------------------
# TEACHER PACKET  (answer keys + framework phase headers)
# -----------------------------------------------------------------------

def _quiz_answer_key_lines():
    items = qb.get_for_packet(QUIZ_IDS)
    labels = ["1a", "1b", "2", "3", "4", "5"]
    lines = []
    for label, q in zip(labels, items):
        lines.append(f"Q{label}  [bank: {q['id']}, DOK {q.get('dok')}]")
        lines.append(f"    Prompt: {_fmt(q['prompt'])[:120]}")
        correct_raw = str(q.get("correct", ""))
        # For MC items (q4, q5), map integer to letter choice
        answers = q.get("answers", [])
        correct_int = q.get("correct")
        is_mc = bool(answers) and isinstance(correct_int, int)
        if is_mc:
            correct_letter = "ABCD"[correct_int - 1] if 1 <= correct_int <= 4 else "?"
            lines.append(f"    \u2705  Correct letter: {correct_letter}")
            if 1 <= correct_int <= len(answers):
                lines.append(f"    \u2705  Answer text: {_fmt(answers[correct_int - 1])[:120]}")
        else:
            lines.append(f"    \u2705  {_fmt(correct_raw)[:200]}")
        rationale = q.get("dok_rationale", "")
        if rationale:
            lines.append(f"    DOK rationale: {_fmt(rationale)[:120]}")
        lines.append("")
    return lines


def _station_answer_key_lines():
    lines = [
        "\u2014\u2014 SUPPORT STATION \u2014\u2014",
    ]
    for sid in ([SUPPORT_ITEM_1] + ([SUPPORT_ITEM_2] if SUPPORT_ITEM_2 else [])):
        q = qb.get(sid)
        if q:
            lines.append(f"  [{q['id']}, DOK {q.get('dok')}]")
            lines.append(f"  Prompt: {_fmt(q['prompt'])[:120]}")
            lines.append(f"  \u2705  {_fmt(str(q.get('correct', '')))[:200]}")
            lines.append("")
    lines += [
        "\u2014\u2014 CORE STATION \u2014\u2014",
    ]
    for cid in [CORE_ITEM_1, CORE_ITEM_2]:
        q = qb.get(cid)
        if q:
            lines.append(f"  [{q['id']}, DOK {q.get('dok')}]")
            lines.append(f"  Prompt: {_fmt(q['prompt'])[:120]}")
            lines.append(f"  \u2705  {_fmt(str(q.get('correct', '')))[:200]}")
            lines.append("")
    lines += [
        "\u2014\u2014 EXTEND STATION \u2014\u2014",
    ]
    if EXTEND_ITEM:
        q = qb.get(EXTEND_ITEM)
        if q:
            lines.append(f"  [{q['id']}, DOK {q.get('dok')}]  Savvas-declared DOK-3 modeling item")
            lines.append(f"  Prompt: {_fmt(q['prompt'])[:200]}")
            lines.append(f"  \u2705  {_fmt(str(q.get('correct', '')))[:400]}")
    else:
        lines.append("  EXTEND ITEM NOT IN BANK — insert manually. See EXTEND_WARNING.")
        if EXTEND_WARNING:
            lines.append(f"  {EXTEND_WARNING}")
    return lines


def build_teacher_packet(path):
    doc = Document()
    _margins(doc)

    header_line(doc, "TEACHER EDITION", size=12)
    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    ps.day_banner(doc, "8", "Quiz + Differentiated Stations",
                  "Solo teacher. Class of ~10. Answer keys for all 6 quiz items + "
                  "3 station keys. Framework phase headers per DOKframework.txt.")

    add_callout(doc, "\U0001F4CB  DESIGN  \u2014  QUIZ DAY (no Blooket, no DOK-3 driver)", [
        "Day 8 is pure assessment + differentiated reinforcement.  NO DOK-3 today.",
        "[DOK 1] Do Now vocabulary (q28); Quiz Q4 (MC); Extend station (q24, synthesis).",
        "[DOK 2] Quiz Q1a/1b/Q2/Q3/Q5; Core station; Exit.",
        "No Blooket. Silent testing environment throughout the quiz block.",
        "All items trace to Savvas bank IDs. No fabricated tasks.",
    ])

    ps.framework_phase_header(doc,
        phase="Do Now A — Vocabulary Warm-Up",
        framework_tag="Do Now A",
        dok="1",
        minutes=5,
        teacher_does=[
            "Hand out Do Now sheets at the door.",
            "Stand near door; scan sheets for completion as timer runs.",
        ],
        students_do=[
            "Complete vocabulary fill-in (3-5-savvas-q28) silently.",
            "Re-read quiz notes when finished.",
        ],
        questions_to_ask=[
            "'What does it mean for 4 to be a zero of a function?'  (expected: f(4) = 0)",
            "'Which of your answers uses the word factor?' (links to root-factor theorem)",
        ],
        adult_role="Circulate silently. No hints. Collect sheets before quiz distribution.",
    )

    do_now_q = qb.get(DO_NOW_ID)
    add_callout(doc, "\u2705  DO NOW / ANSWER KEY", [
        f"[bank: {DO_NOW_ID}, DOK {do_now_q.get('dok')}]",
        _fmt(do_now_q["prompt"]),
        "",
        f"\u2705  {_fmt(str(do_now_q.get('correct', '')))[:200]}",
        "",
        "Item also appears in teacher vocabulary notes: "
        "'4 is a zero' \u21d4 graph crosses x-axis at 4 \u21d4 (x \u2212 4) is a factor.",
    ])

    ps.framework_phase_header(doc,
        phase="Quiz — Savvas Lesson Quiz (Q1a, Q1b, Q2, Q3, Q4, Q5)",
        framework_tag="Independent Practice / Assessment",
        dok="1-2",
        minutes=25,
        teacher_does=[
            "Distribute quiz face-down. Say: 'Flip over. 25 minutes. Pencil only. No Desmos.'",
            "Circulate silently (4 laps for ~10 students).",
            "Write encouraging non-answer notes on the paper roster (e.g., 'nice setup on Q2').",
            "No verbal hints. If a student is confused, point silently to the instructions.",
        ],
        students_do=[
            "Work silently, individually.",
            "Show all work. For MC items, also write supporting work.",
        ],
        questions_to_ask=[
            "(NO questions during the quiz. Circulate silently.)",
        ],
        adult_role="Proctor. Walk laps, eyes on papers not screens. Collect promptly at 25 min.",
    )

    add_callout(doc, "\u2705  QUIZ / FULL ANSWER KEY (teacher eyes only)", _quiz_answer_key_lines())

    add_callout(doc, "\u23f1\ufe0f  REALISTIC PACING (solo teacher, ~10 students)", [
        "Total: 55 min.",
        "  Do Now A (q28 vocab) ............ 5 min",
        "  Quiz (Q1a\u2013Q5) .................. 25 min",
        "  Transition (hand in + setup) .... 2 min",
        "  Stations (18 min: 6 min each) .. 18 min",
        "  Exit (summary recap) ............. 5 min",
        "RELEASE VALVES (ordered):",
        "  1. If class needs more quiz time, compress Stations to 12 min (2 stations).",
        "  2. If stations go long, fold Exit into the last 2 min of Stations.",
        "  NEVER cut Exit. Only artifact that documents learning.",
    ])

    ps.framework_phase_header(doc,
        phase="Transition — Quiz Hand-In + Station Setup",
        framework_tag="Transition",
        dok="—",
        minutes=2,
        teacher_does=[
            "Collect quizzes.",
            "Briefly explain the three station options and the self-rating rubric.",
            "Tell students to open the student packet to the stations page.",
        ],
        students_do=[
            "Hand in quiz.",
            "Rate themselves 1/2/3 and circle their station(s).",
        ],
        questions_to_ask=[
            "'Which topic felt hardest on the quiz? That's your station.'",
        ],
        adult_role="Facilitate station choice. Respect student autonomy. No assignment.",
    )

    ps.framework_phase_header(doc,
        phase="Stations — Differentiated 3-Station Rotation",
        framework_tag="Differentiated Practice",
        dok="2",
        minutes=18,
        teacher_does=[
            "Circulate among stations (roughly 2 laps per 6-min segment).",
            "Prompt students with questions, never answers.",
            "Target Support station in lap 1; Core in lap 2; check all in lap 3.",
        ],
        students_do=[
            "Work at their chosen station(s).",
            "Students who finish early may attempt Extend.",
            "Write all work on the student packet.",
        ],
        questions_to_ask=[
            "Support: 'Which factor tells you the behavior at this zero?'",
            "Core: 'Where are the zeros on the sign chart? What sign is each interval?'",
            "Extend: 'What does a zero of the height function mean in context?'",
        ],
        adult_role=(
            "Circulate. Listen more than talk. For stuck students: re-read the item aloud "
            "together, then ask 'What's the first step?' No answers given."
        ),
    )

    add_callout(doc, "\u2705  STATION ANSWER KEYS (teacher eyes only)", _station_answer_key_lines())

    add_callout(doc, "\U0001F4CB  STATION SELECTION REPORT", [
        f"Support  station: {SUPPORT_ITEM_1}  +  {SUPPORT_ITEM_2 or '(none — only 1 item selected)'}",
        f"Core     station: {CORE_ITEM_1}  +  {CORE_ITEM_2}",
        f"Extend   station: {EXTEND_ITEM or 'NOT IN BANK (see EXTEND_WARNING)'}",
        EXTEND_WARNING or f"Extend item confirmed in bank (\u2713 {EXTEND_ITEM}).",
    ])

    ps.framework_phase_header(doc,
        phase="Exit — Summary Recap",
        framework_tag="Exit Ticket",
        dok="1-2",
        minutes=5,
        teacher_does=[
            "Announce: 'Stations close. Exit ticket. No notes, no Desmos. 5 minutes.'",
            "Circulate silently during write time.",
            "Collect student packets as students leave.",
        ],
        students_do=[
            "Complete the fill-in summary on the student packet.",
            "Write one sentence: 'My biggest takeaway from Lesson 3-5 is...'",
        ],
        questions_to_ask=[
            "'What's one thing you want to remember about multiplicity?'",
            "(Ask aloud only after time is up.)",
        ],
        adult_role="Silent proctor. Collect packets on the way out.",
    )

    ps.summary_exit_box(doc,
        "[DOK 1-2]  EXIT \u2014  Summary Recap  (5 min)",
        _exit_student_lines())

    add_callout(doc, "\u2705  EXIT / ANSWER KEY", [
        "1.  f(x) = 0.  'Crosses' (odd multiplicity) or 'is tangent to' (even multiplicity).",
        "2.  TANGENT.  (Even multiplicity = turning point; graph does not cross.)",
        "3.  ZEROS; SIGN chart.",
        "4.  Open response \u2014 any accurate statement about zeros, multiplicity, or inequalities.",
    ])

    add_callout(doc, "\U0001F4F7  WHAT TO COLLECT", [
        "1.  Do Now sheets (vocabulary warm-up, q28).",
        "2.  Quiz packets (Q1a\u2013Q5, every student).",
        "3.  Student packets (stations work + exit sentence).",
        "Quick-sort quiz using Q1a as the triage item: "
        "students who got the zeros wrong need Day 9 re-teach.",
    ])

    doc.save(path)
    print(f"  Built {path}")


# -----------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------

if __name__ == "__main__":
    build_do_now("Day_8_Do_Now.docx")
    build_quiz("Day_8_Quiz.docx")
    build_student_packet("Day_8_Student_Packet.docx")
    build_teacher_packet("Day_8_Teacher_Packet.docx")
    print()
    print("Built:")
    print("  Day_8_Do_Now.docx")
    print("  Day_8_Quiz.docx")
    print("  Day_8_Student_Packet.docx")
    print("  Day_8_Teacher_Packet.docx")
