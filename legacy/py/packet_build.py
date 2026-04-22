"""Shared packet-assembly primitives for build_L*_P*_packets.py.

The 18 lesson builders all reimplement the same 4 helper functions
(`set_cell`, `add_callout`, `bank_item_block`, `add_header_block`) plus a
margin-setup block at the top of every `build_do_now` / `build_student` /
`build_teacher`. This module centralizes them so each builder stays focused
on lesson content: objectives, phase headers, callouts, item-order.

Lower-level docx primitives (shading, running headers, framework_phase_header,
render_prompt, latex_to_unicode) live in `packet_styles.py` — this module
sits one layer up and handles structural blocks.

Usage in a lesson builder::

    from packet_build import (
        TABLE_STYLE, new_packet_doc,
        set_cell, add_callout, bank_item_block, add_header_block,
    )

    def build_student(path):
        doc = new_packet_doc(top=0.55, bottom=0.55, left=0.7, right=0.7)
        add_header_block(doc,
                         lesson_title=LESSON_TITLE,
                         lesson_label=LESSON_LABEL,
                         objectives=OBJECTIVES,
                         framework_header=FRAMEWORK_HEADER)
        ...
"""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import packet_styles as ps

TABLE_STYLE = "Table Grid"


def new_packet_doc(*, top: float = 0.6, bottom: float = 0.6,
                   left: float = 0.75, right: float = 0.75) -> Document:
    """Return a fresh Document with margins in inches on every section."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
    return doc


def set_cell(cell, lines, bold_first: bool = False) -> None:
    """Write lines into a table cell. If bold_first, first paragraph is bold."""
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex: str = "CFE2F3") -> None:
    """Shaded single-cell table with a bold title and body lines below."""
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label: str | None = None,
                    space_after_inches: float = 1.0) -> None:
    """Render a questionbank item: optional teal label, prompt, answer list."""
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    ps.render_prompt(prompt_p, q["prompt"], font_size=11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            run = line.add_run(f"  ({chr(65 + i)})  " + ps.latex_to_unicode(a))
            run.font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, lesson_title: str, lesson_label: str,
                     objectives, framework_header,
                     for_teacher: bool = False,
                     day_number: int = 2) -> None:
    """Running header + day banner + two-column objectives + framework table.

    objectives / framework_header: iterables of (label, body) tuples.
    for_teacher adds ' · Teacher Edition' to the subtitle.
    day_number is passed through to ps.day_banner (legacy field, most
    builders pass 2 regardless of the actual period).
    """
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, day_number, lesson_title,
                  subtitle=lesson_label + (" · Teacher Edition" if for_teacher else ""))
    objectives = list(objectives)
    t = doc.add_table(rows=len(objectives), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(objectives):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    framework_header = list(framework_header)
    t = doc.add_table(rows=len(framework_header), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(framework_header):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
