"""Shared docx formatting helpers for the daily lesson packet builders.

Adopts the visual language of `lesson35_8day.tex` (teal section labels,
dark-blue day banners, red-bordered exit box, light-gray sentence-frame
callout, running name/block header) into the python-docx toolchain so
the teacher prints DOCX -> PDF and sees the same look.

Colors match the .tex definitions:
    tealheader  #0097A7
    lightgray   #F2F2F2
    warmred     #C0392B
    dayblue     #2B5C8A
"""
import re
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ----------------------------------------------------------------------
# LaTeX-to-Unicode pre-renderer
#
# Converts common inline LaTeX patterns to readable Unicode text so
# Word (docx) packets show math readably instead of raw LaTeX commands.
# Rules are applied in order -- some depend on earlier passes.
# NOT a full LaTeX renderer; \begin{...}/\end{...} environments are left
# intact.  Unknown \commands are left as-is so the teacher can see them.
# ----------------------------------------------------------------------

# Unicode digit maps for superscripts / subscripts
_SUP_DIGITS = str.maketrans("0123456789", "⁰¹²³⁴⁵⁶⁷⁸⁹")
_SUB_DIGITS = str.maketrans("0123456789", "₀₁₂₃₄₅₆₇₈₉")
_SUB_LETTERS = {
    "a": "ₐ", "e": "ₑ", "o": "ₒ", "x": "ₓ",
    "i": "ᵢ", "n": "ₙ", "r": "ᵣ", "u": "ᵤ",
    "v": "ᵥ",
}


def _extract_braced(s: str, pos: int):
    """Starting at pos (which must point at '{'), extract the balanced
    braced content and return (content, end_pos) where end_pos is the
    index after the closing '}'.  Returns (None, pos) on failure."""
    if pos >= len(s) or s[pos] != "{":
        return None, pos
    depth = 0
    i = pos
    while i < len(s):
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
            if depth == 0:
                return s[pos + 1:i], i + 1
        i += 1
    # unbalanced braces -- return whatever we have
    return s[pos + 1:], len(s)


def _parenthesize_if_multichar(s: str) -> str:
    s = s.strip()
    if len(s) <= 1:
        return s
    return "(" + s + ")"


def _digits_to_sup(s: str) -> str:
    """All-digit string -> unicode superscripts.  Otherwise -> ^(s)."""
    if s.isdigit():
        return s.translate(_SUP_DIGITS)
    return "^(" + s + ")"


def _digits_to_sub(s: str) -> str:
    """All-digit string -> unicode subscripts.  Otherwise -> _(s)."""
    if s.isdigit():
        return s.translate(_SUB_DIGITS)
    return "_(" + s + ")"


def latex_to_unicode(s: str) -> str:
    r"""Convert common inline LaTeX patterns to readable Unicode.

    Applied in order:
      1. Fractions  (\dfrac, \frac, \tfrac)
      2. \sqrt[n]{x} and \sqrt{x}
      3. Text wrappers  (\text, \mathrm, \mathbf, ...)
      4. \left / \right delimiters
      5. Spacing macros  (\quad, \,, \!, ...)
      6. {,} thousands-comma
      7. Greek letters
      8. Operators / relations / arrows
      9. Function names  (\log, \ln, \sin, ...)
     10. ^{...} multi-char superscripts
     11. _{...} multi-char subscripts
     12. Single-char ^x / _x
    """

    # ---- 1. Fractions ---------------------------------------------------
    def _expand_fracs(text: str) -> str:
        result = []
        i = 0
        pat = re.compile(r"\\(dfrac|tfrac|frac)")
        while i < len(text):
            m = pat.search(text, i)
            if not m:
                result.append(text[i:])
                break
            result.append(text[i:m.start()])
            j = m.end()
            while j < len(text) and text[j] == " ":
                j += 1
            num, j = _extract_braced(text, j)
            if num is None:
                # Shorthand form: \frac<char><char>  e.g. \frac34 = 3/4
                if j < len(text) and text[j] != "{":
                    num = text[j]
                    j += 1
                    den = text[j] if j < len(text) and text[j] != " " else "?"
                    if den != "?":
                        j += 1
                else:
                    result.append(m.group())
                    i = m.end()
                    continue
            else:
                while j < len(text) and text[j] == " ":
                    j += 1
                den, j = _extract_braced(text, j)
                if den is None:
                    # try single-char denominator without braces
                    if j < len(text) and text[j] not in (" ", "\\", "{"):
                        den = text[j]
                        j += 1
                    else:
                        den = "?"
            num_r = latex_to_unicode(num)
            den_r = latex_to_unicode(den) if den is not None else "?"
            frac_str = (_parenthesize_if_multichar(num_r) + "/" +
                        _parenthesize_if_multichar(den_r))
            # If the character immediately after the fraction is a backslash
            # (another LaTeX command), insert a separating space so the
            # rendered fraction doesn't run into the next token.
            if j < len(text) and text[j] == "\\":
                frac_str += " "
            result.append(frac_str)
            i = j
        return "".join(result)

    s = _expand_fracs(s)

    # ---- 2a. \sqrt[n]{x} ------------------------------------------------
    def _expand_nthroot(text: str) -> str:
        result = []
        i = 0
        pat = re.compile(r"\\sqrt\[([^\]]+)\]")
        while i < len(text):
            m = pat.search(text, i)
            if not m:
                result.append(text[i:])
                break
            result.append(text[i:m.start()])
            n_str = m.group(1)
            j = m.end()
            content, j = _extract_braced(text, j)
            content_r = latex_to_unicode(content) if content is not None else "?"
            # convert n to superscript chars
            if all(c.isdigit() for c in n_str):
                n_sup = "".join(c.translate(_SUP_DIGITS) for c in n_str)
            else:
                n_sup = n_str
            result.append(n_sup + "√(" + content_r + ")")
            i = j
        return "".join(result)

    s = _expand_nthroot(s)

    # ---- 2b. \sqrt{x} ---------------------------------------------------
    def _expand_sqrt(text: str) -> str:
        result = []
        i = 0
        pat = re.compile(r"\\sqrt")
        while i < len(text):
            m = pat.search(text, i)
            if not m:
                result.append(text[i:])
                break
            result.append(text[i:m.start()])
            j = m.end()
            content, j = _extract_braced(text, j)
            content_r = latex_to_unicode(content) if content is not None else "?"
            result.append("√(" + content_r + ")")
            i = j
        return "".join(result)

    s = _expand_sqrt(s)

    # ---- 3. Text/math wrappers ------------------------------------------
    def _strip_wrapper(text: str, cmd: str) -> str:
        result = []
        i = 0
        pat = re.compile(re.escape(cmd))
        while i < len(text):
            m = pat.search(text, i)
            if not m:
                result.append(text[i:])
                break
            result.append(text[i:m.start()])
            j = m.end()
            content, j = _extract_braced(text, j)
            result.append(content if content is not None else "")
            i = j
        return "".join(result)

    for _cmd in (r"\text", r"\mathrm", r"\mathbf", r"\textbf", r"\textit",
                 r"\mathit", r"\mathsf", r"\mathtt", r"\boldsymbol",
                 r"\mbox"):
        s = _strip_wrapper(s, _cmd)

    # ---- 4. \left / \right delimiters -----------------------------------
    s = re.sub(r"\\left\s*\\\{", "{", s)
    s = re.sub(r"\\right\s*\\\}", "}", s)
    s = re.sub(r"\\left\s*\[", "[", s)
    s = re.sub(r"\\right\s*\]", "]", s)
    s = re.sub(r"\\left\s*\(", "(", s)
    s = re.sub(r"\\right\s*\)", ")", s)
    s = re.sub(r"\\left\s*\.", "", s)
    s = re.sub(r"\\right\s*\.", "", s)

    # ---- 5. Spacing macros ----------------------------------------------
    s = re.sub(r"\\qquad\b", "        ", s)
    s = re.sub(r"\\quad\b",  "    ", s)
    s = re.sub(r"\\[!,;: ]", " ", s)
    s = re.sub(r"\\ ", " ", s)

    # ---- 6. Thousands comma: {,} -> , -----------------------------------
    s = s.replace("{,}", ",")

    # ---- 7. Greek letters -----------------------------------------------
    _GREEK = {
        r"\alpha": "α", r"\beta": "β", r"\gamma": "γ",
        r"\delta": "δ", r"\varepsilon": "ε", r"\epsilon": "ε",
        r"\zeta": "ζ", r"\eta": "η", r"\vartheta": "θ",
        r"\theta": "θ", r"\iota": "ι", r"\kappa": "κ",
        r"\lambda": "λ", r"\mu": "μ", r"\nu": "ν",
        r"\xi": "ξ", r"\varpi": "π", r"\pi": "π",
        r"\varrho": "ρ", r"\rho": "ρ",
        r"\varsigma": "ς", r"\sigma": "σ", r"\tau": "τ",
        r"\upsilon": "υ", r"\varphi": "φ", r"\phi": "φ",
        r"\chi": "χ", r"\psi": "ψ", r"\omega": "ω",
        r"\Gamma": "Γ", r"\Delta": "Δ", r"\Theta": "Θ",
        r"\Lambda": "Λ", r"\Xi": "Ξ", r"\Pi": "Π",
        r"\Sigma": "Σ", r"\Upsilon": "Υ", r"\Phi": "Φ",
        r"\Psi": "Ψ", r"\Omega": "Ω",
    }
    for _cmd, _uni in sorted(_GREEK.items(), key=lambda x: -len(x[0])):
        s = re.sub(re.escape(_cmd) + r"(?![a-zA-Z])", _uni, s)

    # ---- 8. Operators / relations / arrows ------------------------------
    _OPS = [
        (r"\Leftrightarrow", "⇔"), (r"\Rightarrow", "⇒"),
        (r"\rightarrow", "→"), (r"\leftarrow", "←"),
        (r"\to", "→"),
        (r"\leq", "≤"), (r"\le", "≤"),
        (r"\geq", "≥"), (r"\ge", "≥"),
        (r"\neq", "≠"), (r"\ne", "≠"),
        (r"\approx", "≈"), (r"\equiv", "≡"),
        (r"\cdot", "·"), (r"\times", "×"), (r"\div", "÷"),
        (r"\pm", "±"), (r"\mp", "∓"), (r"\infty", "∞"),
        (r"\cap", "∩"), (r"\cup", "∪"),
        (r"\in", "∈"), (r"\notin", "∉"),
        (r"\subset", "⊂"), (r"\subseteq", "⊆"),
        (r"\ldots", "…"), (r"\cdots", "⋯"), (r"\dots", "…"),
    ]
    for _cmd, _uni in sorted(_OPS, key=lambda x: -len(x[0])):
        s = re.sub(re.escape(_cmd) + r"(?![a-zA-Z])", _uni, s)

    # ---- 9. Function names ----------------------------------------------
    _FUNCS = [
        "arcsin", "arccos", "arctan", "sinh", "cosh", "tanh",
        "sin", "cos", "tan", "cot", "sec", "csc",
        "log", "ln", "exp",
        "lim", "max", "min", "gcd", "lcm", "det",
        "deg", "mod",
    ]
    for _fn in _FUNCS:
        s = re.sub(r"\\" + _fn + r"(?![a-zA-Z])", _fn, s)

    # ---- 10. ^{...} multi-char superscripts -----------------------------
    def _expand_sup_braced(text: str) -> str:
        result = []
        i = 0
        while i < len(text):
            if text[i] == "^" and i + 1 < len(text) and text[i + 1] == "{":
                content, j = _extract_braced(text, i + 1)
                if content is not None:
                    content_r = latex_to_unicode(content)
                    result.append(_digits_to_sup(content_r))
                    i = j
                    continue
            result.append(text[i])
            i += 1
        return "".join(result)

    s = _expand_sup_braced(s)

    # ---- 11. _{...} multi-char subscripts -------------------------------
    def _expand_sub_braced(text: str) -> str:
        result = []
        i = 0
        while i < len(text):
            if text[i] == "_" and i + 1 < len(text) and text[i + 1] == "{":
                content, j = _extract_braced(text, i + 1)
                if content is not None:
                    content_r = latex_to_unicode(content)
                    result.append(_digits_to_sub(content_r))
                    i = j
                    continue
            result.append(text[i])
            i += 1
        return "".join(result)

    s = _expand_sub_braced(s)

    # ---- 12. Single-char ^x / _x ----------------------------------------
    def _expand_sup_single(text: str) -> str:
        result = []
        i = 0
        while i < len(text):
            if (text[i] == "^" and i + 1 < len(text)
                    and text[i + 1] != "{" and text[i + 1] != "^"):
                ch = text[i + 1]
                if ch.isdigit():
                    result.append(ch.translate(_SUP_DIGITS))
                    i += 2
                    continue
                # letters -- no complete unicode block; leave as ^x
            result.append(text[i])
            i += 1
        return "".join(result)

    s = _expand_sup_single(s)

    def _expand_sub_single(text: str) -> str:
        result = []
        i = 0
        while i < len(text):
            if (text[i] == "_" and i + 1 < len(text)
                    and text[i + 1] != "{" and text[i + 1] != "_"):
                ch = text[i + 1]
                if ch.isdigit():
                    result.append(ch.translate(_SUB_DIGITS))
                    i += 2
                    continue
                if ch in _SUB_LETTERS:
                    result.append(_SUB_LETTERS[ch])
                    i += 2
                    continue
            result.append(text[i])
            i += 1
        return "".join(result)

    s = _expand_sub_single(s)

    return s


# ----------------------------------------------------------------------
# render_prompt
#
# Central helper: write registry prompt text into a Word paragraph
# after running it through latex_to_unicode().
#
# Usage (in any builder's bank_item_block or equivalent):
#     from packet_styles import render_prompt
#     render_prompt(prompt_p, q["prompt"], font_size=11)
#
# TODO: migrate all builder bank_item_block() calls to use render_prompt
#       so new LaTeX patterns only need fixing in one place.
# ----------------------------------------------------------------------

def render_prompt(paragraph, text: str, *, font_size=None):
    """Add *text* to *paragraph* after converting LaTeX to Unicode.

    Returns the run so callers can apply additional formatting if needed.
    """
    run = paragraph.add_run(latex_to_unicode(text))
    if font_size is not None:
        run.font.size = Pt(font_size)
    return run


COLORS = {
    "teal":      "0097A7",
    "lightgray": "F2F2F2",
    "warmred":   "C0392B",
    "warmredlt": "FCEAE8",  # very pale red (title backdrop)
    "dayblue":   "2B5C8A",
    "white":     "FFFFFF",
    "tealpale":  "E6F5F7",  # 10% tint for generalize-box background
}

TABLE_STYLE = "Table Grid"


# ----------------------------------------------------------------------
# Low-level cell shading / border helpers (python-docx OxmlElement)
# ----------------------------------------------------------------------

def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, *, top=None, left=None, bottom=None, right=None,
                     default=None) -> None:
    """Set per-side borders. Each side is a dict like
    {"val": "single", "sz": 4, "color": "C0392B"} or None.
    If `default` is given, any side not specified gets that spec."""
    tc_pr = cell._tc.get_or_add_tcPr()
    # remove any existing tcBorders so we don't double up
    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    sides = {"top": top, "left": left, "bottom": bottom, "right": right}
    for side, spec in sides.items():
        if spec is None and default is not None:
            spec = default
        if spec is None:
            continue
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   spec.get("val", "single"))
        b.set(qn("w:sz"),    str(spec.get("sz", 4)))
        b.set(qn("w:space"), str(spec.get("space", 0)))
        b.set(qn("w:color"), spec.get("color", "auto"))
        borders.append(b)
    tc_pr.append(borders)


def _run(paragraph, text: str, *, bold=False, color=None, size=None,
         italic=False):
    r = paragraph.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


# ----------------------------------------------------------------------
# Day banner -- dark-blue solid block with "Day N | Title" + gray subtitle
# ----------------------------------------------------------------------

def day_banner(doc, day_num: int, title: str, subtitle: str = "") -> None:
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    cell = t.rows[0].cells[0]
    cell.text = ""
    shade_cell(cell, COLORS["dayblue"])
    # thin border so the banner reads as a block
    set_cell_borders(cell, default={"val": "single", "sz": 4,
                                    "color": COLORS["dayblue"]})
    p = cell.paragraphs[0]
    _run(p, f"Day {day_num}  │  {title}",
         bold=True, color=COLORS["white"], size=16)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if subtitle:
        p = doc.add_paragraph()
        _run(p, subtitle, italic=True, size=10, color="555555")
    doc.add_paragraph()


# ----------------------------------------------------------------------
# Teal section -- 2-col table with teal-shaded left label, white bold,
# body lines on the right. Replaces the unstyled add_two_col pattern.
# ----------------------------------------------------------------------

def teal_section(doc, label: str, body_lines) -> None:
    if isinstance(body_lines, str):
        body_lines = [body_lines]
    t = doc.add_table(rows=1, cols=2)
    t.style = TABLE_STYLE
    t.columns[0].width = Inches(1.3)
    t.columns[1].width = Inches(5.2)
    left, right = t.rows[0].cells

    shade_cell(left, COLORS["teal"])
    left.text = ""
    p = left.paragraphs[0]
    _run(p, label, bold=True, color=COLORS["white"], size=11)
    left.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    right.text = ""
    for i, line in enumerate(body_lines):
        p = right.paragraphs[0] if i == 0 else right.add_paragraph()
        p.add_run(line)
    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    doc.add_paragraph()


# ----------------------------------------------------------------------
# Summary-Exit box -- red-bordered block with a titled top-left tab.
# Mirrors the .tex cerbox. Works for CER or for summary-style exits.
# ----------------------------------------------------------------------

def summary_exit_box(doc, title: str, lines) -> None:
    if isinstance(lines, str):
        lines = [lines]
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_borders(cell, default={"val": "single", "sz": 8,
                                    "color": COLORS["warmred"]})
    shade_cell(cell, COLORS["warmredlt"])
    cell.text = ""

    p = cell.paragraphs[0]
    _run(p, title, bold=True, color=COLORS["warmred"], size=11)
    for line in lines:
        cp = cell.add_paragraph()
        cp.add_run(line)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    doc.add_paragraph()


# ----------------------------------------------------------------------
# Sentence-frame callout -- light-gray block with a 3-pt teal left accent.
# Mirrors the .tex generalizebox.
# ----------------------------------------------------------------------

def sentence_frame_box(doc, lines) -> None:
    if isinstance(lines, str):
        lines = [lines]
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    shade_cell(cell, COLORS["lightgray"])
    set_cell_borders(
        cell,
        left={"val": "single", "sz": 24, "color": COLORS["teal"]},
        top={"val": "nil"}, right={"val": "nil"}, bottom={"val": "nil"},
    )
    cell.text = ""
    p = cell.paragraphs[0]
    _run(p, "Sentence frame", bold=True, color=COLORS["teal"], size=10)
    for line in lines:
        cp = cell.add_paragraph()
        cp.add_run(line).italic = True
    doc.add_paragraph()


# ----------------------------------------------------------------------
# Running header -- name / block fields on every page.
# ----------------------------------------------------------------------

def running_header_name_block(doc, label_right: str = "BLOCK") -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    # Tab stop to push right label to the right margin.
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5))
    _run(p, "FULL NAME: ________________________", size=9)
    p.add_run("\t")
    _run(p, f"{label_right}: __________", size=9)


# ----------------------------------------------------------------------
# Framework phase header with explicit "Teacher / Students / Questions
# to ask / Adult role" columns per DOKframework.txt.
# Used in teacher packets so evaluators scanning with the framework
# printout find the exact labels.
# ----------------------------------------------------------------------

def framework_phase_header(doc, *, phase: str, framework_tag: str,
                           dok: str, minutes,
                           teacher_does: list, students_do: list,
                           questions_to_ask: list, adult_role: str) -> None:
    title = f"[{framework_tag}]  [DOK {dok}]  {phase}"
    if minutes is not None:
        title += f"  ({minutes} min)"
    teal_section(doc, "PHASE", [title])

    rows = [
        ("Teachers will…",    teacher_does),
        ("Students will…",    students_do),
        ("Questions to ask",       questions_to_ask),
        ("Adult role",             [adult_role]),
    ]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = TABLE_STYLE
    t.columns[0].width = Inches(1.3)
    t.columns[1].width = Inches(5.2)
    for i, (label, lines) in enumerate(rows):
        left, right = t.rows[i].cells
        shade_cell(left, COLORS["tealpale"])
        left.text = ""
        _run(left.paragraphs[0], label, bold=True,
             color=COLORS["teal"], size=10)
        right.text = ""
        for j, line in enumerate(lines):
            p = right.paragraphs[0] if j == 0 else right.add_paragraph()
            p.add_run("• " + line if len(lines) > 1 else line)
    doc.add_paragraph()


# ---------------------------------------------------------------------
# Visuals checklist hook for packet builders.
#
# Call once per builder (after all docx files are saved) to emit a
# sidecar `<Day>_Visuals_Checklist.md` listing every embedded visual
# the packet relies on. Teacher reviews this before print to confirm
# photos/graphs/tables/maps are embedded cleanly (no answer-key leak,
# no multi-item crops).
#
# Usage (e.g. build_day6_packets.py __main__):
#     import qb
#     from packet_styles import emit_visuals_checklist
#     emit_visuals_checklist(_ALL_IDS, "Day_6_Visuals_Checklist.md",
#                            title="Day 6 - Visuals Checklist")
# ---------------------------------------------------------------------

def emit_visuals_checklist(packet_ids, out_path, *, title=None):
    """Write a visuals checklist for the given packet IDs.

    Thin wrapper around qb.write_visuals_checklist() so each builder
    can emit its own sidecar with one import. Prints a confirmation
    line if any visual items were found, else notes that none were.
    """
    import qb
    from pathlib import Path
    qb.write_visuals_checklist(
        packet_ids,
        out_path,
        title=title or f"{Path(out_path).stem} - Visuals Checklist",
    )
    rows = qb.visuals_for(packet_ids)
    n_flag = sum(1 for r in rows if r["needs_cleanup"])
    if rows:
        flag = f" ({n_flag} need cleanup)" if n_flag else ""
        print(f"  + {out_path}: {len(rows)} visual(s){flag}")
    else:
        print(f"  + {out_path}: (no visuals in this packet)")


# ── Self-test (python packet_styles.py) ─────────────────────────────
if __name__ == "__main__":
    cases = [
        # (input, expected)
        (r"\log_5 6 + \dfrac{1}{2}\log_5 y",
         "log₅ 6 + 1/2 log₅ y"),
        (r"e^{rt}",
         "e^(rt)"),
        (r"\ln 1{,}000 \neq 3",
         "ln 1,000 ≠ 3"),
        (r"\log_{10} x",
         "log₁₀ x"),
        (r"\sqrt[3]{x^2}",
         "³√(x²)"),
        (r"\dfrac{x+1}{x-2}",
         "(x+1)/(x-2)"),
        (r"3^x = 50",
         "3^x = 50"),
        (r"\pi r^2",
         "π r²"),
    ]
    passed = 0
    failed = 0
    for inp, expected in cases:
        got = latex_to_unicode(inp)
        ok = got == expected
        if ok:
            passed += 1
        else:
            failed += 1
        status = "PASS" if ok else "FAIL"
        print(f"  [{status}]  {inp!r}")
        if not ok:
            print(f"         expected: {expected!r}")
            print(f"         got:      {got!r}")
    print(f"\n{passed}/{passed + failed} passed.")
