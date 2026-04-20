"""Generate Combined Day 4-5 materials \u2014 Savvas-only, DOK-3 modeling day.

Design: folds old Day 4 (Real & Complex Zeros) + old Day 5 (Modeling) into
ONE period.  Every student-facing work item traces to a Savvas bank ID.
DOK-3 driver = Savvas Practice #27 (storage box volume, Example 4 anchor) \u2014
the first of Savvas's three declared DOK-3 items in this lesson.

Phases (55 min, fits 55-min Tue A; 65-min F gets slack for full Practice):
  0\u20135    Do Now A : Savvas #28 \u2014 zeros \u21d4 factor vocab             [DOK 1]
  5\u20137    Do Now B : Blooket login                                       [\u2014]
  7\u201314   Do Now C : Blooket rule recall (conjugate pairs + synth div)    [DOK 1]
  14\u201324  Launch    : Savvas Example 3 \u2014 all real + complex zeros       [DOK 2]
  24\u201334  Practice  : Try It 3a + 3b (find all zeros)                    [DOK 2]
  34\u201349  Explore   : Savvas Practice #27 storage box (DOK-3 capstone)  [DOK 3]
  49\u201352  Share/Sum : Synthesis + self-rating                            [DOK 2]
  52\u201355  Exit      : Short recap (Savvas-anchored)                     [DOK 1\u20132]
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

# Every work item routes through the bank.
DO_NOW_ID    = "3-5-savvas-q28"     # zeros \u21d4 factor vocabulary
PRACTICE_IDS = ["3-5-tryit-3a", "3-5-tryit-3b"]
DOK3_ID      = "3-5-savvas-q27"     # storage box \u2014 DOK-3 driver
EXIT_ID      = "3-5-savvas-q17"     # short "all real + complex zeros"

OBJECTIVES = [
    ("Math Objective",
     "Find ALL zeros of a polynomial function (real and complex) using "
     "graphing, synthetic division, and the quadratic formula; interpret "
     "zeros in a real-world modeling context."),
    ("Language Objective",
     "Students will use the frames \u201cThe graph shows ___ real zero(s), "
     "but the polynomial must have ___ total zero(s) because ___\u201d and "
     "\u201cThe zero at x = ___ represents ___ in the context.\u201d"),
    ("Essential Understanding",
     "A polynomial of degree n has n total zeros, counted with multiplicity "
     "and including complex zeros.  Non-real complex zeros occur in "
     "conjugate pairs.  In a modeling context, each zero has a real-world "
     "meaning tied to the quantities the function represents."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Students find all zeros of a polynomial (real + complex), then apply "
     "zero-finding to a volume-modeling context where factors represent "
     "physical dimensions."),
    ("Essential Question",
     "How do synthetic division and the quadratic formula let you find "
     "EVERY zero of a polynomial \u2014 and what do those zeros mean when "
     "the polynomial models a real-world quantity like volume or profit?"),
    ("Materials",
     "Student laptops with Desmos; pencils; Blooket access; printed "
     "synthetic-division reference card (optional). No chart paper."),
]


def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if i == 0 and bold_first:
            run.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_two_col(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = TABLE_STYLE
    t.columns[0].width = Inches(1.8)
    t.columns[1].width = Inches(4.7)
    for i, (label, content) in enumerate(rows):
        set_cell(t.rows[i].cells[0], label, bold_first=True)
        if isinstance(content, str):
            content = [content]
        set_cell(t.rows[i].cells[1], content)
    doc.add_paragraph()


def add_callout(doc, title, lines):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.add_run(title).bold = True
    for line in lines:
        cell.add_paragraph(line)
    doc.add_paragraph()


def header_line(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def blank_lines(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.add_run("______________________________________________________________")


def _format_unicode(prompt: str) -> str:
    prompt = (prompt
              .replace("\u00e2\u2020\u2019", "\u2192")
              .replace("\u00e2\u20ac\u201d", "\u2014")
              .replace("\u00e2\u20ac\u201c", "\u2013"))
    SUPS = {"2": "\u00b2", "3": "\u00b3", "4": "\u2074",
            "5": "\u2075", "6": "\u2076"}
    out, i = [], 0
    while i < len(prompt):
        ch = prompt[i]
        if ch == "^" and i + 1 < len(prompt) and prompt[i+1] in SUPS:
            out.append(SUPS[prompt[i+1]])
            i += 2
            continue
        out.append("\u2212" if ch == "-" else ch)
        i += 1
    return "".join(out)


def _strip_stem(prompt: str) -> str:
    txt = _format_unicode(prompt)
    if ":" in txt:
        txt = txt.split(":", 1)[-1].strip()
    return txt.rstrip(".")


# ----------------------------------------------------------------------
# DO NOW SHEET  \u2014  Savvas Practice #28 vocabulary
# ----------------------------------------------------------------------

def build_do_now(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)

    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    header_line(doc, "Combined Day 4-5  |  DO NOW  \u2014  Zeros \u21d4 Factors", size=13)

    p = doc.add_paragraph()
    r = p.add_run("Name: _____________________________________     Date: _____________")
    r.font.size = Pt(11)
    doc.add_paragraph()

    add_callout(doc, "\u270d\ufe0f  5 minutes. Silent. Pencil only. No Desmos.", [
        "Yesterday you read multiplicity from the factored form.  Today the "
        "connection runs both ways.",
        "Complete each statement below so it means the same as \u201c4 is a "
        "zero of the function.\u201d",
    ])

    header_line(doc, "Savvas Practice #28:", size=12)
    p = doc.add_paragraph()
    p.add_run("Complete each statement so it means the same as \u201c4 is a "
              "zero of the function\u201d:").italic = True
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("(1)  The graph of the function crosses the ").bold = False
    p.add_run("__________").bold = True
    p.add_run("  at 4.")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("(2)  ").bold = False
    p.add_run("__________").bold = True
    p.add_run("  is a factor of the polynomial.")
    doc.add_paragraph()

    header_line(doc, "Now use the same idea:", size=12)
    p = doc.add_paragraph()
    p.add_run("If  \u22122  is a zero of a polynomial, one factor is: ").italic = True
    blank_lines(doc, 1)
    p = doc.add_paragraph()
    p.add_run("If a polynomial has degree 3 and only one real zero, how many "
              "complex zeros must it have?  ").italic = True
    blank_lines(doc, 1)

    doc.save(path)


# ----------------------------------------------------------------------
# STUDENT PACKET
# ----------------------------------------------------------------------

LAUNCH_INVESTIGATION = [
    "SAVVAS EXAMPLE 3 \u2014 All Real + Complex Zeros",
    "",
    "In Desmos, graph  f(x) = x\u00b3 \u2212 2x\u00b2 \u2212 2x \u2212 3.",
    "",
    "STEP 1 \u2014 Read the graph.  How many times does f cross the x-axis?  _______",
    "        What is the real zero?  x = _______",
    "",
    "STEP 2 \u2014 The polynomial has degree 3, so it must have 3 total zeros.",
    "        You found 1 real zero.  How many zeros are still unaccounted for?  _______",
    "",
    "STEP 3 \u2014 Use synthetic division to divide f(x) by (x \u2212 ___).  "
    "        Your quotient is a quadratic:  Q(x) = _________________________",
    "",
    "STEP 4 \u2014 Solve the quadratic Q(x) = 0 using the quadratic formula.  "
    "        If the discriminant is negative, the solutions are COMPLEX CONJUGATES.",
    "",
    "        Complex zeros:  x = _______________  and  x = _______________",
    "",
    "STEP 5 \u2014 List all three zeros of f(x):",
    "        x = _______,  x = _______,  x = _______",
    "",
    "THE RULE (Fundamental Theorem of Algebra \u2014 shortcut version):",
    "\u2022  A polynomial of degree n has exactly n total zeros, counting "
    "multiplicity and complex zeros.",
    "\u2022  Non-real complex zeros ALWAYS come in conjugate pairs "
    "(a + bi and a \u2212 bi together).",
]


def _practice_lines() -> list[str]:
    items = qb.get_for_packet(PRACTICE_IDS)
    lines = [
        "For each polynomial: find ALL real and complex zeros.  Show your "
        "synthetic division work on the right; write the complex zeros "
        "using i.  Use Desmos to verify the real zero before dividing.",
        "",
    ]
    labels = ["A.", "B."]
    for label, q in zip(labels, items):
        lines.append(f"{label}  {_strip_stem(q['prompt'])}")
        lines.append("    Real zero(s):  _____________________________________________")
        lines.append("    Synthetic division quotient:  _____________________________")
        lines.append("    Complex zeros:  ___________________________________________")
        lines.append("")
    lines.append("Bank IDs: " + ", ".join(PRACTICE_IDS) + " (Savvas Try It 3a, 3b).")
    return lines


DOK3_STORAGE_BOX = [
    "THE PROBLEM  \u2014  Savvas Practice #27 (Model With Mathematics)",
    "",
    "The height of a rectangular storage box is less than BOTH its length "
    "and its width. The function",
    "",
    "        f(x) = x\u00b3 + 2x\u00b2 \u2212 3x",
    "",
    "represents the volume of the rectangular box, where x is the width "
    "(in feet).",
    "",
    "THE RULES YOU NEED (everything is on this page)",
    "\u2022  RULE 1: Volume = length \u00d7 width \u00d7 height.",
    "\u2022  RULE 2: If you factor f(x), each factor represents ONE dimension.",
    "\u2022  RULE 3: The height must be the SMALLEST of the three factors "
    "(height < length AND height < width).",
    "",
    "STEP 1 \u2014 Factor the volume expression.",
    "        f(x) = x(x\u00b2 + 2x \u2212 3) = _______________________________",
    "",
    "STEP 2 \u2014 Find the zeros of f(x).",
    "        x = _______,  x = _______,  x = _______",
    "",
    "STEP 3 \u2014 Match each factor to a dimension.",
    "        x represents the WIDTH (given).",
    "        Of the two remaining factors (x + 3) and (x \u2212 1), which is the",
    "        HEIGHT (smaller)?  _______     Which is the LENGTH?  _______",
    "",
    "STEP 4 \u2014 Find the dimensions when volume = 10 ft\u00b3.",
    "        Set f(x) = 10:   x\u00b3 + 2x\u00b2 \u2212 3x = 10",
    "        Rearrange:       x\u00b3 + 2x\u00b2 \u2212 3x \u2212 10 = 0",
    "        Try x = 2 (test with synthetic division).  Does it work?  _______",
    "        So width = _______ ft, height = _______ ft, length = _______ ft.",
    "",
    "STEP 5 \u2014 Verify.  Does length \u00d7 width \u00d7 height = 10?  _______",
    "",
    "STEP 6 \u2014 Explain to your partner.  Use the frame:",
    "\u201cThe zero at x = ___ represents ___ in this context because ___.\u201d",
]

SHARE_SUMMARY_STUDENT = [
    "Self-check  \u2014  circle one for each:",
    "1.  I can find every zero of a polynomial (real AND complex).        \u2713    partly    not yet",
    "2.  I can use synthetic division to reduce a cubic to a quadratic.   \u2713    partly    not yet",
    "3.  I can interpret factors as physical dimensions in a model.       \u2713    partly    not yet",
    "",
    "Preview Day 6:  polynomial inequalities \u2014 where is the graph above / below the x-axis?",
]


def _exit_lines() -> list[str]:
    q = qb.get_for_packet([EXIT_ID])[0]
    return [
        f"Savvas Practice #17:",
        _format_unicode(q["prompt"]),
        "",
        "All real and complex zeros:  _____________________________________________",
        "____________________________________________________________________________",
        "",
        "Degree check: does your total zero count match the degree?  _______",
    ]


def build_student(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    ps.running_header_name_block(doc)
    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    ps.day_banner(doc, "4-5", "Real + Complex Zeros + Modeling",
                  "Find EVERY zero using synthetic division and the "
                  "quadratic formula.  Then interpret zeros as physical "
                  "dimensions in a volume model.  All work items Savvas-anchored.")

    add_two_col(doc, OBJECTIVES)
    add_two_col(doc, FRAMEWORK_HEADER)

    add_callout(doc, "\U0001F4C4  DO NOW  \u2014  separate sheet", [
        "You got the Zeros \u21d4 Factors sheet on entry (Savvas Practice #28 + bridge questions).",
        "Hold onto it \u2014 we\u2019ll return to it during the Launch.",
    ])

    add_callout(doc, "\U0001F3AE  Blooket  \u2014  Rule Recall   [Do Now B+C]  [DOK 1]   (2 + 7 min)", [
        "Pure rule recall: degree \u21d2 total zero count (Fundamental Theorem "
        "shortcut), complex zeros come in CONJUGATE PAIRS, synthetic division "
        "steps, factor theorem (zero \u21d4 factor).  No procedural computation.",
    ])

    add_two_col(doc, [(
        "\U0001F4A1  Launch  \u2014  Savvas Example 3 (All Real + Complex Zeros)   [Launch]  [DOK 2]   (10 min)",
        LAUNCH_INVESTIGATION,
    )])

    add_two_col(doc, [(
        "\u270F\uFE0F  Practice  \u2014  Savvas Try It 3a + 3b   [Practice]  [DOK 2]   (10 min)",
        _practice_lines(),
    )])

    add_two_col(doc, [(
        "\U0001F4E6  Explore  \u2014  Storage Box Volume (Savvas Practice #27)   [Explore]  [DOK 3]   (15 min)",
        DOK3_STORAGE_BOX,
    )])

    add_callout(doc,
        "\U0001F501  SHARE / SUMMARY   [Share/Summary]  [DOK 2]   (3 min)",
        SHARE_SUMMARY_STUDENT)

    ps.summary_exit_box(doc,
        "\U0001F4E4  EXIT TICKET   [Exit Ticket]  [DOK 1\u20132]   (5 min)",
        _exit_lines())

    doc.save(path)


# ----------------------------------------------------------------------
# TEACHER PACKET
# ----------------------------------------------------------------------

CRITERIA_FOR_SUCCESS = [
    "Student-facing, revisited during Share/Summary:",
    "1.  I can find every zero of a polynomial (real AND complex).",
    "2.  I can use synthetic division to reduce a cubic to a quadratic.",
    "3.  I can interpret factors as physical dimensions in a model.",
    "Formative evidence:",
    "\u2022  Do Now sheets (Savvas #28 vocab + bridge questions).",
    "\u2022  Blooket dashboard (conjugate pairs + synthetic division rules).",
    "\u2022  Launch Example 3 work (synthetic division + quadratic formula).",
    "\u2022  Practice Try It 3a/3b (complex-zero computation).",
    "\u2022  Storage box page (DOK-3 evidence \u2014 photograph or collect).",
    "\u2022  Exit ticket (Savvas Practice #17 \u2014 all real + complex zeros).",
]

BLOOKET_SCOPE = [
    "SCOPE (pure rule-recall, DOK 1).  Rules this lesson depends on:",
    "\u2022  Degree n \u21d2 n total zeros (counting multiplicity and complex).",
    "\u2022  Non-real complex zeros come in CONJUGATE PAIRS (a + bi with a \u2212 bi).",
    "\u2022  Factor theorem: x = c is a zero \u21d4 (x \u2212 c) is a factor.",
    "\u2022  Synthetic division: bring down the leading coefficient, multiply by c, "
    "add the next coefficient.",
    "\u2022  Quadratic formula applies to the reduced quadratic after synthetic division.",
    "\u2022  Day 2-3 retention: multiplicity = exponent; EVEN \u2192 touch, ODD \u2192 cross.",
    "NO procedural problems.  Dashboard diagnostic: note the two lowest items; "
    "60-second board reset if a rule is cold.",
    "CSV: Blooket_Day45_RealComplex.csv  (regenerate via regen_blooket_csvs.py "
    "after tagging more bank items with day4-real-complex + blooket-pool).",
    "Questions to ask (after game): \u201cWhich rule had the lowest %?  Say it now.\u201d",
    "Adult role: teacher watches dashboard + runs 60-sec reset; aide "
    "troubleshoots Chromebook issues during login.",
]

LAUNCH_KEY = [
    "Savvas Example 3 target:  f(x) = x\u00b3 \u2212 2x\u00b2 \u2212 2x \u2212 3.",
    "Step 1 \u2014 Graph crosses the x-axis ONCE.  Real zero: x = 3.",
    "Step 2 \u2014 Degree 3 \u2192 3 total zeros.  2 unaccounted for.",
    "Step 3 \u2014 Synthetic division by (x \u2212 3):",
    "    3  |   1   \u22122   \u22122   \u22123",
    "       |       3    3    3",
    "       |_______________________",
    "           1    1    1    0",
    "    Quotient: Q(x) = x\u00b2 + x + 1.",
    "Step 4 \u2014 Solve Q(x) = 0 using the quadratic formula:",
    "    x = (\u22121 \u00b1 \u221a(1 \u2212 4))/2 = (\u22121 \u00b1 \u221a(\u22123))/2 = (\u22121 \u00b1 i\u221a3)/2.",
    "    Complex zeros:  x = (\u22121 + i\u221a3)/2  and  x = (\u22121 \u2212 i\u221a3)/2.",
    "Step 5 \u2014 All three zeros: 3, (\u22121 + i\u221a3)/2, (\u22121 \u2212 i\u221a3)/2.",
    "TEACHER MOVE \u2014 the synthetic-division teach:",
    "Model synthetic division ONCE at the board during Launch.  Keep the "
    "mechanics tight (90 seconds max).  Students will do it themselves in "
    "Practice.  If anyone gets lost, pair them with a student who got it "
    "during Practice (peer teaching works better than re-teach here).",
    "Questions to ask: \u201cHow many zeros MUST this polynomial have?  Why?\u201d / "
    "\u201cIf you found 1 real zero, how many are still missing?\u201d / "
    "\u201cWhat does it mean that the discriminant is negative?\u201d",
    "Adult role: teacher runs script + board work; aide walks to partners "
    "stuck loading Desmos.",
]


def _practice_key_lines() -> list[str]:
    items = qb.get_for_packet(PRACTICE_IDS)
    lines = []
    labels = ["A.", "B."]
    for label, q in zip(labels, items):
        lines.append(f"{label}  {_strip_stem(q['prompt'])}  [bank: {q['id']}, DOK {q.get('dok')}]")
        lines.append(f"    \u2705  {_format_unicode(str(q.get('correct', '')))}")
        rationale = q.get("dok_rationale")
        if rationale:
            lines.append(f"    DOK rationale: {_format_unicode(rationale)}")
        lines.append("")
    lines.append("Questions to ask (circulating): \u201cWhat\u2019s the real zero from "
                 "the graph?\u201d / \u201cWhat\u2019s the quotient after dividing?\u201d / "
                 "\u201cIs the discriminant negative?  What does that tell you?\u201d")
    lines.append("Adult role: teacher circulates with prompts; aide stations "
                 "near students who need the synthetic-division algorithm "
                 "written out step-by-step.")
    return lines


def _dok3_key_lines() -> list[str]:
    q = qb.get_for_packet([DOK3_ID])[0]
    return [
        f"[bank: {q['id']}, DOK {q.get('dok')}] \u2014 SAVVAS-DECLARED DOK-3 ANCHOR",
        _format_unicode(q["prompt"]),
        "",
        f"\u2705  {_format_unicode(str(q.get('correct', '')))}",
        "",
        "DOK rationale: " + _format_unicode(str(q.get("dok_rationale", ""))),
        "",
        "Full working:",
        "(a) f(x) = x(x\u00b2 + 2x \u2212 3) = x(x + 3)(x \u2212 1).",
        "(b) Zeros: x = 0, x = \u22123, x = 1.",
        "(c) x represents WIDTH (given). The height is SMALLER than both "
        "the length and the width, so between (x + 3) and (x \u2212 1), "
        "(x \u2212 1) is the HEIGHT (smaller) and (x + 3) is the LENGTH.",
        "(d) Set f(x) = 10: x\u00b3 + 2x\u00b2 \u2212 3x \u2212 10 = 0.  Try x = 2:",
        "    2\u00b3 + 2(4) \u2212 6 \u2212 10 = 8 + 8 \u2212 6 \u2212 10 = 0. \u2713",
        "    So width = 2 ft.  Height = (2 \u2212 1) = 1 ft.  Length = (2 + 3) = 5 ft.",
        "    Verify: 2 \u00d7 1 \u00d7 5 = 10 \u2713.",
        "",
        "WHY THIS IS DOK 3:",
        "\u2022  Students coordinate four sub-problems where each part depends on a "
        "strategic interpretation of the previous.",
        "\u2022  They must map algebraic factors to physical dimensions using the "
        "constraint \u201cheight is smaller\u201d \u2014 that\u2019s not a memorized step.",
        "\u2022  They solve a non-routine cubic equation in context (set f(x) = 10 "
        "and factor/test-a-value \u2014 no direct formula).",
        "",
        "Stuck-pair hint card: \u201cFactor first.  Then list zeros.  Which of "
        "(x + 3) and (x \u2212 1) gives the SMALLER dimension when x = 2?\u201d",
        "Questions to ask (prompt-only): \u201cWhat\u2019s the factored form?\u201d / "
        "\u201cWhat are the zeros?\u201d / \u201cWhich factor is smaller \u2014 (x+3) or (x\u22121) \u2014 "
        "at the width you choose?\u201d",
        "Adult role: teacher circulates, prompts only; aide stations near "
        "pairs stuck on Step 4 (cubic equation) and redirects them to try "
        "small integer values of x.",
        "Release valve: if time is tight, assign Steps 5-6 as HW; preserve "
        "Steps 1-4 in-class.",
    ]


def _exit_key_lines() -> list[str]:
    q = qb.get_for_packet([EXIT_ID])[0]
    return [
        f"[bank: {q['id']}, DOK {q.get('dok')}]",
        _format_unicode(q["prompt"]),
        "",
        f"\u2705  {_format_unicode(str(q.get('correct', '')))}",
        "",
        "Students who finish early: check the conjugate-pair rule held "
        "(if a + bi is a zero, a \u2212 bi must also be).",
        "Questions to ask: none during write time.  As papers come up: "
        "\u201cDid your total count match the degree?\u201d",
        "Adult role: teacher collects at door; aide scans degree-check "
        "answers for Day 6 warm-up diagnostic.",
    ]


def _do_now_key_lines() -> list[str]:
    q = qb.get_for_packet([DO_NOW_ID])[0]
    return [
        f"[bank: {q['id']}, DOK {q.get('dok')}]",
        _format_unicode(q["prompt"]),
        "",
        f"\u2705  {_format_unicode(str(q.get('correct', '')))}",
        "",
        "Bridge questions:",
        "\u2022  If \u22122 is a zero, one factor is (x + 2). [Watch the sign flip.]",
        "\u2022  Degree 3 with only 1 real zero \u2192 2 complex zeros (conjugate pair).",
        "",
        "DIAGNOSTIC scan on entry:",
        "\u2022  Students who wrote (x \u2212 2) for the \u22122 zero have the sign flip "
        "backward \u2014 cold-call during Launch, ask \u201cIf I plug \u22122 into (x \u2212 2), "
        "do I get 0?\u201d",
        "\u2022  Students who didn\u2019t answer the complex-zeros question are primed "
        "for the Launch: the Example 3 discovery lives there.",
        "Questions to ask (circulating silently): none aloud.",
        "Adult role: teacher hands sheets at door, circulates silently; aide "
        "supports processing-speed accommodations without prompting content.",
    ]


IEP_MODS = [
    "\u2022  Do Now is fill-in-the-blank vocabulary (low reading load).",
    "\u2022  Launch scaffolds synthetic division step-by-step on the page.",
    "\u2022  Practice is structured per-item with separate lines for real zero, "
    "quotient, complex zeros.",
    "\u2022  Storage box Explore has all rules printed on the page \u2014 "
    "students self-prompt.",
    "\u2022  Stuck-pair hint card preserves DOK 3 via justification.",
    "\u2022  Release valve: storage box Steps 5-6 \u2192 homework if time tight.",
]

ELL_SUPPORTS = [
    "\u2022  Two new academic terms today: CONJUGATE PAIR, SYNTHETIC DIVISION.",
    "\u2022  Frayer-style quick-reference on the page: \u201cIf a + bi is a zero, so is a \u2212 bi.\u201d",
    "\u2022  Sentence frames on the page: \u201cThe graph shows ___ real zero(s), "
    "but the polynomial must have ___ total zero(s) because ___.\u201d and "
    "\u201cThe zero at x = ___ represents ___ in the context.\u201d",
    "\u2022  Predict-before-verify routine: students commit degree-count before "
    "doing synthetic division.",
]


def build_teacher(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    header_line(doc, "TEACHER EDITION", size=12)
    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    ps.day_banner(doc, "4-5", "Real + Complex Zeros + Modeling",
                  "Teacher pacing + DOK framework crosswalk + answer "
                  "keys + Questions to ask / Adult role per phase.")

    add_two_col(doc, OBJECTIVES + [("CCSS", "F-IF.C.7c, A-APR.B.3, N-CN.C.8, MP.3, MP.4")])
    add_two_col(doc, FRAMEWORK_HEADER)

    add_callout(doc, "\U0001F4CB  DESIGN \u2014  SAVVAS-ONLY DOK-3 DAY", [
        "Do Now \u2192 Launch (all zeros) \u2192 Practice (Try Its) \u2192 ONE DOK-3 Explore "
        "(storage box, Savvas-declared) \u2192 Share/Summary \u2192 Exit.",
        "[DOK 1] Blooket rule recall; Do Now Q28 vocab; Exit Q17.",
        "[DOK 2] Launch Example 3; Practice Try It 3a/3b; Share/Summary.",
        "[DOK 3] Savvas Practice #27 (storage box) \u2014 the Savvas-declared DOK-3 "
        "anchor for Example 4.  Every work item traces to a Savvas bank ID.",
        "This day folds old Day 4 (real + complex zeros) + old Day 5 (modeling) "
        "into one period by treating complex-zero mechanics as launch material "
        "(Example 3) and using the modeling task as the DOK-3 capstone.",
    ])

    add_callout(doc, "\u2705  CRITERIA FOR SUCCESS  /  FORMATIVE ASSESSMENT", CRITERIA_FOR_SUCCESS)

    add_callout(doc, "\u23f1\ufe0f  REALISTIC PACING", [
        "Total: 55 min.  Fits 55-min Tue A tight; 65-min F has 10 min slack.",
        "  Do Now A (solo paper) ........ 5 min",
        "  Do Now B (Blooket login) ..... 2 min",
        "  Do Now C (Blooket game) ...... 7 min",
        "  Launch (Example 3) .......... 10 min",
        "  Practice (Try It 3a + 3b) ... 10 min",
        "  Explore (storage box) ....... 15 min",
        "  Share/Summary ................ 3 min",
        "  Exit (LQ Q17) ................ 5 min",
        "RELEASE VALVES (ordered):",
        "  1. Cut Practice to Try It 3a only (save 5 min for stuck storage-box pairs).",
        "  2. Hint-card a stuck storage-box pair: reveal factored form, "
        "require them to justify dimension mapping.",
        "  3. Storage box Steps 5-6 \u2192 homework; keep Steps 1-4 in-class.",
        "  NEVER cut Exit.  Only artifact that documents learning.",
        "65-min F slack: full Practice for both Try Its, extended storage-box "
        "justification, deeper Share/Summary self-rating.",
    ])

    add_callout(doc, "[Do Now A]  [DOK 1] Savvas Practice #28 + Bridge  (5 min)", [
        "Students receive the Do Now sheet on entry. Silent, 5 min, pencil only.",
        "Vocab check + two bridge questions that prime the Launch discovery.",
    ])
    add_callout(doc, "\u2705  DO NOW / ANSWER KEY", _do_now_key_lines())

    add_callout(doc, "[Do Now B+C]  [DOK 1] Blooket \u2014 Rule Recall ONLY", BLOOKET_SCOPE)

    add_callout(doc, "[Launch]  [DOK 2] Savvas Example 3 \u2014 All Real + Complex Zeros  (10 min)", [
        "Students graph f(x) = x\u00b3 \u2212 2x\u00b2 \u2212 2x \u2212 3, see one real zero, "
        "use synthetic division to reduce to a quadratic, use the quadratic "
        "formula to find the complex conjugate pair.",
        "Script (10 min):",
        "1.  \u201cGraph it in Desmos.  How many times does it cross the x-axis?\u201d (1 min)",
        "2.  Call on a student: \u201cDegree 3 \u2192 how many zeros total?\u201d (30s)",
        "3.  Model synthetic division at board once, 2 min.",
        "4.  Students do synthetic division on their Launch page (2 min).",
        "5.  Model quadratic formula on the quotient, 90s.",
        "6.  Students finish, verify conjugate pair rule (90s).",
        "7.  Name THE RULE (Fundamental Theorem shortcut) and CONJUGATE PAIR. (1 min)",
    ])
    add_callout(doc, "\u2705  LAUNCH / ANSWER KEY", LAUNCH_KEY)

    add_callout(doc, "[Practice]  [DOK 2] Savvas Try It 3a + 3b  (10 min)", [
        "Two Savvas-anchored complex-zero problems (bank IDs: " +
        ", ".join(PRACTICE_IDS) + ").",
        "Students work on both; fast finishers check conjugate-pair rule held "
        "and move to the Explore (storage box).",
        "Teacher circulates with synthetic-division prompts.",
    ])
    add_callout(doc, "\u270F\uFE0F  PRACTICE / ANSWER KEY", _practice_key_lines())

    add_callout(doc, "[Explore]  [DOK 3] Storage Box Volume (Savvas Practice #27)  (15 min)", [
        "Savvas-declared DOK-3 anchor for this lesson.  All three rules "
        "printed on the student page; 6-step scaffold; self-contained.",
        "Students factor the volume polynomial, map factors to physical "
        "dimensions using the \u201cheight is smaller\u201d constraint, then solve "
        "a cubic equation to find dimensions when V = 10.",
        "Teacher role: circulate, prompt-only.  Answer questions with "
        "questions (\u201cWhich of (x+3) and (x\u22121) is smaller at x = 2?\u201d).",
        "Partner roles: one factors (Steps 1-2), the other does dimension "
        "mapping + cubic equation (Steps 3-4).  Both write Steps 5-6.",
    ])
    add_callout(doc, "\U0001F4E6  STORAGE BOX / ANSWER KEY + WHY DOK 3", _dok3_key_lines())

    add_callout(doc, "[Share/Summary]  [DOK 2] Synthesis + Self-Rating  (3 min)", [
        "Teacher script (3 min flat):",
        "1.  Essential Question callback (30s): \u201cHow do synthetic division and "
        "the quadratic formula let us find EVERY zero?\u201d",
        "2.  Language Objective check (30s): \u201cWho used \u2018the zero at x = ___ "
        "represents ___\u2019?\u201d",
        "3.  Self-rating circle (60s).  Teacher scans.",
        "4.  Preview Day 6 (30s): \u201cTomorrow \u2014 polynomial inequalities.\u201d",
        "5.  Transition (30s): \u201cExit ticket.  5 minutes.\u201d",
        "Questions to ask: \u201cWhich Criterion did you rate \u2018partly\u2019?\u201d",
        "Adult role: teacher leads; aide scans self-ratings.",
    ])
    add_callout(doc,
        "\U0001F501  SHARE / SUMMARY (student-facing)",
        SHARE_SUMMARY_STUDENT)

    add_callout(doc, "[Exit Ticket]  [DOK 1\u20132] Savvas Practice #17  (5 min)", [
        "Students find all real and complex zeros of the given polynomial, "
        "then verify the total zero count matches the degree.",
    ])
    ps.summary_exit_box(doc,
        "\U0001F4E4  EXIT TICKET (student-facing)",
        _exit_lines())
    add_callout(doc, "\u2705  EXIT TICKET / ANSWER KEY", _exit_key_lines())

    add_callout(doc, "\U0001F4F7  WHAT TO COLLECT", [
        "1.  Do Now sheets (Practice #28 + bridge).",
        "2.  Launch pages (synthetic division + quadratic formula work).",
        "3.  Storage box pages (DOK-3 evidence \u2014 photograph or collect).",
        "4.  Exit tickets (Practice #17).",
    ])
    add_callout(doc, "\U0001F9E9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS", IEP_MODS)
    add_callout(doc, "\U0001F30D  SUPPORTS FOR ELL STUDENTS", ELL_SUPPORTS)

    doc.save(path)


if __name__ == "__main__":
    build_do_now("Day_45_Do_Now.docx")
    build_student("Day_45_Student_Packet.docx")
    build_teacher("Day_45_Teacher_Packet.docx")
    print("Built Day_45_Do_Now.docx, Day_45_Student_Packet.docx, Day_45_Teacher_Packet.docx")
