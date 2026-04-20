"""Generate Day 3 materials: Do Now sheet + Student Packet + Teacher Packet.

Day 3 = The Magic of Multiplicity.

DESIGN: single-DOK3 spine (matches Day 2 rework).
  Entrance ticket primes the multiplicity pattern.  Blooket is pure rule
  recall.  Launch names the word + builds the rule via (x\u22121)\u207f
  investigation.  Three Try It items consolidate.  ONE self-contained
  DOK-3 driver (Tonya error analysis).  Summary exit ticket.

Phases (55 min, fits 55-min Tue F on the nose; 65-min Mon A / Tue A get slack):
  0\u20135    Do Now A : Touch-or-Cross prediction (3 prompts)     [DOK 2]
  5\u20137    Do Now B : Blooket login                             [\u2014]
  7\u201314   Do Now C : Blooket rule recall (mult + ZPP)          [DOK 1]
  14\u201326  Launch    : Name multiplicity via (x\u22121)\u207f pattern        [DOK 2]
  26\u201333  Practice  : 2 Try It items \u2014 zeros, mults, behavior   [DOK 2]
  33\u201348  Explore   : Tonya error analysis \u2014 the DOK-3 driver   [DOK 3]
  48\u201351  Share/Sum : Synthesis + self-rating                   [DOK 2]
  51\u201355  Exit      : Summary recap (NOT a CER)                 [DOK 1\u20132]
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

# Practice items pulled from the question bank (Savvas-anchored Try Its).
# A bank-id list keeps the source of truth in registry.jsonl; if the items
# are reworded or renumbered, the packet rebuild picks it up automatically.
PRACTICE_IDS = ["3-5-tryit-2a", "3-5-tryit-2b"]

OBJECTIVES = [
    ("Math Objective",
     "Determine the multiplicity of each zero of a polynomial in factored "
     "form, and use multiplicity to predict whether the graph crosses, "
     "touches, or flattens through the x-axis at each zero."),
    ("Language Objective",
     "Students will use the frame \u201cThe multiplicity is ___, so the graph "
     "___ at x = ___\u201d to justify graph behavior."),
    ("Essential Understanding",
     "A repeated factor creates a zero of higher multiplicity. Even "
     "multiplicity causes the graph to touch and turn; odd multiplicity "
     "causes it to cross."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Students name multiplicity, count it from a factored polynomial, and "
     "predict graph behavior at every zero before verifying in Desmos."),
    ("Essential Question",
     "What does a repeated factor do to the graph of a polynomial, and how "
     "can you predict the behavior at every zero just by looking at the "
     "factored form?"),
    ("Materials",
     "Student laptops with Desmos; Day 2 exit ticket (returned); Blooket "
     "access; pencils. No chart paper."),
]


def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if i == 0 and bold_first:
            run.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_two_col(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = TABLE_STYLE
    t.columns[0].width = Inches(1.8)
    t.columns[1].width = Inches(4.7)
    for i, (label, content) in enumerate(rows):
        set_cell(t.rows[i].cells[0], label, bold_first=True)
        if isinstance(content, str):
            content = [content]
        set_cell(t.rows[i].cells[1], content)
    doc.add_paragraph()


def add_callout(doc, title, lines):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.add_run(title).bold = True
    for line in lines:
        cell.add_paragraph(line)
    doc.add_paragraph()


def header_line(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def blank_lines(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.add_run("______________________________________________________________")


# ==================================================================
# DO NOW SHEET (handed out on entry) \u2014 SLIM
# ==================================================================

def build_do_now(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)

    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    header_line(doc, "Day 3  |  DO NOW  \u2014  Touch or Cross?", size=13)

    p = doc.add_paragraph()
    r = p.add_run("Name: _____________________________________     Date: _____________")
    r.font.size = Pt(11)
    doc.add_paragraph()

    add_callout(doc, "\u270d\ufe0f  5 minutes. Silent. Pencil only. No Desmos. No packet.", [
        "From Day 2 you know:  zero of f \u21d4 a factor equals zero (Zero Product Property).",
        "New question:  when a factor appears MORE THAN ONCE, what does the graph do?",
        "For each zero, predict TOUCH (graph kisses the x-axis and turns back) or "
        "CROSS (graph goes through).  Use your eyes and Day 2 thinking \u2014 no Desmos.",
    ])

    header_line(doc, "1.  f(x) = (x + 2)(x \u2212 3)", size=12)
    p = doc.add_paragraph()
    p.add_run("    At x = \u22122 the graph will:  ___________________________")
    p = doc.add_paragraph()
    p.add_run("    At x =  3 the graph will:  ___________________________")
    doc.add_paragraph()

    header_line(doc, "2.  h(x) = (x \u2212 1)\u00b2 (x + 5)", size=12)
    p = doc.add_paragraph()
    p.add_run("    At x =  1 the graph will:  ___________________________")
    p = doc.add_paragraph()
    p.add_run("    At x = \u22125 the graph will:  ___________________________")
    doc.add_paragraph()

    header_line(doc, "3.  Pattern guess:", size=12)
    p = doc.add_paragraph()
    p.add_run("When the factor appears an EVEN number of times, the graph").italic = True
    blank_lines(doc, 2)
    p = doc.add_paragraph()
    p.add_run("When the factor appears an ODD number of times, the graph").italic = True
    blank_lines(doc, 2)

    doc.save(path)


# ==================================================================
# STUDENT PACKET
# ==================================================================

LAUNCH_INVESTIGATION = [
    "We\u2019re about to name the pattern you predicted on the Do Now.  First, "
    "type these in Desmos one at a time and watch what the graph does at  x = 1:",
    "",
    "    (a)  y = (x \u2212 1)\u00b9        At x = 1 the graph: ____________________",
    "    (b)  y = (x \u2212 1)\u00b2        At x = 1 the graph: ____________________",
    "    (c)  y = (x \u2212 1)\u00b3        At x = 1 the graph: ____________________",
    "    (d)  y = (x \u2212 1)\u2074        At x = 1 the graph: ____________________",
    "",
    "What pattern do you see?",
    "_______________________________________________________________",
    "_______________________________________________________________",
    "",
    "THE WORD:   the number of times a factor appears in the factored form is "
    "called its  MULTIPLICITY.",
    "    (x \u2212 1)\u00b2 has multiplicity 2 at x = 1.    (x \u2212 1)\u2074 has multiplicity 4 at x = 1.",
    "",
    "T-CHART  \u2014  fill in after the Desmos investigation:",
    "   ODD multiplicity (1, 3, 5\u2026)        |   EVEN multiplicity (2, 4, 6\u2026)",
    "   Behavior at the zero:                |   Behavior at the zero:",
    "   ____________________________         |   ____________________________",
    "   ____________________________         |   ____________________________",
]

def _format_polynomial_unicode(prompt: str) -> str:
    """Tidy bank prompts (which use ASCII like 'x^2 + 9') into student-facing
    Unicode (\u00b2, \u2212). The bank stays ASCII-clean for round-tripping."""
    # Repair common mojibake from earlier ingests (cp1252-decoded UTF-8).
    prompt = (prompt
              .replace("\u00e2\u2020\u2019", "\u2192")  # â†’  \u2192 (\u2192)
              .replace("\u00e2\u20ac\u201d", "\u2014")  # â€" \u2192 (\u2014)
              .replace("\u00e2\u20ac\u201c", "\u2013")) # â€" \u2192 (\u2013)
    SUPS = {"2":"\u00b2","3":"\u00b3","4":"\u2074","5":"\u2075","6":"\u2076"}
    out = []
    i = 0
    while i < len(prompt):
        ch = prompt[i]
        if ch == "^" and i + 1 < len(prompt) and prompt[i+1] in SUPS:
            out.append(SUPS[prompt[i+1]])
            i += 2
            continue
        if ch == "-":
            out.append("\u2212")
        else:
            out.append(ch)
        i += 1
    return "".join(out)


def _practice_student_lines() -> list[str]:
    items = qb.get_for_packet(PRACTICE_IDS)
    lines = [
        "For each polynomial: list the zeros, the multiplicity of each, and what "
        "the graph does (cross / touch / flatten-through).",
        "",
    ]
    labels = ["A.", "B.", "C.", "D.", "E."]
    for label, q in zip(labels, items):
        # Strip the bank prompt's stem ("Describe the behavior\u2026 f(x) = ...")
        # and surface only the polynomial for student use.
        prompt = _format_polynomial_unicode(q["prompt"])
        if "f(x) =" in prompt or "g(x) =" in prompt or "h(x) =" in prompt:
            poly = prompt.split(":", 1)[-1].strip().rstrip(".")
        else:
            poly = prompt
        lines.append(f"{label}  {poly}")
        lines.append("    Zeros & multiplicities:  ____________________________________")
        lines.append("    Behavior at each:        ____________________________________")
        lines.append("")
    lines.append("Source: Savvas Try It items (question bank IDs: " +
                 ", ".join(PRACTICE_IDS) + ").")
    return lines


def _practice_teacher_lines() -> list[str]:
    items = qb.get_for_packet(PRACTICE_IDS)
    lines = []
    labels = ["A.", "B.", "C.", "D.", "E."]
    for label, q in zip(labels, items):
        prompt = _format_polynomial_unicode(q["prompt"])
        if ":" in prompt:
            poly = prompt.split(":", 1)[-1].strip().rstrip(".")
        else:
            poly = prompt.rstrip(".")
        correct = _format_polynomial_unicode(str(q.get("correct", "")))
        lines.append(f"{label}  {poly}")
        lines.append(f"    \u2705  {correct}")
        rationale = q.get("dok_rationale")
        if rationale:
            lines.append(f"    DOK {q.get('dok')} \u2014 {_format_polynomial_unicode(rationale)}")
        lines.append(f"    [bank id: {q['id']}]")
        lines.append("")
    lines.append("Fast-finisher extension (write on the board if a pair finishes "
                 "before the rest):  \u201cBuild a polynomial that CROSSES at x = 0, "
                 "TOUCHES at x = 2, and FLATTEN-CROSSES at x = \u22121.\u201d  Sample "
                 "answer: f(x) = x(x \u2212 2)\u00b2 (x + 1)\u00b3.")
    return lines


# Lazily computed at first use so an empty/broken bank surfaces in build_*().
TRY_IT_STUDENT = None  # filled by build_student()

# Self-contained DOK-3 driver: rule card on the page so students can pick it
# up and work without waiting on the teacher.
TONYA_STUDENT = [
    "THE PROBLEM",
    "Tonya sketches  f(x) = (x + 1)\u00b2 (x \u2212 3).",
    "She draws the graph CROSSING the x-axis at  x = \u22121  and CROSSING at  x = 3.",
    "",
    "THE RULES YOU NEED (everything is right here)",
    "\u2022  RULE 1:  The multiplicity of a zero is the EXPONENT on its factor.",
    "\u2022  RULE 2:  EVEN multiplicity \u2192 graph TOUCHES and turns at that zero.",
    "\u2022  RULE 3:  ODD multiplicity \u2192 graph CROSSES at that zero.",
    "",
    "STEP 1  \u2014  What did Tonya get RIGHT?",
    "    (Look at the zero at x = 3.  What\u2019s its multiplicity?  Does crossing match?)",
    "_____________________________________________________________________",
    "_____________________________________________________________________",
    "",
    "STEP 2  \u2014  What did Tonya get WRONG?  Use the word \u201cmultiplicity\u201d.",
    "    (Look at the zero at x = \u22121.  What\u2019s its multiplicity?  What SHOULD the graph do?)",
    "_____________________________________________________________________",
    "_____________________________________________________________________",
    "",
    "STEP 3  \u2014  Defend the CORRECT behavior at  x = \u22121.",
    "    Use the sentence frame:  \u201cThe multiplicity is ___, so the graph ___ at x = ___",
    "    because ___.\u201d",
    "_____________________________________________________________________",
    "_____________________________________________________________________",
    "_____________________________________________________________________",
    "",
    "STEP 4  \u2014  Verify in Desmos.",
    "    Graph  f(x) = (x + 1)\u00b2 (x \u2212 3).  Touch at x = \u22121?  Cross at x = 3?",
    "",
    "STEP 5  \u2014  Explain to your partner using the sentence frame above.",
]

EXIT_STUDENT = [
    "Today we discovered:",
    "",
    "1.  The MULTIPLICITY of a zero is the  ______________  on its factor.",
    "",
    "2.  EVEN multiplicity  \u2192  graph  ______________  at the zero.",
    "    ODD multiplicity   \u2192  graph  ______________  at the zero.",
    "",
    "3.  Quick application: in  p(x) = (x + 2)\u00b3 (x \u2212 4),",
    "      at  x = \u22122  the multiplicity is ____  and the graph will ____________.",
    "      at  x =  4  the multiplicity is ____  and the graph will ____________.",
    "",
    "4.  One sentence \u2014 the most important thing I learned today:",
    "_____________________________________________________________________",
]

SHARE_SUMMARY_STUDENT = [
    "Self-check  \u2014  circle one for each:",
    "1.  I can state the multiplicity of each zero from factored form.   \u2713    partly    not yet",
    "2.  I can predict cross vs. touch from the multiplicity.            \u2713    partly    not yet",
    "3.  I can spot a multiplicity error in someone else\u2019s graph.        \u2713    partly    not yet",
    "",
    "Preview Day 4: what if a polynomial\u2019s zeros aren\u2019t all real numbers?",
]


def build_student(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    ps.running_header_name_block(doc)
    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    ps.day_banner(doc, 3, "The Magic of Multiplicity",
                  "Use Desmos to compare graphs.  Use cross, touch, turn, "
                  "even, odd, and multiplicity in every explanation.")

    add_two_col(doc, OBJECTIVES)
    add_two_col(doc, FRAMEWORK_HEADER)

    add_callout(doc, "\U0001F501  WELCOME BACK", [
        "Day 2: zeros + a point \u2192 equation (you worked BACKWARD).",
        "Today: when a factor appears more than once, the graph does something "
        "different at that zero. We\u2019ll name it.",
        "Materials: laptop (Desmos), pencil.",
    ])

    add_callout(doc, "\U0001F4C4  DO NOW  \u2014  separate sheet", [
        "You got the Touch-or-Cross sheet on entry.",
        "Hold onto it \u2014 we\u2019ll compare predictions during the Launch.",
    ])

    add_callout(doc, "\U0001F3AE  Blooket  \u2014  Rule Recall   [Do Now B+C]  [DOK 1]   (2 + 7 min)", [
        "Teacher puts the code on the screen. Log in (2 min), then 7-minute game.",
        "Focus: counting multiplicity from exponents, EVEN \u2192 touch, ODD \u2192 cross, "
        "plus a couple Day 2 ZPP recall items.",
    ])

    add_two_col(doc, [(
        "\U0001F4A1  Launch  \u2014  Name the Pattern   [Launch]  [DOK 2]   (12 min)",
        LAUNCH_INVESTIGATION,
    )])

    add_two_col(doc, [(
        "\u270F\uFE0F  Practice  \u2014  Zeros + Multiplicities + Behavior   [Practice]  [DOK 2]   (7 min)",
        _practice_student_lines(),
    )])

    add_two_col(doc, [(
        "\U0001F50E  Explore  \u2014  Tonya\u2019s Graph (Error Analysis)   [Explore]  [DOK 3]   (15 min)",
        TONYA_STUDENT,
    )])

    add_callout(doc,
        "\U0001F501  SHARE / SUMMARY   [Share/Summary]  [DOK 2]   (3 min)",
        SHARE_SUMMARY_STUDENT)

    ps.summary_exit_box(doc,
        "\U0001F4E4  EXIT TICKET  \u2014  SUMMARY   [Exit Ticket]  [DOK 1\u20132]   (4 min)",
        EXIT_STUDENT)

    doc.save(path)


# ==================================================================
# TEACHER PACKET
# ==================================================================

CRITERIA_FOR_SUCCESS = [
    "Student-facing, revisited during Share/Summary:",
    "1.  I can state the multiplicity of each zero from factored form.",
    "2.  I can predict cross vs. touch from the multiplicity.",
    "3.  I can spot and correct a multiplicity error in another student\u2019s graph.",
    "Formative evidence:",
    "\u2022  Do Now sheets (scan #3 pattern guesses on entry).",
    "\u2022  Blooket dashboard (DOK 1 rule recall).",
    "\u2022  Try It A\u2013C pages (circulate, photograph).",
    "\u2022  Tonya error analysis page \u2014 collected (DOK 3 evidence).",
    "\u2022  Summary exit ticket.",
]

DO_NOW_KEY = [
    "Expected predictions (accept any clear \u201ctouch\u201d / \u201ccross\u201d):",
    "1.  f(x) = (x + 2)(x \u2212 3)",
    "    \u2022  x = \u22122: CROSS  (multiplicity 1)",
    "    \u2022  x =  3: CROSS  (multiplicity 1)",
    "2.  h(x) = (x \u2212 1)\u00b2 (x + 5)",
    "    \u2022  x =  1: TOUCH  (multiplicity 2)",
    "    \u2022  x = \u22125: CROSS  (multiplicity 1)",
    "3.  Pattern target:",
    "    \u2022  EVEN \u2192 touches and turns",
    "    \u2022  ODD  \u2192 crosses",
    "DIAGNOSTIC scan as students hand sheets in:",
    "\u2022  Students who wrote \u201ccross\u201d on EVERY zero in #2 \u2192 cold-call during Launch.",
    "\u2022  Students who articulated the EVEN/ODD rule in #3 \u2192 have them share first. "
    "Peer-voiced rule beats teacher-voiced rule.",
]

BLOOKET_SCOPE = [
    "SCOPE (pure rule-recall, DOK 1).  Multiplicity rules + Day 2 ZPP retention.",
    "Habits of mind today\u2019s work depends on:",
    "\u2022  Multiplicity = exponent on the factor (e.g., (x \u2212 4)\u00b2 \u2192 mult 2 at x = 4).",
    "\u2022  EVEN multiplicity \u2192 TOUCH and turn.",
    "\u2022  ODD multiplicity \u2192 CROSS.",
    "\u2022  Day 2 retention: zero \u21d4 factor = 0; zero at x = c gives factor (x \u2212 c).",
    "DO NOT include items that require students to factor an expression \u2014 "
    "that\u2019s DOK 2 problem solving and doesn\u2019t belong on flashcards.",
    "Dashboard diagnostic: note the two lowest items. If a rule is cold, "
    "60-second board reset before Launch.",
    "CSV: Blooket_Day3_Multiplicity.csv  (regenerate via regen_blooket_csvs.py).",
    "Questions to ask (after game): \u201cWhich rule had the lowest %?  Say it now.\u201d "
    "(Surface the weak rule aloud before Launch so it primes the \u2212f(x) synthesis.)",
    "Adult role: teacher watches dashboard and runs the 60-second board reset if "
    "needed; aide (if present) troubleshoots any Chromebook issues during the "
    "2-min login window.",
]

LAUNCH_KEY = [
    "Target observations at  x = 1:",
    "(a)  y = (x \u2212 1)\u00b9  \u2192  clean straight-through CROSS (linear slope).",
    "(b)  y = (x \u2212 1)\u00b2  \u2192  TOUCHES (parabolic kiss; turns back).",
    "(c)  y = (x \u2212 1)\u00b3  \u2192  CROSSES but FLATTENS at the zero (S-curve).",
    "(d)  y = (x \u2212 1)\u2074  \u2192  TOUCHES, flatter than (b).",
    "Pattern: EVEN \u2192 touch/turn; ODD \u2192 cross; higher multiplicity \u2192 flatter at the zero.",
    "T-chart fill (canonical):",
    "  ODD: graph crosses through the x-axis (1 = clean cross; 3, 5 = flatten-through).",
    "  EVEN: graph touches the x-axis and turns back (sign of f doesn\u2019t change there).",
    "TEACHER MOVE \u2014 the naming:",
    "After (b) is graphed, pause: \u201cWho predicted touch on the Do Now? Why?\u201d "
    "Surface a student answer.  Then say: \u201cThe word for how many times a factor "
    "appears is MULTIPLICITY.\u201d Write on the board.  Continue with (c) and (d).",
    "DON\u2019T volunteer the FLATTEN-as-multiplicity-grows insight \u2014 let students "
    "notice it.  Bonus: \u201cWhy does (c) flatten compared to (a)?\u201d (Cubing a "
    "small number makes it tiny \u2014 graph stays near the axis longer.)",
]

# TRY_IT_KEY now generated from the bank \u2014 see _practice_teacher_lines().

TONYA_KEY = [
    "Target: Tonya\u2019s f(x) = (x + 1)\u00b2 (x \u2212 3).  She drew BOTH zeros crossing.",
    "\u2705  STEP 1 (RIGHT): x = 3 has multiplicity 1 (odd), so crossing IS correct there.",
    "\u2705  STEP 2 (WRONG): the factor (x + 1) has exponent 2, so x = \u22121 has "
    "multiplicity 2 (even). Even multiplicity means the graph TOUCHES and TURNS, "
    "NOT crosses.",
    "\u2705  STEP 3 (sample defense, sentence frame):",
    "    \u201cThe multiplicity is 2, so the graph touches and turns at x = \u22121, "
    "because (x + 1)\u00b2 stays the same sign just before and just after x = \u22121, "
    "so the overall sign of f(x) can\u2019t flip there.\u201d",
    "WHY THIS IS DOK 3:",
    "\u2022  Students critique a peer\u2019s reasoning (not just compute).",
    "\u2022  They construct an evidence-based argument using domain vocabulary "
    "(multiplicity, even, sign change).",
    "\u2022  Justification \u2014 not procedure execution \u2014 is the deliverable.",
    "Stuck-pair hint card: \u201cAt x = \u22121, write down the exponent on (x + 1). "
    "Now read Rule 2.\u201d  (Preserves DOK 3 because they still must articulate WHY.)",
    "ELL look-for: \u201cbecause\u201d / \u201csince\u201d in the defense; \u201ctouch\u201d / "
    "\u201cturn\u201d / \u201ckisses\u201d all valid academic-language proxies.",
]

EXIT_KEY = [
    "Expected responses:",
    "1.  EXPONENT  (accept: \u201cpower,\u201d \u201cnumber on top of the parenthesis\u201d).",
    "2.  EVEN \u2192 TOUCHES (accept: \u201ckisses,\u201d \u201cturns,\u201d \u201cbounces\u201d).",
    "    ODD  \u2192 CROSSES (accept: \u201cgoes through,\u201d \u201ccuts through\u201d).",
    "3.  Application to p(x) = (x + 2)\u00b3 (x \u2212 4):",
    "    x = \u22122: multiplicity 3, CROSSES (with a flatten).",
    "    x =  4: multiplicity 1, CROSSES (clean).",
    "4.  Open response.  Strong answers connect EXPONENT \u2192 PARITY \u2192 BEHAVIOR.",
    "INTENTIONALLY NOT A CER.  The DOK-3 generative work happened in Tonya.",
    "If you want a CER for HW: \u201cExplain why even multiplicity must touch, "
    "using sign-chart reasoning.\u201d",
]

IEP_MODS = [
    "\u2022  Do Now uses CHECK-or-CROSS predictions (low writing load).",
    "\u2022  Launch (x \u2212 1)\u207f investigation: partner with a strong-grapher; "
    "IEP student records observations in writing.",
    "\u2022  Tonya error analysis: structured Steps 1\u20135 with self-contained rule card. "
    "Sentence frame printed twice (T-chart + Step 3).",
    "\u2022  Stuck-pair hint card: \u201cWrite the exponent. Read Rule 2.\u201d",
    "\u2022  Try It is structured per-zero (zeros + mult + behavior on three lines).",
    "\u2022  Exit ticket is fill-in-the-blank with one open sentence \u2014 low writing load.",
    "\u2022  Preferential seating near board for Launch synthesis.",
]

ELL_SUPPORTS = [
    "\u2022  Sentence frame printed on Tonya page (Step 3) and on the T-chart.",
    "\u2022  \u201cTouches / kisses / turns / bounces\u201d all accepted for even-mult behavior.",
    "\u2022  \u201cCrosses / goes through / cuts through\u201d all accepted for odd-mult behavior.",
    "\u2022  ONE new academic word today: MULTIPLICITY.  No other new vocab.",
    "\u2022  Predict-before-verify routine: students commit in writing before speaking.",
    "\u2022  T-chart side-by-side keeps EVEN and ODD as a visual reference.",
    "\u2022  Summary exit ticket is fill-in-the-blank + one sentence (low writing load).",
]


def build_teacher(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    header_line(doc, "TEACHER EDITION", size=12)
    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    ps.day_banner(doc, 3, "The Magic of Multiplicity",
                  "Teacher pacing + DOK framework crosswalk + answer keys + "
                  "Questions to ask / Adult role per phase.")

    add_two_col(doc, OBJECTIVES + [("CCSS", "F-IF.C.7c, A-APR.B.3, MP.3, MP.7")])
    add_two_col(doc, FRAMEWORK_HEADER)

    add_callout(doc, "\U0001F4CB  DESIGN \u2014  ONE DOK-3 SPINE", [
        "Do Now \u2192 Launch (name the word) \u2192 Practice (apply) \u2192 ONE DOK-3 Explore "
        "(Tonya) \u2192 Summary Exit.",
        "[DOK 1] Blooket rule recall.",
        "[DOK 2] Touch/Cross prediction; Launch (x\u22121)\u207f investigation; Try It; summary exit.",
        "[DOK 3] Tonya error analysis \u2014 the sole DOK-3 task today.",
        "Why this design (Day 1\u20132 evidence): bundling two DOK-3 tasks in one period "
        "stretches the lesson to 1.4\u20131.5\u00d7 the budget. Day 3 keeps Tonya as the spine "
        "and lets the (x\u22121)\u207f investigation stay DOK 2 (pattern observation, not "
        "argument construction).",
    ])

    add_callout(doc, "\u2705  CRITERIA FOR SUCCESS  /  FORMATIVE ASSESSMENT", CRITERIA_FOR_SUCCESS)

    add_callout(doc, "\u23f1\ufe0f  REALISTIC PACING", [
        "Total: 55 min.  Blocks: Tuesday F = 55 min (on the nose), "
        "Monday A / Tuesday A = 65 min (10 min slack).",
        "  Do Now A (solo paper) ........ 5 min",
        "  Do Now B (Blooket login) ..... 2 min",
        "  Do Now C (Blooket game) ...... 7 min",
        "  Launch ((x\u22121)\u207f + naming) .... 12 min",
        "  Practice (Try It A\u2013B) ......... 7 min",
        "  Explore (Tonya) .............. 15 min",
        "  Summary Exit .................. 4 min",
        "RELEASE VALVES (in order of preference):",
        "  1. Cut Practice from 10 \u2192 8 min (assign Try It C as HW).",
        "  2. Hint-card a stuck pair on Tonya (Step 1: write the exponent).",
        "  3. Cut Launch from 12 \u2192 10 min (skip the (x\u22121)\u2074 graph).",
        "  NEVER cut the Summary Exit. Only artifact that documents learning.",
        "65-min slack uses: extended Tonya partner CER, fast-finisher extension "
        "in Practice (build a polynomial with cross/touch/flatten constraints), "
        "deeper Frayer for multiplicity.",
    ])

    # ---- Do Now A ----
    add_callout(doc, "[Do Now A]  [DOK 2] Touch-or-Cross Prediction  (5 min)", [
        "Students receive the Do Now sheet on entry. Silent, 5 min, pencil only.",
        "Three small prediction items + one pattern guess. Breathable.",
        "Bridges from Day 2 (still ZPP) and primes today\u2019s pattern reveal.",
        "Questions to ask (circulating silently): none aloud. Scan for "
        "\u201ccross-everywhere\u201d misconception and note those students for Launch cold-call.",
        "Adult role: teacher hands sheets at door, circulates silently; aide (if present) "
        "supports processing-speed accommodations without prompting content.",
    ])
    add_callout(doc, "\u2705  DO NOW / ANSWER KEY", DO_NOW_KEY)

    # ---- Do Now B+C ----
    add_callout(doc, "[Do Now B+C]  [DOK 1] Blooket \u2014 Rule Recall ONLY", BLOOKET_SCOPE)

    # ---- Launch ----
    add_callout(doc, "[Launch]  [DOK 2] Name the Pattern via (x \u2212 1)\u207f  (12 min)", [
        "Pulls Do Now into a public synthesis, then partners run a fast Desmos "
        "investigation of (x \u2212 1)\u207f for n = 1, 2, 3, 4. Teacher names the word "
        "MULTIPLICITY at the moment the (x \u2212 1)\u00b2 graph touches.",
        "Script:",
        "1.  \u201cPull out your Do Now.\u201d (15s)",
        "2.  Cold-call the student who wrote a clean even/odd rule in #3. (60s)",
        "3.  Cold-call a student who wrote \u201ccross\u201d everywhere in #2 \u2014 "
        "address the misconception publicly. (60s)",
        "4.  Partners type (a)\u2013(d) in Desmos, record behavior at x = 1. (5 min)",
        "5.  Class share: \u201cWhat happened at multiplicity 2 vs. 3 vs. 4?\u201d (3 min)",
        "6.  Name the word, build the T-chart together. (3 min)",
        "Questions to ask: \u201cWhat changed between (x\u22121)\u00b9 and (x\u22121)\u00b2 at x = 1?\u201d / "
        "\u201cHow many times does the factor (x\u22121) appear in each case?\u201d / "
        "\u201cWho predicted touch on the Do Now?  What made you predict that?\u201d",
        "Adult role: teacher runs script + cold-calls; aide (if present) walks to "
        "partners stuck loading Desmos, keeps them in the 5-min window.",
    ])
    add_callout(doc, "\u2705  LAUNCH / ANSWER KEY", LAUNCH_KEY)

    # ---- Practice ----
    add_callout(doc, "[Practice]  [DOK 2] Try It A\u2013B  (7 min)", [
        "Two items pulled from the Savvas bank.  Each: zeros + multiplicity + behavior.",
        "Teacher circulates with sentence-frame prompts, not corrections.",
        "Questions to ask: \u201cWhich factor tells you the behavior at this zero?\u201d / "
        "\u201cIs this multiplicity even or odd?\u201d",
        "Adult role: teacher circulates with prompts; aide (if present) scans for "
        "students stuck on Item B\u2019s irreducible quadratic and redirects to Rule 1.",
        "Fast finishers: write a polynomial with prescribed cross/touch/flatten "
        "behaviors at three zeros (extension on the board).",
    ])
    add_callout(doc, "\u270F\uFE0F  TRY IT / ANSWER KEY", _practice_teacher_lines())

    # ---- Explore (Tonya, the DOK-3 driver) ----
    add_callout(doc, "[Explore]  [DOK 3] Tonya Error Analysis  (15 min)", [
        "ONE driver.  All three rules printed on the student page; 5-step scaffold; "
        "self-contained.  Students critique Tonya\u2019s sketch and defend the correct "
        "behavior using multiplicity language.",
        "Partner roles: one writes Step 1\u20132, the other drafts Step 3 (the defense). "
        "Swap for fast-finisher CER expansion if time allows.",
        "Questions to ask (prompt-only, answer questions with questions): "
        "\u201cWhat\u2019s the exponent on (x + 1)?\u201d / \u201cRead Rule 2 aloud.  "
        "Does that match what Tonya drew?\u201d / \u201cWhich factor appears twice?  "
        "What does that mean?\u201d",
        "Adult role: teacher circulates with prompts only (no answers); aide (if "
        "present) stations near pairs with processing-speed accommodations, prompts "
        "partner swap when Step 3 stalls.",
    ])
    add_callout(doc, "\U0001F50E  TONYA / ANSWER KEY + WHY DOK 3", TONYA_KEY)

    # ---- Share / Summary (distinct phase, per DOKframework.txt) ----
    add_callout(doc, "[Share/Summary]  [DOK 2] Synthesis + Self-Rating  (3 min)", [
        "Teacher script (3 min flat):",
        "1.  \u201cPencils up.\u201d Essential Question callback (45s): \u201cWhat does a "
        "repeated factor do, and how do you predict the behavior at every zero?\u201d",
        "2.  Language Objective check (30s): \u201cWho used \u2018the multiplicity is ___, "
        "so the graph ___ at x = ___\u2019? Finish with what?\u201d",
        "3.  Self-rating circle on packet (45s).  Teacher scans.",
        "4.  Preview Day 4 (30s): \u201cTomorrow: what if the zeros aren\u2019t all real?\u201d",
        "5.  Transition (30s): \u201cSummary exit ticket. 4 minutes.\u201d",
        "Questions to ask: \u201cWhich Criterion for Success did you rate \u2018partly\u2019? "
        "What\u2019s the blocker?\u201d (use answers to target Day 4 warm-up).",
        "Adult role: teacher leads; aide (if present) scans self-ratings as packets "
        "tilt up during transition.",
    ])
    add_callout(doc,
        "\U0001F501  SHARE / SUMMARY (student-facing self-rating)",
        SHARE_SUMMARY_STUDENT)

    # ---- Exit Ticket ----
    add_callout(doc, "[Exit Ticket]  [DOK 1\u20132] Summary (NOT a CER)  (4 min)", [
        "Design intent: SUMMARY of today\u2019s learning, not a new cognitive task.",
        "Three fill-in-the-blank items + one open sentence.",
        "Why not CER here: full CER writing in 4 min is unrealistic; the DOK-3 "
        "generative work already happened in Tonya.",
        "Questions to ask: none during write time.  As papers come up: "
        "\u201cBiggest thing you learned?  Say it in one breath.\u201d",
        "Adult role: teacher collects at the door; aide (if present) totals #3 "
        "responses for next-day diagnostic while students pack up.",
    ])
    ps.summary_exit_box(doc,
        "\U0001F4E4  EXIT TICKET  \u2014  SUMMARY (student-facing)",
        EXIT_STUDENT)
    add_callout(doc, "\u2705  EXIT TICKET / ANSWER KEY", EXIT_KEY)

    # ---- Closers ----
    add_callout(doc, "\U0001F4F7  WHAT TO COLLECT", [
        "1.  Do Now sheets \u2014 collected on entry scan or during Launch.",
        "2.  Tonya page \u2014 photograph or collect (DOK-3 evidence).",
        "3.  Summary exit ticket (every student).",
        "4.  Try It pages \u2014 quick scan for diagnostic.",
    ])
    add_callout(doc, "\U0001F440  LOOK-FORS DURING WALKTHROUGH", [
        "[Do Now A] Silent solo work? Pattern guess attempted in #3?",
        "[Do Now C] Teacher monitoring Blooket dashboard diagnostically?",
        "[Launch] Cold-call used for misconception? Word \u201cmultiplicity\u201d named at the right moment?",
        "[Practice] Sentence-frame prompts? Teacher circulating with questions, not answers?",
        "[Explore] Partners reading the rule card? \u201cBecause\u201d / \u201csince\u201d in Step 3?",
        "[Exit] All four items attempted? Sentence complete?",
    ])
    add_callout(doc, "\U0001F9E9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS", IEP_MODS)
    add_callout(doc, "\U0001F30D  SUPPORTS FOR ELL STUDENTS", ELL_SUPPORTS)

    doc.save(path)


if __name__ == "__main__":
    build_do_now("Day_3_Do_Now.docx")
    build_student("Day_3_Student_Packet.docx")
    build_teacher("Day_3_Teacher_Packet.docx")
    print("Built Day_3_Do_Now.docx, Day_3_Student_Packet.docx, Day_3_Teacher_Packet.docx")
