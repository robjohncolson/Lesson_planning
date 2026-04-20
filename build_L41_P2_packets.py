"""Generate Lesson 4-1 Period 2 materials — Klimsara-adapted, no Blooket.

Design: Inverse variation applications + the lesson's single-DOK-3 spine.
P2 is the conceptual core: students apply y = k/x to real contexts, then
tackle the Ramón performance task (#26) which contrasts DIRECT vs INVERSE
variation — the pedagogical anchor of the whole lesson.

Phases (55 min base; Period F 65-min):
  0–  5   Do Now    : Practice #15 — quick table-check of inverse variation [DOK 1]
  5– 17   Launch    : teacher models Example 3 (bouzouki string / frequency) [DOK 2]
 17– 50   Explore   : Try-It 3 + Practice #12/#17/#23 + ⭐ DOK-3 q26 (Ramón) [DOK 2-3]
 50– 55   Share/Sum : DOK-3 pair presents + Wed bridge to reciprocal function [DOK 2]
 55– 58   Exit      : summary recap fill-in                                  [DOK 1-2]
 (back)   Optional: Practice #25 (SAT/ACT MC DOK-3) — fast-finisher stretch
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_ID   = "4-1-savvas-q15"
EXPLORE_IDS = [
    "4-1-savvas-try-it-3-lesson-4-1-matches-examp",
    "4-1-savvas-q12",        # HOT — inverse square
    "4-1-savvas-q17",        # wavelength/frequency graph
    "4-1-savvas-q23",        # Boyle's Law
    "4-1-savvas-q26",        # ⭐ DOK-3 spine — Ramón (direct vs inverse)
]
DOK3_ID       = "4-1-savvas-q26"
REINFORCE_IDS = ["4-1-savvas-q25"]   # SAT/ACT MC DOK-3 stretch

_ALL_IDS = [DO_NOW_ID] + EXPLORE_IDS + REINFORCE_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 4-1 · Period 2"
LESSON_TITLE = "Inverse Variation Applications + DOK-3 Performance Task"

OBJECTIVES = [
    ("Math Objective",
     "Apply inverse variation to real-world contexts (wavelength, gas pressure, "
     "travel distance), and distinguish DIRECT variation (y = kx) from INVERSE "
     "variation (xy = k) inside a single multi-part problem."),
    ("Language Objective",
     "Students will use the frame \u201cThis is ___ variation because as ___ "
     "___, ___ ___ proportionally.\u201d to justify whether a relationship is "
     "direct or inverse."),
    ("Essential Understanding",
     "Not every covariation is inverse variation. Direct variation has a "
     "constant RATIO (y/x = k); inverse variation has a constant PRODUCT "
     "(xy = k). A single scenario can contain BOTH types across different "
     "variable pairs."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Students model 4 contextual inverse-variation problems, then tackle a "
     "Performance Task (Ramón's road trip) that embeds direct and inverse "
     "variation in one scenario. Every item is Savvas-sourced."),
    ("Essential Question",
     "How does inverse variation model real-world relationships, and how is "
     "it different from direct variation?"),
    ("Materials",
     "Pencils; student packet; calculator. No Chromebooks required."),
]


# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.0):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    prompt_p.add_run(q["prompt"]).font.size = Pt(11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            line.add_run(f"  ({chr(65+i)})  {a}").font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, for_teacher=False):
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 2, LESSON_TITLE,
                  subtitle=LESSON_LABEL + (" · Teacher Edition" if for_teacher else ""))
    t = doc.add_table(rows=len(OBJECTIVES), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(OBJECTIVES):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    t = doc.add_table(rows=len(FRAMEWORK_HEADER), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(FRAMEWORK_HEADER):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()


# ── Do Now ───────────────────────────────────────────────────────────────

def build_do_now(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 2, "Do Now — " + LESSON_TITLE, subtitle=LESSON_LABEL)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from P1",
        dok="1",
        minutes=5,
        teacher_does=["Hand out at door.", "Circulate 3 min.",
                      "Cold-call 1 student: \"Is this table inverse? What's k?\""],
        students_do=["Compute xy for each column.",
                     "If constant → inverse; state k. If not → not inverse."],
        questions_to_ask=[
            "What do you multiply to check for inverse variation?",
            "What would make this NOT inverse? (xy not constant.)",
        ],
        adult_role="Solo teacher: circulate, time 5 min, pull sheets.",
    )

    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label="Do Now — Is this inverse variation?",
                    space_after_inches=2.0)

    add_callout(doc, "🎯  SENTENCE FRAME",
                ["\u201cThe product xy is ___ for every pair, so this ___ (is / is not) an inverse variation. The constant of variation is ___.\u201d"],
                color_hex="FFF2CC")

    doc.save(path)


# ── Student packet ───────────────────────────────────────────────────────

def build_student(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=False)

    ps.teal_section(doc, "LAUNCH — Inverse Variation in Context (teacher models Example 3)", [])
    add_callout(doc, "📘  APPLICATION RULES (keep printed on your page)",
                [
                    "• To model an inverse variation in context:",
                    "     1. Identify the two varying quantities and their constant product k.",
                    "     2. Compute k from one given pair: k = (quantity 1) × (quantity 2).",
                    "     3. Write the equation: y = k/x (or rename variables to fit the context).",
                    "     4. Solve for the unknown quantity.",
                    "• Example (Savvas Ex 3, bouzouki): string length s varies inversely with frequency f.",
                    "     26-inch string → 329.63 cycles/sec → k = 26 × 329.63 = 8,570.38.",
                    "     New length 13 inches → f = 8,570.38 / 13 = 659.26 cycles/sec.",
                ],
                color_hex="D9EAD3")

    add_callout(doc, "⚠️  DIRECT vs INVERSE (both appear in today's DOK-3!)",
                [
                    "DIRECT: y = kx → y/x is constant. As x increases, y INCREASES proportionally.",
                    "INVERSE: y = k/x (xy = k). As x increases, y DECREASES proportionally.",
                    "Test: compute y/x (direct) AND xy (inverse). Whichever is constant tells you the relationship.",
                ],
                color_hex="F4CCCC")
    doc.add_paragraph()

    ps.teal_section(doc, "EXPLORE — Think-Pair-Share (33 min)",
                    ["Work with a partner. Move in order. Practice #26 (Ramón) is the big one — plan ~15 min for it."])

    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 3 — Ice cube melting",
        "Practice #12 — HOT: inverse SQUARE variation",
        "Practice #17 — Radio wavelength (graph)",
        "Practice #23 — Boyle's Law (volleyball vs basketball)",
        "Practice #26  ⭐ PERFORMANCE TASK · Ramón's road trip",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=label, space_after_inches=1.5)

    ps.teal_section(doc, "SHARE / SUMMARY (5 min)", [])
    ps.sentence_frame_box(doc, [
        "Pair share (Ramón task): \u201cPart A is ___ variation because distance and time have a constant RATIO of ___. Part B is ___ variation because gas and time have a constant PRODUCT of ___.\u201d",
        "Whole class: one pair reads the contrast. Class signals thumbs up/sideways.",
    ])

    add_callout(doc, "🔭  BRIDGE TO P3 (Wednesday)",
                ["Tomorrow we graph the RECIPROCAL FUNCTION y = 1/x and its translations. Watch for: the graph of inverse variation has asymptotes and a specific shape — a hyperbola. You've been drawing them already; we'll name the features."],
                color_hex="FCE5CD")

    ps.summary_exit_box(doc, "EXIT TICKET — Summary Recap",
                        [
                            "Today I learned that _____________________________________________________.",
                            "The biggest trap in Ramón's task was _____________________________________________________.",
                            "I distinguish direct from inverse variation by _____________________________________________________.",
                        ])

    doc.add_page_break()
    ps.teal_section(doc, "OPTIONAL REINFORCEMENT (not collected)",
                    ["If you finish Explore early. #25 is a fast SAT-style question that stretches the inverse-variation idea."])
    reinforce = qb.get_for_packet(REINFORCE_IDS)
    for i, q in enumerate(reinforce, 1):
        bank_item_block(doc, q, label=f"Optional #{i}  (Savvas Practice, SAT/ACT style)",
                        space_after_inches=1.8)

    doc.save(path)


# ── Teacher packet ───────────────────────────────────────────────────────

def build_teacher(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=True)

    add_callout(doc, "📌  PRE-FLIGHT — P2 IS THE DOK-3 SPINE DAY",
                [
                    "Today's conceptual core is Practice #26 (Ramón). Part A is direct variation, Part B is inverse — same scenario, different variable pairs. That contrast IS the lesson's deepest content.",
                    "Pace: Try-It 3 (5 min) → #12 (4 min) → #17 (5 min) → #23 (7 min) → #26 (12-15 min). ~33-36 min total Explore.",
                    "If running tight: cut #12 (the HOT inverse-square is tangent, not central). Never cut #26.",
                ],
                color_hex="FFD9E0")

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from P1",
        dok="1",
        minutes=5,
        teacher_does=["Hand out at door.", "Circulate 3 min.",
                      "Cold-call 1 student for xy computation."],
        students_do=["Compute xy for each column.",
                     "State k if constant; explain if not."],
        questions_to_ask=[
            "What do you multiply to check inverse variation?",
            "Is xy constant here? What does that tell you?",
        ],
        adult_role="Solo teacher: circulate silently. No hints.",
    )
    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label=f"Do Now ({q['id']})", space_after_inches=0.3)
    add_callout(doc, "✅  ANSWER KEY",
                [q.get("notes") or "(verify)"],
                color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Launch",
        framework_tag="teacher models Example 3",
        dok="2",
        minutes=12,
        teacher_does=[
            "Project Example 3 (bouzouki). Think aloud: extract the two quantities (s = string length, f = frequency). Compute k = sf.",
            "\"Now I can find f for any s. Plug in the new length, solve.\"",
            "Flag the Application Rules card — print it in student minds.",
            "30-sec pause: \"What's k for YOUR Do Now table?\" (partner check)",
            "Do NOT solve Try-It 3 or Practice items — release at minute 17.",
        ],
        students_do=[
            "Watch the bouzouki example silently.",
            "During pause: state k from Do Now to partner.",
            "Copy Application Rules card into margin.",
            "Note the DIRECT vs INVERSE warning: this comes back in #26.",
        ],
        questions_to_ask=[
            "What two quantities are related? What's constant?",
            "\"Every radio wave with frequency f has wavelength w such that wf = constant.\" Why is that inverse?",
            "If frequency doubles, what happens to wavelength? Why?",
        ],
        adult_role="Project. Think aloud. Release promptly at minute 17.",
    )
    add_callout(doc, "🎤  TEACHER SCRIPT",
                [
                    "1. \u201cWatch me set up a real-world inverse variation. Bouzouki string: length s varies inversely with frequency f.\u201d",
                    "2. \u201cStep 1: Identify the quantities and find k. k = sf = 26 × 329.63 = 8,570.38.\u201d",
                    "3. \u201cStep 2: Write the equation. s = k/f, so s = 8570.38/f.\u201d",
                    "4. \u201cStep 3: Plug in the new value. For s = 13: f = 8570.38/13 = 659.26 cycles/sec.\u201d",
                    "5. \u201cTonight's Performance Task has a twist — ONE scenario contains BOTH direct and inverse variation. Be ready to tell them apart.\u201d",
                    "6. Release to Explore.",
                ],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Explore",
        framework_tag="TPS + DOK-3 driver",
        dok="2-3",
        minutes=33,
        teacher_does=[
            "5 laps across 33 min.",
            "Lap 1 (5 min): Try-It 3 — ice cube. Press pairs to identify k = Temp × Time.",
            "Lap 2 (4 min): #12 — inverse square. This is a THINKING item; let pairs struggle.",
            "Lap 3 (5 min): #17 — graph reading + context. Confirm k = 300,000.",
            "Lap 4 (7 min): #23 — Boyle's Law. Check that pairs compute volleyball's k = 1350.",
            "Lap 5 (12 min): ⭐ #26 Ramón. PART A is direct; PART B is inverse. Press hard on that contrast.",
            "Do NOT give #26 answers. Pairs present Part A/B to each other; one pair presents at Share.",
        ],
        students_do=[
            "5 items in order: Try-It 3 → #12 → #17 → #23 → #26.",
            "For #26: BEFORE solving, identify which part is direct and which is inverse. Write it down.",
            "Justify with sentence frame.",
        ],
        questions_to_ask=[
            "What's k in this problem? What units?",
            "#12: y = k/x². If x → 4x, what happens to y? (Divide by 16.)",
            "#17: from the graph, what's k? How do you read it?",
            "#23: Whose k do we use — the volleyball's, or the basketball's? (Volleyball's; the basketball only gives the target PRESSURE.)",
            "#26 Part A: is distance/time constant? → direct variation.",
            "#26 Part B: is gas × time constant? → inverse variation.",
        ],
        adult_role="Solo teacher: 5 laps. Press DIRECT vs INVERSE distinction on #26.",
    )
    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 3", "Practice #12", "Practice #17", "Practice #23",
        "Practice #26 ⭐ DOK-3 DRIVER",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=f"{label} ({q['id']})", space_after_inches=0.3)
        add_callout(doc, "✅  ANSWER / ANTICIPATED TALK",
                    [q.get("notes") or "(no answer key — verify before class)"],
                    color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Share / Summary",
        framework_tag="academic conversation + Wed bridge",
        dok="2",
        minutes=5,
        teacher_does=[
            "Cold-call 1 pair to present Ramón's Part A (direct) + Part B (inverse) using sentence frame.",
            "Class signals thumbs. Write the direct-vs-inverse test on board.",
            "Wed bridge: \"Tomorrow we graph these inverse functions and name the features.\"",
        ],
        students_do=[
            "Presenting pair uses sentence frame for Parts A and B.",
            "Rest of class signals and writes takeaway in margin.",
        ],
        questions_to_ask=[
            "Summarize: how do you test direct vs inverse?",
            "Wednesday preview: what shape do you think y = k/x graphs as?",
        ],
        adult_role="Cold-call a pair that struggled on #26 — hold them accountable to the direct-vs-inverse distinction.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit Ticket",
        framework_tag="summary recap (not graded)",
        dok="1-2",
        minutes=3,
        teacher_does=[
            "Collect at door.",
            "Scan #3 (how to distinguish direct from inverse). Plan P3 Do Now if many students are still fuzzy.",
        ],
        students_do=["3 sentence stems, honest."],
        questions_to_ask=[
            "Did students articulate the test (y/x constant = direct; xy constant = inverse)?",
        ],
        adult_role="Collect. No grade.",
    )
    add_callout(doc, "📋  EXIT TICKET — Student Form",
                [
                    "Today I learned that ___.",
                    "The biggest trap in Ramón's task was ___.",
                    "I distinguish direct from inverse variation by ___.",
                ],
                color_hex="FFF2CC")

    doc.add_page_break()
    add_callout(doc, "🕐  PERIOD A / PERIOD F — 65-MIN CLASS NOTES",
                [
                    "Both sections should have 65 min for P2. Use the extra 10 min on Explore #26 (Ramón) — the richer discussion, the better the DOK-3 payoff.",
                    "If finished early: hand out Practice #25 (SAT/ACT MC) as fast-finisher.",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "🧩  MODIFICATIONS / SUPPORT FOR IEP STUDENTS",
                [
                    "• Provide the Application Rules card as a sticker/cut-out.",
                    "• For #26, provide pre-formatted work space: \"PART A: distance = ___ × time\" etc.",
                    "• Accept verbal sentence-frame completion in lieu of written.",
                ],
                color_hex="E2F0D9")
    add_callout(doc, "🌍  SUPPORTS FOR ELL STUDENTS",
                [
                    "• Sentence frames on student page (don't skip).",
                    "• Word cards: \u201cfrequency\u201d = how often; \u201cvolume\u201d = space inside; \u201cpressure\u201d = push.",
                    "• \u201cPerformance Task\u201d = an assessment task with multiple parts.",
                    "• Pair ELL students with English-dominant partners for TPS.",
                ],
                color_hex="DEEBF7")

    doc.save(path)


# ── Main ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_do_now("L41_P2_Do_Now.docx")
    build_student("L41_P2_Student_Packet.docx")
    build_teacher("L41_P2_Teacher_Packet.docx")
    print("Built L41_P2_Do_Now.docx, L41_P2_Student_Packet.docx, L41_P2_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L41_P2_Visuals_Checklist.md",
                              title="L41 P2 — Visuals Checklist")
