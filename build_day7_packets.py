"""Generate Combined Day 7 materials — Savvas-only, synthesis + performance task day.

Design: Day 7 folds synthesis review + the Venetta Performance Task (Savvas
Practice #30) as the sole DOK-3 driver.  Every student-facing work item
traces to a Savvas bank ID.

Phases (57 min):
  0–5    Do Now A : Savvas Practice #29 (SAT-style: which graph is f(x)?)   [DOK 2]
  5–7    Do Now B : Blooket login                                            [—]
  7–13   Do Now C : Blooket synthesis review (whole-lesson rules)            [DOK 1]
  13–14  Blooket Wrap : candy + lowest rule aloud                            [DOK 1]
  14–22  Launch    : 3-concept rapid review (multiplicity, complex, ineq)    [DOK 2]
  22–44  Explore   : Savvas Practice #30 — Venetta Performance Task          [DOK 3]
  44–48  Share/Sum : cold-call pairs + self-rating                           [DOK 2]
  48–53  Exit      : summary recap fill-in + sentence                        [DOK 1–2]
"""
import io
import sys

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

# Every work item routes through the bank.
DO_NOW_ID   = "3-5-savvas-q29"   # SAT-style graph ID, Day 2-3 bridge
LAUNCH_IDS  = ["3-5-savvas-q17", "3-5-tryit-6a"]  # complex zeros + inequality
DOK3_ID     = "3-5-savvas-q30"   # Venetta performance task — DOK-3 driver

OBJECTIVES = [
    ("Math Objective",
     "Synthesize the key ideas of Lesson 3-5 — zeros, multiplicity, complex "
     "zeros, and polynomial inequalities — and apply them to a multi-part "
     "real-world performance task (Savvas Practice #30, Venetta deli profit)."),
    ("Language Objective",
     "Students will use the frames \u201cThe zero at t = ___ represents ___ "
     "in the context because ___\u201d and \u201cThe profit is negative on the "
     "interval ___ because ___ .\u201d"),
    ("Essential Understanding",
     "A polynomial\u2019s zeros, sign behavior, and graph are three views of the "
     "same structure.  In a modeling context each zero has a real-world "
     "meaning, and intervals of positive/negative values answer questions "
     "about profit, height, or volume."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Students pull together multiplicity, complex zeros, and sign-chart "
     "reasoning to tackle a complete multi-part performance task anchored "
     "in a real-world profit model."),
    ("Essential Question",
     "How do zeros, sign charts, and the factored form work together to "
     "answer questions about when a business is profitable, when it breaks "
     "even, and when it loses money?"),
    ("Materials",
     "Student laptops with Desmos; pencils; Blooket access; "
     "printed Day 7 packet.  No chart paper."),
]


def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if i == 0 and bold_first:
            run.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_two_col(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = TABLE_STYLE
    t.columns[0].width = Inches(1.8)
    t.columns[1].width = Inches(4.7)
    for i, (label, content) in enumerate(rows):
        set_cell(t.rows[i].cells[0], label, bold_first=True)
        if isinstance(content, str):
            content = [content]
        set_cell(t.rows[i].cells[1], content)
    doc.add_paragraph()


def add_callout(doc, title, lines):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.add_run(title).bold = True
    if isinstance(lines, str):
        lines = [lines]
    for line in lines:
        cell.add_paragraph(line)
    doc.add_paragraph()


def header_line(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def blank_lines(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.add_run("______________________________________________________________")


def _format_unicode(prompt: str) -> str:
    prompt = (prompt
              .replace("\u00e2\u2020\u2019", "\u2192")
              .replace("\u00e2\u20ac\u201d", "\u2014")
              .replace("\u00e2\u20ac\u201c", "\u2013")
              .replace("\u00c2\u00b1", "\u00b1")
              .replace("\u00e2\u02c6\u0161", "\u221a"))
    SUPS = {"2": "\u00b2", "3": "\u00b3", "4": "\u2074",
            "5": "\u2075", "6": "\u2076"}
    out, i = [], 0
    while i < len(prompt):
        ch = prompt[i]
        if ch == "^" and i + 1 < len(prompt) and prompt[i+1] in SUPS:
            out.append(SUPS[prompt[i+1]])
            i += 2
            continue
        out.append("\u2212" if ch == "-" else ch)
        i += 1
    return "".join(out)


def _strip_stem(prompt: str) -> str:
    txt = _format_unicode(prompt)
    if ":" in txt:
        txt = txt.split(":", 1)[-1].strip()
    return txt.rstrip(".")


# ----------------------------------------------------------------------
# DO NOW SHEET — Savvas Practice #29 (SAT-style graph ID)
# ----------------------------------------------------------------------

def build_do_now(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)

    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    header_line(doc, "Day 7  |  DO NOW  \u2014  SAT-Style Graph Identification", size=13)

    p = doc.add_paragraph()
    r = p.add_run("Name: _____________________________________     Date: _____________")
    r.font.size = Pt(11)
    doc.add_paragraph()

    add_callout(doc, "\u270d\ufe0f  5 minutes. Silent. Pencil only. No graphing calculator.", [
        "This is a SAT/ACT-style problem.  Use what you know about zeros, "
        "factored form, and end behavior to identify the correct graph "
        "WITHOUT a calculator.",
        "Think about: zeros \u2192 behavior \u2192 sketch.",
    ])

    _do_now_q = qb.get_for_packet([DO_NOW_ID])[0]
    header_line(doc, f"Savvas Practice #29  [bank: {DO_NOW_ID}]:", size=12)
    p = doc.add_paragraph()
    p.add_run(_format_unicode(_do_now_q["prompt"])).italic = True
    doc.add_paragraph()

    _answers = _do_now_q.get("answers", ["Graph A", "Graph B", "Graph C", "Graph D"])
    p = doc.add_paragraph()
    p.add_run("Circle your answer:     " + "          ".join(_answers))
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("In one sentence, explain how you know:").bold = True
    blank_lines(doc, 1)
    blank_lines(doc, 1)

    add_callout(doc, "Check your reasoning (use these steps):", [
        "1.  Can you factor f(x)?  Factor out the GCF, then look for more factors.",
        "2.  What are the zeros?  Do they match the x-intercepts in the graph?",
        "3.  Check end behavior: odd degree, positive leading coefficient \u2192 "
        "down-left, up-right.",
    ])

    doc.save(path)


# ----------------------------------------------------------------------
# STUDENT PACKET
# ----------------------------------------------------------------------

_venetta_q = qb.get_for_packet([DOK3_ID])[0]
_venetta_prompt = _format_unicode(_venetta_q["prompt"])

VENETTA_STUDENT = [
    f"THE PROBLEM  \u2014  Savvas Practice #30 (Performance Task)  [bank: {DOK3_ID}]",
    "",
    _venetta_prompt,
    "",
    "THE RULES YOU NEED (everything is on this page)",
    "\u2022  RULE 1 (Break-even):   Break-even means P(t) = 0.  "
    "Find the zeros of P(t).",
    "\u2022  RULE 2 (Sign chart):   A sign chart tells you where P(t) > 0 "
    "(profitable) and where P(t) < 0 (losing money).",
    "     \u2022  Pick a test value in each interval.",
    "     \u2022  Plug into the factored form.  Negative \u00d7 negative \u00d7 positive "
    "= what sign?",
    "\u2022  RULE 3 (Context):     t = 0 means the year 2000.  "
    "Only t \u2265 0 is meaningful (franchises can\u2019t open before 2000).",
    "",
    "PART A \u2014 Sketch a graph of P(t).",
    "  Factor P(t):  P(t) = t(\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f\u202f) = ______________________",
    "  Zeros:   t = _______,   t = _______,   t = _______",
    "  Sketch (label the zeros and axes; show end behavior):",
    "",
    "  [                                                             ]",
    "  [                                                             ]",
    "  [                                                             ]",
    "  [                                                             ]",
    "  [                                                             ]",
    "",
    "PART B \u2014 During which year(s) did Venetta\u2019s franchises break even?",
    "  Break-even means P(t) = 0.  Which zeros are in the range t \u2265 0?",
    "  Break-even at t = _______ (year _______) and t = _______ (year _______)",
    "",
    "PART C \u2014 Interpret years of negative profit.",
    "  Use your sign chart.  Between which values of t is P(t) < 0?",
    "  t is between _______ and _______  \u2192  years _______ to _______",
    "  What does this mean in context?",
    "  \u201cThe franchises lost money during ________________________________.\u201d",
    "",
    "PART D \u2014 During which year range was the franchise profitable?",
    "  P(t) > 0 when t > _______  (for t \u2265 0)",
    "  Profitable years: _______ and beyond.",
    "  Justify using your sign chart:",
    "  ___________________________________________________________________",
    "  ___________________________________________________________________",
    "",
    "Bank ID: 3-5-savvas-q30 (Savvas Practice #30 \u2014 Performance Task, DOK 3).",
]

SHARE_SUMMARY_STUDENT = [
    "Self-check  \u2014  circle one for each:",
    "1.  I can factor a cubic polynomial and find its zeros.             \u2713    partly    not yet",
    "2.  I can use a sign chart to find where P(t) is positive/negative. \u2713    partly    not yet",
    "3.  I can interpret zeros and intervals in a real-world context.    \u2713    partly    not yet",
]


def _exit_lines() -> list[str]:
    return [
        "Fill in each blank:",
        "",
        "1.  A zero of a polynomial function means the graph __________________ the x-axis",
        "    (if multiplicity is odd) or __________________ to the x-axis (if even).",
        "",
        "2.  Non-real complex zeros always come in __________________________ pairs.",
        "",
        "3.  A sign chart tells you where a polynomial is __________________ (> 0)",
        "    or __________________ (< 0).",
        "",
        "4.  In the Venetta model, the franchises were profitable for years _________ ",
        "    and beyond because ___________________________________________________.",
        "",
        "Biggest thing I learned today across the whole lesson:",
        "_______________________________________________________________________",
        "_______________________________________________________________________",
    ]


def build_student(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    ps.running_header_name_block(doc)
    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    ps.day_banner(doc, "7", "Synthesis + Venetta Performance Task",
                  "Pull together zeros, multiplicity, complex zeros, and "
                  "inequalities.  DOK-3 driver = Savvas Practice #30 "
                  "(Venetta deli profit).  All items Savvas-anchored.")

    add_two_col(doc, OBJECTIVES)
    add_two_col(doc, FRAMEWORK_HEADER)

    add_callout(doc, "\U0001F4C4  DO NOW  \u2014  separate sheet", [
        "You got the SAT-style graph sheet on entry (Savvas Practice #29).",
        "Hold onto it \u2014 we\u2019ll return to it during Launch.",
    ])

    add_callout(doc, "\U0001F3AE  Blooket  \u2014  Synthesis Rule Recall   [Do Now B+C]  [DOK 1]   (2 + 6 min)", [
        "Whole-lesson synthesis round.  Rules: zero \u21d4 factor; multiplicity "
        "\u2192 cross or tangent; complex conjugate pairs; sign chart / inequality.  "
        "No procedural computation.",
    ])

    _launch_q2 = qb.get_for_packet(["3-5-savvas-q17"])[0]
    _launch_q3 = qb.get_for_packet(["3-5-tryit-6a"])[0]
    add_two_col(doc, [(
        "\U0001F4A1  Launch  \u2014  3-Concept Rapid Review   [Launch]  [DOK 2]   (8 min)",
        [
            "Teacher will walk through 3 quick verbal checks.  Call out answers.",
            "",
            "1.  MULTIPLICITY CHECK (150s):",
            "    [Verbal rapid-review \u2014 teacher will project the example. Listen and call out the answer.]",
            "",
            "2.  COMPLEX-PAIR CHECK (150s)  [bank: 3-5-savvas-q17]:",
            "    " + _format_unicode(_launch_q2["prompt"]),
            "    Answer aloud: _____________________________________________",
            "",
            "3.  INEQUALITY CHECK (150s)  [bank: 3-5-tryit-6a]:",
            "    " + _format_unicode(_launch_q3["prompt"]),
            "    Answer aloud: _____________________________________________",
        ],
    )])

    add_two_col(doc, [(
        "\U0001F969  Explore  \u2014  Venetta Performance Task (Savvas #30)   [Explore]  [DOK 3]   (22 min)",
        VENETTA_STUDENT,
    )])

    add_callout(doc,
        "\U0001F501  SHARE / SUMMARY   [Share/Summary]  [DOK 2]   (4 min)",
        SHARE_SUMMARY_STUDENT)

    ps.summary_exit_box(doc,
        "\U0001F4E4  EXIT TICKET   [Exit Ticket]  [DOK 1\u20132]   (5 min)",
        _exit_lines())

    doc.save(path)


# ----------------------------------------------------------------------
# TEACHER PACKET
# ----------------------------------------------------------------------

CRITERIA_FOR_SUCCESS = [
    "Student-facing criteria (revisited during Share/Summary):",
    "1.  I can factor a cubic polynomial and find its zeros.",
    "2.  I can use a sign chart to find where P(t) is positive/negative.",
    "3.  I can interpret zeros and intervals in a real-world context.",
    "Formative evidence:",
    "\u2022  Do Now sheet (Savvas #29 SAT-style identification + reasoning sentence).",
    "\u2022  Blooket dashboard (synthesis rule recall).",
    "\u2022  Launch rapid-review answers (verbal cold-calls).",
    "\u2022  Venetta packet (DOK-3 evidence \u2014 photograph or collect).",
    "\u2022  Exit ticket (fill-in recap + \u201cbiggest thing\u201d sentence).",
]

BLOOKET_SCOPE = [
    "SCOPE (pure rule-recall, DOK 1).  Synthesis of whole lesson:",
    "\u2022  Zero \u21d4 factor theorem (zero at c \u2194 (x \u2212 c) is a factor).",
    "\u2022  Odd multiplicity \u2192 graph crosses the x-axis.",
    "\u2022  Even multiplicity \u2192 graph is tangent to the x-axis (turning point).",
    "\u2022  Non-real complex zeros come in conjugate pairs (a + bi with a \u2212 bi).",
    "\u2022  Sign chart: find zeros, pick test values, determine sign in each interval.",
    "\u2022  Polynomial inequality: solution = union of intervals where sign is correct.",
    "NO procedural factoring problems.  Dashboard diagnostic: note the two lowest.",
    "Mode: Fishing Frenzy (default) or Racing (if repeated).",
    "Blooket Wrap at 1 min: candy pass, class repeats lowest-scored rule aloud.",
    "Questions to ask (Wrap): \u201cWhich rule had the lowest %?  Everyone say it now.\u201d",
    "Adult role: teacher watches dashboard; runs Wrap transition.",
]

DO_NOW_KEY = [
    "Savvas Practice #29  [bank: 3-5-savvas-q29, DOK 1]",
    "",
    "f(x) = x\u00b3 + x\u00b2 \u2212 4x.  Factor: f(x) = x(x\u00b2 + x \u2212 4).",
    "The quadratic x\u00b2 + x \u2212 4 has discriminant 1 + 16 = 17 > 0, "
    "so two irrational real zeros. Total: three real zeros.",
    "",
    "\u2705  Correct answer: Graph C.",
    "   (The correct graph shows a cubic, positive leading coefficient, "
    "three x-intercepts: one at x = 0 and two symmetric near \u00b12.)",
    "",
    "DIAGNOSTIC scan on entry:",
    "\u2022  Students who chose A or B likely missed the zero at x = 0 "
    "(forgot to factor out x).",
    "\u2022  Students who chose D likely got the end behavior reversed.",
    "\u2022  Use cold-call during Launch to surface the factoring step.",
    "",
    "Questions to ask (circulating silently): none aloud.",
    "Adult role: teacher hands sheets at door, walks one silent lap.",
]

LAUNCH_KEY = [
    "Three rapid-review checks (8 min, DOK 2).  Teacher projects; class calls out.",
    "",
    "1.  MULTIPLICITY CHECK (150s)  [bank: implicit, multiplicity pattern]",
    "    f(x) = (x \u2212 1)\u00b2(x + 3).",
    "    \u2022  At x = 1: multiplicity 2 (EVEN) \u2192 graph is TANGENT to x-axis (turning point).",
    "    \u2022  At x = \u22123: multiplicity 1 (ODD) \u2192 graph CROSSES the x-axis.",
    "    Teacher: \u201cEven multiplicity \u2014 tangent.  Odd multiplicity \u2014 crosses.\u201d",
    "    Correct pronunciation: \u201ctouches and turns back\u201d or \u201ctangent\u201d for even.",
    "",
    "2.  COMPLEX-PAIR CHECK (150s)  [bank: 3-5-savvas-q17]",
    "    f(x) = x\u00b3 \u2212 x\u00b2 \u2212 4x \u2212 6.  Real zero at x = 3.",
    "    Synthetic division by (x \u2212 3) \u2192 quotient x\u00b2 + 2x + 2.",
    "    Quadratic formula: x = (\u22122 \u00b1 \u221a(4 \u2212 8))/2 = \u22121 \u00b1 i.",
    "    \u2705  All zeros: x = 3, x = \u22121 + i, x = \u22121 \u2212 i.",
    "    Teacher: \u201cConjugate pairs \u2014 if a + bi is a zero, a \u2212 bi must be too.\u201d",
    "",
    "3.  INEQUALITY CHECK (150s)  [bank: 3-5-tryit-6a]",
    "    2x\u00b3 + 12x\u00b2 + 12x < 0.",
    "    Factor: 2x(x\u00b2 + 6x + 6) < 0.",
    "    Zeros: x = 0, x = \u22123 \u00b1 \u221a3 (approx \u22124.73 and \u22121.27).",
    "    Sign chart: negative on (\u2212\u221e, \u22123\u2212\u221a3) and (\u22123+\u221a3, 0).",
    "    \u2705  Solution: x < \u22123 \u2212 \u221a3  OR  \u22123 + \u221a3 < x < 0.",
    "",
    "Questions to ask (rapid fire):",
    "\u2022  \u201cEven or odd multiplicity?\u201d",
    "\u2022  \u201cConjugate of \u22121 + i?\u201d",
    "\u2022  \u201cHow do you know which interval is negative on the sign chart?\u201d",
    "Adult role: teacher leads rapid-fire, cold-calls 2\u20133 students per check.",
]

VENETTA_KEY = [
    "[bank: 3-5-savvas-q30, DOK 3] \u2014 SAVVAS-DECLARED DOK-3 ANCHOR (Performance Task)",
    "",
    "P(t) = t\u00b3 + t\u00b2 \u2212 6t.",
    "",
    "\u2705  FULL WORKING:",
    "Factor:  P(t) = t(t\u00b2 + t \u2212 6) = t(t + 3)(t \u2212 2).",
    "Zeros:   t = 0, t = \u22123, t = 2.",
    "",
    "PART A \u2014 Graph:",
    "  Cubic, positive leading coefficient (end: down-left, up-right).",
    "  Crosses x-axis at t = \u22123 (mult 1), t = 0 (mult 1), t = 2 (mult 1).",
    "  For t \u2265 0: starts at origin, dips below in (0, 2), crosses at t = 2, then climbs.",
    "",
    "PART B \u2014 Break-even (t \u2265 0 context):",
    "  Meaningful zeros: t = 0 (year 2000) and t = 2 (year 2002).",
    "  t = \u22123 is discarded (before franchise opened).",
    "  \u2705  Break-even at year 2000 and year 2002.",
    "",
    "PART C \u2014 Negative profit interval:",
    "  Sign chart for t \u2265 0: test t = 1 \u2192 P(1) = 1 + 1 \u2212 6 = \u22124 < 0.",
    "  P(t) < 0 on (0, 2).",
    "  \u2705  Franchises lost money during the first two years (2000\u20132002).",
    "",
    "PART D \u2014 Profitable years:",
    "  Test t = 3 \u2192 P(3) = 27 + 9 \u2212 18 = 18 > 0.",
    "  P(t) > 0 for t > 2.",
    "  \u2705  Profitable from year 2002 onward (t > 2).",
    "",
    "WHY THIS IS DOK 3:",
    "\u2022  Students produce a graph (Part A) and read it in context.",
    "\u2022  They discard the negative zero as contextually meaningless (Part B) \u2014 "
    "not a memorized step.",
    "\u2022  They build and read a sign chart to answer Part C and Part D.",
    "\u2022  The whole task requires coordinating four sub-problems across multiple "
    "representations (algebraic, graphical, contextual).",
    "",
    "Stuck-pair hint card:",
    "\u201cFactor first.  Then list all three zeros.  Which zeros have t \u2265 0?  "
    "Now plug in t = 1 \u2014 is P positive or negative?\u201d",
    "",
    "Questions to ask (prompt-only, never deliver):",
    "\u2022  \u201cWhat does P(t) = 0 mean in the context of the franchise?\u201d",
    "\u2022  \u201cCan the franchise lose money before it opens?  Which zeros are meaningful?\u201d",
    "\u2022  \u201cWhat sign does P(t) have between t = 0 and t = 2?  How do you know?\u201d",
    "\u2022  \u201cWhat does \u2018profitable\u2019 mean in terms of the sign of P?\u201d",
    "Adult role: teacher circulates with hint cards, prompts only; single teacher walking ~5 pairs.",
    "Release valve: if time tight, skip Part D written justification; "
    "preserve Parts A\u2013C in-class.",
]

EXIT_KEY = [
    "Fill-in answers:",
    "1.  \u201ccrosses\u201d / \u201ctangent\u201d",
    "2.  \u201cconjugate\u201d",
    "3.  \u201cpositive\u201d / \u201cnegative\u201d",
    "4.  t > 2  (year 2002 onward);  because P(t) > 0 for t > 2 "
    "(sign chart confirms positive interval after the last zero).",
    "",
    "Scan the \u2018biggest thing\u2019 sentence tonight.  Students who wrote only "
    "\u2018we did the deli problem\u2019 (no vocabulary) \u2192 target for Day 8 warm-up.",
    "Questions to ask: none during write time.  At door: \u201cBiggest thing \u2014 "
    "one word. Go.\u201d",
    "Adult role: teacher collects at door; scans criterion-3 self-ratings "
    "for Day 8 station assignments.",
]

IEP_MODS = [
    "\u2022  Do Now is multiple-choice + one-sentence (low writing load).",
    "\u2022  Launch rapid-review is verbal \u2014 no writing required.",
    "\u2022  Venetta task has all rules printed on the page; student can "
    "self-prompt without waiting for teacher.",
    "\u2022  Factored form is scaffolded with blanks, not a blank page.",
    "\u2022  Sign chart template is provided in the student packet.",
    "\u2022  Stuck-pair hint card preserves DOK 3 via the justification demand.",
    "\u2022  Release valve: Part D justification \u2192 homework if time tight.",
]

ELL_SUPPORTS = [
    "\u2022  Three new application terms today: BREAK-EVEN, PROFITABLE, CONTEXTUAL.",
    "\u2022  Sentence frames on the page: \u201cThe zero at t = ___ represents ___ "
    "in the context because ___\u201d and \u201cThe profit is negative on the "
    "interval ___ because ___ .\u201d",
    "\u2022  t = 0 \u2192 year 2000 translation box printed on the student packet.",
    "\u2022  Predict-before-verify: students commit to zeros and sign before "
    "checking with Desmos.",
]


def build_teacher(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    header_line(doc, "TEACHER EDITION", size=12)
    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    ps.day_banner(doc, "7", "Synthesis + Venetta Performance Task",
                  "Teacher pacing + DOK framework crosswalk + answer keys + "
                  "Questions to ask / Adult role per phase.")

    add_two_col(doc, OBJECTIVES + [("CCSS",
        "F-IF.C.7c, A-APR.B.3, N-CN.C.8, A-REI.D, MP.2, MP.4")])
    add_two_col(doc, FRAMEWORK_HEADER)

    add_callout(doc, "\U0001F4CB  DESIGN \u2014  SAVVAS-ONLY SYNTHESIS + DOK-3 DAY", [
        "Do Now \u2192 Blooket (synthesis) \u2192 Blooket Wrap \u2192 "
        "Launch (rapid review) \u2192 ONE DOK-3 Explore (Venetta, Savvas-declared) "
        "\u2192 Share/Summary \u2192 Exit.",
        "[DOK 1] Blooket synthesis recall; Blooket Wrap.",
        "[DOK 2] Do Now #29; Launch rapid-review; Share/Summary.",
        "[DOK 3] Savvas Practice #30 (Venetta) \u2014 the ONLY DOK-3 driver. "
        "Do NOT add a second.",
        "Every work item traces to a Savvas bank ID.",
    ])

    add_callout(doc, "\u2705  CRITERIA FOR SUCCESS  /  FORMATIVE ASSESSMENT", CRITERIA_FOR_SUCCESS)

    add_callout(doc, "\u23f1\ufe0f  REALISTIC PACING", [
        "Total: 57 min.  Fits 65-min F with slack; tight for 55-min A.",
        "  Do Now A (Savvas #29) ......... 5 min",
        "  Do Now B (Blooket login) ....... 2 min",
        "  Do Now C (Blooket game) ........ 6 min",
        "  Blooket Wrap ................... 1 min",
        "  Launch (3-concept review) ...... 8 min",
        "  Explore (Venetta DOK-3) ....... 26 min  (includes 4-min circulation buffer)",
        "  Share/Summary .................. 4 min",
        "  Exit ........................... 5 min",
        "RELEASE VALVES (ordered):",
        "  1. Trim Launch to 2 checks instead of 3 (save 2 min).",
        "  2. Hint-card a stuck Venetta pair: reveal factored form, "
        "require them to justify the sign chart.",
        "  3. Venetta Part D justification \u2192 homework; preserve A\u2013C in-class.",
        "  NEVER cut Exit.  Only artifact that documents learning.",
        "55-min A tight: drop Launch inequality check (use only multiplicity + complex pair).",
    ])

    ps.framework_phase_header(doc,
        phase="Do Now  \u2014  Savvas Practice #29 (SAT-style graph ID)",
        framework_tag="Do Now A",
        dok="2",
        minutes=5,
        teacher_does=[
            "Hand sheet at the door.  Walk one silent lap.  "
            "Scan what they circled: A or B \u2192 missed zero at 0; "
            "D \u2192 end behavior reversed.",
        ],
        students_do=[
            "Identify the correct graph of f(x) = x\u00b3 + x\u00b2 \u2212 4x without a calculator.",
            "Write one sentence explaining their reasoning.",
        ],
        questions_to_ask=[
            "None aloud during silent work.",
        ],
        adult_role="Teacher hands sheets at door; circulates silently for 5 min.")
    add_callout(doc, "\u2705  DO NOW / ANSWER KEY", DO_NOW_KEY)

    ps.framework_phase_header(doc,
        phase="Blooket  \u2014  Synthesis Rule Recall",
        framework_tag="Do Now B+C",
        dok="1",
        minutes=8,
        teacher_does=[
            "Display Blooket code.  Confirm synthesis CSV is loaded.  "
            "Run one 6-min game.  Watch dashboard for two lowest rules.",
            "Blooket Wrap (1 min): candy pass, class repeats lowest rule aloud.",
        ],
        students_do=[
            "Log in (2 min).  Play whole-lesson synthesis game (6 min).",
        ],
        questions_to_ask=[
            "\u201cWhich rule had the lowest %?  Everyone say it now.\u201d",
        ],
        adult_role="Teacher watches dashboard, runs Wrap transition.")
    add_callout(doc, "[Do Now B+C]  Blooket scope", BLOOKET_SCOPE)

    ps.framework_phase_header(doc,
        phase="Launch  \u2014  3-Concept Rapid Review",
        framework_tag="Launch",
        dok="2",
        minutes=8,
        teacher_does=[
            "Project three quick-check examples one at a time.  Cold-call "
            "2\u20133 students per check.  Deliver correct Savvas vocabulary aloud.",
        ],
        students_do=[
            "Call out answers (verbal; no new writing needed).",
            "Bridge Do Now graph reasoning to the Venetta task.",
        ],
        questions_to_ask=[
            "\u201cEven or odd multiplicity?\u201d",
            "\u201cConjugate of \u22121 + i?\u201d",
            "\u201cHow do you know which interval is negative on the sign chart?\u201d",
        ],
        adult_role="Teacher leads rapid-fire; solo teacher, cold-calls low-stakes.")
    add_callout(doc, "\u2705  LAUNCH / ANSWER KEY", LAUNCH_KEY)

    ps.framework_phase_header(doc,
        phase="Explore  \u2014  Venetta Performance Task (Savvas Practice #30)",
        framework_tag="Explore",
        dok="3",
        minutes=22,
        teacher_does=[
            "Circulate with hint cards.  Prompt only \u2014 never deliver answers.",
            "Four laps for ~10 students; 2\u20133 min per lap.",
            "Reveal hint-card answer only when a pair is genuinely stuck "
            "after 8+ min; require written justification.",
        ],
        students_do=[
            "Work Parts A\u2013D on the Venetta task.  All rules are on the page.",
            "Factor, find zeros, build sign chart, interpret context.",
        ],
        questions_to_ask=[
            "\u201cWhat does P(t) = 0 mean in the context of the franchise?\u201d",
            "\u201cWhich zeros have t \u2265 0?  Why does that matter?\u201d",
            "\u201cWhat sign does P have between t = 0 and t = 2?\u201d",
            "\u201cWhat does \u2018profitable\u2019 mean in terms of the sign of P?\u201d",
        ],
        adult_role="Solo teacher circulates with prompts.")
    add_callout(doc, "\U0001F969  VENETTA / ANSWER KEY + WHY DOK 3", VENETTA_KEY)

    ps.framework_phase_header(doc,
        phase="Share / Summary  \u2014  Synthesis + Self-Rating",
        framework_tag="Share/Summary",
        dok="2",
        minutes=4,
        teacher_does=[
            "Cold-call Pair 1 to describe the sign chart.",
            "Cold-call Pair 2 to justify the profitable interval.",
            "Self-rating circle on packet.  Teacher scans.",
            "Preview Day 8 (30s): \u201cTomorrow \u2014 quiz + differentiated stations.\u201d",
        ],
        students_do=[
            "Listen to cold-calls; rate themselves on three criteria.",
        ],
        questions_to_ask=[
            "\u201cWho can describe the sign chart for P(t)?\u201d",
            "\u201cWho can justify why P is positive for t > 2?\u201d",
            "\u201cWhich criterion did you rate \u2018partly\u2019?\u201d",
        ],
        adult_role="Solo teacher leads; cold-call phrasing is low-stakes.")
    add_callout(doc, "\U0001F501  SHARE / SUMMARY (student-facing)", SHARE_SUMMARY_STUDENT)

    ps.framework_phase_header(doc,
        phase="Exit Ticket  \u2014  Summary Recap",
        framework_tag="Exit Ticket",
        dok="1\u20132",
        minutes=5,
        teacher_does=[
            "Silent write time.  Collect at door.",
            "Scan \u2018biggest thing\u2019 sentences that evening for Day 8 diagnostics.",
        ],
        students_do=[
            "Complete fill-in recap.  Write one \u2018biggest thing\u2019 sentence.",
        ],
        questions_to_ask=[
            "None during write time.  At door: \u201cBiggest thing \u2014 one word. Go.\u201d",
        ],
        adult_role="Teacher collects at door; scans self-ratings for Day 8 station assignments.")
    ps.summary_exit_box(doc,
        "\U0001F4E4  EXIT TICKET (student-facing)",
        _exit_lines())
    add_callout(doc, "\u2705  EXIT TICKET / ANSWER KEY", EXIT_KEY)

    add_callout(doc, "\U0001F4F7  WHAT TO COLLECT", [
        "1.  Do Now sheets (Practice #29 + reasoning sentence).",
        "2.  Venetta packet (DOK-3 evidence \u2014 photograph or collect).",
        "3.  Exit tickets (fill-in recap + \u2018biggest thing\u2019 sentence).",
    ])
    add_callout(doc, "\U0001F9E9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS", IEP_MODS)
    add_callout(doc, "\U0001F30D  SUPPORTS FOR ELL STUDENTS", ELL_SUPPORTS)

    doc.save(path)


if __name__ == "__main__":
    build_do_now("Day_7_Do_Now.docx")
    build_student("Day_7_Student_Packet.docx")
    build_teacher("Day_7_Teacher_Packet.docx")
    print("Built Day_7_Do_Now.docx, Day_7_Student_Packet.docx, Day_7_Teacher_Packet.docx")
