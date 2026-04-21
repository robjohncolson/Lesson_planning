"""Generate Lesson 5-4 Period 2 materials — Klimsara-adapted, no Blooket.

Design: Radical equations with rational exponents + the lesson's single-DOK-3 spine.
P2 is the conceptual core: students solve equations with rational exponents and tackle
the Half-Life Model task (#41) which asks them to apply y = a·(0.5)^(t/h) to find the
amount of soft drink remaining after 16 hours — the pedagogical anchor of Topic 5.

Phases (55 min base; Period F 65-min):
  0–  5   Do Now    : Practice #21 — bridge: solve a radical equation from P1 [DOK 1]
  5– 17   Launch    : teacher models Example 4 (rational exponents) [DOK 2]
 17– 50   Explore   : Try-It 4 + Practice #12/#33/#34 + ⭐ DOK-3 q41 (Half-Life Model) [DOK 2-3]
 50– 55   Share/Sum : DOK-3 pair presents half-life computation + bridge to P3 [DOK 2]
 55– 58   Exit      : summary recap fill-in                              [DOK 1-2]
 (back)   Optional: Practice #19 (extraneous-solution argument) — conceptual stretch
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_ID   = "5-4-savvas-q21"
LAUNCH_IDS  = ["5-4-savvas-example-4-lesson-5-4"]
EXPLORE_IDS = [
    "5-4-savvas-try-it-4-lesson-5-4",
    "5-4-savvas-q12",
    "5-4-savvas-q33",
    "5-4-savvas-q34",
    "5-4-savvas-q41",        # ⭐ DOK-3 spine — Half-Life Model
]
DOK3_ID       = "5-4-savvas-q41"
REINFORCE_IDS = ["5-4-savvas-q19"]   # extraneous-solution argument stretch

_ALL_IDS = [DO_NOW_ID] + LAUNCH_IDS + EXPLORE_IDS + REINFORCE_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 5-4 · Period 2"
LESSON_TITLE = "Radical Equations With Rational Exponents + DOK-3 Modeling"

OBJECTIVES = [
    ("Math Objective",
     "Solve equations with rational exponents by converting to radical form and raising "
     "both sides to the reciprocal power; apply rational-exponent models to real-world "
     "half-life and decay problems."),
    ("Language Objective",
     "Translate a real-world half-life scenario into an equation with a rational exponent, "
     "compute, and interpret the result in context."),
    ("Essential Understanding",
     "Rational exponents are the bridge between algebraic power operations and real-world "
     "exponential/radical modeling — the half-life formula y = a·(1/2)^(t/h) uses "
     "fractional powers directly."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Solve rational-exponent equations; model real-world decay with fractional-power formulas."),
    ("Essential Question",
     "How do rational-exponent equations model real-world decay, and when do extraneous "
     "solutions arise?"),
    ("Materials",
     "Student packet, pencil, calculator for Q#41. Teacher models Ex 4; DOK-3 Practice #41 "
     "is the topic's capstone modeling task."),
]


# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.0):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    prompt_p.add_run(q["prompt"]).font.size = Pt(11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            line.add_run(f"  ({chr(65+i)})  {a}").font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, for_teacher=False):
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 2, LESSON_TITLE,
                  subtitle=LESSON_LABEL + (" · Teacher Edition" if for_teacher else ""))
    t = doc.add_table(rows=len(OBJECTIVES), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(OBJECTIVES):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    t = doc.add_table(rows=len(FRAMEWORK_HEADER), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(FRAMEWORK_HEADER):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()


# ── Do Now ───────────────────────────────────────────────────────────────

def build_do_now(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 2, "Do Now — " + LESSON_TITLE, subtitle=LESSON_LABEL)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from P1",
        dok="1",
        minutes=5,
        teacher_does=["Hand out at door.", "Circulate 3 min.",
                      "Cold-call 1 student: \"Isolate the radical, then square both sides. What do you get?\""],
        students_do=["Solve the radical equation by isolating the radical and raising both sides to the appropriate power.",
                     "Check for extraneous solutions."],
        questions_to_ask=[
            "What operation undoes a cube root?",
            "Why do we check our solution in the original equation?",
        ],
        adult_role="Solo teacher: circulate, time 5 min, pull sheets.",
    )

    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label="Do Now — Solve this radical equation",
                    space_after_inches=2.0)

    add_callout(doc, "\U0001f3af  SENTENCE FRAME",
                ["“I isolate the radical, giving ___. I raise both sides to the power ___, getting ___. I check by substituting back: the left side equals ___ and the right side equals ___. The solution is ___.”"],
                color_hex="FFF2CC")

    doc.save(path)


# ── Student packet ───────────────────────────────────────────────────────

def build_student(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=False)

    ps.teal_section(doc, "LAUNCH — Rational-Exponent Equations (teacher models Example 4)", [])
    add_callout(doc, "\U0001f4d8  RATIONAL-EXPONENT RULES (keep printed on your page)",
                [
                    "RATIONAL-EXPONENT RULES",
                    "1. Rewrite x^(m/n) as (ⁿ√x)^m — bridge to radicals you already know.",
                    "2. To solve (expr)^(m/n) = k, raise both sides to the reciprocal power n/m.",
                    "3. Simplify, then CHECK — even rational-exponent equations can have extraneous solutions when the denominator of the exponent is EVEN.",
                    "4. For real-world decay y = a·(0.5)^(t/h): half-life h tells you the exponent's denominator; plug in t, evaluate.",
                ],
                color_hex="D9EAD3")

    add_callout(doc, "⚠️  EXTRANEOUS SOLUTIONS ARISE WHEN THE DENOMINATOR IS EVEN",
                [
                    "(expr)^(m/n) = k → raise both sides to n/m",
                    "If n is EVEN, squaring can introduce solutions that don't satisfy the original equation.",
                    "ALWAYS substitute back into the ORIGINAL equation to verify.",
                ],
                color_hex="F4CCCC")
    doc.add_paragraph()

    ps.teal_section(doc, "EXPLORE — Think-Pair-Share (33 min)",
                    ["Work with a partner. Move in order. Practice #41 (Half-Life Model) is the big one — plan ~12 min for it."])

    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 4 — Rational-exponent equation",
        "Practice #12 — Solve with rational exponent",
        "Practice #33 — Rational-exponent equation",
        "Practice #34 — Rational-exponent equation",
        "Practice #41  ⭐ DOK-3 HALF-LIFE MODEL",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=label, space_after_inches=1.5)

    ps.teal_section(doc, "SHARE / SUMMARY (5 min)", [])
    ps.sentence_frame_box(doc, [
        "Pair share (Half-Life Model): “This equation has exponent ___/___, so I raise both sides to the power ___/___. After simplifying I get ___. In the half-life formula, a = ___ (start), h = ___ (half-life period), t = ___ (elapsed time), giving y ≈ ___.”",
        "Whole class: one pair shares their computation and interpretation. Class signals thumbs up/sideways.",
    ])

    add_callout(doc, "\U0001f52d  BRIDGE TO P3",
                ["Next period we apply rational-exponent skills to assessment-critical items. Bring your check-for-extraneous habit."],
                color_hex="FCE5CD")

    ps.summary_exit_box(doc, "EXIT TICKET — Summary Recap",
                        [
                            "To solve (expr)^(m/n) = k I raise both sides to power ___.",
                            "Half-life tells me the ___ of the exponent.",
                            "After 16 hours, 50 mL of soft drink decays to ___ mL.",
                            "One biggest thing I learned today: ________________________",
                        ])

    doc.add_page_break()
    ps.teal_section(doc, "OPTIONAL REINFORCEMENT (not collected)",
                    ["If you finish Explore early. #19 is a conceptual extraneous-solution argument that extends the half-life skill."])
    reinforce = qb.get_for_packet(REINFORCE_IDS)
    for i, q in enumerate(reinforce, 1):
        bank_item_block(doc, q, label=f"Optional #{i}  (Savvas Practice, extraneous-solution argument)",
                        space_after_inches=1.8)

    doc.save(path)


# ── Teacher packet ───────────────────────────────────────────────────────

def build_teacher(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=True)

    add_callout(doc, "\U0001f4cc  PRE-FLIGHT — P2 IS THE DOK-3 SPINE DAY",
                [
                    "Today's conceptual core is Practice #41 (Half-Life Model). Students must (a) identify a=50, h=5, t=16, (b) compute exponent 16/5 = 3.2, (c) evaluate 50·(0.5)^3.2 ≈ 5.44 mL, (d) interpret: 'about 5.44 mL of the drink remains after 16 hours.' This is the single most important item in Topic 5 — all three assessment items chain off this skill.",
                    "Pace: Try-It 4 (4 min) → #12 (4 min) → #33 (4 min) → #34 (4 min) → #41 (12 min). ~35 min total Explore. (Adjust: Try-It 4 + #12 = 8 min block, #33 + #34 = 8 min block.)",
                    "If running tight: cut #33 or #34. Never cut #41.",
                ],
                color_hex="FFD9E0")

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from P1",
        dok="1",
        minutes=5,
        teacher_does=["Hand out at door.", "Circulate 3 min.",
                      "Cold-call 1 student for radical isolation and solution check."],
        students_do=["Solve by isolating the radical and raising to the reciprocal power.",
                     "Check for extraneous solutions."],
        questions_to_ask=[
            "What operation undoes a cube root?",
            "Did you verify by substituting back into the original?",
        ],
        adult_role="Solo teacher: circulate silently. No hints.",
    )
    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label=f"Do Now ({q['id']})", space_after_inches=0.3)
    add_callout(doc, "✅  ANSWER KEY",
                [q.get("notes") or "(verify)"],
                color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Launch",
        framework_tag="teacher models Example 4",
        dok="2",
        minutes=12,
        teacher_does=[
            "Project Example 4. Think aloud: rewrite the rational exponent as a radical, then raise both sides to the reciprocal power.",
            "\"The exponent m/n means I raise to the n/m to undo it — the reciprocal flips the fraction.\"",
            "Flag the Rational-Exponent Rules card — print it in student minds.",
            "30-sec pause: \"What is the reciprocal of the exponent? Can an extraneous solution appear?\" (partner check)",
            "Do NOT solve Try-Its or Practice items — release at minute 17.",
        ],
        students_do=[
            "Watch Example 4 silently.",
            "During pause: state the reciprocal exponent to partner.",
            "Copy Rules card into margin.",
            "Note the half-life formula — this comes back in #41.",
        ],
        questions_to_ask=[
            "What must I do BEFORE raising both sides to the reciprocal power?",
            "Why do I need to check my solution even after correct algebra?",
            "In y = a·(0.5)^(t/h), which value is the half-life? (h)",
        ],
        adult_role="Project. Think aloud. Release promptly at minute 17.",
    )
    example_items = qb.get_for_packet(LAUNCH_IDS)
    add_callout(doc, "\U0001f3a4  TEACHER SCRIPT",
                [
                    "1. “Watch me solve an equation with a rational exponent. I rewrite x^(m/n) as the nth root of x, raised to the m.”",
                    "2. “Step 1: Isolate the expression with the rational exponent on one side.”",
                    "3. “Step 2: Raise BOTH sides to the reciprocal power n/m to eliminate the exponent.”",
                    "4. “Step 3: Simplify and solve for the variable.”",
                    "5. “Step 4: CHECK by substituting back into the original — even perfect algebra can produce extraneous solutions when the denominator is even.”",
                    "6. “Today’s DOK-3 task applies this exact skill to a real-world half-life model. The exponent t/h IS a rational exponent.”",
                    "7. Release to Explore.",
                ],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Explore",
        framework_tag="TPS + DOK-3 driver",
        dok="2-3",
        minutes=33,
        teacher_does=[
            "4 laps across 33 min.",
            "Lap 1 (8 min): Try-It 4 + #12 — rational-exponent equations. Press pairs to write the reciprocal power explicitly.",
            "Lap 2 (8 min): #33 + #34 — rational-exponent equations. Watch for pairs who skip the check step.",
            "Lap 3 (12 min): ⭐ #41 Half-Life Model. Students identify a, h, t; compute exponent; evaluate; interpret.",
            "Do NOT give #41 answer. Pairs present computation and interpretation at Share.",
        ],
        students_do=[
            "5 items in order: Try-It 4 → #12 → #33 → #34 → #41.",
            "For #41: BEFORE computing, identify a=50 (start amount), h=5 (half-life period in hours), t=16 (elapsed hours).",
            "Justify with sentence frame.",
        ],
        questions_to_ask=[
            "What is the reciprocal of the exponent? Write it before you raise both sides.",
            "#12: does the denominator of the exponent tell you anything about extraneous solutions?",
            "#33 / #34: did you check your solution in the original equation?",
            "#41: what does a=50 represent? What does h=5 represent?",
            "#41: compute 16/5. What decimal is that? Now evaluate (0.5)^3.2.",
            "#41: interpret your answer — what does 5.44 mL mean in context?",
        ],
        adult_role="Solo teacher: 4 laps. Press INTERPRETATION on #41 — decimal answer alone is not enough.",
    )
    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 4", "Practice #12", "Practice #33", "Practice #34",
        "⭐ DOK-3 HALF-LIFE MODEL — Practice #41 (Topic 5 spine)",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=f"{label} ({q['id']})", space_after_inches=0.3)
        if "DOK-3" in label or "#41" in label:
            add_callout(doc, "⭐ DOK-3 HALF-LIFE MODEL — Practice #41 (Topic 5 spine)",
                        [
                            "This is the single most important item in the topic. Students should (a) identify a=50, h=5, t=16, (b) compute exponent 16/5 = 3.2, (c) evaluate 50·(0.5)^3.2 ≈ 5.44 mL, (d) interpret: 'about 5.44 mL of the drink remains after 16 hours.' All three assessment items chain off this skill.",
                        ],
                        color_hex="FFD9E0")
        else:
            add_callout(doc, "✅  ANSWER / ANTICIPATED TALK",
                        [q.get("notes") or "(no answer key — verify before class)"],
                        color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Share / Summary",
        framework_tag="academic conversation + P3 bridge",
        dok="2",
        minutes=5,
        teacher_does=[
            "Cold-call 1 pair to present their Half-Life computation and interpretation using sentence frame.",
            "Class signals thumbs. Write the half-life formula y = a·(0.5)^(t/h) on board with a, h, t labeled.",
            "P3 bridge: \"Next period we apply these skills to assessment-critical items — bring your check-for-extraneous habit.\"",
        ],
        students_do=[
            "Presenting pair reads their computation and interpretation.",
            "Rest of class signals and writes takeaway in margin.",
        ],
        questions_to_ask=[
            "Summarize: what does the exponent t/h represent in the half-life formula?",
            "Does a rational exponent always produce a unique solution? (No — check for extraneous.)",
        ],
        adult_role="Cold-call a pair that struggled on #41 — hold them accountable to the full interpretation step.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit Ticket",
        framework_tag="summary recap (not graded)",
        dok="1-2",
        minutes=3,
        teacher_does=[
            "Collect at door.",
            "Scan stem 3 (50 → 5.44 mL). Plan P3 Do Now if many students got the computation wrong or skipped the interpretation.",
        ],
        students_do=["4 sentence stems, honest."],
        questions_to_ask=[
            "Did students correctly compute (0.5)^3.2 and interpret the result in context?",
        ],
        adult_role="Collect. No grade.",
    )
    add_callout(doc, "\U0001f4cb  EXIT TICKET — Student Form",
                [
                    "To solve (expr)^(m/n) = k I raise both sides to power ___.",
                    "Half-life tells me the ___ of the exponent.",
                    "After 16 hours, 50 mL of soft drink decays to ___ mL.",
                    "One biggest thing I learned today: ________________________",
                ],
                color_hex="FFF2CC")

    doc.add_page_break()
    add_callout(doc, "\U0001f550  PERIOD A / PERIOD F — 65-MIN CLASS NOTES",
                [
                    "Both sections should have 65 min for P2. Use the extra 10 min on Explore #41 (Half-Life Model) — press students to write the full interpretation sentence, not just the decimal.",
                    "If finished early: hand out Practice #19 (extraneous-solution argument) as fast-finisher.",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "\U0001f9e9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS",
                [
                    "• Provide the Rational-Exponent Rules card as a sticker/cut-out.",
                    "• For #41, provide the formula y = a·(0.5)^(t/h) pre-labeled with a=50, h=5, t=16.",
                    "• Accept verbal interpretation in lieu of written sentence.",
                ],
                color_hex="E2F0D9")
    add_callout(doc, "\U0001f30d  SUPPORTS FOR ELL STUDENTS",
                [
                    "• Sentence frames on student page (don't skip).",
                    "• Word cards: “reciprocal power” = flip the exponent fraction; “half-life” = time for half the amount to remain; “extraneous solution” = an answer the algebra gives but the original equation rejects.",
                    "• “Half-Life Model” = a formula showing how something decreases by half every h hours.",
                    "• Pair ELL students with English-dominant partners for TPS.",
                ],
                color_hex="DEEBF7")

    doc.save(path)


# ── Main ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_do_now("L54_P2_Do_Now.docx")
    build_student("L54_P2_Student_Packet.docx")
    build_teacher("L54_P2_Teacher_Packet.docx")
    print("Built L54_P2_Do_Now.docx, L54_P2_Student_Packet.docx, L54_P2_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L54_P2_Visuals_Checklist.md",
                              title="L54 P2 — Visuals Checklist")
