"""Generate Lesson 6-4 Period 1 materials — single-period quick treatment.

Single-period lesson — condensed Klimsara. If 45-min Wed F variant, cut last
DOK-2 practice item (q9) from Explore — graph-reading partially covered by q32.

NOTE: This is a COMPRESSED lesson — one Example modeled in Launch (Ex 3B:
inverse of g(x)=log_7(x+5) + composition verification), which aligns with
the standard Klimsara one-example rule. Lesson 6-4 content density (graphing
logs + finding inverses) is handled by Do Now coverage, allowing a tighter
Launch model.

Design: Logarithmic Functions + DOK-3 Sales Revenue Inverse (Item 28).
This single-period treatment covers graphing log functions, key features
(asymptote, domain, range, intercepts), and finding inverses of exponential
and logarithmic functions. Ends with Practice #28 (Sales Revenue R = 12·log(a+1)+25),
the DOK-3 driver: find the inverse, then judge which form is easier for
specific R values.

Phases (50 min base; 45-min Wed F variant cuts q9):
  0–  7   Do Now    : 3 items — q31 (translation multi-select), q5 (graph log_4 x),
                      q32 (SAT/ACT inverse of f(x)=5^(x+1))               [DOK 1-2]
  7– 15   Launch    : teacher models Ex 3B (inverse of log_7(x+5),
                      composition verification; contrast with Ex 3A 90s)  [DOK 2]
 15– 40   Explore   : ⭐ DOK-3 q28 (Sales Revenue) + DOK-2 q21/q24/q13/q9 [DOK 2-3]
 40– 45   Share/Sum : DOK-3 pair presents function-form choice + 6-5 bridge[DOK 2]
 45– 50   Exit      : find inverse of f(x)=3·log_5(x−2)+1; asymptotes     [DOK 1-2]
 (back)   Optional: q33 (telescope PT) — stretch, not required
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_IDS  = [
    "6-4-savvas-q31",    # translation identification multi-select: h(x)=ln(x+2)−1
    "6-4-savvas-q5",     # graph y=log_4 x; state domain/range/intercepts/end behavior
    "6-4-savvas-q32",    # SAT/ACT: inverse of f(x)=5^(x+1)
]
LAUNCH_ID   = "6-4-savvas-example-3-lesson-6-4"   # Ex 3B: inverse of log_7(x+5) + comp check
EXPLORE_IDS = [
    "6-4-savvas-q28",    # ⭐ DOK-3 spine — Sales Revenue inverse
    "6-4-savvas-q21",    # Form B Q7 rehearsal (inverse of 5^(x-3))
    "6-4-savvas-q24",    # Form B Q12 rehearsal (inverse of log_2(8x))
    "6-4-savvas-q13",    # Form B Q11 rehearsal (asymptote of translated log)
    "6-4-savvas-q9",     # Form B Q9/Q10 rehearsal (are graphs inverses; inverse from graph)
]
DOK3_ID       = "6-4-savvas-q28"
REINFORCE_IDS = ["6-4-savvas-q33"]  # telescope PT — stretch, not required

_ALL_IDS = DO_NOW_IDS + [LAUNCH_ID] + EXPLORE_IDS + REINFORCE_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 6-4 · Quick"
LESSON_TITLE = "Logarithmic Functions + DOK-3 Sales Revenue Inverse (Item 28)"

OBJECTIVES = [
    ("Math Objective",
     "Graph logarithmic functions, identify key features (asymptote, domain, range, "
     "intercepts), and find equations of inverses of exponential and logarithmic "
     "functions."),
    ("Language Objective",
     "Describe how transforming a log function shifts its asymptote and key points; "
     "explain why the inverse of a log is an exponential."),
    ("Essential Understanding",
     "Log and exponential functions are inverses — swap x and y, solve. "
     "The asymptote of a log is the y-axis mirror of an exp function's "
     "horizontal asymptote."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Fluency with graphing log functions and finding inverses; apply to the "
     "Sales Revenue DOK-3 task (Item 28 rehearsal for Form B model-then-invert "
     "structure)."),
    ("Essential Question",
     "How do the graphs of logarithmic and exponential functions relate, "
     "and how do we find one from the other?"),
    ("Materials",
     "Student packet, pencil, calculator. Rules card: log_b 1 = 0; log_b b = 1; "
     "y = log_b x has VA at x = 0; y = log_b(x − h) + k has VA at x = h; "
     "inverse of y = a·b^x is y = log_b(x/a); "
     "inverse of y = log_b(x − h) + k is y = b^(x−k) + h."),
]

RULES_TEXT = [
    "LOG FUNCTION KEY FEATURES (y = log_b x, b > 0, b ≠ 1):",
    "  • Domain: x > 0         • Range: all real numbers",
    "  • VA: x = 0 (y-axis)    • x-intercept: (1, 0)",
    "  • Passes through (b, 1)",
    "TRANSLATION (y = log_b(x − h) + k):",
    "  • VA shifts to x = h        • Domain: x > h",
    "  • Point (b + h, 1 + k)",
    "INVERSE RULES:",
    "  • y = a·b^x  ⇒  inverse y = log_b(x/a)",
    "  • y = log_b(x − h) + k  ⇒  inverse y = b^(x−k) + h",
    "COMPOSITION CHECK: f(f⁻¹(x)) = x",
]


# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.0):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    ps.render_prompt(prompt_p, q["prompt"], font_size=11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            run = line.add_run(f"  ({chr(65+i)})  " + ps.latex_to_unicode(a))
            run.font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, for_teacher=False):
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, LESSON_TITLE,
                  subtitle=LESSON_LABEL + (" · Teacher Edition" if for_teacher else ""))
    t = doc.add_table(rows=len(OBJECTIVES), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(OBJECTIVES):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    t = doc.add_table(rows=len(FRAMEWORK_HEADER), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(FRAMEWORK_HEADER):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()


# ── Do Now ───────────────────────────────────────────────────────────────

def build_do_now(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, "Do Now — " + LESSON_TITLE, subtitle=LESSON_LABEL)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge to log inverses launch",
        dok="1-2",
        minutes=7,
        teacher_does=["Hand out at door.", "Circulate 4 min.",
                      "Cold-call 1 student: \"For q31 — which transformation shifts the "
                      "vertical asymptote? How does h(x)=ln(x+2)−1 compare to y=ln x?\"",
                      "Pull Do Now sheets before Launch."],
        students_do=["q31: identify translations of h(x)=ln(x+2)−1 (multi-select).",
                     "q5: graph y=log_4 x; label domain, range, VA, intercept, end behavior.",
                     "q32: find the inverse of f(x)=5^(x+1) (SAT/ACT format)."],
        questions_to_ask=[
            "q31: what does replacing x with (x+2) do to the vertical asymptote?",
            "q5: what is the domain? Where is the vertical asymptote?",
            "q32: how do you undo an exponential — what operation is the inverse of 5^x?",
        ],
        adult_role="Solo teacher: circulate, time 7 min, pull sheets.",
    )

    dn_items = qb.get_for_packet(DO_NOW_IDS)
    labels = [
        "Do Now 1 — Translation identification: h(x)=ln(x+2)−1",
        "Do Now 2 — Graph y=log₄ x; state key features",
        "Do Now 3 — SAT/ACT: inverse of f(x)=5^(x+1)",
    ]
    for label, q in zip(labels, dn_items):
        bank_item_block(doc, q, label=label, space_after_inches=1.5)

    add_callout(doc, "\U0001f3af  SENTENCE FRAME",
                ["“The graph of h(x) = ln(x+2)−1 is the graph of y = ln x shifted "
                 "___ units left and ___ units down. "
                 "The vertical asymptote moves from x = ___ to x = ___.”"],
                color_hex="FFF2CC")

    doc.save(path)


# ── Student packet ───────────────────────────────────────────────────────

def build_student(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=False)

    ps.teal_section(doc, "LAUNCH — Inverse of a Log Function (teacher models Ex 3B)", [])
    add_callout(doc, "\U0001f4d8  LOG FUNCTION KEY FEATURES + INVERSE RULES (keep printed on your page)",
                RULES_TEXT,
                color_hex="D9EAD3")

    add_callout(doc, "⚠️  ONE EXAMPLE IN LAUNCH — Ex 3B: inverse of g(x)=log₇(x+5)",
                [
                    "Teacher models Example 3B: find the inverse of g(x) = log_7(x+5), "
                    "then verify using composition f(f⁻¹(x)) = x.",
                    "Contrast: Ex 3A (inverse of an exponential) covered in 90 sec "
                    "— same swap-and-solve process, direction reversed.",
                    "Key move: swap x and y → 7^(x) = y + 5 → solve for y.",
                ],
                color_hex="CFE2F3")
    doc.add_paragraph()

    ps.teal_section(doc, "EXPLORE — Think-Pair-Share (25 min)",
                    ["Work with a partner. Start with ⭐ Practice #28 (DOK-3 Sales Revenue). "
                     "Then DOK-2 items in order.",
                     "45-min Wed F: q9 is dropped. Skip to Exit after q13."])

    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Practice #28  ⭐ DOK-3 SALES REVENUE INVERSE — Form B model-then-invert rehearsal",
        "Practice #21 — Inverse of exponential 5^(x-3)  [Form B Q7 rehearsal]",
        "Practice #24 — Inverse of log_2(8x), token drag  [Form B Q12 rehearsal]",
        "Practice #13 — Asymptote of translated log from graph  [Form B Q11 rehearsal]",
        "Practice #9 — Are graphs inverses? Inverse from graph  [Form B Q9/Q10 rehearsal]",
    ]
    for label, q in zip(labels, items):
        space = 2.0 if "DOK-3" in label else 1.5
        bank_item_block(doc, q, label=label, space_after_inches=space)

    ps.teal_section(doc, "SHARE / SUMMARY (5 min)", [])
    ps.sentence_frame_box(doc, [
        "DOK-3 pair share: “For Practice #28, I found the inverse by ___. "
        "For revenue R = ___ dollars, the original form R = 12·log(a+1)+25 is easier "
        "because ___. The inverse form a = ___ is easier when ___.”",
        "Bridge to 6-5: log properties let us rewrite and simplify log expressions — "
        "that makes inverses and equations easier.",
    ])

    add_callout(doc, "\U0001f50d  BRIDGE TO LESSON 6-5",
                ["In 6-5 you will learn log properties (product, quotient, power rules). "
                 "Those properties let you expand or condense log expressions — "
                 "which makes solving equations like the one in #28 much faster."],
                color_hex="FCE5CD")

    ps.summary_exit_box(doc, "EXIT TICKET — Inverse + Asymptote",
                        [
                            "Find the inverse of f(x) = 3·log₅(x − 2) + 1.",
                            "State the vertical asymptote of f.",
                            "State the horizontal asymptote of f⁻¹.",
                            "One biggest thing I learned today: ________________________",
                        ])

    doc.add_page_break()
    ps.teal_section(doc, "OPTIONAL REINFORCEMENT (not collected)",
                    ["If you finish Explore early. #33 is a Savvas performance task "
                     "(telescope) that extends inverse/log fluency to a modeling context."])
    reinforce = qb.get_for_packet(REINFORCE_IDS)
    for i, q in enumerate(reinforce, 1):
        bank_item_block(doc, q, label=f"Optional #{i}  (Savvas Practice #33 — Telescope PT)",
                        space_after_inches=1.8)

    doc.save(path)


# ── Teacher packet ───────────────────────────────────────────────────────

def build_teacher(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=True)

    add_callout(doc, "\U0001f4cc  PRE-FLIGHT — SINGLE-PERIOD COMPRESSED LESSON + DOK-3 CAPSTONE",
                [
                    "This is a single-period quick treatment of Lesson 6-4. ONE example modeled "
                    "in Launch (Ex 3B: inverse of g(x)=log_7(x+5) + composition verification). "
                    "Contrast with Ex 3A (inverse of exponential) in 90 sec only — same "
                    "swap-and-solve process, direction reversed. Release to Explore at minute 15.",
                    "DOK-3 driver: Practice #28 (Sales Revenue R = 12·log(a+1)+25). "
                    "Students (a) find the inverse function a = 10^((R−25)/12) − 1, "
                    "(b) judge which form is easier for specific R values. "
                    "Real Savvas values from SE source.",
                    "Pace: ⭐ q28 (10 min) → q21 (4 min) → q24 (4 min) "
                    "→ q13 (4 min) → q9 (3 min). ~25 min total Explore.",
                    "45-min Wed F variant: cut q9. Never cut q28.",
                    "Assessment alignment: q21→Form B Q7, q24→Q12, q13→Q11, "
                    "q9→Q9/Q10, q28→DOK-3 rehearsal.",
                ],
                color_hex="FFD9E0")

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge to log inverses launch",
        dok="1-2",
        minutes=7,
        teacher_does=["Hand out at door.", "Circulate 4 min.",
                      "Cold-call on q31: ask which transformation shifts the VA.",
                      "Pull sheets at minute 7."],
        students_do=["q31: translation multi-select for h(x)=ln(x+2)−1.",
                     "q5: graph y=log_4 x with all key features.",
                     "q32: find inverse of f(x)=5^(x+1) (SAT/ACT)."],
        questions_to_ask=[
            "q31: how does (x+2) inside the log change the VA?",
            "q5: domain? VA? x-intercept? End behavior as x→0⁺ and x→∞?",
            "q32: undo 5^x using log base 5 — what is the inverse operation?",
        ],
        adult_role="Solo teacher: circulate silently. No hints on q32.",
    )
    dn_items = qb.get_for_packet(DO_NOW_IDS)
    for i, q in enumerate(dn_items, 1):
        bank_item_block(doc, q, label=f"Do Now {i} ({q['id']})", space_after_inches=0.3)
    add_callout(doc, "✅  ANSWER KEY",
                [q.get("notes") or "(verify before class)" for q in dn_items],
                color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Launch",
        framework_tag="teacher models Ex 3B (inverse of log) + 90s Ex 3A contrast",
        dok="2",
        minutes=8,
        teacher_does=[
            "Project Ex 3B. Think aloud: g(x) = log_7(x+5). Step 1: swap x and y. "
            "Step 2: write 7^x = y+5. Step 3: solve — y = 7^x − 5.",
            "\"Now verify with composition: g(g⁻¹(x)) = log_7((7^x−5)+5) = log_7(7^x) = x. ✔\"",
            "90-sec contrast: Ex 3A (inverse of exponential, e.g. y = 3·5^x). "
            "\"Swap x and y → x = 3·5^y → x/3 = 5^y → y = log_5(x/3). Same swap-and-solve.\"",
            "Flag the Rules card — students copy VA shift rule and inverse rules into margin.",
            "30-sec pause: partner repeats the inverse steps in their own words.",
            "Do NOT solve Explore items — release at minute 15.",
        ],
        students_do=[
            "Watch Ex 3B silently. Note swap-then-solve sequence.",
            "Copy composition verification step.",
            "During 90-sec contrast: note that inverting an exponential gives a log, "
            "and inverting a log gives an exponential.",
            "Copy Rules card VA and inverse rules into margin.",
        ],
        questions_to_ask=[
            "In g(x)=log_7(x+5), what is the domain? VA? (x=−5)",
            "After swapping x and y, how do we isolate y when 7^x = y+5?",
            "What is g⁻¹(x)? What is its domain? (x ∈ ℝ)",
            "Why does the VA of g become the HA of g⁻¹?",
        ],
        adult_role="Project Ex 3B. Model composition check. 90-sec Ex 3A contrast. "
                   "Release promptly at minute 15.",
    )
    launch_item = qb.get_for_packet([LAUNCH_ID])[0]
    bank_item_block(doc, launch_item, label=f"Example 3B ({launch_item['id']})",
                    space_after_inches=0.3)
    add_callout(doc, "\U0001f3a4  TEACHER SCRIPT",
                [
                    "1. \"Watch me find the inverse of a LOG function.\"",
                    "2. \"Step 1: replace g(x) with y. Step 2: swap x and y.\"",
                    "3. \"Step 3: rewrite in exponential form — 7^x = y+5.\"",
                    "4. \"Step 4: solve for y — y = 7^x − 5. That’s g⁻¹(x).\"",
                    "5. \"Composition check: g(g⁻¹(x)) = log_7(7^x−5+5) = x. ✔\"",
                    "6. \"Now 90 seconds: Ex 3A — inverting an exponential gives a log. "
                    "Same swap-and-solve.\"",
                    "7. Release to Explore.",
                ],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Explore",
        framework_tag="TPS + DOK-3 driver",
        dok="2-3",
        minutes=25,
        teacher_does=[
            "4 laps across 25 min.",
            "Lap 1 (10 min): ⭐ q28 DOK-3. Press pairs on inverse setup and form comparison. "
            "DO NOT give the inverse formula — redirect with 'what does it mean to undo R?'",
            "Lap 2 (4 min): q21. Press notation: what is the inverse of 5^(x-3)?",
            "Lap 3 (4 min): q24. Press: inverse of log_2(8x) — watch for domain.",
            "Lap 4 (4 min): q13. Press: how does the horizontal shift move the VA?",
            "Lap 5 (3 min): q9. Are graphs inverses? Key test: reflective symmetry over y=x.",
            "45-min Wed F: skip q9.",
        ],
        students_do=[
            "5 items in order: ⭐ q28 → q21 → q24 → q13 → q9.",
            "For q28: BEFORE computing, write down R = 12·log(a+1)+25 and plan how "
            "to isolate a.",
            "For q28 part 2: decide which form (original or inverse) is easier for "
            "R = 49 and R = 61. Write a justification sentence.",
        ],
        questions_to_ask=[
            "q28: how do you undo +25 first? Then undo ×12?",
            "q28: after (R−25)/12 = log(a+1), what is the next step?",
            "q28: for R=49, would you use R=12·log(a+1)+25 or a=10^((R−25)/12)−1? Why?",
            "q21: identify base and exponent — what is the log that undoes 5^x?",
            "q24: inverse of log_2(8x) — swap x and y first. Then rewrite.",
            "q13: find the VA of the translated log from the graph.",
            "q9: if two graphs are inverses, what visual test do they pass?",
        ],
        adult_role="Solo teacher: 4 laps. On q28, press the judgment step (which form and why) "
                   "— that is the DOK-3 performance criterion.",
    )
    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Practice #28 ⭐ DOK-3 SALES REVENUE INVERSE",
        "Practice #21", "Practice #24", "Practice #13", "Practice #9",
    ]
    form_b = {
        "6-4-savvas-q21": "Form B Q7: inverse of 5^(x-3)",
        "6-4-savvas-q24": "Form B Q12: inverse of log_2(8x), token drag",
        "6-4-savvas-q13": "Form B Q11: asymptote of translated log from graph",
        "6-4-savvas-q9":  "Form B Q9/Q10: are graphs inverses; inverse from graph",
        "6-4-savvas-q28": "DOK-3 rehearsal for Form B model-then-invert structure",
    }
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=f"{label} ({q['id']})", space_after_inches=0.3)
        if "DOK-3" in label:
            add_callout(doc, "⭐ DOK-3 SALES REVENUE — Practice #28",
                        [
                            "Students must (a) isolate a from R = 12·log(a+1)+25, "
                            "yielding a = 10^((R−25)/12) − 1, "
                            "(b) judge which form is easier for specific R values.",
                            "Key criteria: (1) correct algebraic isolation steps shown, "
                            "(2) inverse formula correct (base 10, not natural log), "
                            "(3) justification for form preference cites a specific computation.",
                            "Assessment alignment: " + form_b[q["id"]],
                        ],
                        color_hex="FFD9E0")
        else:
            add_callout(doc, "✅  ANSWER / ANTICIPATED TALK",
                        [q.get("notes") or "(no answer key — verify before class)",
                         "Assessment alignment: " + form_b.get(q["id"], "—")],
                        color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Share / Summary",
        framework_tag="academic conversation + 6-5 bridge",
        dok="2",
        minutes=5,
        teacher_does=[
            "Cold-call 1 pair to present q28 solution and form-choice justification.",
            "Class signals thumbs. Write a = 10^((R−25)/12) − 1 on board.",
            "6-5 bridge: \"Log properties (product/quotient/power rules) will let you "
            "expand and condense log expressions — making inverses and equations faster.\"",
        ],
        students_do=[
            "Presenting pair: read inverse formula and cite which form is easier for R=49 and R=61.",
            "Rest of class: signal thumbs and write bridge note in margin.",
        ],
        questions_to_ask=[
            "How did you isolate a from R = 12·log(a+1)+25?",
            "For R=49, which form is computationally easier? Why?",
        ],
        adult_role="Cold-call a pair that struggled on the judgment step — "
                   "hold them accountable to citing a computation, not just a preference.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit Ticket",
        framework_tag="inverse + asymptote (not graded)",
        dok="1-2",
        minutes=5,
        teacher_does=[
            "Collect at door.",
            "Scan: (1) Did students swap x/y correctly? "
            "(2) Did students solve for y in exponential form? "
            "(3) Did they state VA of f as x=2 and HA of f⁻¹ as y=2?",
        ],
        students_do=["Two-part custom exit: find inverse, then state both asymptotes."],
        questions_to_ask=[
            "Inverse of f(x) = 3·log_5(x−2)+1: swap → exponential form → solve.",
            "VA of f is x=___ (vertical asymptote of original log).",
            "HA of f⁻¹ is y=___ (the VA of f becomes the HA of the inverse).",
        ],
        adult_role="Collect. No grade. Use responses to plan any re-teach before 6-5.",
    )
    add_callout(doc, "\U0001f4cb  EXIT TICKET — Student Form",
                [
                    "Find the inverse of f(x) = 3·log₅(x − 2) + 1.",
                    "State the vertical asymptote of f.",
                    "State the horizontal asymptote of f⁻¹.",
                    "One biggest thing I learned today: ________________________",
                    "",
                    "ANSWER KEY (TE only): f⁻¹(x) = 5^((x−1)/3) + 2. "
                    "VA of f: x = 2. HA of f⁻¹: y = 2.",
                ],
                color_hex="FFF2CC")

    doc.add_page_break()
    add_callout(doc, "\U0001f550  PERIOD A / PERIOD F — TIMING NOTES",
                [
                    "50-min base: standard pacing above.",
                    "45-min Wed F variant: drop q9 from Explore. Never cut q28.",
                    "If finished early: hand out Practice #33 (Telescope PT — stretch).",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "\U0001f9e9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS",
                [
                    "• Provide the Log Function Key Features + Inverse Rules card as a sticker/cut-out.",
                    "• For q28, provide the setup: (R−25)/12 = log(a+1). Students solve from that point.",
                    "• Accept verbal justification for form-choice in lieu of written sentence.",
                ],
                color_hex="E2F0D9")
    add_callout(doc, "\U0001f30d  SUPPORTS FOR ELL STUDENTS",
                [
                    "• Sentence frames on student page (don’t skip).",
                    "• Word cards: ‘vertical asymptote’ = a line the graph never touches; "
                    "‘inverse function’ = reverses the input/output; "
                    "‘composition’ = plugging one function inside another.",
                    "• ‘log base b of x’ = the exponent you put on b to get x.",
                    "• Pair ELL students with English-dominant partners for TPS.",
                ],
                color_hex="DEEBF7")

    doc.save(path)


# ── Main ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_do_now("L64_P1_Do_Now.docx")
    build_student("L64_P1_Student_Packet.docx")
    build_teacher("L64_P1_Teacher_Packet.docx")
    print("Built L64_P1_Do_Now.docx, L64_P1_Student_Packet.docx, L64_P1_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L64_P1_Visuals_Checklist.md",
                              title="L64 P1 — Visuals Checklist")
