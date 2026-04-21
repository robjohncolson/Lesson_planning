"""Generate Lesson 3-5 Period 3 (Wednesday) materials — Klimsara-adapted, no Blooket.

Design: real + complex zeros + modeling application. Single-DOK-3 spine uses
Savvas Practice #27 (Storage Box) as the modeling driver. All student-facing
items trace to Savvas bank IDs.

Phases (A: 65 min; F: 45 min short — see Period F variant notes in teacher pkt):
  0–  5   Do Now    : bridge prompt (conjugate pair coefficients)             [DOK 1-2]
  5– 17   Launch    : teacher models Example 3 + Example 4 briefly            [DOK 2]
 17– 55   Explore   : Try-It 3a/3b + Practice #17 + DOK-3 driver q27 TPS      [DOK 2-3]
 55– 62   Share/Sum : multi-step strategy articulation                        [DOK 2]
 62– 65   Exit      : summary recap fill-in + one sentence                    [DOK 1-2]
 (back)   Optional Reinforcement: Practice #28, #30 (not collected)           [DOK 2-3]

Period F 45-min variant:
  - Skip Example 4 in Launch (stay on Example 3 only).
  - Explore cut to 25 min: Try-It 3a + Practice #27 (core only). Skip Try-It 3b + #17.
  - Share/Exit stay at 7+3. Framework phases never cut.
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

# The Do Now for P3 is a teacher-authored bridge prompt (conjugate pairs) posted
# on the board from Tuesday's Share phase. It is NOT a Savvas item — it's a
# 1-minute connect prompt, not a packet problem. So no bank ID is needed.
# For formal compliance we use a Savvas-sourced warm-up as the written Do Now
# item and post the bridge prompt on the board as a verbal lead-in.
DO_NOW_ID      = "3-5-savvas-q11"   # intersection problem — real + complex framing
PRACTICE_IDS   = ["3-5-tryit-3a", "3-5-tryit-3b",
                  "3-5-savvas-q17", "3-5-savvas-q27"]
DOK3_ID        = "3-5-savvas-q27"   # Storage Box — the single DOK-3 spine
REINFORCE_IDS  = ["3-5-savvas-q28", "3-5-savvas-q30"]

_ALL_IDS = [DO_NOW_ID] + PRACTICE_IDS + REINFORCE_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 3-5 · Period 3"
LESSON_TITLE = "Real + Complex Zeros, and Modeling"

OBJECTIVES = [
    ("Math Objective",
     "Find all zeros (real and complex) of a polynomial function, and use the "
     "factored form to model a volume-constrained application."),
    ("Language Objective",
     "Students will use the frame \u201cThe zeros are ___; because ___, the "
     "complex conjugate ___ must also be a zero.\u201d to justify complex-zero "
     "reasoning."),
    ("Essential Understanding",
     "Complex zeros of a polynomial with real coefficients come in conjugate "
     "pairs. A cubic model of a physical quantity (volume, profit) can be "
     "solved for specific values by factoring and solving."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Students find all zeros (real + complex) of two cubics, then model a "
     "Storage Box volume problem end-to-end. Every item is Savvas-sourced."),
    ("Essential Question",
     "How do real and complex zeros together describe a polynomial, and how "
     "do you use that structure to answer a real-world model?"),
    ("Materials",
     "Pencils; student packet; calculator (optional). Desmos allowed for Try-"
     "It 3 check only, NOT for Practice #27."),
]


# ── Low-level helpers (mirrored from L35_P2) ─────────────────────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    lines = [title] + list(body_lines)
    set_cell(cell, lines, bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.0):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    ps.render_prompt(prompt_p, q["prompt"], font_size=11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            run = line.add_run(f"  ({chr(65+i)})  " + ps.latex_to_unicode(a))
            run.font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, for_teacher=False):
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 3, LESSON_TITLE,
                  subtitle=LESSON_LABEL + (" · Teacher Edition" if for_teacher else ""))
    t = doc.add_table(rows=len(OBJECTIVES), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(OBJECTIVES):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    t = doc.add_table(rows=len(FRAMEWORK_HEADER), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(FRAMEWORK_HEADER):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()


# ── Do Now ───────────────────────────────────────────────────────────────

def build_do_now(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 3, "Do Now — " + LESSON_TITLE, subtitle=LESSON_LABEL)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from Tuesday",
        dok="1-2",
        minutes=5,
        teacher_does=["Hand out at door.",
                      "Post Tuesday's bridge prompt on the board as they work."],
        students_do=["Read the graph and list intersection points.",
                     "Jot: are the zeros rational, irrational, or complex?"],
        questions_to_ask=[
            "Where does the graph cross the given line?",
            "Reminder from Tuesday: if a cubic has zeros 2, 1+√3, and 1−√3 — what's the nature of its coefficients?",
        ],
        adult_role="Solo teacher: circulate, time 5 min, then pull sheets.",
    )
    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label="Do Now — Intersection Points", space_after_inches=2.0)

    add_callout(doc, "🔗  BRIDGE FROM TUESDAY (verbal, posted on board)",
                ["A cubic has zeros 2, 1 + √3, and 1 − √3. What is the nature of its coefficients — rational, irrational, or complex? Why? (This is Topic 3 Assessment Q#2 — no surprises on Thursday.)"],
                color_hex="FCE5CD")

    doc.save(path)


# ── Student packet ───────────────────────────────────────────────────────

def build_student(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=False)

    ps.teal_section(doc, "LAUNCH — Complex Zeros + Modeling (teacher models on screen)", [])
    add_callout(doc, "📘  COMPLEX-ZERO RULES (keep printed on your page)",
                [
                    "• A polynomial with real coefficients that has complex zeros has them in CONJUGATE PAIRS.",
                    "• If a + bi is a zero, then a − bi must also be a zero.",
                    "• A degree-n polynomial has exactly n zeros (counting multiplicity). Some may be real, some complex.",
                    "• Example: if a cubic has a real zero at x = 4 and a complex zero 2 + 3i, its third zero must be 2 − 3i.",
                ],
                color_hex="D9EAD3")
    add_callout(doc, "📐  MODELING RULES (for the Storage Box)",
                [
                    "• Set up a factored-form volume function V(x) = (expr₁)(expr₂)(expr₃).",
                    "• Each factor is a physical dimension — label each: length, width, height.",
                    "• Use the physical constraint (e.g., \u201clength is greatest\u201d) to decide which factor is which dimension.",
                    "• To find dimensions at V = 10 ft³, solve the cubic equation — show all work.",
                ],
                color_hex="FFF2CC")
    doc.add_paragraph()

    ps.teal_section(doc, "EXPLORE — Think-Pair-Share (38 min)",
                    ["Work with a partner. Use the rule cards above. Storage Box (Practice #27) is the big one — plan 20 min for it."])

    items = qb.get_for_packet(PRACTICE_IDS)
    labels = ["Try It 3a", "Try It 3b", "Practice #17", "Practice #27 ⭐ MODELING"]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=label, space_after_inches=1.5)

    ps.teal_section(doc, "SHARE / SUMMARY (7 min)", [])
    ps.sentence_frame_box(doc, [
        "Pair share: \u201cFor the Storage Box, the factored form was ___. I chose ___ as the length because ___. To find V = 10 ft³, I ___.\u201d",
        "Whole class: one pair at random presents Part (a)-(b)-(c)-(d) of #27.",
    ])

    ps.summary_exit_box(doc, "EXIT TICKET — Summary Recap",
                        [
                            "Today I learned that _____________________________________________________.",
                            "For the Storage Box, the hardest step was _____________________________________________________.",
                            "Thursday's assessment will test _____________________________________________________.",
                        ])

    doc.add_page_break()
    ps.teal_section(doc, "OPTIONAL REINFORCEMENT (not collected)",
                    ["If you finish early or want extra practice before Thursday's assessment."])
    reinforce = qb.get_for_packet(REINFORCE_IDS)
    for i, q in enumerate(reinforce, 1):
        bank_item_block(doc, q, label=f"Optional #{i}  (Savvas Practice)",
                        space_after_inches=1.8)

    doc.save(path)


# ── Teacher packet ───────────────────────────────────────────────────────

def build_teacher(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=True)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from Tuesday",
        dok="1-2",
        minutes=5,
        teacher_does=["Hand out at door.", "Post bridge prompt on board.",
                      "Circulate 3 min; cold-call 1 student for one intersection point."],
        students_do=["List the intersection points.",
                     "Respond to the bridge prompt verbally if called on."],
        questions_to_ask=[
            "How did you find where the graphs cross?",
            "Bridge: nature of the cubic's coefficients — why? (Conjugate irrational pair → rational coefficients.)",
        ],
        adult_role="Solo teacher: stand at door, circulate 3 min, cold-call.",
    )
    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label=f"Do Now ({q['id']})", space_after_inches=0.3)
    add_callout(doc, "✅  ANSWER KEY",
                [q.get("notes") or "(verify from bank)"],
                color_hex="D9EAD3")
    add_callout(doc, "🔗  BRIDGE PROMPT — Expected Answer",
                ["\u201cRational coefficients. The zeros 1 + √3 and 1 − √3 are irrational conjugates — when multiplied, (x − (1+√3))(x − (1−√3)) = x² − 2x − 2, which has rational coefficients. Multiplied by (x − 2), the full cubic has all rational coefficients.\u201d",
                 "This is Topic 3 Assessment Q#2, answer C."],
                color_hex="FCE5CD")

    ps.framework_phase_header(
        doc,
        phase="Launch",
        framework_tag="teacher models Example 3 + brief Example 4",
        dok="2",
        minutes=12,
        teacher_does=[
            "Project Example 3 (real + complex zeros cubic). Think aloud: find one real zero, divide, solve the quadratic.",
            "Brief Example 4 (modeling): 90 sec walk-through of Savvas's storage box — just the setup, not the solution.",
            "Pause once for 30-sec partner-check on conjugate pair.",
        ],
        students_do=[
            "Watch silently through Example 3.",
            "Turn-and-talk: \u201cWhy does the quadratic produce the complex pair?\u201d",
            "Write the complex-zero rule in packet margin.",
        ],
        questions_to_ask=[
            "If one real zero is x = 4, what's the next step?",
            "Why does the discriminant being negative produce complex zeros?",
            "How do you know the complex zero's conjugate is also a zero?",
        ],
        adult_role="Project. Think aloud. Do NOT solve Try-Its. Release promptly at minute 17.",
    )
    add_callout(doc, "🎤  TEACHER SCRIPT",
                [
                    "1. \u201cWatch me find all zeros of this cubic. Step 1: guess a real zero by the rational root theorem. Step 2: synthetic divide. Step 3: solve the leftover quadratic.\u201d",
                    "2. When quadratic's discriminant is negative → \u201cNow we get complex zeros. They come in conjugate pairs — so if I get 2 + 3i, I also have 2 − 3i. That's not a coincidence; it's a rule.\u201d",
                    "3. Example 4 flash: project the storage box picture. \u201cNotice the factored form V(x) = x(x+3)(x−1). We'll unpack this in Explore on the Practice #27 item.\u201d",
                    "4. Release: \u201cYou have Try-It 3a, 3b, then Practice #17, then the big one — #27. Manage your time.\u201d",
                ],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Explore",
        framework_tag="Think-Pair-Share + DOK-3 modeling",
        dok="2-3",
        minutes=38,
        teacher_does=[
            "5 laps of circulation across 38 min.",
            "Lap 1: Try-Its 3a/3b. Lap 2: Practice #17. Laps 3-5: Practice #27 (⭐).",
            "Do not give #27 answers. Pairs present parts (a-d) to each other; then you pick one pair for Share.",
        ],
        students_do=[
            "Pairs work through Try-It 3a/3b (~12 min), then #17 (~6 min), then the full 20 min on #27.",
            "For #27, write each factor's physical meaning BEFORE solving part (d).",
            "Use sentence frame for the Share.",
        ],
        questions_to_ask=[
            "What are ALL zeros of this polynomial — real and complex?",
            "How do the factors of V(x) relate to the box's physical dimensions?",
            "What does the constraint \u201clength is greatest\u201d tell you?",
            "For V = 10 ft³, what cubic equation must you solve, and how?",
        ],
        adult_role="Solo teacher: 5 laps of circulation. Press pairs on #27 part (c) — physical dimension interpretation is the DOK-3 pivot.",
    )
    items = qb.get_for_packet(PRACTICE_IDS)
    labels = ["Try It 3a", "Try It 3b", "Practice #17", "Practice #27 ⭐ DOK-3 DRIVER"]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=f"{label} ({q['id']})", space_after_inches=0.3)
        add_callout(doc, "✅  ANSWER / ANTICIPATED TALK",
                    [q.get("notes") or "(no answer key — verify before class)"],
                    color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Share / Summary",
        framework_tag="academic conversation + assessment preview",
        dok="2",
        minutes=7,
        teacher_does=[
            "Call 1 pair to the board to present #27 parts (a)-(d).",
            "Add 30-sec connection: \u201cTomorrow's assessment Q#6 (Lucy's tray) uses the same factored-volume idea.\u201d",
            "Post the exit ticket cue.",
        ],
        students_do=[
            "Presenting pair walks through each part with sentence frame.",
            "Class signals thumbs up/sideways after each part.",
            "Write the Lucy-tray preview in margin.",
        ],
        questions_to_ask=[
            "Summarize the modeling strategy in one sentence.",
            "What's different between #27 and tomorrow's Lucy-tray problem? (Different physical setup, same factoring strategy.)",
        ],
        adult_role="Cold-call a pair that struggled earlier (not the strongest). Hold others accountable to listen.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit Ticket",
        framework_tag="summary recap (not graded)",
        dok="1-2",
        minutes=3,
        teacher_does=[
            "Collect at door.",
            "Scan responses to #3 (assessment prediction). Use tonight to adjust Thursday Do Now emphasis.",
        ],
        students_do=[
            "Complete 3 sentence stems.",
            "Be honest about #2 — the hardest step tells us what to hit Thursday morning.",
        ],
        questions_to_ask=[
            "Did students name a specific step in #2?",
            "Did they name Topic 3 content (factoring, multiplicity, conjugate pairs, modeling) in #3?",
        ],
        adult_role="Collect. No grade.",
    )
    add_callout(doc, "📋  EXIT TICKET — Student Form",
                [
                    "Today I learned that ___.",
                    "For the Storage Box, the hardest step was ___.",
                    "Thursday's assessment will test ___.",
                ],
                color_hex="FFF2CC")

    doc.add_page_break()

    # Period F short-variant panel — critical for Wed
    add_callout(doc, "🕐  PERIOD F — 45-MIN SHORT VARIANT (Wed only)",
                [
                    "Period F has 45 min on Wednesday. Apply these cuts:",
                    "• Launch: skip Example 4 brief (stay on Example 3). Still 12 min.",
                    "• Explore: 25 min max. DROP Try-It 3b and Practice #17. Keep Try-It 3a + Practice #27 only. #27 is non-negotiable (DOK-3 spine + Thursday-assessment prep).",
                    "• Share: 5 min (cut from 7). One pair presents #27 parts (a)+(b) only; (c)+(d) as HW hint.",
                    "• Exit: 3 min (unchanged).",
                    "• Optional reinforcement: announce but don't expect completion.",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "🧩  MODIFICATIONS / SUPPORT FOR IEP STUDENTS",
                [
                    "• Provide Complex-Zero Rules as a half-sheet cut-out.",
                    "• For #27, provide a pre-labeled box sketch (length / width / height slots).",
                    "• Accept partial solutions on #27(d) with verbal explanation of which factor is which dimension.",
                ],
                color_hex="E2F0D9")
    add_callout(doc, "🌍  SUPPORTS FOR ELL STUDENTS",
                [
                    "• Sentence frames on student page (do not skip).",
                    "• Word cards: \u201cconjugate\u201d = pair that together gives real numbers; \u201cimaginary\u201d = contains i = √−1; \u201cmodel\u201d = equation that describes a real situation.",
                    "• Pair ELL students with English-dominant partners for Share.",
                    "• On #27 part (c), allow verbal responses in students' strongest language before writing.",
                ],
                color_hex="DEEBF7")

    doc.save(path)


# ── Main ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_do_now("L35_P3_Do_Now.docx")
    build_student("L35_P3_Student_Packet.docx")
    build_teacher("L35_P3_Teacher_Packet.docx")
    print("Built L35_P3_Do_Now.docx, L35_P3_Student_Packet.docx, L35_P3_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L35_P3_Visuals_Checklist.md",
                              title="L35 P3 — Visuals Checklist")
