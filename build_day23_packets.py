"""Generate Combined Day 2-3 materials \u2014 Savvas-only multiplicity practice day.

Design: folds old Day 2 (Graphing from Factored Form) + old Day 3 (Multiplicity)
into ONE period. Every student-facing work item traces to a Savvas bank ID.
Savvas Example 2 teacher content + ELL scaffolds are injected from
`questionbank/teacher_prompts/3-5.jsonl` via `qb.teacher_prompts()`.

Context: solo teacher, class of up to ten. No adult-aide language.

Phases (55 min):
  0\u20135    Do Now A : LQ Q4 \u2014 zeros + multiplicities (graph shown)     [DOK 1]
  5\u20137    Do Now B : Blooket login                                       [\u2014]
  7\u201314   Do Now C : Blooket rule recall                                 [DOK 1]
  14\u201326  Launch    : Savvas Example 2 \u2014 name MULTIPLICITY              [DOK 2]
  26\u201346  Practice  : Try It 2a, 2b + Savvas #14/#15/#16 (choose 3)      [DOK 2]
  46\u201350  Share/Sum : Synthesis + self-rating                            [DOK 2]
  50\u201355  Exit      : LQ Q5 \u2014 zeros + behavior (no Desmos)             [DOK 2]
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_ID    = "3-5-lq-q4"
PRACTICE_IDS = ["3-5-tryit-2a", "3-5-tryit-2b",
                "3-5-savvas-q14", "3-5-savvas-q15", "3-5-savvas-q16"]
EXIT_ID      = "3-5-lq-q5"

OBJECTIVES = [
    ("Math Objective",
     "Identify the zeros of a polynomial function from factored form, "
     "determine the multiplicity of each zero, and use multiplicity to "
     "predict whether the graph crosses the x-axis or is tangent to it "
     "(turning point)."),
    ("Language Objective",
     "Students will use the frame \u201cThe multiplicity is ___, so the "
     "graph ___ at x = ___\u201d to justify graph behavior."),
    ("Essential Understanding",
     "The zeros of a polynomial function can be determined using factoring. "
     "A repeated factor creates a zero of higher multiplicity. Even "
     "multiplicity causes the graph to touch and turn; odd multiplicity "
     "causes it to cross."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Students factor a polynomial, read multiplicity from the factored "
     "form, and predict graph behavior at each zero before verifying in Desmos."),
    ("Essential Question",
     "How does the factored form of a polynomial tell you everything you "
     "need to sketch its graph \u2014 zeros AND behavior at each zero?"),
    ("Materials",
     "Student laptops with Desmos; pencils; Blooket access. No chart paper."),
]

# --- Low-level helpers -------------------------------------------------

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if i == 0 and bold_first:
            run.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_two_col(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = TABLE_STYLE
    t.columns[0].width = Inches(1.8)
    t.columns[1].width = Inches(4.7)
    for i, (label, content) in enumerate(rows):
        set_cell(t.rows[i].cells[0], label, bold_first=True)
        if isinstance(content, str):
            content = [content]
        set_cell(t.rows[i].cells[1], content)
    doc.add_paragraph()


def add_callout(doc, title, lines):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    p.add_run(title).bold = True
    for line in lines:
        cell.add_paragraph(line)
    doc.add_paragraph()


def add_phase_callout(doc, phase_label, dok, minutes, title, body):
    """Callout that leads with bold DOK badge + phase + time."""
    if isinstance(body, str):
        body = [body]
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    # Badge-style first line
    dok_run = p.add_run(f"  {dok}  ")
    dok_run.bold = True
    dok_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    ps.shade_cell(cell, ps.COLORS["teal"])  # whole first cell is tinted; we'll undo with per-run coloring
    # Actually reset cell color and apply only to phase header run via whole line styling
    # Simpler: the cell gets a teal left bar via border, not a full teal fill.
    # Revert full-cell shade:
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), ps.COLORS["tealpale"])
    tc_pr.append(shd)
    ps.set_cell_borders(cell,
        left={"val": "single", "sz": 24, "color": ps.COLORS["teal"]},
        top={"val": "nil"}, right={"val": "nil"}, bottom={"val": "nil"})
    # Re-write the paragraph cleanly
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(f"[{dok}]  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(ps.COLORS["teal"])
    r = p.add_run(f"{phase_label}  ")
    r.bold = True
    r = p.add_run(f"({minutes} min)")
    r.font.color.rgb = RGBColor(0x55, 0x66, 0x77)
    p = cell.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    for line in body:
        cp = cell.add_paragraph()
        cp.add_run(line)
    doc.add_paragraph()


def header_line(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def blank_lines(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.add_run("______________________________________________________________")


def add_image_from_bank(doc, qid, width=Inches(3.5)):
    """Embed a bank item's visual (question image) at the given width."""
    q = qb.get(qid)
    if not q or not q.get("image"):
        return False
    path = q["image"]
    if not os.path.exists(path):
        return False
    doc.add_picture(path, width=width)
    return True


def add_blank_grid(doc, width=Inches(3.2)):
    if os.path.exists("assets/blank_grid.png"):
        doc.add_picture("assets/blank_grid.png", width=width)
        return True
    return False


# --- Bank prompt formatting -------------------------------------------

def _format_unicode(prompt):
    prompt = (prompt
              .replace("\u00e2\u2020\u2019", "\u2192")
              .replace("\u00e2\u20ac\u201d", "\u2014")
              .replace("\u00e2\u20ac\u201c", "\u2013"))
    SUPS = {"2": "\u00b2", "3": "\u00b3", "4": "\u2074",
            "5": "\u2075", "6": "\u2076"}
    out, i = [], 0
    while i < len(prompt):
        ch = prompt[i]
        if ch == "^" and i + 1 < len(prompt) and prompt[i+1] in SUPS:
            out.append(SUPS[prompt[i+1]])
            i += 2
            continue
        out.append("\u2212" if ch == "-" else ch)
        i += 1
    return "".join(out)


def _strip_stem(prompt):
    txt = _format_unicode(prompt)
    if ":" in txt:
        txt = txt.split(":", 1)[-1].strip()
    return txt.rstrip(".")


# ======================================================================
# DO NOW SHEET
# ======================================================================

def build_do_now(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    header_line(doc, "Combined Day 2-3  |  DO NOW  \u2014  Zeros + Multiplicities", size=13)

    p = doc.add_paragraph()
    r = p.add_run("Name: _____________________________________     Date: _____________")
    r.font.size = Pt(11)

    # Prominent DOK badge + instructions
    add_callout(doc, "[DOK 1]  \u2022  5 minutes  \u2022  Silent  \u2022  Pencil only  \u2022  No Desmos", [
        "Savvas Lesson Quiz Q4 (multiple-choice).  Circle ONE letter.",
        "Then in your own words, explain how you know.",
    ])

    header_line(doc, "Savvas Lesson Quiz Q4:", size=12)
    p = doc.add_paragraph()
    p.add_run("What are the zeros and their multiplicities for the function").italic = True
    p = doc.add_paragraph()
    p.add_run("    y = x\u00b3 + 3x\u00b2 + x + 3,  shown in the graph?").italic = True
    doc.add_paragraph()

    # Embed the graph + MC choices directly
    added = add_image_from_bank(doc, DO_NOW_ID, width=Inches(3.3))
    if not added:
        p = doc.add_paragraph()
        r = p.add_run("[ graph image missing \u2014 see Savvas Lesson Quiz Q4 in teacher materials ]")
        r.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x00, 0x00)

    doc.add_paragraph()
    header_line(doc, "Your answer: A  /  B  /  C  /  D", size=12)
    doc.add_paragraph()

    header_line(doc, "Explain how you know (1 sentence):", size=12)
    blank_lines(doc, 3)

    doc.save(path)


# ======================================================================
# STUDENT PACKET
# ======================================================================

LAUNCH_INVESTIGATION = [
    "SAVVAS EXAMPLE 2 \u2014 Multiplicity",
    "",
    "With your partner, graph both functions in Desmos.  Watch  x = \u22122:",
    "",
    "    (1)  y = (x + 2)\u00b2 (x \u2212 1)",
    "    (2)  y = (x + 2)\u00b3 (x \u2212 1)",
    "",
    "What did the graph DO at  x = \u22122  in each case?",
    "",
    "    Graph (1):  ______________________________________________________",
    "    Graph (2):  ______________________________________________________",
    "",
    "THE WORD:  the number of times a factor appears in the factored form is "
    "called its  MULTIPLICITY.",
    "    (x + 2)\u00b2  has multiplicity 2 at x = \u22122.   (x + 2)\u00b3  has multiplicity 3 at x = \u22122.",
    "",
    "T-CHART  \u2014  fill in after the Desmos investigation:",
    "   ODD multiplicity (1, 3, 5\u2026)            EVEN multiplicity (2, 4, 6\u2026)",
    "   Behavior at the zero:                    Behavior at the zero:",
    "   _______________________________          _______________________________",
    "",
    "Sentence frame:  \u201cThe multiplicity is ___, so the graph ___ at x = ___.\u201d",
]


def _practice_lines():
    items = qb.get_for_packet(PRACTICE_IDS)
    lines = [
        "For each polynomial: list the zeros, the multiplicity of each, and "
        "describe the behavior of the graph at each zero using Savvas\u2019s "
        "language \u2014 CROSSES the x-axis (odd multiplicity) or is TANGENT to "
        "the x-axis / has a TURNING POINT (even multiplicity).  Use Desmos "
        "only to CHECK your prediction, not to make it.",
        "",
        "REQUIRED:  A and B.   PICK ONE of C, D, E.",
        "",
    ]
    labels = ["A.", "B.", "C.", "D.", "E."]
    for label, q in zip(labels, items):
        lines.append(f"{label}  {_strip_stem(q['prompt'])}")
        lines.append("    Zeros & multiplicities:  ____________________________________")
        lines.append("    Behavior at each:        ____________________________________")
        lines.append("")
    lines.append("(All five items from Savvas.  Bank IDs: " +
                 ", ".join(PRACTICE_IDS) + ".)")
    return lines


SHARE_SUMMARY_STUDENT = [
    "Self-check  \u2014  circle one for each:",
    "1.  I can read the multiplicity of each zero from the factored form.   \u2713    partly    not yet",
    "2.  I can predict cross vs. touch from the multiplicity.               \u2713    partly    not yet",
    "3.  I can factor a polynomial in standard form before using the rule.  \u2713    partly    not yet",
    "",
    "Preview Day 4-5:  real + complex zeros + Savvas storage-box volume (DOK 3).",
]


def _exit_lines():
    q = qb.get_for_packet([EXIT_ID])[0]
    return [
        "[DESMOS ALLOWED:  yes, to verify zeros.  You must still write the "
        "behavior in your own words.]",
        "",
        "Savvas Lesson Quiz Q5:",
        _format_unicode(q["prompt"]),
        "",
        "Zeros:  _____________________________________________________",
        "",
        "Behavior at each zero (crosses the x-axis, OR tangent / turning point):",
        "____________________________________________________________",
        "____________________________________________________________",
        "____________________________________________________________",
    ]


def build_student(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    ps.running_header_name_block(doc)
    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    ps.day_banner(doc, "2-3", "Zeros + Multiplicity (Savvas-anchored)",
                  "Read multiplicity from factored form; predict graph "
                  "behavior at every zero. All work comes from Savvas.")

    add_two_col(doc, OBJECTIVES)
    add_two_col(doc, FRAMEWORK_HEADER)

    add_phase_callout(doc, "DO NOW", "DOK 1", "5",
        "Savvas Lesson Quiz Q4 \u2014 separate sheet",
        [
            "You got the Zeros + Multiplicities sheet on entry.  Hold onto it \u2014 "
            "we\u2019ll return to it during the Launch.",
        ])

    add_phase_callout(doc, "BLOOKET", "DOK 1", "2 + 7",
        "Rule Recall (on the board, not the leaderboard)",
        [
            "Pure rule recall: ZPP, multiplicity = exponent, EVEN \u2192 touch, "
            "ODD \u2192 cross, highest-power \u2192 end behavior. No procedural factoring.",
        ])

    add_phase_callout(doc, "LAUNCH", "DOK 2", "12",
        "Name the Pattern \u2014 Savvas Example 2", [])
    add_callout(doc, "", LAUNCH_INVESTIGATION)

    add_phase_callout(doc, "PRACTICE", "DOK 2", "20",
        "Savvas Try Its + Practice Items", [])
    add_callout(doc, "", _practice_lines())
    # Sketch space for the brave: one small blank grid for optional sketching
    p = doc.add_paragraph()
    r = p.add_run("Extra sketch space (use if you want to check a shape by hand):")
    r.italic = True
    r.font.size = Pt(10)
    add_blank_grid(doc, width=Inches(3.0))
    doc.add_paragraph()

    add_phase_callout(doc, "SHARE / SUMMARY", "DOK 2", "4",
        "Synthesis + self-rating",
        SHARE_SUMMARY_STUDENT)

    ps.summary_exit_box(doc,
        "[DOK 2]  EXIT TICKET  \u2014  5 min  \u2014  Savvas Lesson Quiz Q5",
        _exit_lines())

    doc.save(path)


# ======================================================================
# TEACHER PACKET
# ======================================================================

CRITERIA_FOR_SUCCESS = [
    "Student-facing, revisited during Share/Summary:",
    "1.  I can read the multiplicity of each zero from the factored form.",
    "2.  I can predict cross vs. touch from the multiplicity.",
    "3.  I can factor a polynomial in standard form before using the rule.",
    "Formative evidence:",
    "\u2022  Do Now (Savvas LQ Q4 multiple choice \u2014 collect on entry or after Launch).",
    "\u2022  Blooket per-student percentages on the dashboard.",
    "\u2022  Practice pages (Try It 2a, 2b + Practice #14/#15/#16) \u2014 circulate + photograph.",
    "\u2022  Summary exit ticket (Savvas LQ Q5).",
]

BLOOKET_SCOPE = [
    "SCOPE (DOK 1 rule-recall only).  Rules today\u2019s work depends on:",
    "\u2022  Zero Product Property: zero \u21d4 factor = 0.",
    "\u2022  Sign direction: zero at x = c gives factor (x \u2212 c).",
    "\u2022  Multiplicity = exponent on the factor.",
    "\u2022  EVEN multiplicity \u2192 TOUCH and turn.",
    "\u2022  ODD multiplicity \u2192 CROSS.",
    "\u2022  Highest power \u2192 end behavior.",
    "NO procedural factoring items. Dashboard diagnostic: note the two "
    "lowest items; 60-second board reset if a rule is cold.",
    "DISPLAY NOTE: project the PowerPoint slide 3 (rule-recall list) during "
    "play, not the Blooket leaderboard. Students see the rules they are "
    "being tested on.",
    "CSV: Blooket_Day3_Multiplicity.csv (regenerate via regen_blooket_csvs.py).",
    "Questions to ask (after game): \u201cWhich rule had the lowest percent? Say it now.\u201d",
]

# Pull teacher_prompts for Example 2 into LAUNCH_KEY dynamically.


def _launch_key_lines():
    prompts = qb.teacher_prompts("3-5", anchor_example=2)
    lines = [
        "SAVVAS EXAMPLE 2 TARGET OBSERVATIONS at x = \u22122:",
        "    y = (x + 2)\u00b2 (x \u2212 1)  \u2192  graph is TANGENT to the x-axis at x = \u22122 (turning point; multiplicity 2).",
        "    y = (x + 2)\u00b3 (x \u2212 1)  \u2192  graph CROSSES the x-axis at x = \u22122 (multiplicity 3).",
        "    Shared between the two: zeros at x = \u22122 and x = 1.  End behavior is",
        "    opposite (cubic vs. quartic-like behavior near the zero).",
        "",
        "T-CHART  \u2014  what it is + how to build it",
        "    A T-chart in this lesson is a two-column reference students draw on",
        "    their packet (or you draw on the board): ODD multiplicity on one side,",
        "    EVEN multiplicity on the other. Each column lists the behavior rule and",
        "    a small sketch. It is the physical artifact of the multiplicity rule.",
        "    ODD column:  graph CROSSES the x-axis.  (Multiplicity 1 = a clean",
        "                 linear crossing; higher odd multiplicities still cross,",
        "                 just with visible curvature near the zero.)  Savvas does",
        "                 not give the curvature a separate name \u2014 it is still",
        "                 \u2018crosses.\u2019",
        "    EVEN column: graph is TANGENT to the x-axis / has a TURNING POINT",
        "                 (sign of f does not change at the zero).",
        "",
        "LAUNCH STEP-BY-STEP (12 minutes, solo teacher):",
        "    0:00 \u2013 0:15  \u201cPull out your Do Now.\u201d",
        "    0:15 \u2013 2:00  TURN-AND-TALK with a partner: \u201cWhat did you circle? Why?\u201d",
        "                Students discuss for ~90 seconds. Teacher walks two laps of",
        "                the room (10 students \u2192 ~5 pairs, easy to monitor).",
        "    2:00 \u2013 3:00  COLD-CALL a student who wrote \u201ccross everywhere\u201d on an earlier",
        "                prediction (from the Do Now entry scan). Low-stakes ask:",
        "                \u201c[Name], can you read me your one-sentence explanation?\u201d",
        "                No grade, no candy \u2014 just a sincere \u201cthanks\u201d and",
        "                \u201canyone want to build on that?\u201d. Normalize wrong answers as",
        "                the class finding the edge of what we know.",
        "    3:00 \u2013 8:00  Partners graph (x + 2)\u00b2 (x \u2212 1)  and  (x + 2)\u00b3 (x \u2212 1) in",
        "                Desmos. Teacher DOES NOT draw on the board or doc-cam.",
        "                Teacher walks the room reading student screens over their",
        "                shoulders. This is the fastest way to check 10 students.",
        "    8:00 \u2013 10:00 CLASS SHARE: ask any student (or call on someone who got",
        "                it first) to describe what (x + 2)\u00b2 did at x = \u22122 and what",
        "                (x + 2)\u00b3 did. \u201cNAME MULTIPLICITY\u201d = at the moment the class",
        "                agrees on \u201ctangent / turning point\u201d vs \u201ccrosses,\u201d the",
        "                teacher SAYS the word out loud: \u201cThis thing we\u2019re calling",
        "                a repeated factor has a name: multiplicity.\u201d Advance to",
        "                PowerPoint slide 5 (definition + example). That slide IS",
        "                the board write \u2014 you do not need to write the definition",
        "                on the whiteboard yourself.",
        "    10:00\u201312:00 Build the T-chart TOGETHER on the packet. Project the",
        "                PowerPoint slide showing the T-chart template; students fill",
        "                it in on their packet as you narrate: \u201cODD \u2192 ___? EVEN \u2192 ___?\u201d",
    ]
    lines += [
        "",
        "SAVVAS TEACHER-EDITION PROMPTS (inject during Launch, from teacher_prompts bank):",
    ]
    for tp in prompts:
        t = tp.get("type", "")
        if t in ("purposeful_question", "elicit_evidence", "habits_of_mind_reason"):
            lines.append(f"  \u2022  [{t}]  {tp.get('prompt','')}")
            lines.append(f"       Expected: {tp.get('expected_response','')}")
            lines.append(f"       When to use: {tp.get('use','')}")
    lines += [
        "",
        "ELL SCAFFOLDS (inject between Launch and Practice, from teacher_prompts):",
    ]
    for tp in prompts:
        t = tp.get("type", "")
        if t in ("ell_speaking", "ell_vocabulary"):
            lines.append(f"  \u2022  [{t}]  {tp.get('prompt','')[:240]}")
            lines.append(f"       Expected: {tp.get('expected_response','')[:180]}")
            lines.append(f"       Use: {tp.get('use','')}")
    return lines


def _practice_key_lines():
    items = qb.get_for_packet(PRACTICE_IDS)
    lines = []
    labels = ["A.", "B.", "C.", "D.", "E."]
    for label, q in zip(labels, items):
        lines.append(f"{label}  {_strip_stem(q['prompt'])}  [bank: {q['id']}, DOK {q.get('dok')}]")
        correct = _format_unicode(str(q.get("correct", "")))
        lines.append(f"    \u2705  {correct}")
        rationale = q.get("dok_rationale")
        if rationale:
            lines.append(f"    DOK rationale: {_format_unicode(rationale)}")
        lines.append("")
    lines += [
        "CIRCULATION PATTERN (solo teacher, ~10 students, ~5 pairs):",
        "  Lap 1 (min 0\u20135):  check that every pair has A started with zeros written.",
        "  Lap 2 (min 5\u201310): target pairs stuck on Item B\u2019s x\u00b2 + 9 \u2014 remind them that",
        "                     x\u00b2 + positive has NO real zero (Day 1 rule).",
        "  Lap 3 (min 10\u201315): nudge faster pairs to pick their C/D/E and start it.",
        "  Lap 4 (min 15\u201320): check the chosen item. Stuck \u2192 hint, not answer.",
        "Questions to ask (circulating, NEVER deliver the answer):",
        "  \u2022  \u201cWhich factor tells you the behavior at this zero?\u201d",
        "  \u2022  \u201cIs this multiplicity even or odd?\u201d",
        "  \u2022  \u201cWrite the exponent. Now read Rule 2 on your packet.\u201d",
    ]
    return lines


def _exit_key_lines():
    q = qb.get_for_packet([EXIT_ID])[0]
    return [
        f"[bank: {q['id']}, DOK {q.get('dok')}]",
        "DESMOS POLICY: students may use Desmos to VERIFY zeros. The behavior",
        "description (crosses / tangent / turning point) must be written in their own words.",
        "",
        _format_unicode(q["prompt"]),
        "",
        f"\u2705  {_format_unicode(str(q.get('correct', '')))}",
        "",
        "DOK rationale: " + _format_unicode(str(q.get("dok_rationale", ""))),
        "",
        "Full credit = zeros correct AND behavior using multiplicity language",
        "(\u201ccrosses\u201d for odd multiplicity; \u201ctangent\u201d, \u201cturning point\u201d, or "
        "\u201ctouches and turns back\u201d all accepted for even multiplicity \u2014 "
        "matches Savvas language).",
    ]


def _do_now_key_lines():
    q = qb.get_for_packet([DO_NOW_ID])[0]
    answers = q.get("answers", [])
    correct_idx = q.get("correct")
    correct_letter = "ABCD"[correct_idx - 1] if isinstance(correct_idx, int) and 1 <= correct_idx <= 4 else "?"
    lines = [
        f"[bank: {q['id']}, DOK {q.get('dok')}] \u2014 multiple choice (A/B/C/D)",
        _format_unicode(q["prompt"]),
        "",
    ]
    for i, ans in enumerate(answers):
        mark = "\u2705" if (i + 1) == correct_idx else "   "
        lines.append(f"  {mark}  ({'ABCD'[i]}) {_format_unicode(ans)}")
    lines += [
        "",
        f"CORRECT: {correct_letter}.  Factor by grouping: x\u00b3 + 3x\u00b2 + x + 3 = "
        "x\u00b2(x + 3) + 1(x + 3) = (x + 3)(x\u00b2 + 1).  Real zero at x = \u22123 with "
        "multiplicity 1. The factor x\u00b2 + 1 contributes no real zero.",
        "",
        "DIAGNOSTIC SCAN (on entry or during Blooket login):",
        "\u2022  Students who picked A or C \u2192 misread the graph; cold-call them in Launch.",
        "\u2022  Students who picked D \u2192 confused multiplicity with number of crossings.",
        "\u2022  Students who picked B and wrote a clean explanation \u2192 have them share",
        "  first during Launch class-share.",
    ]
    return lines


IEP_MODS = [
    "\u2022  Do Now uses a Savvas-printed graph (low reading load).",
    "\u2022  Launch is Savvas\u2019s own side-by-side comparison \u2014 visual, no standard-form factoring.",
    "\u2022  Practice is structured per-item (zeros + mult + behavior on three lines each).",
    "\u2022  Choice in Practice: A + B required, C/D/E pick one. Reduces overwhelm.",
    "\u2022  Exit structure mirrors Do Now structure \u2014 same verbs, same shape.",
    "\u2022  Desmos is explicitly allowed on the exit to verify zeros; behavior must still be written.",
]

ELL_SUPPORTS = [
    "\u2022  Sentence frame printed on the Launch T-chart: \u201cThe multiplicity is ___, "
    "so the graph ___ at x = ___.\u201d",
    "\u2022  \u201cTouches / kisses / turns / bounces\u201d all accepted for even multiplicity.",
    "\u2022  \u201cCrosses / goes through / cuts through\u201d all accepted for odd multiplicity.",
    "\u2022  ONE new academic word today: MULTIPLICITY. Bridge to multiplication via",
    "  the Savvas etymology prompt (see Launch answer key).",
    "\u2022  Savvas ELL speaking scaffold (from teacher_prompts bank) uses a 6-question",
    "  sequence: even/odd \u2192 behavior at x = \u22122 \u2192 generalize. Run in the last",
    "  3 minutes of Launch for any ELL who needs it, or during Practice as a 1:1.",
]


def build_teacher(path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    header_line(doc, "TEACHER EDITION", size=12)
    header_line(doc, "ALGEBRA 2  |  LESSON 3-5", size=14)
    ps.day_banner(doc, "2-3", "Zeros + Multiplicity (Savvas-anchored)",
                  "Solo teacher. Class of ~10. Teacher pacing + DOK "
                  "framework crosswalk + answer keys + Questions to ask "
                  "per phase. Savvas Example 2 teacher-edition content "
                  "inlined from the teacher_prompts bank.")

    add_two_col(doc, OBJECTIVES + [("CCSS", "F-IF.C.7c, A-APR.B.3, MP.3, MP.7")])
    add_two_col(doc, FRAMEWORK_HEADER)

    add_callout(doc, "\U0001F4CB  DESIGN \u2014  SAVVAS-ONLY PRACTICE DAY (no DOK 3)", [
        "Do Now \u2192 Launch \u2192 Practice \u2192 Share/Summary \u2192 Exit.",
        "[DOK 1] Blooket; Do Now Q4 (MC); Try It 2a/2b.",
        "[DOK 2] Launch investigation; Practice #14/#15/#16; Share/Summary; Exit Q5.",
        "[DOK 3] NOT on this day. Savvas\u2019s editorial choice: DOK-3 in this lesson is "
        "modeling-only (Practice #25, #27, #30). DOK-3 capstone \u2192 Combined Day 4-5.",
        "Every work item on the student packet traces to a bank ID. No fabricated tasks.",
    ])

    add_callout(doc, "\u2705  CRITERIA FOR SUCCESS  /  FORMATIVE ASSESSMENT", CRITERIA_FOR_SUCCESS)

    add_callout(doc, "\u23f1\ufe0f  REALISTIC PACING (solo teacher, ~10 students)", [
        "Total: 55 min.  Fits 55-min Tue A; 65-min F has 10 min slack.",
        "  Do Now A (LQ Q4 MC) .......... 5 min",
        "  Do Now B (Blooket login) ..... 2 min",
        "  Do Now C (Blooket game) ...... 7 min",
        "  Launch (Example 2) .......... 12 min",
        "  Practice (A, B + 1 of C/D/E). 20 min",
        "  Share/Summary ................ 4 min",
        "  Exit (LQ Q5) ................. 5 min",
        "RELEASE VALVES (ordered):",
        "  1. Cut Practice to A + B only (save 8 min for stuck pairs).",
        "  2. Cut Launch from 12 \u2192 10 min (skip extended class share).",
        "  3. Cut Share/Summary to 3 min.",
        "  NEVER cut Exit.  Only artifact that documents learning.",
        "65-min F slack: all five Practice items, run ELL speaking scaffold during Launch.",
    ])

    add_callout(doc, "[Do Now A]  [DOK 1] Savvas Lesson Quiz Q4 (MC)  (5 min)", [
        "Students receive the Do Now sheet (Q4 stem + graph + choices) on entry. "
        "Silent, 5 min, pencil only. Multiple-choice with one short written explanation.",
    ])
    add_callout(doc, "\u2705  DO NOW / ANSWER KEY", _do_now_key_lines())

    add_callout(doc, "[Do Now B+C]  [DOK 1] Blooket \u2014 Rule Recall ONLY", BLOOKET_SCOPE)

    add_callout(doc, "[Launch]  [DOK 2] Savvas Example 2 \u2014 Name Multiplicity  (12 min)", [
        "Solo teacher. Students graph on Desmos \u2014 teacher does NOT draw on board.",
        "Teacher walks the room reading screens; leverages the small class size.",
    ])
    add_callout(doc, "\u2705  LAUNCH / ANSWER KEY + SAVVAS TE CONTENT + ELL SCAFFOLDS",
                _launch_key_lines())

    add_callout(doc, "[Practice]  [DOK 2] Savvas Try Its + Practice  (20 min)", [
        "Five Savvas items (bank IDs: " + ", ".join(PRACTICE_IDS) + ").",
        "Required: A + B. Choice: one of C/D/E. Fast finishers: all five, plus",
        "the \u201cbuild a polynomial with a simple crossing, a tangent/turning "
        "point, and a higher-odd-multiplicity crossing at prescribed zeros\u201d",
        "board extension.",
    ])
    add_callout(doc, "\u270F\uFE0F  PRACTICE / ANSWER KEY + CIRCULATION PATTERN",
                _practice_key_lines())

    add_callout(doc, "[Share/Summary]  [DOK 2] Synthesis + Self-Rating  (4 min)", [
        "Teacher script (4 min flat):",
        "1.  Essential Question callback (45s): \u201cHow does factored form tell you "
        "everything you need to sketch?\u201d",
        "2.  Language Objective check (30s): \u201cWho used \u2018the multiplicity is ___, "
        "so the graph ___\u2019?\u201d Call on two students.",
        "3.  Self-rating circle (60s). Teacher scans.",
        "4.  Preview Day 4-5 (30s): \u201cTomorrow \u2014 complex zeros + Savvas storage-box "
        "volume problem (DOK 3).\u201d",
        "5.  Transition (30s): \u201cExit ticket. LQ Q5. 5 minutes. Desmos OK for zeros.\u201d",
        "Questions to ask: \u201cWhich Criterion did you rate \u2018partly\u2019? What\u2019s the blocker?\u201d",
    ])
    add_callout(doc, "\U0001F501  SHARE / SUMMARY (student-facing)", SHARE_SUMMARY_STUDENT)

    add_callout(doc, "[Exit Ticket]  [DOK 2] Savvas Lesson Quiz Q5  (5 min)", [
        "DESMOS POLICY: YES, students may use Desmos to VERIFY the zeros.",
        "Behavior at each zero must still be written in their own words using",
        "multiplicity language.",
        "Full credit = zeros correct AND behavior description correct.",
    ])
    ps.summary_exit_box(doc, "\U0001F4E4  EXIT TICKET (student-facing)", _exit_lines())
    add_callout(doc, "\u2705  EXIT TICKET / ANSWER KEY", _exit_key_lines())

    add_callout(doc, "\U0001F4F7  WHAT TO COLLECT", [
        "1.  Do Now sheets (Q4 MC + 1-sentence explanation).",
        "2.  Practice pages (circulate + photograph).",
        "3.  Exit tickets (Q5, every student).",
        "4.  Share/Summary self-ratings.",
    ])
    add_callout(doc, "\U0001F9E9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS", IEP_MODS)
    add_callout(doc, "\U0001F30D  SUPPORTS FOR ELL STUDENTS", ELL_SUPPORTS)

    doc.save(path)


if __name__ == "__main__":
    build_do_now("Day_23_Do_Now.docx")
    build_student("Day_23_Student_Packet.docx")
    build_teacher("Day_23_Teacher_Packet.docx")
    print("Built Day_23_Do_Now.docx, Day_23_Student_Packet.docx, Day_23_Teacher_Packet.docx")
