"""Generate Lesson 3-5 Period 4 (Thursday) assessment-day wrapper materials.

The Topic 3 assessment itself is `a2topic3assess.docx` (not generated here).
This script produces the surrounding class-period artifacts:

  - L35_P4_Do_Now.docx        — mixed-review warm-up (NOT a review lesson; just a
                                 gentle warm-up to prime the brain for the test)
  - L35_P4_Exit_Ticket.docx   — post-assessment self-reflection sheet
  - L35_P4_Teacher_Notes.docx — timing + proctoring notes for the period

Phases (55 min base / 65 min Period F):
  0– 3  Do Now warm-up (gentle, no new content)                       [DOK 1]
  3– 5  Launch: test format reminder (calculator, time, flag Qs)
  5–40  Assessment: 11 questions                                      [varies]
 40–55  Share/Reflect + Exit self-reflection + buffer
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_IDS = ["3-5-savvas-q13", "3-5-savvas-q14"]
_ALL_IDS = DO_NOW_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 3-5 · Period 4"
LESSON_TITLE = "Topic 3 Assessment Day"


def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.2):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    ps.render_prompt(prompt_p, q["prompt"], font_size=11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            run = line.add_run(f"  ({chr(65+i)})  " + ps.latex_to_unicode(a))
            run.font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def build_do_now(path):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.75); s.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 4, "Do Now — Gentle Warm-Up", subtitle=LESSON_LABEL)

    add_callout(doc, "🌤️  BEFORE THE ASSESSMENT",
                ["This is just a warm-up. Don't overthink it. Pick one problem you feel confident about and write it clearly. Breathe. You know this."],
                color_hex="FFF2CC")

    for i, qid in enumerate(DO_NOW_IDS, 1):
        q = qb.get_for_packet([qid])[0]
        bank_item_block(doc, q, label=f"Warm-Up #{i}", space_after_inches=1.6)

    doc.save(path)


def build_exit(path):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.75); s.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 4, "Exit Ticket — Post-Assessment Reflection",
                  subtitle=LESSON_LABEL)

    ps.summary_exit_box(doc, "SELF-REFLECTION (not graded)",
                        [
                            "1. Which question was hardest? Q#___. Why was it hard?",
                            "",
                            "2. Which question was easiest? Q#___. What made it click?",
                            "",
                            "3. One thing I want to strengthen before Topic 4 starts:",
                            "",
                            "4. On a scale 1-5, how confident do I feel about Topic 3 overall?",
                            "   (1 = very unsure   5 = very confident)       ___ / 5",
                        ])

    add_callout(doc, "📣  HONESTY HELPS",
                ["Your self-reflection tells the teacher what to review before Topic 4. No grade, no shame — just real info. Be honest."],
                color_hex="D9EAD3")

    doc.save(path)


def build_teacher_notes(path):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.5); s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)
    ps.day_banner(doc, 4, LESSON_TITLE, subtitle=LESSON_LABEL + " · Teacher Notes")

    ps.framework_phase_header(
        doc,
        phase="Do Now (Gentle Warm-Up)",
        framework_tag="low-stakes prime",
        dok="1",
        minutes=3,
        teacher_does=["Hand out warm-up sheet at door.",
                      "Quick no-stakes check; do not stop to re-teach.",
                      "Collect warm-up sheets (or leave on desks) at minute 3."],
        students_do=["Pick one warm-up problem. Write confidently.",
                     "Breathe. Settle in."],
        questions_to_ask=["(None — this is priming, not teaching.)"],
        adult_role="Solo teacher: low-key, calm energy. Don't hover.",
    )

    ps.framework_phase_header(
        doc,
        phase="Launch (Test Format Reminder)",
        framework_tag="logistics",
        dok="—",
        minutes=2,
        teacher_does=[
            "Announce: 11 questions, calculator allowed, expect 35 min.",
            "\"If stuck, flag it and move on. Come back at the end.\"",
            "No phones. Pencil only.",
        ],
        students_do=["Clear desk except pencil + calculator.",
                     "Raise hand if a question is ambiguous (wording, not math)."],
        questions_to_ask=["Any logistical questions before we start?"],
        adult_role="Solo teacher: reset room tone. Start timer visibly.",
    )

    ps.framework_phase_header(
        doc,
        phase="Assessment (Topic 3)",
        framework_tag="proctoring",
        dok="1-3 mix",
        minutes=35,
        teacher_does=[
            "Circulate silently every 10 min.",
            "At minute 20: \"15 min remaining. Flag anything you need to skip.\"",
            "At minute 30: \"5 min. Finalize your answers.\"",
        ],
        students_do=[
            "Work independently. No partner talk.",
            "If stuck >2 min on a question, flag and move on.",
        ],
        questions_to_ask=["(Silent proctoring — no teaching during assessment.)"],
        adult_role="Solo teacher: proctor, time, say nothing content-related.",
    )

    add_callout(doc, "📋  ASSESSMENT COVERAGE MAP (for your reference only)",
                [
                    "Q#1 Remainder Theorem — Lesson 3-4",
                    "Q#2 Cubic with conjugate irrational pair — Lesson 3-5 (Tue bridge prompt)",
                    "Q#3 Binomial Theorem — Lesson 3-3",
                    "Q#4 Simplify to standard form — Lesson 3-2/3",
                    "Q#5 Zeros + end behavior from graph — Lesson 3-5 (Tue)",
                    "Q#6 Lucy's tray volume model — Lesson 3-5 (Wed, parallel to #27)",
                    "Q#7 Polynomial identities factoring — Lesson 3-3",
                    "Q#8 Cube vs sphere growth — Lesson 3-1/2",
                    "Q#9 Synthetic division — Lesson 3-4",
                    "Q#10 Identify zeros — Lesson 3-5 (Tue)",
                    "Q#11 Graph transformations of polynomial — Lesson 3-1/2",
                ],
                color_hex="DEEBF7")

    ps.framework_phase_header(
        doc,
        phase="Share / Reflect",
        framework_tag="brief debrief",
        dok="—",
        minutes=5,
        teacher_does=[
            "Collect tests. Keep tone positive regardless of class vibes.",
            "Hand out self-reflection sheet.",
            "\"One sentence: how did that feel?\" — volunteer only.",
        ],
        students_do=["Complete self-reflection sheet.",
                     "Optional: share one-sentence vibe check."],
        questions_to_ask=["How did that feel overall? (Volunteer only.)"],
        adult_role="Solo teacher: warm, non-judgmental. Do not give away answers.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit",
        framework_tag="self-reflection",
        dok="—",
        minutes=5,
        teacher_does=["Collect self-reflection at door.",
                      "Tonight: scan confidence scores. Plan next-day Do Now if class averaged <3."],
        students_do=["Complete all 4 reflection items honestly."],
        questions_to_ask=["Which Q was hardest/easiest?",
                          "What do you want to strengthen?"],
        adult_role="Solo teacher: collect, don't read in front of them.",
    )

    add_callout(doc, "🕐  PERIOD F (65-MIN) NOTE",
                [
                    "Period F has 10 extra minutes on Thursday. Use:",
                    "• 5 min: post-assessment math-chat (students compare strategies on ONE problem they found interesting; teacher picks, no right-answer pressure).",
                    "• 5 min: Topic 4 preview (pencil out: \"Next week we start inverse variation. Think about: what does 'as x doubles, y halves' mean?\").",
                    "• Do NOT reteach assessment content. Period A doesn't get this extension.",
                ],
                color_hex="FFD9E0")

    doc.save(path)


if __name__ == "__main__":
    build_do_now("L35_P4_Do_Now.docx")
    build_exit("L35_P4_Exit_Ticket.docx")
    build_teacher_notes("L35_P4_Teacher_Notes.docx")
    print("Built L35_P4_Do_Now.docx, L35_P4_Exit_Ticket.docx, L35_P4_Teacher_Notes.docx")
