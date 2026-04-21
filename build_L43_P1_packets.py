"""Generate Lesson 4-3 Period 1 materials — Klimsara-adapted, no Blooket.

Design: Introduce rational expressions (equivalent forms & simplification). Period 1
lays the conceptual foundation — factor, cancel common factors only, state domain
restrictions. Assessment items for 4-3 come from later periods; P1 is foundation,
not assessment-critical.

Phases (55 min base; Period F 65-min on Fri; Period A 65-min on Mon):
  0–  5   Do Now    : 4-3 Model & Discuss starter — notice on a graph with a  [DOK 2]
                     "skipped" point (domain restriction bridge)
  5– 17   Launch    : teacher models Example 1 (equivalent rational expressions) [DOK 1-2]
                     + Example 2 (simplify by factoring and canceling)
 17– 50   Explore   : Try-It 1 + Try-It 2 + Practice #20/#22/#24/#17/#19 TPS   [DOK 1-2]
 50– 55   Share/Sum : sentence-frame consolidation on domain restrictions        [DOK 2]
 55– 58   Exit      : summary recap fill-in                                     [DOK 1-2]
 (back)   Optional Reinforcement: Practice #25
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_ID      = "4-3-savvas-model-discuss-lesson-4-3-launch"
# Explore sequence: Try-It 1 -> #20 -> #22 -> #24 -> Try-It 2 -> #17 -> #19.
EXPLORE_IDS = [
    "4-3-savvas-try-it-1-lesson-4-3",
    "4-3-savvas-q20",
    "4-3-savvas-q22",
    "4-3-savvas-q24",
    "4-3-savvas-try-it-2-lesson-4-3",
    "4-3-savvas-q17",
    "4-3-savvas-q19",
]
REINFORCE_IDS  = ["4-3-savvas-q25"]

_ALL_IDS = [DO_NOW_ID] + EXPLORE_IDS + REINFORCE_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 4-3 · Period 1"
LESSON_TITLE = "Rational Expressions — Equivalent & Simplify"

OBJECTIVES = [
    ("Math Objective",
     "Write equivalent rational expressions over a restricted domain, identify "
     "domain restrictions from factored denominators, and simplify rational "
     "expressions by canceling common factors."),
    ("Language Objective",
     "Explain, in writing and orally using sentence frames, why a value is "
     "excluded from the domain and why we cancel factors but not terms."),
    ("Essential Understanding",
     "A rational expression equals an equivalent simpler form ONLY when both "
     "share the same domain — canceling a factor never erases the restriction "
     "it imposed."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Write equivalent rational expressions, simplify, and identify domain "
     "restrictions from polynomial denominators."),
    ("Essential Question",
     "Why does canceling a common factor not remove its domain restriction?"),
    ("Materials",
     "Student packet, pencil, highlighter. Teacher displays worked Example 1 "
     "then gradual-release into Example 2."),
]


# ── Helpers ──────────────────────────────────────────────────────────────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.0):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    prompt_p.add_run(q["prompt"]).font.size = Pt(11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            line.add_run(f"  ({chr(65+i)})  {a}").font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, for_teacher=False):
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, LESSON_TITLE,
                  subtitle=LESSON_LABEL + (" · Teacher Edition" if for_teacher else ""))
    t = doc.add_table(rows=len(OBJECTIVES), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(OBJECTIVES):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    t = doc.add_table(rows=len(FRAMEWORK_HEADER), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(FRAMEWORK_HEADER):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()


# ── Do Now ─────────────────────────────────────────────────────────────────────────

def build_do_now(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 1, "Do Now — " + LESSON_TITLE, subtitle=LESSON_LABEL)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="notice & wonder",
        dok="2",
        minutes=5,
        teacher_does=["Project the graph on screen or sketch on board.",
                      "Don't answer questions — this is exploration.",
                      "Collect sheets at 5 min sharp."],
        students_do=["Look at the graph. Notice a point seems to be missing.",
                     "Write 1 thing you NOTICE and 1 thing you WONDER.",
                     "Bonus: write one sentence about what the equation might be doing at that x-value."],
        questions_to_ask=[
            "What is happening at the missing point?",
            "What must be true about the equation at that x-value?",
            "Could the graph ever have a y-value at that x?",
        ],
        adult_role="Solo teacher: circulate silently. No hints. Students must struggle productively for a moment.",
    )

    add_callout(doc, "\U0001f534  STARTER — A Graph with a Missing Point",
                [
                    "Starter: The graph below shows a line that suddenly ‘skips’ a point. What must be true about the equation for that x-value? Write one sentence.",
                    "",
                    "A) Notice: Write something you notice about this graph.",
                    "",
                    "B) Wonder: Write something you wonder about what causes a graph to skip a point.",
                    "",
                    "C) Sentence: What must be true about the equation at the skipped x-value?",
                ],
                color_hex="FFF2CC")

    doc.save(path)


# ── Student packet ───────────────────────────────────────────────────────────────────────────────

def build_student(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=False)

    ps.teal_section(doc, "LAUNCH — Rational Expressions (teacher models on screen)", [])
    add_callout(doc, "\U0001f4d8  EQUIVALENT + SIMPLIFY RULES (keep printed on your page)",
                [
                    "1. Factor the numerator AND denominator completely.",
                    "2. Cancel only common FACTORS (never terms across + or −).",
                    "3. State domain restrictions: every zero of the ORIGINAL denominator is excluded — including factors you canceled.",
                ],
                color_hex="D9EAD3")

    add_callout(doc, "⚠️  COMMON ERROR",
                [
                    "Do NOT cancel terms that are added or subtracted across a fraction bar.",
                    "You can only cancel factors that multiply across the entire numerator or denominator.",
                    "Test: factor completely FIRST, then cancel only identical factor pairs.",
                ],
                color_hex="F4CCCC")
    doc.add_paragraph()

    ps.teal_section(doc, "EXPLORE — Think-Pair-Share (33 min)",
                    ["Work with a partner. Use the Equivalent + Simplify Rules above for every problem. Move through items in order — each builds on the last."])

    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 1 — Write an equivalent rational expression",
        "Practice #20 — Identify equivalent expressions",
        "Practice #22 — Identify domain restrictions",
        "Practice #24 — Apply domain restrictions",
        "Try It 2 — Simplify by factoring and canceling",
        "Practice #17 — Simplify",
        "Practice #19 — Simplify with domain restrictions (extension)",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=label, space_after_inches=1.4)

    # Share / Summary
    ps.teal_section(doc, "SHARE / SUMMARY (5 min)", [])
    ps.sentence_frame_box(doc, [
        "Pair share: “The factored form is ___, so the domain excludes x = ___. After canceling the factor ___, the simplified expression is ___ with the same domain restriction.”",
        "Whole class: one pair reads their sentence. Class signals thumbs up/sideways.",
    ])

    ps.summary_exit_box(doc, "EXIT TICKET — Summary Recap",
                        [
                            "To simplify (x²+3x−10)/(x²−4) I first factor to ___.",
                            "The restricted domain values are x ≠ ___.",
                            "After canceling the factor ___, my simplified expression is ___.",
                            "One biggest thing I learned today: ________________________",
                        ])

    doc.add_page_break()
    ps.teal_section(doc, "OPTIONAL REINFORCEMENT (not collected)",
                    ["If you finish Explore early or want extra practice before Period 2."])

    # Model & Discuss extension (demoted from Explore to optional)
    add_callout(doc, "\U0001f52d  OPTIONAL · Model & Discuss",
                [
                    "Return to the Do Now graph. With your partner:",
                    "",
                    "B) Use Structure. Why can the graph never have a point at the skipped x-value, no matter what simplified form we write?",
                    "",
                    "C) Examine at least two rational expressions you simplified in Explore. For each, explain in your own words why the domain restriction from the canceled factor still applies.",
                ],
                color_hex="CFE2F3")

    reinforce = qb.get_for_packet(REINFORCE_IDS)
    for i, q in enumerate(reinforce, 1):
        bank_item_block(doc, q, label=f"Optional #{i}  (Savvas Practice)",
                        space_after_inches=1.8)

    doc.save(path)


# ── Teacher packet ──────────────────────────────────────────────────────────────────────────────────────

def build_teacher(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=True)

    add_callout(doc, "\U0001f4cc  PRE-FLIGHT NOTE (read before class)",
                [
                    "Explore is 7 bank items: Try-It 1 → #20 → #22 → #24 → Try-It 2 → #17 → #19 (extension). Paced for ~4-5 min per item.",
                    "If pairs finish early, push them to the Optional Reinforcement: Model & Discuss Parts B + C (verbal/TPS), then Practice #25.",
                    "If pace slows, cut #19 (extension) — it's the DOK-2 stretch; the core DOK-1 arc (Try-It 1 → #20/#22/#24 → Try-It 2 → #17) is the minimum viable set.",
                ],
                color_hex="FFD9E0")

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="notice & wonder",
        dok="2",
        minutes=5,
        teacher_does=["Project the graph (or use printed Do Now sheet).",
                      "Circulate silently. Do NOT teach during this phase.",
                      "Collect at 5 min sharp."],
        students_do=["Write 1 notice + 1 wonder.",
                     "Write one sentence about what the equation must be doing at the skipped x-value."],
        questions_to_ask=["(None during the phase — teacher harvests answers in Launch.)"],
        adult_role="Solo teacher: circulate silently. Resist the urge to give hints.",
    )
    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label=f"Do Now ({q['id']})", space_after_inches=0.3)
    add_callout(doc, "✅  ANTICIPATED STUDENT RESPONSES",
                [
                    "NOTICE (typical): \"There's a hole in the graph.\" \"The line has a gap.\" \"One point is missing.\"",
                    "WONDER (typical): \"Why is there a hole?\" \"What happens at that x?\" \"Can you have an equation that does this?\"",
                    "SENTENCES: students might write \"the denominator equals zero at that x\" or \"the equation is undefined there.\"",
                    "KEY MOVE in Launch: surface the insight that when the denominator = 0, the expression is undefined — this is a domain restriction.",
                ],
                color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Launch",
        framework_tag="teacher models Example 1 + Example 2",
        dok="1-2",
        minutes=12,
        teacher_does=[
            "Bridge from Do Now: \"You noticed the graph skips a point. That happens because the denominator equals zero at that x-value — we call it a domain restriction.\"",
            "Project Example 1 (Savvas Topic 4-3). Think aloud: factor numerator and denominator, identify what x-values make the denominator zero, write the equivalent expression with restriction stated.",
            "Project Example 2. Think aloud: factor completely, cancel the common factor, but keep the domain restriction from the canceled factor in the final answer.",
            "Pause 1× for 30-sec partner-check: \"What x-value is excluded in Example 1?\"",
            "Do NOT solve Try-Its. Release promptly at minute 17.",
        ],
        students_do=[
            "Watch silently through Example 1.",
            "During partner-check: identify the excluded x-value.",
            "Copy the Equivalent + Simplify Rules card into their page margin.",
            "Note the COMMON ERROR warning: cancel only factors, never terms.",
        ],
        questions_to_ask=[
            "\"What values make the denominator zero?\"",
            "\"Can I cancel this factor? Is it a factor of the entire numerator?\"",
            "\"After canceling, does the domain restriction disappear?\"",
        ],
        adult_role="Project, think aloud. Do NOT give students Try-Its to work until minute 17.",
    )
    add_callout(doc, "\U0001f3a4  TEACHER SCRIPT",
                [
                    "1. “From the Do Now, you noticed the graph skips a point. That x-value is excluded from the domain because the denominator equals zero there.”",
                    "2. Example 1 walk-through: “Watch me. I factor top and bottom. Any x that makes the denominator zero is restricted — I write x ≠ that value next to my answer.”",
                    "3. Example 2 walk-through: “Now I simplify. I factor, find the matching factor in numerator and denominator, and cancel it. But — critical — the restriction from that canceled factor STAYS. Canceling doesn’t erase the hole.”",
                    "4. Common-error flag: “Don’t cancel terms that are added or subtracted. Factor FIRST, then cancel only matching factor pairs.”",
                    "5. Release: “Now you try. Try-It 1 asks for an equivalent expression. Try-It 2 asks you to simplify completely.”",
                ],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Explore",
        framework_tag="Think-Pair-Share + Practice",
        dok="1-2",
        minutes=33,
        teacher_does=[
            "5 laps circulation across 33 min (~6-7 min each).",
            "Lap 1: Try-It 1 — press pairs to factor before writing the equivalent form.",
            "Lap 2: #20 + #22 — check that pairs identify domain restrictions correctly.",
            "Lap 3: #24 + Try-It 2 — press pairs to factor completely before canceling.",
            "Lap 4: #17 — verify pairs state the domain restriction alongside the simplified form.",
            "Lap 5: #19 (extension) — stretch. Cut if pace slows.",
            "Do not give answers. Use sentence frame to press justification.",
        ],
        students_do=[
            "Pairs move in order through 7 items: 2 Try-Its + 5 Practice.",
            "For each item: factor completely FIRST, then cancel, then state restriction.",
            "For Practice #19 (simplify with domain restrictions): show all steps.",
            "Justify with sentence frame before writing.",
        ],
        questions_to_ask=[
            "Did you factor completely before canceling?",
            "What x-values make the denominator zero?",
            "After simplifying, what is the domain restriction?",
            "Can you use the sentence frame to explain your answer?",
            "For #19: what is the domain restriction from the canceled factor?",
        ],
        adult_role="Solo teacher: 5 laps. Press pairs to justify using the rule card. No answers.",
    )
    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 1",
        "Practice #20 — Identify equivalent",
        "Practice #22",
        "Practice #24",
        "Try It 2",
        "Practice #17 — Simplify",
        "Practice #19 — Simplify with domain restrictions",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=f"{label} ({q['id']})", space_after_inches=0.3)
        add_callout(doc, "✅  ANSWER / ANTICIPATED TALK",
                    [q.get("notes") or "(no answer key — verify before class)"],
                    color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Share / Summary",
        framework_tag="academic conversation",
        dok="2",
        minutes=5,
        teacher_does=[
            "Cold-call 2 pairs to read their sentence-frame sentence.",
            "Write one student-generated sentence on board.",
            "Preview P2: \"Next class we’ll multiply and divide rational expressions — the same factor-and-cancel logic applies.\"",
        ],
        students_do=[
            "Presenting pair reads their sentence.",
            "Class signals thumbs up/sideways.",
            "Note: next class extends this to multiplication and division.",
        ],
        questions_to_ask=[
            "What’s a sentence that captures today’s rule about domain restrictions?",
            "Why can’t we just ignore the restriction after we cancel?",
        ],
        adult_role="Cold-call 2 pairs. Keep it brief — 5 min.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit Ticket",
        framework_tag="summary recap (not graded)",
        dok="1-2",
        minutes=2,
        teacher_does=[
            "Collect at door.",
            "Scan the fill-in items — misconceptions about domain restrictions tell you what to re-hit in P2 Do Now.",
        ],
        students_do=[
            "Complete the fill-in stems.",
            "Be honest about confusion.",
        ],
        questions_to_ask=[
            "Did students correctly identify domain restrictions from the original denominator?",
        ],
        adult_role="Collect. No grade.",
    )
    add_callout(doc, "\U0001f4cb  EXIT TICKET — Student Form",
                [
                    "To simplify (x²+3x−10)/(x²−4) I first factor to ___.",
                    "The restricted domain values are x ≠ ___.",
                    "After canceling the factor ___, my simplified expression is ___.",
                    "One biggest thing I learned today: ________________________",
                ],
                color_hex="FFF2CC")

    doc.add_page_break()
    add_callout(doc, "\U0001f550  PERIOD A & F — 65-MIN CLASS NOTES",
                [
                    "Both sections get 65 min for P1. Use the extra 10 min on Explore, NOT Launch.",
                    "Suggested add: project Savvas Practice #22 and #24 on screen for 2 more TPS rounds if pairs finish fast.",
                    "Or: extend Model & Discuss Parts B + C with \"write one more rational expression that has the same restriction — then simplify it.\"",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "\U0001f9e9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS",
                [
                    "• Provide the Equivalent + Simplify Rules card as a sticker/cut-out on their desk.",
                    "• For Try-It 1, allow students to factor with a reference list of factoring patterns.",
                    "• Accept verbal sentence-frame completions in lieu of written.",
                ],
                color_hex="E2F0D9")
    add_callout(doc, "\U0001f30d  SUPPORTS FOR ELL STUDENTS",
                [
                    "• Sentence frames printed on student packet (don’t skip).",
                    "• Word card: “domain restriction” = an x-value that makes the denominator zero; excluded from the domain.",
                    "• “cancel” = divide out a common factor from numerator and denominator.",
                    "• “factor” = rewrite as a product of simpler expressions.",
                    "• Pair ELL students with English-dominant partners for TPS.",
                ],
                color_hex="DEEBF7")

    doc.save(path)


# ── Main ─────────────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_do_now("L43_P1_Do_Now.docx")
    build_student("L43_P1_Student_Packet.docx")
    build_teacher("L43_P1_Teacher_Packet.docx")
    print("Built L43_P1_Do_Now.docx, L43_P1_Student_Packet.docx, L43_P1_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L43_P1_Visuals_Checklist.md",
                              title="L43 P1 — Visuals Checklist")
