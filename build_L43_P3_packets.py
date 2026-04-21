"""Generate Lesson 4-3 Period 3 materials — Klimsara-adapted, no Blooket.

Design: Rational Expression Modeling + Assessment-Critical Simplify = the
**assessment-critical** period. Topic 4 8-Q assessment item #2 (identify domain
restrictions from a cubic denominator) is directly rehearsed by Practice #19.
Ex 2 (simplify with domain restriction) and Ex 6 (SA/V modeling) are paired in
Launch. P3 is pure DOK-2 mastery (no DOK-3 driver — P2 owned the spine).
Students leave P3 able to simplify rational expressions with explicit domain
restrictions and interpret simplified rational models in context.

Phases (55 min base; Period F 65-min):
  0–  5   Do Now    : Practice #11 — identity under domain restriction (P2 bridge) [DOK 2]
  5– 17   Launch    : teacher models Example 2 (simplify) + Example 6 (SA/V)       [DOK 2]
 17– 52   Explore   : Try-It 6 + #14 + #22 + #34 + #37 + ⭐ #19                    [DOK 2]
                     (#19 is assessment-critical: domain restrictions, LEHS Q#2)
 52– 57   Share/Sum : Concept Summary return + LEHS 8-Q assessment preview         [DOK 2]
 57– 58   Exit      : pre-assessment confidence check                              [DOK 1-2]
 (back)   Optional  : Practice #39 (SAT/ACT MC) — extra stretch
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor

import qb
import packet_styles as ps

TABLE_STYLE = "Table Grid"

DO_NOW_ID   = "4-3-savvas-q11"
LAUNCH_IDS  = [
    "4-3-savvas-example-2-lesson-4-3",   # Ex 2 — assessment-critical simplify reprise (LEHS Q#4)
    "4-3-savvas-example-6-lesson-4-3",   # Ex 6 — modeling anchor (SA/V ratio)
]
EXPLORE_IDS = [
    "4-3-savvas-try-it-6-lesson-4-3",
    "4-3-savvas-q14",
    "4-3-savvas-q22",
    "4-3-savvas-q34",
    "4-3-savvas-q37",
    "4-3-savvas-q19",        # ⭐ ASSESSMENT-CRITICAL: domain restrictions — LEHS Q#2
]
REINFORCE_IDS = ["4-3-savvas-q39"]

_ALL_IDS = [DO_NOW_ID] + LAUNCH_IDS + EXPLORE_IDS + REINFORCE_IDS
qb.get_for_packet(_ALL_IDS)

LESSON_LABEL = "Lesson 4-3 · Period 3"
LESSON_TITLE = "Rational Expression Modeling + Assessment-Critical Simplify"

OBJECTIVES = [
    ("Math Objective",
     "Model real-world ratios (surface area to volume, area ratios) as rational "
     "expressions; re-master simplification with domain restrictions on "
     "assessment-aligned items."),
    ("Language Objective",
     "Interpret a simplified rational model in context, naming what x, the "
     "numerator, and the denominator each represent."),
    ("Essential Understanding",
     "The same simplify-and-restrict procedure handles both abstract and modeling "
     "problems — units and context never change the math, only the interpretation."),
]

FRAMEWORK_HEADER = [
    ("Topic Goals",
     "Simplify rational expressions with explicit domain restrictions "
     "(assessment-critical); model ratios of geometric quantities."),
    ("Essential Question",
     "How do simplified rational expressions help us compare or scale real-world "
     "ratios, and what domain values are never allowed?"),
    ("Materials",
     "Student packet, pencil. Teacher pairs Example 2 (simplify) with Example 6 "
     "(SA/V) in Launch. Practice #19 is the assessment rehearsal item — "
     "Topic 4 LEHS Q#2."),
]


# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell(cell, lines, bold_first=False):
    if isinstance(lines, str):
        lines = [lines]
    cell.text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        if bold_first and i == 0:
            run.bold = True


def add_callout(doc, title, body_lines, color_hex="CFE2F3"):
    t = doc.add_table(rows=1, cols=1)
    t.style = TABLE_STYLE
    cell = t.cell(0, 0)
    ps.shade_cell(cell, color_hex)
    set_cell(cell, [title] + list(body_lines), bold_first=True)
    doc.add_paragraph()


def bank_item_block(doc, q, *, label=None, space_after_inches=1.0):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0B, 0x5C, 0x7B)
    prompt_p = doc.add_paragraph()
    prompt_p.add_run(q["prompt"]).font.size = Pt(11)
    if q.get("answers"):
        for i, a in enumerate(q["answers"]):
            line = doc.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.3)
            line.add_run(f"  ({chr(65+i)})  {a}").font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Inches(space_after_inches)


def add_header_block(doc, *, for_teacher=False):
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 3, LESSON_TITLE,
                  subtitle=LESSON_LABEL + (" · Teacher Edition" if for_teacher else ""))
    t = doc.add_table(rows=len(OBJECTIVES), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(OBJECTIVES):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()
    t = doc.add_table(rows=len(FRAMEWORK_HEADER), cols=2)
    t.style = TABLE_STYLE
    for r, (label, body) in enumerate(FRAMEWORK_HEADER):
        t.cell(r, 0).width = Inches(1.6)
        set_cell(t.cell(r, 0), label, bold_first=True)
        set_cell(t.cell(r, 1), body)
    doc.add_paragraph()


# ── Do Now ───────────────────────────────────────────────────────────────

def build_do_now(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    ps.running_header_name_block(doc, label_right="BLOCK")
    ps.day_banner(doc, 3, "Do Now — " + LESSON_TITLE, subtitle=LESSON_LABEL)

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from P2",
        dok="2",
        minutes=5,
        teacher_does=["Hand out at door.", "Circulate 3 min.",
                      "Cold-call: \"What does the domain restriction tell us about the original expression?\""],
        students_do=["Explain why an expression equals 1 over its own domain.",
                     "Write the restriction in set notation."],
        questions_to_ask=[
            "Why can't x equal those excluded values?",
            "Today we model ratios — predict: how does domain restriction apply to SA/V problems?",
        ],
        adult_role="Solo teacher: circulate silently.",
    )

    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label="Do Now — Identity under Domain Restriction",
                    space_after_inches=2.0)

    add_callout(doc, "\U0001f52d  PREDICTION (spend 30 sec)",
                ["Before we model SA/V together, write what you think the ratio (Surface Area)/(Volume) will simplify to for a cube with side x. (No wrong answers — just predict.) We’ll compare in Launch."],
                color_hex="FCE5CD")

    doc.save(path)


# ── Student packet ───────────────────────────────────────────────────────

def build_student(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=False)

    ps.teal_section(doc, "LAUNCH — Simplify + Modeling (teacher models Ex 2 + Ex 6)", [])

    add_callout(doc, "\U0001f4d8  DOMAIN RULES",
                [
                    "1. Domain = all reals EXCEPT zeros of every denominator in every step.",
                    "2. Write the simplified form AND the restriction list together.",
                    "3. Canceled factors still exclude their zeros.",
                ],
                color_hex="D9EAD3")

    add_callout(doc, "\U0001f4d8  MODELING RULES",
                [
                    "1. Identify the ratio in words first (what / what).",
                    "2. Translate each quantity to a rational expression, with units.",
                    "3. Simplify, restrict domain, then interpret the simplified ratio in context.",
                ],
                color_hex="FFF2CC")

    doc.add_paragraph()

    ps.teal_section(doc, "EXPLORE — Think-Pair-Share (35 min)",
                    ["Work with a partner. Practice #19 is assessment-critical (domain restrictions) — plan ~8 min for it."])

    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 6 — Simplify a rational expression",
        "Practice #14 — Simplify with factoring",
        "Practice #22 — Simplify: identify domain",
        "Practice #34 — Modeling ratio problem",
        "Practice #37 — Modeling: interpret in context",
        "Practice #19  ⭐ ASSESSMENT CRITICAL — Domain restrictions (LEHS Q#2)",
    ]
    for label, q in zip(labels, items):
        bank_item_block(doc, q, label=label, space_after_inches=1.5)

    ps.teal_section(doc, "SHARE / SUMMARY (5 min)", [])
    ps.sentence_frame_box(doc, [
        "Pair share (#19): “This ratio compares ___ to ___. The rational model is ___ / ___. "
        "After factoring and canceling, it simplifies to ___ with x ≠ ___, which means in "
        "context that ___.”",
        "Whole class: one pair reads #19. Class signals thumbs up/sideways.",
    ])

    add_callout(doc, "\U0001f4ca  CONCEPT SUMMARY (your reference for Topic 4 assessment)",
                [
                    "SIMPLIFYING RATIONAL EXPRESSIONS:",
                    "  — FACTOR numerator and denominator completely.",
                    "  — DOMAIN: exclude ALL zeros of any denominator (original AND intermediate steps).",
                    "  — CANCEL common factors (they still exclude their zeros from the domain).",
                    "  — Write simplified form WITH restriction: e.g., -(x+2)/(x+5), x ≠ 2, x ≠ -5.",
                    "",
                    "MODELING WITH RATIONAL EXPRESSIONS:",
                    "  — Name the ratio in words first.",
                    "  — Build each quantity as a polynomial, form the fraction, then simplify.",
                    "  — Interpret the simplified ratio: what do the numerator and denominator mean?",
                    "  — State domain restriction in context (e.g., “x > 0 because length is positive”).",
                ],
                color_hex="DEEBF7")

    ps.summary_exit_box(doc, "EXIT TICKET — Pre-Assessment Confidence Check",
                        [
                            "A rational model needs factoring because ___.",
                            "My simplified ratio is ___, valid when x ≠ ___.",
                            "In context this means ___.",
                            "One biggest thing I learned today: ________________________",
                        ])

    doc.add_page_break()
    ps.teal_section(doc, "OPTIONAL REINFORCEMENT (not collected)",
                    ["Extra stretch for fast finishers. Try without a calculator first."])
    reinforce = qb.get_for_packet(REINFORCE_IDS)
    for i, q in enumerate(reinforce, 1):
        bank_item_block(doc, q, label=f"Optional #{i}  (Savvas SAT/ACT style)",
                        space_after_inches=1.8)

    doc.save(path)


# ── Teacher packet ───────────────────────────────────────────────────────

def build_teacher(path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    add_header_block(doc, for_teacher=True)

    add_callout(doc, "\U0001f4cc  PRE-FLIGHT — P3 IS ASSESSMENT-CRITICAL",
                [
                    "Topic 4 8-Q assessment Q#2 (identify domain restrictions from a cubic denominator) is directly rehearsed by Practice #19.",
                    "Ex 2 (simplify with domain restriction) and Ex 6 (SA/V modeling) are paired in Launch — students see both abstract and applied forms back-to-back.",
                    "No DOK-3 driver today (P2 owned the spine). P3 is pure DOK-2 mastery.",
                    "Pace: Try-It 6 (5) → #14 (5) → #22 (6) → #34 (6) → #37 (5) → #19 (8). Total ~35 min.",
                    "If pace slows: cut #37 (modeling) first. Never cut #19.",
                ],
                color_hex="FFD9E0")

    ps.framework_phase_header(
        doc,
        phase="Do Now",
        framework_tag="bridge from P2",
        dok="2",
        minutes=5,
        teacher_does=["Hand out at door.", "Circulate 3 min.",
                      "Cold-call for domain restriction reasoning."],
        students_do=["Explain the identity over its domain.", "Write restriction in set notation."],
        questions_to_ask=[
            "WHY is the domain restricted there?",
            "Based on yesterday: predict what SA/V simplifies to for a cube.",
        ],
        adult_role="Solo teacher: circulate silently.",
    )
    q = qb.get_for_packet([DO_NOW_ID])[0]
    bank_item_block(doc, q, label=f"Do Now ({q['id']})", space_after_inches=0.3)
    add_callout(doc, "✅  ANSWER KEY",
                [q.get("notes") or "(verify)"],
                color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Launch",
        framework_tag="teacher models Ex 2 + Ex 6 back-to-back",
        dok="2",
        minutes=12,
        teacher_does=[
            "FIRST 6 min: Example 2 — simplify (4−x²)/(x²+3x−10). Factor both, cancel (x−2), write domain: x ≠ 2, x ≠ −5. Emphasize: canceled factor still restricts.",
            "NEXT 5 min: Example 6 — SA/V ratio for prism vs. cylinder. Show ratio setup → simplify → both equal 3/x. Interpret: “x is the side length; ratio shrinks as x grows.”",
            "30 sec: point at packet rule cards. \"Mark them. These are your assessment reference.\"",
            "Do NOT solve Try-Its. Release at minute 17.",
        ],
        students_do=[
            "Watch Ex 2. Copy each factoring step. Note domain restriction next to canceled factor.",
            "Watch Ex 6. Write the ratio in words before algebra: SA ÷ Volume.",
            "Copy BOTH rule cards into margins.",
        ],
        questions_to_ask=[
            "Why does the canceled (x−2) factor STILL appear in the domain restriction?",
            "What does the domain restriction x ≠ 0 mean in the SA/V context?",
            "For Ex 6: if x doubles, what happens to the SA/V ratio? (It halves.)",
            "How is modeling with rational expressions different from just simplifying?",
        ],
        adult_role="Project both examples. Think aloud on domain restriction reasoning.",
    )
    launch_items = qb.get_for_packet(LAUNCH_IDS)
    launch_labels = [
        "Example 2 — Simplify (4−x²)/(x²+3x−10)",
        "Example 6 — SA/V Modeling (prism vs. cylinder)",
    ]
    for label, lq in zip(launch_labels, launch_items):
        bank_item_block(doc, lq, label=f"{label} ({lq['id']})", space_after_inches=0.3)
        add_callout(doc, "✅  ANSWER KEY",
                    [lq.get("notes") or "(no answer key — verify before class)"],
                    color_hex="D9EAD3")

    add_callout(doc, "\U0001f3a4  TEACHER SCRIPT",
                [
                    "1. “Yesterday we simplified with one factoring move. Today we do multi-step AND interpret in context.”",
                    "2. Ex 2: “Factor top: difference of squares → (2−x)(2+x). Factor bottom: (x−2)(x+5). Now — watch: (2−x) = −(x−2). Cancel. Domain: x ≠ 2, x ≠ −5. The 2 came from the canceled factor — it STILL blocks.”",
                    "3. “Ex 6: what IS efficiency? It’s SA divided by Volume. Build each one, simplify, both land on 3/x. Interpretation: as x grows, the ratio 3/x shrinks — bigger packages are more efficient.”",
                    "4. “Same procedure, new meaning. The math doesn’t change — just what we say about the answer.”",
                    "5. Release.",
                ],
                color_hex="CFE2F3")

    ps.framework_phase_header(
        doc,
        phase="Explore",
        framework_tag="TPS · pure DOK-2 mastery · assessment rehearsal",
        dok="2",
        minutes=35,
        teacher_does=[
            "6 laps across 35 min.",
            "Laps 1-3: Try-It 6 + #14 + #22. Abstract simplification with domain. (5-6 min each.)",
            "Laps 4-5: #34 + #37 — modeling items. Build ratio → simplify → interpret.",
            "Lap 6: #19 — the assessment rehearsal. Press on domain restriction from cubic denominator.",
            "Do not give answers. Use sentence frame to press interpretation.",
        ],
        students_do=[
            "6 items in order. Abstract forms first (Try-It 6, #14, #22), then modeling (#34, #37), then assessment rehearsal (#19).",
            "For #19: BEFORE simplifying, state all denominator zeros. Then simplify. Then write restriction.",
            "Justify with sentence frame.",
        ],
        questions_to_ask=[
            "What are ALL the denominators across every step? (Not just the final one.)",
            "For modeling items: what does the numerator represent? The denominator?",
            "For #14: did you catch both restrictions? (Factor completely first.)",
            "For #19: what’s special about a cubic denominator? (May have up to 3 zeros.)",
            "For #37: interpret your answer — what does x ≠ ___ mean in context?",
        ],
        adult_role="Solo teacher: 6 laps. Press domain restriction reasoning on every item. #19 IS the assessment rehearsal.",
    )
    items = qb.get_for_packet(EXPLORE_IDS)
    labels = [
        "Try It 6", "Practice #14", "Practice #22", "Practice #34", "Practice #37",
        "Practice #19 ⭐ ASSESSMENT-CRITICAL",
    ]
    for label, qi in zip(labels, items):
        bank_item_block(doc, qi, label=f"{label} ({qi['id']})", space_after_inches=0.3)
        add_callout(doc, "✅  ANSWER / ANTICIPATED TALK",
                    [qi.get("notes") or "(no answer key — verify before class)"],
                    color_hex="D9EAD3")

    ps.framework_phase_header(
        doc,
        phase="Share / Summary",
        framework_tag="Concept Summary + LEHS 8-Q preview",
        dok="2",
        minutes=5,
        teacher_does=[
            "Cold-call 1 pair to present #19 (domain restrictions from cubic denominator).",
            "Project Concept Summary (from student packet).",
            "\"Topic 4 8-Q assessment: expect Q#2 on THIS content — domain ID from a cubic. You just rehearsed it.\"",
        ],
        students_do=[
            "Presenting pair uses sentence frame for #19.",
            "Rest of class signals and annotates their Concept Summary.",
        ],
        questions_to_ask=[
            "What are the TWO big ideas from Lesson 4-3?",
            "If you see a cubic denominator on the assessment, what’s the FIRST thing you do?",
        ],
        adult_role="Frame today as assessment rehearsal. Students leave confident on domain restrictions.",
    )

    ps.framework_phase_header(
        doc,
        phase="Exit Ticket",
        framework_tag="pre-assessment confidence check",
        dok="1-2",
        minutes=3,
        teacher_does=[
            "Collect at door.",
            "Tonight: scan for students who can’t complete the simplified ratio or restriction. Insert a 5-min Do Now recap when the 8-Q lands.",
        ],
        students_do=[
            "Complete the three fill-in stems.",
            "Write one biggest thing learned.",
        ],
        questions_to_ask=[
            "Did any student leave the restriction blank?",
            "Did any student interpret correctly but miss the domain?",
        ],
        adult_role="Collect. No grade. Use as data for assessment prep.",
    )
    add_callout(doc, "\U0001f4cb  EXIT TICKET — Student Form",
                [
                    "A rational model needs factoring because ___.",
                    "My simplified ratio is ___, valid when x ≠ ___.",
                    "In context this means ___.",
                    "One biggest thing I learned today: ________________________",
                ],
                color_hex="FFF2CC")

    doc.add_page_break()
    add_callout(doc, "\U0001f550  PERIOD A / PERIOD F — 65-MIN CLASS NOTES",
                [
                    "Both sections should have 65 min. Use extra 10 min on Explore #19 (assessment-critical).",
                    "If finished early: Practice #39 (SAT/ACT MC) as fast-finisher stretch.",
                    "Alternative: project Practice #22 as a second domain-identification prompt — more restriction practice.",
                ],
                color_hex="FFD9E0")

    add_callout(doc, "\U0001f9e9  MODIFICATIONS / SUPPORT FOR IEP STUDENTS",
                [
                    "• Provide both Rule cards as half-sheet cut-outs on their desks.",
                    "• For #19, provide a pre-factored scaffold showing the denominator already factored so they focus on identifying zeros.",
                    "• Accept verbal domain restriction identification in lieu of written.",
                ],
                color_hex="E2F0D9")
    add_callout(doc, "\U0001f30d  SUPPORTS FOR ELL STUDENTS",
                [
                    "• Sentence frames on student page (don’t skip).",
                    "• Word cards: “domain” = allowed x-values; “restriction” = x-values blocked by division by zero.",
                    "• “simplify” = reduce to lowest terms by canceling common factors.",
                    "• “ratio” = one quantity divided by another (e.g., SA ÷ Volume).",
                    "• Pair ELL students with English-dominant partners for TPS.",
                ],
                color_hex="DEEBF7")

    doc.save(path)


# ── Main ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_do_now("L43_P3_Do_Now.docx")
    build_student("L43_P3_Student_Packet.docx")
    build_teacher("L43_P3_Teacher_Packet.docx")
    print("Built L43_P3_Do_Now.docx, L43_P3_Student_Packet.docx, L43_P3_Teacher_Packet.docx")
    ps.emit_visuals_checklist(_ALL_IDS, "L43_P3_Visuals_Checklist.md",
                              title="L43 P3 — Visuals Checklist")
